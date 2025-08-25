import sys
import os
import pandas as pd
import numpy as np
from datetime import datetime, timedelta
import json
import smtplib
from email.mime.text import MIMEText
from email.mime.multipart import MIMEMultipart
import requests
import zipfile
import shutil
import traceback
import logging
import time
from PyQt5.QtWidgets import *
from PyQt5.QtCore import *
from PyQt5.QtGui import *
from PyQt5.QtPrintSupport import *
import matplotlib.pyplot as plt
import matplotlib
matplotlib.use('Qt5Agg')
from matplotlib.backends.backend_qt5agg import FigureCanvasQTAgg as FigureCanvas
from matplotlib.figure import Figure
import seaborn as sns
import warnings
warnings.filterwarnings('ignore')
import qtawesome as qta
import openpyxl
from openpyxl.styles import Font


# QTableView modelini Excel'e aktaran fonksiyon

from PyQt5.QtWidgets import QProgressDialog

def export_table_to_excel(model, proxy_model, parent):
    try:
        file_path, _ = QFileDialog.getSaveFileName(parent, "Excel Olarak Kaydet", "", "Excel Dosyası (*.xlsx)")
        if not file_path:
            return
        rows = proxy_model.rowCount()
        cols = proxy_model.columnCount()
        data = []
        headers = []
        for col in range(cols):
            headers.append(str(model.headerData(col, Qt.Orientation.Horizontal)))
        # Yükleniyor penceresi
        progress = QProgressDialog("Veriler Excel'e aktarılıyor...", None, 0, rows, parent)
        progress.setWindowTitle("Lütfen Bekleyin")
        progress.setWindowModality(Qt.WindowModal)
        progress.setMinimumDuration(0)
        for row in range(rows):
            row_data = []
            for col in range(cols):
                index = proxy_model.index(row, col)
                row_data.append(proxy_model.data(index))
            data.append(row_data)
            progress.setValue(row)
            if progress.wasCanceled():
                QMessageBox.warning(parent, "İptal Edildi", "Excel aktarımı kullanıcı tarafından iptal edildi.")
                return
        progress.setValue(rows)
        df = pd.DataFrame(data, columns=headers)
        df.to_excel(file_path, index=False)
        QMessageBox.information(parent, "Başarılı", "Excel dosyası başarıyla kaydedildi.")
    except Exception as e:
        logger.error(f"Excel'e aktarma hatası: {str(e)}", exc=e)
        QMessageBox.critical(parent, "Hata", "Excel'e aktarırken bir hata oluştu. Lütfen dosyanın açık olmadığından ve yeterli disk alanı olduğundan emin olun. Teknik detaylar log dosyasına kaydedildi.")

from PyQt5.QtWidgets import (
    QApplication, QMainWindow, QWidget, QHBoxLayout, QVBoxLayout, QPushButton,
    QLabel, QListWidget, QListWidgetItem, QStackedWidget, QTableWidget, QTableWidgetItem,
    QDialog, QTextEdit, QLineEdit, QMessageBox, QFormLayout, QFileDialog, QComboBox, QDateEdit,
    QMenuBar, QMenu, QAction, QGridLayout, QFrame, QGroupBox, QScrollArea, QTabWidget,
    QHeaderView, QAbstractItemView, QSplitter, QToolBar, QStatusBar, QProgressBar,
    QTableView, QCheckBox
)
from PyQt5.QtGui import QIcon, QFont, QPixmap, QPalette, QColor, QStandardItemModel, QStandardItem, QKeySequence
from PyQt5.QtCore import QSize, QDate, Qt, QTimer, pyqtSignal, QThread, pyqtSignal, QSortFilterProxyModel
from PyQt5.QtWidgets import QShortcut
import pandas as pd
import json
import shutil

def save_vehicle_photo(selected_photo_path, plaka):
    target_dir = os.path.join("veri", "arac_fotograflari")
    os.makedirs(target_dir, exist_ok=True)
    ext = os.path.splitext(selected_photo_path)[1].lower()
    target_path = os.path.join(target_dir, f"{plaka}{ext}")
    # Eğer kaynak ve hedef aynıysa tekrar kopyalama!
    if os.path.abspath(selected_photo_path) == os.path.abspath(target_path):
        return target_path
    shutil.copyfile(selected_photo_path, target_path)
    return target_path  # Bunu veri tabanında/foto_path olarak saklayın
# =============================================================================
# KONFİGÜRASYON VE AYARLAR
# =============================================================================

class Config:
    """Uygulama konfigürasyonu"""
    
    # Uygulama bilgileri
    APP_NAME = "Araç Filo Yönetim Sistemi"
    VERSION = "22.07.24.01"
    DEVELOPER = "Ertuğrul Kamil ŞAHİN"
    EMAIL = "ertugrul.yazilim@gmail.com"
    
    # Güncelleme ayarları
    UPDATE_CHECK_URL = "https://api.github.com/repos/ertugrul-yazilim/arac.filo/contents/version.txt"
    GITHUB_REPO_URL = "https://github.com/ertugrul-yazilim/arac.filo"
    UPDATE_CHECK_INTERVAL = 24 * 60 * 60  # 24 saat (saniye cinsinden)
    AUTO_UPDATE_CHECK = True  # Otomatik güncelleme kontrolü
    SHOW_BETA_UPDATES = False  # Beta sürümleri göster
    
    # Mail ayarları (Geri bildirim için)
    MAIL_SMTP_SERVER = "smtp.gmail.com"
    MAIL_SMTP_PORT = 587
    MAIL_USERNAME = "ertugrul.yazilim@gmail.com"
    MAIL_PASSWORD = "kmrp noyx hxlc ntcy"  # Gmail uygulama şifresi
    MAIL_FROM_NAME = "Araç Filo Yönetim Sistemi"
    

    
    # Dosya yolları
    DATA_DIR = "veri"
    BACKUP_DIR = "veri/yedekler"
    LOG_DIR = "veri/loglar"
    
    # Excel dosya isimleri
    EXCEL_FILES = {
        'araclar': 'araclar.xlsx',
        'giderler': 'giderler.xlsx',
        'suruculer': 'suruculer.xlsx', 
        'bakimlar': 'bakimlar.xlsx',
        'yakitlar': 'yakitlar.xlsx',
        'cezalar': 'cezalar.xlsx',
        'kazalar': 'kazalar.xlsx',
        'belgeler': 'belgeler.xlsx',
        'hatirlatmalar': 'hatirlatmalar.xlsx',
        'silinen_araclar': 'silinen_araclar.xlsx',
        'silinen_suruculer': 'silinen_suruculer.xlsx',
        'silinen_bakimlar': 'silinen_bakimlar.xlsx',
        'silinen_giderler': 'silinen_giderler.xlsx',
        'silinen_cezalar': 'silinen_cezalar.xlsx',
        'silinen_kazalar': 'silinen_kazalar.xlsx'
    }
    
    # Menü yapısı - Modern ikonlarla
    MENU_ITEMS = [
        # Ana Yönetim
        {"id": "dashboard", "title": "Ana Sayfa", "icon": "fa.home", "emoji_icon": "🏠", "group": "main", "shortcut": "Ctrl+1"},
        
        # Araç Yönetimi
        {"id": "araclar", "title": "Araçlar", "icon": "fa.car", "emoji_icon": "🚗", "group": "vehicles", "shortcut": "Ctrl+2"},
        {"id": "suruculer", "title": "Sürücüler", "icon": "fa.user-tie", "emoji_icon": "👨‍✈️", "group": "vehicles", "shortcut": "Ctrl+3"},
        
        # Operasyonel İşlemler
        {"id": "bakimlar", "title": "Bakım & Onarımlar", "icon": "fa.wrench", "emoji_icon": "🔧", "group": "operations", "shortcut": "Ctrl+4"},
        {"id": "yakitlar", "title": "Yakıtlar", "icon": "fa.gas-pump", "emoji_icon": "⛽", "group": "operations", "shortcut": "Ctrl+5"},
        {"id": "trafik", "title": "Trafik Cezaları & Kazalar", "icon": "fa.exclamation-triangle", "emoji_icon": "🚨", "group": "operations", "shortcut": "Ctrl+6"},
        
        # Raporlama & Planlama
        {"id": "raporlar", "title": "Raporlar", "icon": "fa.chart-bar", "emoji_icon": "📊", "group": "reports", "shortcut": "Ctrl+7"},
        
        # Sistem
        {"id": "belgeler", "title": "Belgeler", "icon": "fa.folder", "emoji_icon": "📁", "group": "system", "shortcut": "Ctrl+8"},
        {"id": "ayarlar", "title": "Ayarlar", "icon": "fa.cog", "emoji_icon": "⚙️", "group": "system", "shortcut": "Ctrl+9"}
    ]
    
    # Menü grupları
    MENU_GROUPS = {
        "main": {"title": "Ana Yönetim", "color": "#3498db"},
        "vehicles": {"title": "Araç Yönetimi", "color": "#e74c3c"},
        "operations": {"title": "Operasyonel İşlemler", "color": "#f39c12"},
        "reports": {"title": "Raporlama & Planlama", "color": "#27ae60"},
        "system": {"title": "Sistem", "color": "#9b59b6"}
    }

class SimpleLogger:
    def __init__(self):
        self.log_file = os.path.join(Config.LOG_DIR, "app.log")
        self._ensure_log_dir()
    def _ensure_log_dir(self):
        os.makedirs(Config.LOG_DIR, exist_ok=True)
    def _write_log(self, level, msg):
        timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        log_entry = f"[{timestamp}] [{level}] {msg}\n"
        print(log_entry.strip())
        try:
            with open(self.log_file, 'a', encoding='utf-8') as f:
                f.write(log_entry)
        except Exception as e:
            print(f"Log dosyası yazma hatası: {e}")
    def info(self, msg): 
        self._write_log("INFO", msg)
    def error(self, msg, exc=None): 
        if exc:
            msg += f" - Exception: {str(exc)}"
        self._write_log("ERROR", msg)
    def warning(self, msg): 
        self._write_log("WARNING", msg)
    def debug(self, msg): 
        self._write_log("DEBUG", msg)
    def critical(self, msg, exc=None): 
        if exc:
            msg += f" - Exception: {str(exc)}"
        self._write_log("CRITICAL", msg)
    def log_system_event(self, event, details=""): 
        self._write_log("SYSTEM", f"{event}: {details}")
    def log_user_action(self, action, details=""): 
        self._write_log("USER", f"{action}: {details}")

logger = SimpleLogger()

# =============================================================================
# GELİŞMİŞ GÜNCELLEME SİSTEMİ
# =============================================================================

class UpdateChecker(QThread):
    """Gelişmiş güncelleme kontrol sistemi"""
    
    update_available = pyqtSignal(dict)  # update_info dict
    no_update = pyqtSignal()
    error_occurred = pyqtSignal(str)
    check_progress = pyqtSignal(str)  # İlerleme durumu
    
    def __init__(self):
        super().__init__()
        self.last_check_time = None
        self.check_interval = 24 * 60 * 60  # 24 saat (saniye cinsinden)
    
    def run(self):
        """Güncelleme kontrolü yap"""
        try:
            self.check_progress.emit("GitHub sunucusuna bağlanılıyor...")
            
            # GitHub API'den en son release bilgilerini al
            response = requests.get(Config.UPDATE_CHECK_URL, timeout=15)
            
            if response.status_code == 200:
                self.check_progress.emit("Sürüm bilgileri alınıyor...")
                data = response.json()
                
                # Detaylı sürüm bilgileri
                latest_version = data.get('tag_name', '').lstrip('v')
                release_name = data.get('name', '')
                release_body = data.get('body', '')
                download_url = data.get('html_url', '')
                published_at = data.get('published_at', '')
                prerelease = data.get('prerelease', False)
                
                # Asset'leri kontrol et (exe dosyası var mı?)
                assets_url = data.get('assets_url', '')
                if assets_url:
                    assets_response = requests.get(assets_url, timeout=10)
                    if assets_response.status_code == 200:
                        assets = assets_response.json()
                        exe_assets = [asset for asset in assets if asset.get('name', '').endswith('.exe')]
                        if exe_assets:
                            download_url = exe_assets[0].get('browser_download_url', download_url)
                
                self.check_progress.emit("Sürüm karşılaştırması yapılıyor...")
                
                if self.compare_versions(latest_version, Config.VERSION):
                    # Güncelleme mevcut
                    update_info = {
                        'version': latest_version,
                        'name': release_name,
                        'body': release_body,
                        'download_url': download_url,
                        'published_at': published_at,
                        'prerelease': prerelease,
                        'current_version': Config.VERSION
                    }
                    self.update_available.emit(update_info)
                else:
                    self.no_update.emit()
                    
            elif response.status_code == 404:
                self.error_occurred.emit("GitHub repository bulunamadı. Lütfen repository URL'sini kontrol edin.")
            elif response.status_code == 403:
                self.error_occurred.emit("GitHub API limiti aşıldı. Lütfen daha sonra tekrar deneyin.")
            else:
                self.error_occurred.emit(f"GitHub sunucusuna ulaşılamadı (HTTP {response.status_code})")
                
        except requests.exceptions.Timeout:
            self.error_occurred.emit("Bağlantı zaman aşımına uğradı. İnternet bağlantınızı kontrol edin.")
        except requests.exceptions.ConnectionError:
            self.error_occurred.emit("İnternet bağlantısı yok. Lütfen bağlantınızı kontrol edin.")
        except requests.exceptions.RequestException as e:
            self.error_occurred.emit(f"Güncelleme sunucusuna ulaşılamadı: {str(e)}")
        except Exception as e:
            self.error_occurred.emit(f"Güncelleme kontrolü hatası: {str(e)}")
    
    def compare_versions(self, latest, current):
        """Gelişmiş versiyon karşılaştırması"""
        try:
            # Versiyon formatını temizle
            latest = latest.strip().lower()
            current = current.strip().lower()
            
            # Alpha, beta, rc gibi özel sürümleri işle
            latest_parts = self._parse_version(latest)
            current_parts = self._parse_version(current)
            
            # Ana versiyon numaralarını karşılaştır
            for i in range(max(len(latest_parts['numbers']), len(current_parts['numbers']))):
                latest_num = latest_parts['numbers'][i] if i < len(latest_parts['numbers']) else 0
                current_num = current_parts['numbers'][i] if i < len(current_parts['numbers']) else 0
                
                if latest_num > current_num:
                    return True
                elif latest_num < current_num:
                    return False
            
            # Ana versiyonlar eşitse, özel sürüm bilgilerini kontrol et
            if latest_parts['suffix'] and not current_parts['suffix']:
                return True  # Mevcut sürüm stable, yeni sürüm pre-release
            elif not latest_parts['suffix'] and current_parts['suffix']:
                return True  # Yeni sürüm stable, mevcut sürüm pre-release
            
            return False
            
        except Exception as e:
            logger.error(f"Versiyon karşılaştırma hatası: {str(e)}")
            return False
    
    def _parse_version(self, version_str):
        """Versiyon string'ini parse et"""
        import re
        
        # Versiyon numaralarını ayır
        numbers = []
        suffix = ""
        
        # Sayısal kısımları bul
        number_matches = re.findall(r'\d+', version_str)
        numbers = [int(x) for x in number_matches]
        
        # Özel sürüm bilgilerini bul (alpha, beta, rc, vb.)
        suffix_match = re.search(r'[a-zA-Z]+', version_str)
        if suffix_match:
            suffix = suffix_match.group().lower()
        
        return {
            'numbers': numbers,
            'suffix': suffix
        }
    
    def should_check_for_updates(self):
        """Güncelleme kontrolü yapılmalı mı?"""
        if not self.last_check_time:
            return True
        
        time_diff = time.time() - self.last_check_time
        return time_diff >= self.check_interval
    
    def mark_check_completed(self):
        """Kontrol tamamlandı olarak işaretle"""
        self.last_check_time = time.time()

class VersionInfo:
    """Versiyon bilgileri sınıfı"""
    
    def __init__(self, version, name="", body="", download_url="", published_at="", prerelease=False):
        self.version = version
        self.name = name
        self.body = body
        self.download_url = download_url
        self.published_at = published_at
        self.prerelease = prerelease
    
    def get_formatted_date(self):
        """Tarihi formatla"""
        try:
            if self.published_at:
                dt = datetime.fromisoformat(self.published_at.replace('Z', '+00:00'))
                return dt.strftime("%d.%m.%Y %H:%M")
            return ""
        except:
            return ""
    
    def get_changelog_preview(self, max_lines=5):
        """Değişiklik notlarının önizlemesini al"""
        if not self.body:
            return "Değişiklik notu bulunamadı."
        
        lines = self.body.split('\n')
        preview_lines = []
        
        for line in lines[:max_lines]:
            line = line.strip()
            if line:
                # Markdown formatını temizle
                line = line.lstrip('#').lstrip('*').lstrip('-').lstrip()
                preview_lines.append(line)
        
        if len(lines) > max_lines:
            preview_lines.append("...")
        
        return '\n'.join(preview_lines)

class UpdateDialog(QDialog):
    """Gelişmiş güncelleme dialog"""
    
    def __init__(self, parent=None, update_info=None):
        super().__init__(parent)
        self.update_info = update_info or {}
        self.init_ui()
    
    def init_ui(self):
        self.setWindowTitle("🔄 Güncelleme Bulundu")
        self.setMinimumSize(600, 500)
        self.setModal(True)
        
        layout = QVBoxLayout(self)
        layout.setSpacing(15)
        layout.setContentsMargins(20, 20, 20, 20)
        
        # Başlık
        title_label = QLabel("🎉 Yeni Sürüm Mevcut!")
        title_label.setStyleSheet("""
            font-size: 18px;
            font-weight: bold;
            color: #2c3e50;
            margin-bottom: 10px;
        """)
        title_label.setAlignment(Qt.AlignmentFlag.AlignCenter)
        layout.addWidget(title_label)
        
        # Sürüm bilgileri
        self.create_version_info_section(layout)
        
        # Değişiklik notları
        self.create_changelog_section(layout)
        
        # Butonlar
        self.create_buttons_section(layout)
    
    def create_version_info_section(self, parent_layout):
        """Sürüm bilgileri bölümü"""
        info_frame = QFrame()
        info_frame.setStyleSheet("""
            QFrame {
                background-color: #f8f9fa;
                border: 1px solid #dee2e6;
                border-radius: 8px;
                padding: 15px;
            }
        """)
        
        info_layout = QVBoxLayout(info_frame)
        info_layout.setSpacing(10)
        
        # Sürüm karşılaştırması
        version_layout = QHBoxLayout()
        
        # Mevcut sürüm
        current_version_frame = QFrame()
        current_version_frame.setStyleSheet("""
            QFrame {
                background-color: #e9ecef;
                border-radius: 6px;
                padding: 10px;
            }
        """)
        current_layout = QVBoxLayout(current_version_frame)
        
        current_label = QLabel("Mevcut Sürüm")
        current_label.setStyleSheet("font-size: 11px; color: #6c757d; font-weight: bold;")
        current_version = QLabel(f"v{self.update_info.get('current_version', Config.VERSION)}")
        current_version.setStyleSheet("font-size: 16px; font-weight: bold; color: #495057;")
        
        current_layout.addWidget(current_label)
        current_layout.addWidget(current_version)
        
        # Ok işareti
        arrow_label = QLabel("→")
        arrow_label.setStyleSheet("font-size: 20px; color: #6c757d; margin: 0 15px;")
        arrow_label.setAlignment(Qt.AlignmentFlag.AlignCenter)
        
        # Yeni sürüm
        new_version_frame = QFrame()
        new_version_frame.setStyleSheet("""
            QFrame {
                background-color: #d4edda;
                border-radius: 6px;
                padding: 10px;
            }
        """)
        new_layout = QVBoxLayout(new_version_frame)
        
        new_label = QLabel("Yeni Sürüm")
        new_label.setStyleSheet("font-size: 11px; color: #155724; font-weight: bold;")
        new_version = QLabel(f"v{self.update_info.get('version', '')}")
        new_version.setStyleSheet("font-size: 16px; font-weight: bold; color: #155724;")
        
        new_layout.addWidget(new_label)
        new_layout.addWidget(new_version)
        
        version_layout.addWidget(current_version_frame)
        version_layout.addWidget(arrow_label)
        version_layout.addWidget(new_version_frame)
        version_layout.addStretch()
        
        info_layout.addLayout(version_layout)
        
        # Ek bilgiler
        details_layout = QHBoxLayout()
        
        # Yayın tarihi
        if self.update_info.get('published_at'):
            date_label = QLabel(f"📅 Yayın Tarihi: {self.get_formatted_date()}")
            date_label.setStyleSheet("font-size: 12px; color: #6c757d;")
            details_layout.addWidget(date_label)
        
        # Pre-release uyarısı
        if self.update_info.get('prerelease', False):
            prerelease_label = QLabel("⚠️ Pre-release Sürüm")
            prerelease_label.setStyleSheet("font-size: 12px; color: #856404; background-color: #fff3cd; padding: 2px 6px; border-radius: 3px;")
            details_layout.addWidget(prerelease_label)
        
        details_layout.addStretch()
        info_layout.addLayout(details_layout)
        
        parent_layout.addWidget(info_frame)
    
    def create_changelog_section(self, parent_layout):
        """Değişiklik notları bölümü"""
        changelog_frame = QFrame()
        changelog_frame.setStyleSheet("""
            QFrame {
                background-color: white;
                border: 1px solid #dee2e6;
                border-radius: 8px;
                padding: 15px;
            }
        """)
        
        changelog_layout = QVBoxLayout(changelog_frame)
        
        # Başlık
        changelog_title = QLabel("📝 Değişiklik Notları")
        changelog_title.setStyleSheet("font-size: 14px; font-weight: bold; color: #2c3e50; margin-bottom: 10px;")
        changelog_layout.addWidget(changelog_title)
        
        # Değişiklik notları
        changelog_text = QTextEdit()
        changelog_text.setReadOnly(True)
        changelog_text.setMaximumHeight(200)
        changelog_text.setStyleSheet("""
            QTextEdit {
                background-color: #f8f9fa;
                border: 1px solid #dee2e6;
                border-radius: 4px;
                padding: 10px;
                font-size: 12px;
                line-height: 1.4;
            }
        """)
        
        changelog_content = self.update_info.get('body', 'Değişiklik notu bulunamadı.')
        if changelog_content:
            # Markdown formatını temizle
            lines = changelog_content.split('\n')
            cleaned_lines = []
            for line in lines:
                line = line.strip()
                if line:
                    # Markdown başlıklarını temizle
                    line = line.lstrip('#').lstrip('*').lstrip('-').lstrip()
                    cleaned_lines.append(line)
            
            changelog_text.setPlainText('\n'.join(cleaned_lines))
        else:
            changelog_text.setPlainText("Bu sürüm için değişiklik notu bulunamadı.")
        
        changelog_layout.addWidget(changelog_text)
        parent_layout.addWidget(changelog_frame)
    
    def create_buttons_section(self, parent_layout):
        """Butonlar bölümü"""
        btn_layout = QHBoxLayout()
        btn_layout.setSpacing(10)
        
        # İndir butonu
        btn_download = QPushButton("⬇️ Güncellemeyi İndir")
        btn_download.setStyleSheet("""
            QPushButton {
                background-color: #28a745;
                color: white;
                border: none;
                padding: 12px 24px;
                border-radius: 6px;
                font-size: 14px;
                font-weight: bold;
            }
            QPushButton:hover {
                background-color: #218838;
            }
            QPushButton:pressed {
                background-color: #1e7e34;
            }
        """)
        btn_download.clicked.connect(self.download_update)
        
        # Daha sonra butonu
        btn_later = QPushButton("⏰ Daha Sonra")
        btn_later.setStyleSheet("""
            QPushButton {
                background-color: #6c757d;
                color: white;
                border: none;
                padding: 12px 24px;
                border-radius: 6px;
                font-size: 14px;
                font-weight: bold;
            }
            QPushButton:hover {
                background-color: #5a6268;
            }
        """)
        btn_later.clicked.connect(self.remind_later)
        
        # İptal butonu
        btn_cancel = QPushButton("❌ İptal")
        btn_cancel.setStyleSheet("""
            QPushButton {
                background-color: #dc3545;
                color: white;
                border: none;
                padding: 12px 24px;
                border-radius: 6px;
                font-size: 14px;
                font-weight: bold;
            }
            QPushButton:hover {
                background-color: #c82333;
            }
        """)
        btn_cancel.clicked.connect(self.reject)
        
        btn_layout.addWidget(btn_download)
        btn_layout.addWidget(btn_later)
        btn_layout.addWidget(btn_cancel)
        
        parent_layout.addLayout(btn_layout)
    
    def get_formatted_date(self):
        """Tarihi formatla"""
        try:
            published_at = self.update_info.get('published_at', '')
            if published_at:
                dt = datetime.fromisoformat(published_at.replace('Z', '+00:00'))
                return dt.strftime("%d.%m.%Y %H:%M")
            return ""
        except:
            return ""
    
    def download_update(self):
        """Güncellemeyi indir"""
        try:
            download_url = self.update_info.get('download_url', '')
            if not download_url:
                QMessageBox.warning(self, "Uyarı", "İndirme linki bulunamadı.")
                return
            
            import webbrowser
            webbrowser.open(download_url)
            
            QMessageBox.information(self, "✅ İndirme Başlatıldı", 
                "Tarayıcıda indirme sayfası açıldı.\n\n"
                "📋 Kurulum Adımları:\n"
                "1. Dosyayı indirin\n"
                "2. Mevcut uygulamayı kapatın\n"
                "3. İndirilen dosyayı çalıştırın\n"
                "4. Kurulumu tamamlayın\n\n"
                "⚠️ Önemli: Kurulum öncesi verilerinizi yedekleyin!")
            
            self.accept()
            
        except Exception as e:
            QMessageBox.critical(self, "❌ İndirme Hatası", f"İndirme başlatılamadı:\n{str(e)}")
    
    def remind_later(self):
        """Daha sonra hatırlat"""
        QMessageBox.information(self, "⏰ Hatırlatma", 
            "Güncelleme hatırlatması 24 saat sonra tekrar gösterilecek.")
        self.reject()

# =============================================================================
# GERİ BİLDİRİM SİSTEMİ
# =============================================================================



class FeedbackDialog(QDialog):
    """Geri bildirim dialog"""
    
    def __init__(self, parent=None):
        super().__init__(parent)
        self.init_ui()
    
    def init_ui(self):
        self.setWindowTitle("Geri Bildirim Gönder")
        self.setMinimumSize(500, 400)
        self.setModal(True)
        
        layout = QVBoxLayout(self)
        
        # Başlık
        title = QLabel("Geri Bildirim")
        title.setStyleSheet("font-size: 18px; font-weight: bold; color: #2c3e50; margin-bottom: 10px;")
        title.setAlignment(Qt.AlignmentFlag.AlignCenter)
        layout.addWidget(title)
        
        # Form
        form_layout = QFormLayout()
        
        self.name = QLineEdit()
        self.name.setPlaceholderText("Adınız (opsiyonel)")
        
        self.email = QLineEdit()
        self.email.setPlaceholderText("E-posta adresiniz (opsiyonel)")
        
        self.category = QComboBox()
        self.category.addItems([
            "Öneri", "Hata Bildirimi", "Yeni Özellik", "Genel", "Diğer"
        ])
        
        self.subject = QLineEdit()
        self.subject.setPlaceholderText("Konu başlığı")
        
        self.message = QTextEdit()
        self.message.setPlaceholderText("Mesajınızı buraya yazın...")
        self.message.setMinimumHeight(150)
        
        form_layout.addRow("Ad:", self.name)
        form_layout.addRow("E-posta:", self.email)
        form_layout.addRow("Kategori:", self.category)
        form_layout.addRow("Konu:", self.subject)
        form_layout.addRow("Mesaj:", self.message)
        
        layout.addLayout(form_layout)
        
        # Butonlar
        btn_layout = QHBoxLayout()
        
        btn_send = QPushButton("📧 Gönder")
        btn_send.setStyleSheet("""
            QPushButton {
                background-color: #3498db;
                color: white;
                border: none;
                padding: 10px 20px;
                border-radius: 5px;
                font-weight: bold;
            }
            QPushButton:hover {
                background-color: #2980b9;
            }
        """)
        btn_send.clicked.connect(self.send_feedback)
        
        btn_cancel = QPushButton("İptal")
        btn_cancel.setStyleSheet("""
            QPushButton {
                background-color: #95a5a6;
                color: white;
                border: none;
                padding: 10px 20px;
                border-radius: 5px;
                font-weight: bold;
            }
            QPushButton:hover {
                background-color: #7f8c8d;
            }
        """)
        btn_cancel.clicked.connect(self.reject)
        
        btn_layout.addWidget(btn_send)
        btn_layout.addWidget(btn_cancel)
        layout.addLayout(btn_layout)
    
    def send_feedback(self):
        """Geri bildirim gönder"""
        if not self.subject.text().strip():
            QMessageBox.warning(self, "Uyarı", "Lütfen konu başlığı girin.")
            return
        
        if not self.message.toPlainText().strip():
            QMessageBox.warning(self, "Uyarı", "Lütfen mesajınızı yazın.")
            return
        
        try:
            # Mail içeriği oluştur
            msg = MIMEMultipart()
            msg['From'] = f"{Config.MAIL_FROM_NAME} <{Config.MAIL_USERNAME}>"
            msg['To'] = Config.MAIL_USERNAME
            msg['Subject'] = f"[Geri Bildirim] {self.subject.text().strip()}"
            
            # Mail gövdesi
            body = f"""
Geri Bildirim Detayları:
========================

Kategori: {self.category.currentText()}
Ad: {self.name.text().strip() or 'Belirtilmemiş'}
E-posta: {self.email.text().strip() or 'Belirtilmemiş'}
Konu: {self.subject.text().strip()}

Mesaj:
{self.message.toPlainText().strip()}

---
Bu mesaj Araç Filo Yönetim Sistemi v{Config.VERSION} tarafından gönderilmiştir.
            """
            
            msg.attach(MIMEText(body, 'plain', 'utf-8'))
            
            # Mail gönder
            server = smtplib.SMTP(Config.MAIL_SMTP_SERVER, Config.MAIL_SMTP_PORT)
            server.starttls()
            server.login(Config.MAIL_USERNAME, Config.MAIL_PASSWORD)
            server.send_message(msg)
            server.quit()
            
            QMessageBox.information(self, "Başarılı", 
                "Geri bildiriminiz başarıyla gönderildi.\nTeşekkür ederiz!")
            self.accept()
            
        except Exception as e:
            QMessageBox.critical(self, "Hata", f"Mail gönderme hatası: {str(e)}")

# =============================================================================
# OTOMASYON SİSTEMİ
# =============================================================================

class AutomationSystem:
    """Otomasyon sistemi - hatırlatmalar ve kontroller"""
    
    def __init__(self, data_manager):
        self.data_manager = data_manager
    
    def check_expiring_documents(self):
        """Süresi Yaklaşan İşlemleri kontrol et"""
        try:
            vehicles = self.data_manager.load_data('araclar')
            current_date = datetime.now().date()
            expiring_items = []
            
            for _, vehicle in vehicles.iterrows():
                plaka = vehicle.get('Plaka', '')
                
                # Sigorta kontrolü
                sigorta_date = self._parse_date(vehicle.get('Sigorta Bitiş', ''))
                if sigorta_date:
                    days_left = (sigorta_date - current_date).days
                    if 0 <= days_left <= 30:
                        expiring_items.append({
                            'type': 'Sigorta',
                            'plaka': plaka,
                            'date': sigorta_date,
                                'days_left': days_left,
                                'color': '#e74c3c' if days_left <= 7 else '#f39c12'
                            })
                
                # Muayene kontrolü
                muayene_date = self._parse_date(vehicle.get('Muayene Tarihi', ''))
                if muayene_date:
                    days_left = (muayene_date - current_date).days
                    if 0 <= days_left <= 30:
                        expiring_items.append({
                            'type': 'Muayene',
                            'plaka': plaka,
                            'date': muayene_date,
                            'days_left': days_left,
                            'color': '#e74c3c' if days_left <= 7 else '#f39c12'
                        })
            
            return expiring_items
            
        except Exception as e:
            print(f"Belge kontrolü hatası: {e}")
            return []
    
    def _parse_date(self, date_str):
        """Tarih string'ini parse et"""
        try:
            if pd.isna(date_str) or not date_str:
                return None
            return pd.to_datetime(date_str).date()
        except:
            return None
    
    def _parse_date(self, date_str):
        """Tarih string'ini parse et - Türkçe format (dd.mm.yyyy) için dayfirst=True"""
        try:
            if pd.isna(date_str) or not date_str:
                return None
            return pd.to_datetime(date_str, dayfirst=True).date()
        except:
            return None
    
    def create_reminder(self, vehicle_plaka, reminder_type, expiry_date):
        """Hatırlatma oluştur"""
        try:
            reminders = self.data_manager.load_data('hatirlatmalar')
            
            # Mevcut hatırlatma kontrolü
            existing = reminders[
                (reminders['Başlık'].str.contains(vehicle_plaka, na=False)) &
                (reminders['Açıklama'].str.contains(reminder_type, na=False))
            ]
            
            if not existing.empty:
                return  # Zaten hatırlatma var
            
            # Yeni hatırlatma
            new_reminder = {
                'ID': len(reminders) + 1 if not reminders.empty else 1,
                'Başlık': f"{vehicle_plaka} - {reminder_type} Hatırlatması",
                'Açıklama': f"{vehicle_plaka} plakalı aracın {reminder_type} süresi {expiry_date} tarihinde dolacak.",
                'Tarih': expiry_date.strftime("%Y-%m-%d"),
                'Tür': reminder_type,
                'Durum': 'Bekliyor',
                'Oluşturma Tarihi': datetime.now().strftime("%Y-%m-%d %H:%M:%S")
            }
            
            new_df = pd.DataFrame([new_reminder])
            reminders = pd.concat([reminders, new_df], ignore_index=True)
            
            return self.data_manager.save_data('hatirlatmalar', reminders)
            
        except Exception as e:
            print(f"Hatırlatma oluşturma hatası: {e}")
            return False



# =============================================================================
# EXCEL VERİ YÜKLEME SİSTEMİ
# =============================================================================

class ExcelTemplateManager:
    """Excel şablon yönetim sistemi"""
    
    def __init__(self, data_manager):
        self.data_manager = data_manager
        self.template_descriptions = {
            'araclar': {
                'title': 'Araçlar Şablonu',
                'description': 'Araç bilgilerini içeren Excel şablonu. Plaka, marka, model, yıl, şasi no, ruhsat seri no, yakıt tipi, araç tipi, son km, durum, sigorta bitiş, muayene tarihi bilgilerini içerir.',
                                     'columns': ['ID', 'Plaka', 'Marka', 'Model', 'Yıl', 'Şasi No', 'Ruhsat Seri No', 'Yakıt Tipi', 'Araç Tipi', 'Son KM', 'Durum', 'Sigorta Bitiş', 'Muayene Tarihi', 'KM Geçmişi', 'Oluşturma Tarihi', 'Güncelleme Tarihi', 'Son Hatırlatma']
            },
            'suruculer': {
                'title': 'Sürücüler Şablonu',
                'description': 'Sürücü bilgilerini içeren Excel şablonu. Ad soyad, TC kimlik, telefon, ehliyet sınıfı, ehliyet tarihi, atanan araç, atama tarihi, durum bilgilerini içerir.',
                'columns': ['ID', 'Ad Soyad', 'TC Kimlik', 'Telefon', 'Ehliyet Sınıfı', 'Ehliyet Tarihi', 'Atanan Araç', 'Atama Tarihi', 'Durum', 'Oluşturma Tarihi', 'Güncelleme Tarihi']
            },
            'bakimlar': {
                'title': 'Bakım & Onarımlar Şablonu',
                'description': 'Bakım ve onarım kayıtlarını içeren Excel şablonu. Araç plakası, bakım tarihi, km, işlem türü, servis adı, tutar, açıklama, fatura dosyası bilgilerini içerir.',
                                     'columns': ['ID', 'Araç Plakası', 'Bakım Tarihi', 'KM', 'İşlem Türü', 'Servis Adı', 'Tutar', 'Açıklama', 'Oluşturma Tarihi']
            },
            'yakitlar': {
                'title': 'Yakıtlar Şablonu',
                'description': 'Yakıt alım kayıtlarını içeren Excel şablonu. Araç plakası, yakıt tipi, tarih, litre, tutar, bayi, açıklama bilgilerini içerir.',
                'columns': ['ID', 'Araç Plakası', 'Yakıt Tipi', 'Tarih', 'Litre', 'Tutar', 'Bayi', 'Açıklama', 'Oluşturma Tarihi']
            },
            'giderler': {
                'title': 'Giderler Şablonu',
                'description': 'Araç giderlerini içeren Excel şablonu. Araç plakası, tarih, gider türü, tutar, açıklama, fatura dosyası bilgilerini içerir.',
                                     'columns': ['ID', 'Araç Plakası', 'Tarih', 'Gider Türü', 'Tutar', 'Açıklama', 'Oluşturma Tarihi']
            },
            'cezalar': {
                'title': 'Trafik Cezaları Şablonu',
                'description': 'Trafik cezalarını içeren Excel şablonu. Araç plakası, sürücü, ceza tarihi, ceza türü, ceza tutarı, ceza yeri, ceza nedeni, ödeme durumu, ödeme tarihi, ceza dosyası bilgilerini içerir.',
                                     'columns': ['ID', 'Araç Plakası', 'Sürücü', 'Ceza Tarihi', 'Ceza Türü', 'Ceza Tutarı', 'Ceza Yeri', 'Ceza Nedeni', 'Ödeme Durumu', 'Ödeme Tarihi', 'Oluşturma Tarihi']
            },
            'kazalar': {
                'title': 'Kazalar Şablonu',
                'description': 'Kaza kayıtlarını içeren Excel şablonu. Araç plakası, sürücü, kaza tarihi, kaza yeri, kaza türü, hasar durumu, hasar tutarı, sigorta şirketi, sigorta dosya no, kaza açıklaması, kaza dosyası bilgilerini içerir.',
                                     'columns': ['ID', 'Araç Plakası', 'Sürücü', 'Kaza Tarihi', 'Kaza Yeri', 'Kaza Türü', 'Hasar Durumu', 'Hasar Tutarı', 'Sigorta Şirketi', 'Sigorta Dosya No', 'Kaza Açıklaması', 'Oluşturma Tarihi']
            },
            
        }
    
    def create_template(self, data_type, save_path=None):
        """Belirtilen veri türü için Excel şablonu oluştur"""
        if data_type not in self.template_descriptions:
            raise ValueError(f"Geçersiz veri türü: {data_type}")
        
        if save_path is None:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"{data_type}_sablonu_{timestamp}.xlsx"
            save_path = os.path.join(os.getcwd(), filename)
        
        # Örnek verilerle DataFrame oluştur
        df = self._create_template_with_examples(data_type)
        
        # Excel dosyasına kaydet
        with pd.ExcelWriter(save_path, engine='openpyxl') as writer:
            df.to_excel(writer, sheet_name='Veri', index=False)
            
            # Açıklama sayfası oluştur
            self._create_instruction_sheet(writer, data_type)
        
        return save_path
    
    def _create_template_with_examples(self, data_type):
        """Örnek verilerle şablon oluştur"""
        if data_type == 'araclar':
            df = pd.DataFrame({
                'ID': [1, 2],
                'Plaka': ['34ABC123', '06XYZ789'],
                'Marka': ['Renault', 'Ford'],
                'Model': ['Megane', 'Focus'],
                'Yıl': ['2020', '2019'],
                'Şasi No': ['VF1KZ0E0000000001', 'WF0AXXGAF5K123456'],
                'Ruhsat Seri No': ['123456789', '987654321'],
                'Yakıt Tipi': ['Benzin', 'Dizel'],
                'Araç Tipi': ['Binek', 'SUV'],
                'Son KM': ['45000', '32000'],
                'Durum': ['Aktif', 'Aktif'],
                'Sigorta Bitiş': ['31.12.2024', '15.06.2024'],
                'Muayene Tarihi': ['30.09.2024', '20.03.2024'],
                
                'KM Geçmişi': ['2020:0, 2021:15000, 2022:30000, 2023:45000', '2019:0, 2020:12000, 2021:25000, 2022:32000'],
                'Oluşturma Tarihi': ['01.01.2024 10:00:00', '01.01.2024 10:00:00'],
                'Güncelleme Tarihi': ['01.01.2024 10:00:00', '01.01.2024 10:00:00'],
                'Son Hatırlatma': ['', '']
            })
        elif data_type == 'suruculer':
            df = pd.DataFrame({
                'ID': [1, 2],
                'Ad Soyad': ['Ahmet Yılmaz', 'Fatma Demir'],
                'TC Kimlik': ['12345678901', '98765432109'],
                'Telefon': ['0532 123 45 67', '0533 987 65 43'],
                'Ehliyet Sınıfı': ['B', 'B'],
                'Ehliyet Tarihi': ['15.03.2015', '20.07.2018'],
                'Atanan Araç': ['34ABC123', '06XYZ789'],
                'Atama Tarihi': ['01.01.2024 10:00:00', '01.01.2024 10:00:00'],
                'Durum': ['Aktif', 'Aktif'],
                'Oluşturma Tarihi': ['01.01.2024 10:00:00', '01.01.2024 10:00:00'],
                'Güncelleme Tarihi': ['01.01.2024 10:00:00', '01.01.2024 10:00:00']
            })
        elif data_type == 'bakimlar':
            df = pd.DataFrame({
                'ID': [1, 2],
                'Araç Plakası': ['34ABC123', '06XYZ789'],
                'Bakım Tarihi': ['15.01.2024', '20.02.2024'],
                'KM': ['45000', '32000'],
                'İşlem Türü': ['Periyodik Bakım', 'Yağ Değişimi'],
                'Servis Adı': ['Oto Servis Merkezi', 'Ford Yetkili Servis'],
                'Tutar': ['1.250,00', '850,00'],
                'Açıklama': ['Motor yağı, filtre değişimi', 'Yağ ve filtre değişimi'],
                
                'Oluşturma Tarihi': ['15.01.2024 10:00:00', '20.02.2024 10:00:00']
            })
        elif data_type == 'yakitlar':
            df = pd.DataFrame({
                'ID': [1, 2],
                'Araç Plakası': ['34ABC123', '06XYZ789'],
                'Yakıt Tipi': ['Benzin', 'Dizel'],
                'Tarih': ['10.01.2024', '15.01.2024'],
                'Litre': ['50', '45'],
                'Tutar': ['750,00', '675,00'],
                'Bayi': ['BP İstasyonu', 'Shell İstasyonu'],
                'Açıklama': ['Tam doldurma', 'Yarı doldurma'],
                'Oluşturma Tarihi': ['10.01.2024 10:00:00', '15.01.2024 10:00:00']
            })
        elif data_type == 'giderler':
            df = pd.DataFrame({
                'ID': [1, 2],
                'Araç Plakası': ['34ABC123', '06XYZ789'],
                'Tarih': ['05.01.2024', '12.01.2024'],
                'Gider Türü': ['Yıkama', 'Park Ücreti'],
                'Tutar': ['50,00', '25,00'],
                'Açıklama': ['Detaylı yıkama', 'Şehir merkezi park'],
                
                'Oluşturma Tarihi': ['05.01.2024 10:00:00', '12.01.2024 10:00:00']
            })
        elif data_type == 'cezalar':
            df = pd.DataFrame({
                'ID': [1, 2],
                'Araç Plakası': ['34ABC123', '06XYZ789'],
                'Sürücü': ['Ahmet Yılmaz', 'Fatma Demir'],
                'Ceza Tarihi': ['10.01.2024', '15.01.2024'],
                'Ceza Türü': ['Hız İhlali', 'Park İhlali'],
                'Ceza Tutarı': ['500,00', '150,00'],
                'Ceza Yeri': ['İstanbul, Kadıköy', 'Ankara, Çankaya'],
                'Ceza Nedeni': ['50 km/h sınırda 70 km/h', 'Yasak yerde park'],
                'Ödeme Durumu': ['Ödendi', 'Beklemede'],
                'Ödeme Tarihi': ['15.01.2024', ''],
                
                'Oluşturma Tarihi': ['10.01.2024 10:00:00', '15.01.2024 10:00:00']
            })
        elif data_type == 'kazalar':
            df = pd.DataFrame({
                'ID': [1, 2],
                'Araç Plakası': ['34ABC123', '06XYZ789'],
                'Sürücü': ['Ahmet Yılmaz', 'Fatma Demir'],
                'Kaza Tarihi': ['05.01.2024', '12.01.2024'],
                'Kaza Yeri': ['İstanbul, Beşiktaş', 'Ankara, Kızılay'],
                'Kaza Türü': ['Hafif Hasar', 'Orta Hasar'],
                'Hasar Durumu': ['Ön tampon çizik', 'Yan kapı göçük'],
                'Hasar Tutarı': ['2.500,00', '8.000,00'],
                'Sigorta Şirketi': ['Anadolu Sigorta', 'Axa Sigorta'],
                'Sigorta Dosya No': ['AS2024001', 'AX2024001'],
                'Kaza Açıklaması': ['Park halindeyken çarpma', 'Kavşakta çarpışma'],
                
                'Oluşturma Tarihi': ['05.01.2024 10:00:00', '12.01.2024 10:00:00']
            })
        
        else:
            # Boş DataFrame oluştur
            self.data_manager._create_empty_excel(save_path, data_type)
            return None
        
        return df
    
    def _create_instruction_sheet(self, writer, data_type):
        """Açıklama sayfası oluştur"""
        instructions = {
            'araclar': {
                'title': 'Araçlar Şablonu - Kullanım Talimatları',
                'instructions': [
                    'Bu şablon araç bilgilerini içerir.',
                    'Plaka: Araç plakası (örn: 34ABC123)',
                    'Marka: Araç markası (örn: Renault, Ford)',
                    'Model: Araç modeli (örn: Megane, Focus)',
                    'Yıl: Araç üretim yılı (örn: 2020)',
                    'Şasi No: Araç şasi numarası',
                    'Ruhsat Seri No: Ruhsat seri numarası',
                    'Yakıt Tipi: Benzin, Dizel, LPG, Elektrik',
                    'Araç Tipi: Binek, SUV, Kamyon, vb.',
                    'Son KM: Araçtaki son kilometre',
                    'Durum: Aktif, Pasif, Arızalı',
                    'Sigorta Bitiş: Sigorta bitiş tarihi (gg.aa.yyyy)',
                    'Muayene Tarihi: Muayene tarihi (gg.aa.yyyy)',
                    
                    'KM Geçmişi: Yıllık km geçmişi (2020:0, 2021:15000)',
                    'ÖNEMLİ: Tarih formatı gg.aa.yyyy şeklinde olmalıdır!',
                    'ÖNEMLİ: Tutar formatı 1.250,00 şeklinde olmalıdır!'
                ]
            },
            'suruculer': {
                'title': 'Sürücüler Şablonu - Kullanım Talimatları',
                'instructions': [
                    'Bu şablon sürücü bilgilerini içerir.',
                    'Ad Soyad: Sürücünün tam adı',
                    'TC Kimlik: 11 haneli TC kimlik numarası',
                    'Telefon: İletişim telefonu',
                    'Ehliyet Sınıfı: B, C, D, E vb.',
                    'Ehliyet Tarihi: Ehliyet alma tarihi (gg.aa.yyyy)',
                    'Atanan Araç: Sürücünün atandığı araç plakası',
                    'Atama Tarihi: Atama tarihi',
                    'Durum: Aktif, Pasif',
                    'ÖNEMLİ: Tarih formatı gg.aa.yyyy şeklinde olmalıdır!'
                ]
            },
            'bakimlar': {
                'title': 'Bakım & Onarımlar Şablonu - Kullanım Talimatları',
                'instructions': [
                    'Bu şablon bakım ve onarım kayıtlarını içerir.',
                    'Araç Plakası: Bakım yapılan araç plakası',
                    'Bakım Tarihi: Bakım tarihi (gg.aa.yyyy)',
                    'KM: Bakım sırasındaki kilometre',
                    'İşlem Türü: Periyodik Bakım, Yağ Değişimi, vb.',
                    'Servis Adı: Bakım yapılan servis adı',
                    'Tutar: Bakım tutarı (1.250,00 formatında)',
                    'Açıklama: Bakım detayları',
                    'ÖNEMLİ: Tarih formatı gg.aa.yyyy şeklinde olmalıdır!',
                    'ÖNEMLİ: Tutar formatı 1.250,00 şeklinde olmalıdır!'
                ]
            },
            'yakitlar': {
                'title': 'Yakıtlar Şablonu - Kullanım Talimatları',
                'instructions': [
                    'Bu şablon yakıt alım kayıtlarını içerir.',
                    'Araç Plakası: Yakıt alınan araç plakası',
                    'Yakıt Tipi: Benzin, Dizel, LPG',
                    'Tarih: Yakıt alma tarihi (gg.aa.yyyy)',
                    'Litre: Alınan yakıt miktarı',
                    'Tutar: Yakıt tutarı (750,00 formatında)',
                    'Bayi: Yakıt alınan istasyon',
                    'Açıklama: Ek açıklamalar',
                    'ÖNEMLİ: Tarih formatı gg.aa.yyyy şeklinde olmalıdır!',
                    'ÖNEMLİ: Tutar formatı 750,00 şeklinde olmalıdır!'
                ]
            },
            'giderler': {
                'title': 'Giderler Şablonu - Kullanım Talimatları',
                'instructions': [
                    'Bu şablon araç giderlerini içerir.',
                    'Araç Plakası: Gider yapılan araç plakası',
                    'Tarih: Gider tarihi (gg.aa.yyyy)',
                    'Gider Türü: Yıkama, Park, Otopark, vb.',
                    'Tutar: Gider tutarı (50,00 formatında)',
                    'Açıklama: Gider detayları',
                    'ÖNEMLİ: Tarih formatı gg.aa.yyyy şeklinde olmalıdır!',
                    'ÖNEMLİ: Tutar formatı 50,00 şeklinde olmalıdır!'
                ]
            },
            'cezalar': {
                'title': 'Trafik Cezaları Şablonu - Kullanım Talimatları',
                'instructions': [
                    'Bu şablon trafik cezalarını içerir.',
                    'Araç Plakası: Ceza alan araç plakası',
                    'Sürücü: Ceza alan sürücü adı',
                    'Ceza Tarihi: Ceza tarihi (gg.aa.yyyy)',
                    'Ceza Türü: Hız İhlali, Park İhlali, vb.',
                    'Ceza Tutarı: Ceza tutarı (500,00 formatında)',
                    'Ceza Yeri: Ceza alınan yer',
                    'Ceza Nedeni: Ceza nedeni',
                    'Ödeme Durumu: Ödendi, Beklemede',
                    'Ödeme Tarihi: Ödeme tarihi (gg.aa.yyyy)',
                    'ÖNEMLİ: Tarih formatı gg.aa.yyyy şeklinde olmalıdır!',
                    'ÖNEMLİ: Tutar formatı 500,00 şeklinde olmalıdır!'
                ]
            },
            'kazalar': {
                'title': 'Kazalar Şablonu - Kullanım Talimatları',
                'instructions': [
                    'Bu şablon kaza kayıtlarını içerir.',
                    'Araç Plakası: Kaza yapan araç plakası',
                    'Sürücü: Kaza yapan sürücü adı',
                    'Kaza Tarihi: Kaza tarihi (gg.aa.yyyy)',
                    'Kaza Yeri: Kaza yeri',
                    'Kaza Türü: Hafif Hasar, Orta Hasar, Ağır Hasar',
                    'Hasar Durumu: Hasar detayları',
                    'Hasar Tutarı: Hasar tutarı (2.500,00 formatında)',
                    'Sigorta Şirketi: Sigorta şirketi adı',
                    'Sigorta Dosya No: Sigorta dosya numarası',
                    'Kaza Açıklaması: Kaza detayları',
                    'ÖNEMLİ: Tarih formatı gg.aa.yyyy şeklinde olmalıdır!',
                    'ÖNEMLİ: Tutar formatı 2.500,00 şeklinde olmalıdır!'
                ]
            },
            
        }
        
        if data_type in instructions:
            info = instructions[data_type]
            
            # Açıklama sayfası oluştur
            instruction_df = pd.DataFrame({
                'Açıklama': info['instructions']
            })
            
            instruction_df.to_excel(writer, sheet_name='Kullanım Talimatları', index=False)
            
            # Başlık ekle
            workbook = writer.book
            worksheet = writer.sheets['Kullanım Talimatları']
            worksheet.insert_rows(1)
            worksheet['A1'] = info['title']
            worksheet['A1'].font = openpyxl.styles.Font(bold=True, size=14)
    
    def create_all_templates(self, output_dir=None):
        """Tüm şablonları oluştur"""
        if output_dir is None:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            output_dir = os.path.join(os.getcwd(), f"excel_sablonlari_{timestamp}")
        
        os.makedirs(output_dir, exist_ok=True)
        created_files = []
        
        for data_type in self.template_descriptions.keys():
            if data_type not in ['silinen_araclar', 'silinen_suruculer', 'silinen_bakimlar', 'silinen_giderler', 'silinen_cezalar', 'silinen_kazalar', 'belgeler', 'hatirlatmalar']:
                filename = f"{data_type}_sablonu.xlsx"
                filepath = os.path.join(output_dir, filename)
                self.create_template(data_type, filepath)
                created_files.append(filepath)
        
        return output_dir, created_files
    
    def get_template_info(self, data_type):
        """Şablon bilgilerini döndür"""
        return self.template_descriptions.get(data_type, {})


class ExcelImporter:
    """Excel veri yükleme sistemi"""
    
    def __init__(self, data_manager):
        self.data_manager = data_manager
    
    def import_vehicles(self, file_path):
        """Araç verilerini import et"""
        try:
            df = pd.read_excel(file_path)
            vehicles = self.data_manager.load_data('araclar')
            
            for _, row in df.iterrows():
                # Plaka kontrolü
                plaka = str(row.get('Plaka', '')).strip()
                if not plaka:
                    continue
                
                # Mevcut araç kontrolü
                existing = vehicles[vehicles['Plaka'] == plaka]
                if not existing.empty:
                    # Güncelleme
                    vehicles = vehicles[vehicles['Plaka'] != plaka]
                
                # Yeni veri
                new_vehicle = {
                    'ID': len(vehicles) + 1 if len(vehicles) > 0 else 1,
                    'Plaka': plaka,
                    'Marka': str(row.get('Marka', '')).strip(),
                    'Model': str(row.get('Model', '')).strip(),
                    'Yıl': str(row.get('Yıl', '')).strip(),
                    'Şasi No': str(row.get('Şasi No', '')).strip(),
                    'Motor No': str(row.get('Motor No', '')).strip(),
                    'Araç Tipi': str(row.get('Araç Tipi', 'Binek')).strip(),
                    'Son KM': str(row.get('Son KM', '')).strip(),
                    'Durum': str(row.get('Durum', 'Aktif')).strip(),
                    'Sigorta Bitiş': str(row.get('Sigorta Bitiş', '')).strip(),
                    'Muayene Tarihi': str(row.get('Muayene Tarihi', '')).strip(),
                    'Evrak Yolu': str(row.get('Evrak Yolu', '')).strip(),
                    'KM Geçmişi': str(row.get('KM Geçmişi', '')).strip(),
                    'Oluşturma Tarihi': datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
                    'Güncelleme Tarihi': datetime.now().strftime("%Y-%m-%d %H:%M:%S")
                }
                
                new_df = pd.DataFrame([new_vehicle])
                vehicles = pd.concat([vehicles, new_df], ignore_index=True)  # type: ignore
            
            return self.data_manager.save_data('araclar', vehicles)
            
        except Exception as e:
            raise Exception(f"Araç import hatası: {str(e)}")
    
    def import_drivers(self, file_path):
        """Sürücü verilerini import et"""
        try:
            df = pd.read_excel(file_path)
            drivers = self.data_manager.load_data('suruculer')
            
            for _, row in df.iterrows():
                # TC kimlik kontrolü
                tc_kimlik = str(row.get('TC Kimlik', '')).strip()
                if not tc_kimlik:
                    continue
                
                # Mevcut sürücü kontrolü
                existing = drivers[drivers['TC Kimlik'] == tc_kimlik]
                if not existing.empty:
                    # Güncelleme
                    drivers = drivers[drivers['TC Kimlik'] != tc_kimlik]
                
                # Yeni veri
                new_driver = {
                    'ID': len(drivers) + 1 if len(drivers) > 0 else 1,
                    'Ad Soyad': str(row.get('Ad Soyad', '')).strip(),
                    'TC Kimlik': tc_kimlik,
                    'Telefon': str(row.get('Telefon', '')).strip(),
                    'Ehliyet Sınıfı': str(row.get('Ehliyet Sınıfı', 'B')).strip(),
                    'Ehliyet Tarihi': str(row.get('Ehliyet Tarihi', '')).strip(),
                    'Atanan Araç': str(row.get('Atanan Araç', '')).strip(),
                    'Atama Tarihi': datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
                    'Durum': str(row.get('Durum', 'Aktif')).strip(),
                    'Oluşturma Tarihi': datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
                    'Güncelleme Tarihi': datetime.now().strftime("%Y-%m-%d %H:%M:%S")
                }
                
                new_df = pd.DataFrame([new_driver])
                drivers = pd.concat([drivers, new_df], ignore_index=True)  # type: ignore
            
            return self.data_manager.save_data('suruculer', drivers)
            
        except Exception as e:
            raise Exception(f"Sürücü import hatası: {str(e)}")
    
    def import_maintenance(self, file_path):
        """Bakım verilerini import et"""
        try:
            df = pd.read_excel(file_path)
            maintenance = self.data_manager.load_data('bakimlar')
            
            for _, row in df.iterrows():
                # Yeni bakım kaydı
                new_maintenance = {
                    'ID': len(maintenance) + 1 if len(maintenance) > 0 else 1,
                    'Araç Plakası': str(row.get('Araç Plakası', '')).strip(),
                    'Bakım Tarihi': str(row.get('Bakım Tarihi', '')).strip(),
                    'KM': str(row.get('KM', '')).strip(),
                    'İşlem Türü': str(row.get('İşlem Türü', '')).strip(),
                    'Servis Adı': str(row.get('Servis Adı', '')).strip(),
                    'Tutar': str(row.get('Tutar', '')).strip(),
                    'Açıklama': str(row.get('Açıklama', '')).strip(),
                    'Fatura Dosyası': str(row.get('Fatura Dosyası', '')).strip(),
                    'Oluşturma Tarihi': datetime.now().strftime("%Y-%m-%d %H:%M:%S")
                }
                
                new_df = pd.DataFrame([new_maintenance])
                maintenance = pd.concat([maintenance, new_df], ignore_index=True)
            
            return self.data_manager.save_data('bakimlar', maintenance)
            
        except Exception as e:
            raise Exception(f"Bakım import hatası: {str(e)}")
    
    def import_expenses(self, file_path):
        """Gider verilerini import et"""
        try:
            df = pd.read_excel(file_path)
            expenses = self.data_manager.load_data('giderler')
            
            for _, row in df.iterrows():
                # Yeni gider kaydı
                new_expense = {
                    'ID': len(expenses) + 1 if len(expenses) > 0 else 1,
                    'Araç Plakası': str(row.get('Araç Plakası', '')).strip(),
                    'Gider Türü': str(row.get('Gider Türü', '')).strip(),
                    'Tarih': str(row.get('Tarih', '')).strip(),
                    'Tutar': str(row.get('Tutar', '')).strip(),
                    'Açıklama': str(row.get('Açıklama', '')).strip(),
                    'Fiş Dosyası': str(row.get('Fiş Dosyası', '')).strip(),
                    'Oluşturma Tarihi': datetime.now().strftime("%Y-%m-%d %H:%M:%S")
                }
                
                new_df = pd.DataFrame([new_expense])
                expenses = pd.concat([expenses, new_df], ignore_index=True)
            
            return self.data_manager.save_data('giderler', expenses)
            
        except Exception as e:
            raise Exception(f"Gider import hatası: {str(e)}")

# =============================================================================
# VERİ YÖNETİMİ (Excel Tabanlı)
# =============================================================================

class DataManager:
    """Excel tabanlı veri yöneticisi"""
    
    def __init__(self):
        self._ensure_directories()
        self._init_excel_files()
    
    def _ensure_directories(self):
        """Gerekli klasörleri oluştur"""
        directories = [Config.DATA_DIR, Config.BACKUP_DIR, Config.LOG_DIR]
        for directory in directories:
            if not os.path.exists(directory):
                os.makedirs(directory)
    
    def _init_excel_files(self):
        """Excel dosyalarını başlat"""
        for key, filename in Config.EXCEL_FILES.items():
            filepath = os.path.join(Config.DATA_DIR, filename)
            if not os.path.exists(filepath):
                self._create_empty_excel(filepath, key)
    
    def _create_empty_excel(self, filepath, data_type):
        """Boş Excel dosyası oluştur"""
        if data_type == 'araclar':
            df = pd.DataFrame({
                'ID': [],
                'Plaka': [],
                'Marka': [],
                'Model': [],
                'Yıl': [],
                'Şasi No': [],
                'Ruhsat Seri No': [],
                'Yakıt Tipi': [],
                'Araç Tipi': [],
                'Son KM': [],
                'Durum': [],
                'Sigorta Bitiş': [],
                'Muayene Tarihi': [],
                'Evrak Yolu': [],
                'KM Geçmişi': [],
                'Oluşturma Tarihi': [],
                'Güncelleme Tarihi': [],
                'Son Hatırlatma': []
            })
        elif data_type == 'suruculer':
            df = pd.DataFrame({
                'ID': [],
                'Ad Soyad': [],
                'TC Kimlik': [],
                'Telefon': [],
                'Ehliyet Sınıfı': [],
                'Ehliyet Tarihi': [],
                'Atanan Araç': [],
                'Atama Tarihi': [],
                'Durum': [],
                'Oluşturma Tarihi': [],
                'Güncelleme Tarihi': []
            })
        elif data_type == 'bakimlar':
            df = pd.DataFrame({
                'ID': [],
                'Araç Plakası': [],
                'Bakım Tarihi': [],
                'KM': [],
                'İşlem Türü': [],
                'Servis Adı': [],
                'Tutar': [],
                'Açıklama': [],
                'Fatura Dosyası': [],
                'Oluşturma Tarihi': []
            })
        elif data_type == 'yakitlar':
            df = pd.DataFrame({
                'ID': [],
                'Araç Plakası': [],
                'Yakıt Tipi': [],
                'Tarih': [],
                'Litre': [],
                'Tutar': [],
                'Bayi': [],
                'Açıklama': [],
                'Oluşturma Tarihi': []
            })
        elif data_type == 'cezalar':
            df = pd.DataFrame({
                'ID': [],
                'Araç Plakası': [],
                'Sürücü': [],
                'Ceza Tarihi': [],
                'Ceza Türü': [],
                'Ceza Tutarı': [],
                'Ceza Yeri': [],
                'Ceza Nedeni': [],
                'Ödeme Durumu': [],
                'Ödeme Tarihi': [],
                'Ceza Dosyası': [],
                'Oluşturma Tarihi': []
            })
        elif data_type == 'kazalar':
            df = pd.DataFrame({
                'ID': [],
                'Araç Plakası': [],
                'Sürücü': [],
                'Kaza Tarihi': [],
                'Kaza Yeri': [],
                'Kaza Türü': [],
                'Hasar Durumu': [],
                'Hasar Tutarı': [],
                'Sigorta Şirketi': [],
                'Sigorta Dosya No': [],
                'Kaza Açıklaması': [],
                'Kaza Dosyası': [],
                'Oluşturma Tarihi': []
            })
        elif data_type == 'hatirlatmalar':
            df = pd.DataFrame({
                'ID': [],
                'Başlık': [],
                'Açıklama': [],
                'Tarih': [],
                'Tür': [],
                'Durum': [],
                'Oluşturma Tarihi': []
            })
        elif data_type == 'silinen_araclar':
            df = pd.DataFrame({
                'ID': [],
                'Plaka': [],
                'Marka': [],
                'Model': [],
                'Yıl': [],
                'Şasi No': [],
                'Motor No': [],
                'Araç Tipi': [],
                'Son KM': [],
                'Durum': [],
                'Sigorta Bitiş': [],
                'Muayene Tarihi': [],
                'Birim': [],
                'Hizmet Kişisi': [],
                'Fotoğraf Yolu': [],
                'Evrak Yolu': [],
                'KM Geçmişi': [],
                'Oluşturma Tarihi': [],
                'Güncelleme Tarihi': [],
                'Silme Tarihi': []
            })
        elif data_type == 'silinen_suruculer':
            df = pd.DataFrame({
                'ID': [],
                'Ad Soyad': [],
                'TC Kimlik': [],
                'Telefon': [],
                'Ehliyet Sınıfı': [],
                'Ehliyet Tarihi': [],
                'Atanan Araç': [],
                'Atama Tarihi': [],
                'Oluşturma Tarihi': [],
                'Güncelleme Tarihi': [],
                'Silme Tarihi': []
            })
        elif data_type == 'silinen_bakimlar':
            df = pd.DataFrame({
                'ID': [],
                'Araç Plakası': [],
                'Bakım Tarihi': [],
                'KM': [],
                'İşlem Türü': [],
                'Açıklama': [],
                'Tutar': [],
                'Servis': [],
                'Oluşturma Tarihi': [],
                'Güncelleme Tarihi': [],
                'Silme Tarihi': []
            })
        elif data_type == 'silinen_giderler':
            df = pd.DataFrame({
                'ID': [],
                'Araç Plakası': [],
                'Tarih': [],
                'Gider Türü': [],
                'Tutar': [],
                'Açıklama': [],
                'Oluşturma Tarihi': [],
                'Güncelleme Tarihi': [],
                'Silme Tarihi': []
            })
        elif data_type == 'silinen_cezalar':
            df = pd.DataFrame({
                'ID': [],
                'Araç Plakası': [],
                'Sürücü': [],
                'Ceza Tarihi': [],
                'Ceza Türü': [],
                'Ceza Tutarı': [],
                'Ceza Yeri': [],
                'Ceza Nedeni': [],
                'Ödeme Durumu': [],
                'Ödeme Tarihi': [],
                'Ceza Dosyası': [],
                'Oluşturma Tarihi': [],
                'Silme Tarihi': []
            })
        elif data_type == 'silinen_kazalar':
            df = pd.DataFrame({
                'ID': [],
                'Araç Plakası': [],
                'Sürücü': [],
                'Kaza Tarihi': [],
                'Kaza Yeri': [],
                'Kaza Türü': [],
                'Hasar Durumu': [],
                'Hasar Tutarı': [],
                'Sigorta Şirketi': [],
                'Sigorta Dosya No': [],
                'Kaza Açıklaması': [],
                'Kaza Dosyası': [],
                'Oluşturma Tarihi': [],
                'Silme Tarihi': []
            })
        elif data_type == 'giderler':
            df = pd.DataFrame({
                'ID': [],
                'Araç Plakası': [],
                'Tarih': [],
                'Gider Türü': [],
                'Tutar': [],
                'Açıklama': [],
                'Fatura Dosyası': [],
                'Oluşturma Tarihi': []
            })
        elif data_type == 'belgeler':
            df = pd.DataFrame({
                'ID': [],
                'Araç Plakası': [],
                'Belge Adı': [],
                'Belge Türü': [],
                'Geçerlilik Tarihi': [],
                'Dosya Yolu': [],
                'Açıklama': [],
                'Oluşturma Tarihi': []
            })
        
        df.to_excel(filepath, index=False)
    
    def load_data(self, data_type):
        filepath = os.path.join(Config.DATA_DIR, Config.EXCEL_FILES[data_type])
        if not os.path.exists(filepath):
            self._create_empty_excel(filepath, data_type)
        df = pd.read_excel(filepath)
        # --- Giderler için eksik sütunları tamamla ---
        if data_type == 'giderler':
            for col in ['Araç Plakası', 'Gider Türü', 'Başlangıç', 'Bitiş', 'Şirket', 'Tutar']:
                if col not in df.columns:
                    df[col] = ''
        return df
    
    def save_data(self, data_type, data):
        """Veriyi Excel dosyasına kaydet"""
        try:
            filepath = os.path.join(Config.DATA_DIR, Config.EXCEL_FILES[data_type])
            data.to_excel(filepath, index=False)
            self._create_backup(data_type)
            logger.info(f"Veri kaydedildi: {data_type} - {len(data)} kayıt")
            return True
        except Exception as e:
            logger.error(f"Veri kaydetme hatası ({data_type}): {str(e)}", e)
            return False
    
    def _create_backup(self, data_type):
        """Yedek oluştur"""
        try:
            source = os.path.join(Config.DATA_DIR, Config.EXCEL_FILES[data_type])
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            backup_name = f"{data_type}_{timestamp}.xlsx"
            backup_path = os.path.join(Config.BACKUP_DIR, backup_name)
            shutil.copy2(source, backup_path)
        except Exception as e:
            self._log_error(f"Yedek oluşturma hatası: {str(e)}")
    
    def _log_error(self, message):
        """Hata logla"""
        timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        log_message = f"[{timestamp}] {message}\n"
        log_file = os.path.join(Config.LOG_DIR, "hata_log.txt")
        with open(log_file, "a", encoding="utf-8") as f:
            f.write(log_message)
    


# =============================================================================
# ANA PENCERE
# =============================================================================

class MainWindow(QMainWindow):
    """Ana uygulama penceresi"""
    
    def __init__(self):
        super().__init__()
        
        # Log sistemi başlat
        logger.log_system_event("Ana pencere başlatılıyor", f"v{Config.VERSION}")
        
        self.data_manager = DataManager()
        self.current_panel = None
        self.automation_system = AutomationSystem(self.data_manager)
        self.last_used_panel = 'dashboard'  # Son kullanılan panel
        
        self.init_ui()
        self.setup_menu()
        self.setup_status_bar()
        self.setup_keyboard_shortcuts()  # Klavye kısayollarını ayarla
        self.load_dashboard()
        
        # Otomatik hatırlatma timer'ı
        self.reminder_timer = QTimer()
        self.reminder_timer.timeout.connect(self.check_automated_reminders)
        self.reminder_timer.start(3600000)  # Her saat kontrol et (1 saat = 3600000 ms)
        
        # Otomatik güncelleme kontrolü timer'ı
        self.update_timer = QTimer()
        self.update_timer.timeout.connect(self.check_for_updates_silent)
        self.update_timer.start(24 * 3600000)  # Her 24 saat kontrol et
        
        # İlk kontrolleri hemen yap
        self.check_automated_reminders()
        
        # Program başlarken GitHub versiyon kontrolü yap (daha hızlı)
        QTimer.singleShot(1000, self.check_startup_update)
        
        logger.log_system_event("Ana pencere başlatıldı", "Başarılı")
    
    def init_ui(self):
        """UI başlat"""
        self.setWindowTitle(f"{Config.APP_NAME} v{Config.VERSION}")
        self.setGeometry(100, 100, 1400, 900)
        self.setMinimumSize(1200, 800)
        
        # Ana widget
        central_widget = QWidget()
        self.setCentralWidget(central_widget)
        
        # Ana layout
        main_layout = QHBoxLayout(central_widget)
        main_layout.setContentsMargins(0, 0, 0, 0)
        main_layout.setSpacing(0)
        
        # Sol menü
        self.create_sidebar()
        
        # Sağ panel
        self.create_main_panel()
        
        # Layout'a ekle
        main_layout.addWidget(self.sidebar, 1)
        main_layout.addWidget(self.main_panel, 4)
        
        # Stil uygula
        self.apply_styles()
    
    def create_sidebar(self):
        """Klasik Windows sidebar oluştur"""
        self.sidebar = QFrame()
        self.sidebar.setMaximumWidth(220)
        self.sidebar.setMinimumWidth(180)
        self.sidebar.setFrameStyle(QFrame.Box)
        self.sidebar.setStyleSheet("""
            QFrame {
                background-color: #f0f0f0;
                border-right: 1px solid #c0c0c0;
            }
        """)
        
        # Ana layout
        main_layout = QVBoxLayout(self.sidebar)
        main_layout.setContentsMargins(0, 0, 0, 0)
        main_layout.setSpacing(0)
        
        # Üst kısım - Logo ve daraltma butonu
        top_frame = QFrame()
        top_frame.setStyleSheet("background-color: #0078d7; border: none;")
        top_layout = QHBoxLayout(top_frame)
        top_layout.setContentsMargins(10, 10, 10, 10)
        
        # Logo/başlık
        title_label = QLabel("Araç Filo Yönetim")
        title_label.setStyleSheet("""
            QLabel {
                font-size: 14px;
                font-weight: bold;
                color: white;
            }
        """)
        
        # Sidebar collapse butonu kaldırıldı - sidebar hep açık kalacak
        
        top_layout.addWidget(title_label)
        top_layout.addStretch()
        
        main_layout.addWidget(top_frame)
        
        # Scroll area için widget
        scroll_widget = QWidget()
        scroll_layout = QVBoxLayout(scroll_widget)
        scroll_layout.setContentsMargins(5, 5, 5, 5)
        scroll_layout.setSpacing(1)
        
        # Menü butonları - Tablo benzeri düzen
        self.menu_buttons = {}
        self.menu_containers = {}
        
        for item in Config.MENU_ITEMS:
            item_id = item['id']
            
            # Menü butonu
            btn_container = self.create_menu_button(item)
            self.menu_buttons[item_id] = btn_container
            self.menu_containers[item_id] = btn_container
            scroll_layout.addWidget(btn_container)
        
        # Alt boşluk
        scroll_layout.addStretch()
        
        # Alt kısım - Sadece versiyon bilgisi
        bottom_frame = QFrame()
        bottom_frame.setStyleSheet("background-color: #ecf0f1; border-top: 1px solid #bdc3c7;")
        bottom_layout = QVBoxLayout(bottom_frame)
        bottom_layout.setContentsMargins(10, 10, 10, 10)
        
        # Versiyon bilgisi
        version_label = QLabel(f"v{Config.VERSION}")
        version_label.setAlignment(Qt.AlignmentFlag.AlignCenter)
        version_label.setStyleSheet("color: #7f8c8d; font-size: 11px; font-weight: bold;")
        bottom_layout.addWidget(version_label)
        
        main_layout.addWidget(bottom_frame)
        
        # Scroll area
        self.scroll_area = QScrollArea()
        self.scroll_area.setWidget(scroll_widget)
        self.scroll_area.setWidgetResizable(True)
        self.scroll_area.setHorizontalScrollBarPolicy(Qt.ScrollBarAlwaysOff)
        self.scroll_area.setVerticalScrollBarPolicy(Qt.ScrollBarAsNeeded)
        self.scroll_area.setStyleSheet("QScrollArea { border: none; }")
        
        main_layout.insertWidget(1, self.scroll_area)
        
        # Sidebar durumu - hep açık kalacak
    
    def create_menu_button(self, item):
        """Klasik Windows menü butonu oluştur"""
        # Ana buton container
        btn_container = QFrame()
        btn_container.setStyleSheet("""
            QFrame {
                background-color: transparent;
                border: none;
                margin: 1px 0px;
            }
        """)
        
        btn_layout = QHBoxLayout(btn_container)
        btn_layout.setContentsMargins(2, 2, 2, 2)
        btn_layout.setSpacing(3)
        
        # Ana buton - Emoji ikonları kullan
        btn = QPushButton(f"{item.get('emoji_icon', '🏠')} {item['title']}")
        
        btn.setMinimumHeight(32)
        btn.setObjectName(f"menu_btn_{item['id']}")
        
        btn.setStyleSheet("""
            QPushButton {
                text-align: left;
                padding: 4px 8px;
                border: none;
                background-color: transparent;
                color: #000000;
                font-size: 11px;
                font-weight: normal;
            }
            QPushButton:hover {
                background-color: #e1e1e1;
                color: #000000;
            }
            QPushButton:pressed {
                background-color: #d4d4d4;
                color: #000000;
            }
            QPushButton:checked {
                background-color: #3498db;
                color: white;
                font-weight: bold;
            }
        """)
        
        # Kısayol tuşu göster
        if 'shortcut' in item:
            shortcut_label = QLabel(item['shortcut'])
            shortcut_label.setStyleSheet("""
                QLabel {
                    color: #666666;
                    font-size: 8px;
                    font-weight: normal;
                }
            """)
            btn_layout.addWidget(btn)
            btn_layout.addWidget(shortcut_label)
            btn_layout.addStretch()
        else:
            btn_layout.addWidget(btn)
            btn_layout.addStretch()
        
        btn.clicked.connect(lambda: self.show_panel(item['id']))
        
        # Tooltip ekle
        tooltip_text = item['title']
        if 'shortcut' in item:
            tooltip_text += f" ({item['shortcut']})"
        btn.setToolTip(tooltip_text)
        
        return btn_container
    
    def create_main_panel(self):
        """Ana panel oluştur"""
        self.main_panel = QStackedWidget()
        self.main_panel.setStyleSheet("""
            QStackedWidget {
                background-color: #f0f0f0;
                border-left: 1px solid #c0c0c0;
            }
        """)
    
    def setup_menu(self):
        """Menü barı oluştur"""
        menubar = self.menuBar()
        if menubar is None:
            return
        
        # Dosya menüsü
        file_menu = menubar.addMenu('Dosya')
        
        # Veri yedekleme
        backup_action = QAction('Veri Yedekle', self)
        backup_action.triggered.connect(self.backup_data)
        file_menu.addAction(backup_action)
        
        # Excel export
        export_action = QAction('Excel\'e Aktar', self)
        export_action.triggered.connect(self.export_to_excel)
        file_menu.addAction(export_action)
        
        file_menu.addSeparator()
        
        # Çıkış
        exit_action = QAction('Çıkış', self)
        exit_action.triggered.connect(lambda: self.close())
        file_menu.addAction(exit_action)
        
        # Yardım menüsü
        help_menu = menubar.addMenu('Yardım')
        
        # Güncelleme kontrolü
        update_action = QAction('Güncelleme Kontrolü', self)
        update_action.triggered.connect(self.check_for_updates)
        help_menu.addAction(update_action)
        
        help_menu.addSeparator()
        
        # Geri bildirim
        feedback_action = QAction('Geri Bildirim Gönder', self)
        feedback_action.triggered.connect(self.show_feedback)
        help_menu.addAction(feedback_action)
        
        help_menu.addSeparator()
        
        about_action = QAction('Hakkında', self)
        about_action.triggered.connect(self.show_about)
        help_menu.addAction(about_action)
    
    def setup_status_bar(self):
        """Durum çubuğu oluştur"""
        self.status_bar = QStatusBar()
        self.setStatusBar(self.status_bar)
        
        # Durum
        self.status_bar.addPermanentWidget(QLabel("Hazır"))
    
    def apply_styles(self):
        """Klasik Windows teması uygula"""
        self.setStyleSheet("""
            QMainWindow {
                background-color: #f0f0f0;
            }
            QMenuBar {
                background-color: #f0f0f0;
                border-bottom: 1px solid #c0c0c0;
                color: #000000;
            }
            QMenuBar::item {
                background-color: transparent;
                padding: 4px 8px;
            }
            QMenuBar::item:selected {
                background-color: #0078d7;
                color: white;
            }
            QMenu {
                background-color: #f0f0f0;
                border: 1px solid #c0c0c0;
                padding: 2px;
            }
            QMenu::item {
                padding: 4px 20px;
            }
            QMenu::item:selected {
                background-color: #0078d7;
                color: white;
            }
            QStatusBar {
                background-color: #f0f0f0;
                border-top: 1px solid #c0c0c0;
                color: #000000;
            }
            QToolBar {
                background-color: #f0f0f0;
                border-bottom: 1px solid #c0c0c0;
                spacing: 2px;
                padding: 2px;
            }
            QToolButton {
                background-color: #f0f0f0;
                border: 1px solid transparent;
                padding: 4px;
                margin: 1px;
            }
            QToolButton:hover {
                background-color: #e1e1e1;
                border: 1px solid #c0c0c0;
            }
            QToolButton:pressed {
                background-color: #d4d4d4;
                border: 1px solid #a0a0a0;
            }
        """)
    
    def show_panel(self, panel_id):
        """Panel göster - Gelişmiş versiyon"""
        # Menü butonlarını güncelle
        for btn_id, btn_container in self.menu_buttons.items():
            btn = btn_container.findChild(QPushButton)
            if btn:
                if btn_id == panel_id:
                    # Seçili buton stilini güncelle
                    btn.setChecked(True)
                    group_color = "#3498db"
                    if 'group' in next((item for item in Config.MENU_ITEMS if item['id'] == btn_id), {}):
                        group = next((item for item in Config.MENU_ITEMS if item['id'] == btn_id), {})['group']
                        if group in Config.MENU_GROUPS:
                            group_color = Config.MENU_GROUPS[group]['color']
                    
                    btn.setStyleSheet(f"""
                        QPushButton {{
                        text-align: left;
                        padding: 10px 15px;
                        border: none;
                            border-radius: 8px;
                            background-color: {group_color};
                        color: white;
                        font-size: 13px;
                            font-weight: 500;
                            border-left: 4px solid {group_color};
                        }}
                """)
                else:
                    # Normal buton stilini güncelle
                    btn.setChecked(False)
                    group_color = "#3498db"
                    if 'group' in next((item for item in Config.MENU_ITEMS if item['id'] == btn_id), {}):
                        group = next((item for item in Config.MENU_ITEMS if item['id'] == btn_id), {})['group']
                        if group in Config.MENU_GROUPS:
                            group_color = Config.MENU_GROUPS[group]['color']
                    
                    btn.setStyleSheet(f"""
                        QPushButton {{
                        text-align: left;
                        padding: 10px 15px;
                        border: none;
                            border-radius: 8px;
                            background-color: #f8f9fa;
                        color: #2c3e50;
                        font-size: 13px;
                            font-weight: 500;
                            border-left: 4px solid transparent;
                        }}
                        QPushButton:hover {{
                            background-color: {group_color};
                        color: white;
                            border-left: 4px solid {group_color};
                        }}
                """)
        
        # Son kullanılan paneli kaydet
        self.last_used_panel = panel_id
        
        # Panel içeriğini yükle
        if panel_id == 'dashboard':
            self.load_dashboard()
        elif panel_id == 'araclar':
            self.load_vehicles_panel()
        elif panel_id == 'suruculer':
            self.load_drivers_panel()
        elif panel_id == 'bakimlar':
            self.load_maintenance_panel()
        elif panel_id == 'yakitlar':
            self.load_fuel_panel()
        elif panel_id == 'trafik':
            self.load_traffic_panel()
        elif panel_id == 'raporlar':
            self.load_reports_panel()
        elif panel_id == 'belgeler':
            self.load_documents_panel()
        elif panel_id == 'ayarlar':
            self.load_settings_panel()
    
    # Sidebar toggle fonksiyonu kaldırıldı - sidebar hep açık kalacak
    
    def setup_keyboard_shortcuts(self):
        """Klavye kısayollarını ayarla"""
        for item in Config.MENU_ITEMS:
            if 'shortcut' in item:
                shortcut = QShortcut(QKeySequence(item['shortcut']), self)
                shortcut.activated.connect(lambda pid=item['id']: self.show_panel(pid))
        
        # Genel kısayollar
        refresh_shortcut = QShortcut(QKeySequence("F5"), self)
        refresh_shortcut.activated.connect(self.refresh_current_panel)
        

        
        # Sidebar kısayolu kaldırıldı - sidebar hep açık kalacak
    
    def refresh_current_panel(self):
        """Mevcut paneli yenile"""
        if hasattr(self, 'current_panel') and self.current_panel:
            if hasattr(self.current_panel, 'load_data'):
                self.current_panel.load_data()
            if hasattr(self.current_panel, 'refresh_dashboard'):
                self.current_panel.refresh_dashboard()
    

    
    def show_usage_statistics(self):
        """Kullanım istatistiklerini göster"""
        # Bu fonksiyon gelecekte kullanım istatistiklerini gösterecek
        QMessageBox.information(self, "Kullanım İstatistikleri", 
            "Bu özellik gelecekte eklenecek.\n"
            "Hangi panellerin daha çok kullanıldığını görebileceksiniz.")
    
    def load_dashboard(self):
        """Ana sayfa yükle"""
        # Mevcut paneli temizle
        if self.current_panel:
            self.main_panel.removeWidget(self.current_panel)
        
        # Yeni dashboard paneli oluştur
        self.current_panel = DashboardPanel(self.data_manager)
        self.main_panel.addWidget(self.current_panel)
        self.main_panel.setCurrentWidget(self.current_panel)
    
    def refresh_dashboard(self):
        """Dashboard'ı yenile"""
        if isinstance(self.current_panel, DashboardPanel):
            self.current_panel.load_data()
            if hasattr(self.current_panel, 'update_maintenance_card'):
                self.current_panel.update_maintenance_card()
            if hasattr(self.current_panel, 'update_fuel_details_label'):
                self.current_panel.update_fuel_details_label()
        

    
    def load_vehicles_panel(self):
        """Araçlar paneli yükle"""
        if self.current_panel:
            self.main_panel.removeWidget(self.current_panel)
        
        self.current_panel = VehiclesPanel(self.data_manager)
        self.main_panel.addWidget(self.current_panel)
        self.main_panel.setCurrentWidget(self.current_panel)
    
    def load_drivers_panel(self):
        """Sürücüler paneli yükle"""
        if self.current_panel:
            self.main_panel.removeWidget(self.current_panel)
        
        self.current_panel = DriversPanel(self.data_manager)
        self.main_panel.addWidget(self.current_panel)
        self.main_panel.setCurrentWidget(self.current_panel)
    
    def load_maintenance_panel(self):
        """Bakım paneli yükle"""
        if self.current_panel:
            self.main_panel.removeWidget(self.current_panel)
        
        self.current_panel = MaintenancePanel(self.data_manager)
        self.main_panel.addWidget(self.current_panel)
        self.main_panel.setCurrentWidget(self.current_panel)
    
    def load_fuel_panel(self):
        if self.current_panel:
            self.main_panel.removeWidget(self.current_panel)
        self.current_panel = YakıtlarPanel(self.data_manager)
        self.main_panel.addWidget(self.current_panel)
        self.main_panel.setCurrentWidget(self.current_panel)
    
    def load_traffic_panel(self):
        """Trafik paneli yükle"""
        if self.current_panel:
            self.main_panel.removeWidget(self.current_panel)
        
        self.current_panel = TrafficPanel(self.data_manager)
        self.main_panel.addWidget(self.current_panel)
        self.main_panel.setCurrentWidget(self.current_panel)
    
    def load_reports_panel(self):
        """Raporlar paneli yükle"""
        if self.current_panel:
            self.main_panel.removeWidget(self.current_panel)
        
        self.current_panel = ReportsPanel(self.data_manager)
        self.main_panel.addWidget(self.current_panel)
        self.main_panel.setCurrentWidget(self.current_panel)
    
    def load_calendar_panel(self):
        """Takvim paneli yükle"""
        if self.current_panel:
            self.main_panel.removeWidget(self.current_panel)
        
        self.current_panel = CalendarPanel(self.data_manager)
        self.main_panel.addWidget(self.current_panel)
        self.main_panel.setCurrentWidget(self.current_panel)
    
    def load_documents_panel(self):
        """Belgeler paneli yükle"""
        if self.current_panel:
            self.main_panel.removeWidget(self.current_panel)
        
        self.current_panel = DocumentsPanel(self.data_manager)
        self.main_panel.addWidget(self.current_panel)
        self.main_panel.setCurrentWidget(self.current_panel)
    
    def load_settings_panel(self):
        """Ayarlar paneli yükle"""
        if self.current_panel:
            self.main_panel.removeWidget(self.current_panel)
        
        self.current_panel = SettingsPanel(self.data_manager)
        self.main_panel.addWidget(self.current_panel)
        self.main_panel.setCurrentWidget(self.current_panel)
    
    def backup_data(self):
        """Veri yedekle"""
        try:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            backup_dir = os.path.join(Config.BACKUP_DIR, f"yedek_{timestamp}")
            os.makedirs(backup_dir, exist_ok=True)
            
            for data_type, filename in Config.EXCEL_FILES.items():
                source = os.path.join(Config.DATA_DIR, filename)
                if os.path.exists(source):
                    shutil.copy2(source, backup_dir)
            
            QMessageBox.information(self, "Başarılı", f"Veriler yedeklendi:\n{backup_dir}")
        except Exception as e:
            QMessageBox.critical(self, "Hata", f"Yedekleme hatası: {str(e)}")
    
    def export_to_excel(self):
        """Excel'e aktar"""
        try:
            folder = QFileDialog.getExistingDirectory(self, "Excel dosyalarını kaydet")
            if folder:
                for data_type, filename in Config.EXCEL_FILES.items():
                    source = os.path.join(Config.DATA_DIR, filename)
                    if os.path.exists(source):
                        dest = os.path.join(folder, filename)
                        shutil.copy2(source, dest)
                
                QMessageBox.information(self, "Başarılı", f"Excel dosyaları kaydedildi:\n{folder}")
        except Exception as e:
            QMessageBox.critical(self, "Hata", f"Export hatası: {str(e)}")
    
    def check_for_updates(self):
        """Gelişmiş güncelleme kontrolü"""
        # Progress dialog göster
        self.progress_dialog = QProgressDialog("Güncelleme kontrol ediliyor...", None, 0, 0, self)
        self.progress_dialog.setWindowTitle("🔄 Güncelleme Kontrolü")
        self.progress_dialog.setWindowModality(Qt.WindowModal)
        self.progress_dialog.setCancelButton(None)
        self.progress_dialog.setMinimumDuration(0)
        self.progress_dialog.show()
        
        # Güncelleme kontrolcüsünü başlat
        self.update_checker = UpdateChecker()
        self.update_checker.update_available.connect(self.show_update_dialog)
        self.update_checker.no_update.connect(self.show_no_update)
        self.update_checker.error_occurred.connect(self.show_update_error)
        self.update_checker.check_progress.connect(self.update_progress)
        self.update_checker.finished.connect(self.update_check_finished)
        self.update_checker.start()
    
    def update_progress(self, message):
        """İlerleme durumunu güncelle"""
        if hasattr(self, 'progress_dialog'):
            self.progress_dialog.setLabelText(message)
    
    def update_check_finished(self):
        """Güncelleme kontrolü tamamlandı"""
        if hasattr(self, 'progress_dialog'):
            self.progress_dialog.close()
        
        # Kontrol tamamlandı olarak işaretle
        if hasattr(self, 'update_checker'):
            self.update_checker.mark_check_completed()
    
    def show_update_dialog(self, update_info):
        """Gelişmiş güncelleme dialog göster"""
        dialog = UpdateDialog(self, update_info)
        result = dialog.exec_()
        
        # Kullanıcı güncellemeyi indirmeyi seçtiyse log kaydı
        if result == QDialog.Accepted:
            logger.log_system_event("Güncelleme indirme başlatıldı", 
                f"v{update_info.get('version', '')} -> v{Config.VERSION}")
    
    def show_no_update(self):
        """Güncelleme yok mesajı"""
        QMessageBox.information(self, "✅ Güncelleme Kontrolü", 
            f"🎉 Tebrikler! Güncel sürümü kullanmaktasınız.\n\n"
            f"📋 Sürüm Bilgileri:\n"
            f"• Mevcut Sürüm: v{Config.VERSION}\n"
            f"• Son Kontrol: {datetime.now().strftime('%d.%m.%Y %H:%M')}\n\n"
            f"💡 İpucu: Güncellemeler otomatik olarak kontrol edilir.")
    
    def show_update_error(self, error_message):
        """Gelişmiş güncelleme hatası mesajı"""
        QMessageBox.warning(self, "⚠️ Güncelleme Hatası", 
            f"Güncelleme kontrolü sırasında bir hata oluştu:\n\n"
            f"❌ Hata: {error_message}\n\n"
            f"🔧 Çözüm Önerileri:\n"
            f"• İnternet bağlantınızı kontrol edin\n"
            f"• Güvenlik duvarı ayarlarınızı kontrol edin\n"
            f"• Daha sonra tekrar deneyin\n\n"
            f"📞 Destek: {Config.EMAIL}")
        
        # Hata logunu kaydet
        logger.error(f"Güncelleme kontrolü hatası: {error_message}")
    
    def check_for_updates_silent(self):
        """Sessiz güncelleme kontrolü (otomatik)"""
        if not hasattr(self, 'update_checker') or not self.update_checker.isRunning():
            self.update_checker = UpdateChecker()
            self.update_checker.update_available.connect(self.show_update_dialog)
            self.update_checker.no_update.connect(lambda: None)  # Sessiz
            self.update_checker.error_occurred.connect(lambda msg: logger.error(f"Otomatik güncelleme hatası: {msg}"))
            self.update_checker.start()
    
    def check_startup_update(self):
        """Program başlarken GitHub versiyon kontrolü"""
        try:
            logger.log_system_event("Başlangıç versiyon kontrolü", "Başlatılıyor")
            
            # Kullanıcıya bilgilendirme mesajı göster
            self.show_update_check_notification()
            
            # GitHub API'den version.txt dosyasını al
            response = requests.get(Config.UPDATE_CHECK_URL, timeout=10)
            
            if response.status_code == 200:
                data = response.json()
                
                # Base64 encoded content'i decode et
                import base64
                content = data.get('content', '')
                if content:
                    # Base64 decode
                    decoded_content = base64.b64decode(content).decode('utf-8')
                    # Satırları ayır ve versiyon numarasını bul
                    lines = decoded_content.strip().split('\n')
                    latest_version = None
                    
                    for line in lines:
                        if line.startswith('Version:'):
                            latest_version = line.replace('Version:', '').strip()
                            break
                    
                    if latest_version:
                        # Debug: API yanıtını logla
                        logger.log_system_event("GitHub API Yanıtı", f"Version.txt içeriği: {decoded_content.strip()}")
                        
                        # Versiyon karşılaştırması
                        if self.compare_versions(latest_version, Config.VERSION):
                            # Yeni sürüm mevcut - kullanıcıya sor
                            download_url = f"{Config.GITHUB_REPO_URL}/releases/latest"
                            self.show_startup_update_dialog(latest_version, download_url)
                        else:
                            logger.log_system_event("Başlangıç versiyon kontrolü", f"Güncel sürüm kullanılıyor (v{Config.VERSION}) - GitHub: {latest_version}")
                            # Güncel sürüm bilgisi göster
                            self.show_current_version_notification()
                    else:
                        logger.error("Version.txt dosyasında versiyon numarası bulunamadı")
                        self.show_update_error_notification("Version.txt dosyası okunamadı")
                else:
                    logger.error("Version.txt dosyası boş")
                    self.show_update_error_notification("Version.txt dosyası boş")
            else:
                logger.error(f"GitHub API hatası: HTTP {response.status_code}")
                self.show_update_error_notification("GitHub sunucusuna ulaşılamadı")
                
        except requests.exceptions.Timeout:
            logger.error("Başlangıç versiyon kontrolü zaman aşımı")
            self.show_update_error_notification("Bağlantı zaman aşımına uğradı")
        except requests.exceptions.ConnectionError:
            logger.error("Başlangıç versiyon kontrolü bağlantı hatası")
            self.show_update_error_notification("İnternet bağlantısı yok")
        except Exception as e:
            logger.error(f"Başlangıç versiyon kontrolü hatası: {str(e)}")
            self.show_update_error_notification(f"Kontrol hatası: {str(e)}")
    
    def show_update_check_notification(self):
        """Güncelleme kontrolü başladığında bilgilendirme"""
        try:
            msg_box = QMessageBox(self)
            msg_box.setWindowTitle("🔄 Güncelleme Kontrolü")
            msg_box.setIcon(QMessageBox.Information)
            msg_box.setText("🔄 GitHub'dan güncelleme kontrol ediliyor...\n\n"
                          "📡 Sunucuya bağlanılıyor...\n"
                          "🔍 Sürüm bilgileri alınıyor...\n"
                          "⚡ Lütfen bekleyin...")
            msg_box.setStandardButtons(QMessageBox.Ok)
            msg_box.setModal(False)  # Modal olmayan dialog
            msg_box.show()
            
            # 2 saniye sonra otomatik kapat
            QTimer.singleShot(2000, msg_box.close)
            
        except Exception as e:
            logger.error(f"Güncelleme kontrolü bilgilendirme hatası: {str(e)}")
    
    def show_current_version_notification(self):
        """Güncel sürüm bilgisi göster"""
        try:
            msg_box = QMessageBox(self)
            msg_box.setWindowTitle("✅ Güncel Sürüm")
            msg_box.setIcon(QMessageBox.Information)
            msg_box.setText(f"🎉 Tebrikler! Güncel sürümü kullanıyorsunuz.\n\n"
                          f"📋 Sürüm: v{Config.VERSION}\n"
                          f"📅 Kontrol: {datetime.now().strftime('%d.%m.%Y %H:%M')}\n"
                          f"🔗 Repository: {Config.GITHUB_REPO_URL}\n\n"
                          f"💡 İpucu: Güncellemeler otomatik olarak kontrol edilir.")
            msg_box.setStandardButtons(QMessageBox.Ok)
            msg_box.setModal(False)
            msg_box.show()
            
            # 4 saniye sonra otomatik kapat
            QTimer.singleShot(4000, msg_box.close)
            
        except Exception as e:
            logger.error(f"Güncel sürüm bilgilendirme hatası: {str(e)}")
    
    def show_update_error_notification(self, error_message):
        """Güncelleme hatası bilgilendirmesi"""
        try:
            msg_box = QMessageBox(self)
            msg_box.setWindowTitle("⚠️ Güncelleme Hatası")
            msg_box.setIcon(QMessageBox.Warning)
            msg_box.setText(f"Güncelleme kontrolü sırasında hata oluştu:\n\n"
                          f"❌ {error_message}\n\n"
                          f"Program normal şekilde çalışmaya devam edecek.")
            msg_box.setStandardButtons(QMessageBox.Ok)
            msg_box.setModal(False)
            msg_box.show()
            
            # 4 saniye sonra otomatik kapat
            QTimer.singleShot(4000, msg_box.close)
            
        except Exception as e:
            logger.error(f"Güncelleme hatası bilgilendirme hatası: {str(e)}")
    
    def compare_versions(self, latest, current):
        """Versiyon karşılaştırması"""
        try:
            latest_parts = [int(x) for x in latest.split('.')]
            current_parts = [int(x) for x in current.split('.')]
            
            for i in range(max(len(latest_parts), len(current_parts))):
                latest_part = latest_parts[i] if i < len(latest_parts) else 0
                current_part = current_parts[i] if i < len(current_parts) else 0
                
                if latest_part > current_part:
                    return True
                elif latest_part < current_part:
                    return False
            
            return False
        except:
            return False
    
    def show_startup_update_dialog(self, latest_version, download_url):
        """Program başlarken güncelleme dialog'u"""
        try:
            # Basit bir dialog oluştur
            msg_box = QMessageBox(self)
            msg_box.setWindowTitle("🔄 Yeni Güncelleme Bulundu")
            msg_box.setIcon(QMessageBox.Information)
            
            msg_box.setText(f"GitHub deposunda yeni bir sürüm bulundu!\n\n"
                          f"Mevcut Sürüm: v{Config.VERSION}\n"
                          f"Yeni Sürüm: v{latest_version}\n\n"
                          f"Güncellemeyi indirmek istiyor musunuz?")
            
            msg_box.setStandardButtons(QMessageBox.Yes | QMessageBox.No)
            msg_box.setDefaultButton(QMessageBox.Yes)
            
            # Buton metinlerini özelleştir
            yes_button = msg_box.button(QMessageBox.Yes)
            yes_button.setText("Evet, İndir")
            yes_button.setStyleSheet("""
                QPushButton {
                    background-color: #28a745;
                    color: white;
                    border: none;
                    padding: 8px 16px;
                    border-radius: 4px;
                    font-weight: bold;
                }
                QPushButton:hover {
                    background-color: #218838;
                }
            """)
            
            no_button = msg_box.button(QMessageBox.No)
            no_button.setText("Hayır, Devam Et")
            no_button.setStyleSheet("""
                QPushButton {
                    background-color: #6c757d;
                    color: white;
                    border: none;
                    padding: 8px 16px;
                    border-radius: 4px;
                    font-weight: bold;
                }
                QPushButton:hover {
                    background-color: #5a6268;
                }
            """)
            
            # Dialog'u göster
            result = msg_box.exec_()
            
            if result == QMessageBox.Yes:
                # GitHub linkine yönlendir
                import webbrowser
                webbrowser.open(download_url)
                
                # Bilgilendirme mesajı
                QMessageBox.information(self, "✅ İndirme Başlatıldı", 
                    "GitHub sayfası tarayıcıda açıldı.\n\n"
                    "📋 Kurulum Adımları:\n"
                    "1. Dosyayı indirin\n"
                    "2. Mevcut uygulamayı kapatın\n"
                    "3. İndirilen dosyayı çalıştırın\n"
                    "4. Kurulumu tamamlayın\n\n"
                    "⚠️ Önemli: Kurulum öncesi verilerinizi yedekleyin!")
                
                logger.log_system_event("Başlangıç güncelleme indirme", f"v{latest_version}")
            else:
                logger.log_system_event("Başlangıç güncelleme reddedildi", f"v{latest_version}")
                
        except Exception as e:
            logger.error(f"Başlangıç güncelleme dialog hatası: {str(e)}")
    

    
    def show_feedback(self):
        """Geri bildirim dialog göster"""
        dialog = FeedbackDialog(self)
        dialog.exec_()
    
    def check_automated_reminders(self):
        """Otomatik hatırlatmaları kontrol et"""
        try:
            expiring_items = self.automation_system.check_expiring_documents()
            current_date = datetime.now().date()
            
            for item in expiring_items:
                # 30 gün kala hatırlatma oluştur
                if item['days_left'] == 30:
                    self.automation_system.create_reminder(
                        item['plaka'], 
                        item['type'], 
                        item['date']
                    )
                
                # 7 gün kala tekrar hatırlatma
                elif item['days_left'] == 7:
                    self.automation_system.create_reminder(
                        item['plaka'], 
                        f"{item['type']} - ACİL", 
                        item['date']
                    )
                
                # Bugün dolacak olanlar için acil hatırlatma
                elif item['days_left'] == 0:
                    self.automation_system.create_reminder(
                        item['plaka'], 
                        f"{item['type']} - BUGÜN DOLUYOR!", 
                        item['date']
                    )
                    
        except Exception as e:
            print(f"Otomatik hatırlatma kontrolü hatası: {e}")
    
    def show_about(self):
        """Hakkında dialog"""
        QMessageBox.about(self, "Hakkında", 
            f"{Config.APP_NAME} v{Config.VERSION}\n\n"
            f"Geliştirici: {Config.DEVELOPER}\n"
            f"E-posta: {Config.EMAIL}\n\n"
            "Araç filo yönetim sistemi")

# =============================================================================
# PANEL SINIFLARI
# =============================================================================

class DashboardPanel(QWidget):
    """Ana sayfa paneli"""
    
    def __init__(self, data_manager):
        super().__init__()
        self.data_manager = data_manager
        self.init_ui()
        self.load_data()
    
    def init_ui(self):
        """Modern ve çekici UI başlat"""
        layout = QVBoxLayout(self)
        layout.setContentsMargins(15, 15, 15, 15)
        layout.setSpacing(12)
        
        # Hoş geldin bölümü
        self.create_welcome_section(layout)
        
        # Özet kartları
        self.create_summary_cards(layout)
        
        # Alt bölüm - Yaklaşan işlemler tam genişlik
        self.create_expiring_documents(layout)
    
    def create_welcome_section(self, parent_layout):
        """Hoş geldin bölümü"""
        welcome_frame = QFrame()
        welcome_frame.setStyleSheet("""
            QFrame {
                background: qlineargradient(x1:0, y1:0, x2:1, y2:0, 
                    stop:0 #667eea, stop:1 #764ba2);
                border-radius: 10px;
                padding: 12px;
            }
        """)
        
        welcome_layout = QHBoxLayout(welcome_frame)
        
        # Sol taraf - Başlık ve açıklama
        left_layout = QVBoxLayout()
        
        # Ana başlık
        welcome_title = QLabel("🚗 Araç Filo Yönetim Sistemi")
        welcome_title.setStyleSheet("""
            font-size: 16px; 
            font-weight: bold; 
            color: white; 
            margin-bottom: 4px;
        """)
        
        # Alt başlık
        welcome_subtitle = QLabel("Filonuzu profesyonelce yönetin")
        welcome_subtitle.setStyleSheet("""
            font-size: 12px; 
            color: rgba(255,255,255,0.9);
            margin-bottom: 2px;
        """)
        
        left_layout.addWidget(welcome_title)
        left_layout.addWidget(welcome_subtitle)
        
        # Sağ taraf - Tarih, saat ve versiyon
        right_layout = QVBoxLayout()
        right_layout.setAlignment(Qt.AlignmentFlag.AlignRight | Qt.AlignmentFlag.AlignTop)
        
        # Tarih
        current_date = QLabel(datetime.now().strftime("%d.%m.%Y"))
        current_date.setStyleSheet("""
            font-size: 12px; 
            font-weight: bold; 
            color: white;
            text-align: right;
            margin-bottom: 2px;
        """)
        
        # Saat
        current_time = QLabel(datetime.now().strftime("%H:%M"))
        current_time.setStyleSheet("""
            font-size: 10px; 
            color: rgba(255,255,255,0.8);
            text-align: right;
            margin-bottom: 2px;
        """)
        
        # Versiyon
        version_info = QLabel(f"v{Config.VERSION}")
        version_info.setStyleSheet("""
            font-size: 9px; 
            color: rgba(255,255,255,0.7);
            text-align: right;
        """)
        
        right_layout.addWidget(current_date)
        right_layout.addWidget(current_time)
        right_layout.addWidget(version_info)
        
        welcome_layout.addLayout(left_layout)
        welcome_layout.addStretch()  # Boşluk ekle
        welcome_layout.addLayout(right_layout)
        
        parent_layout.addWidget(welcome_frame)
    
    def create_quick_actions(self, parent_layout):
        """Hızlı işlem butonları"""
        actions_frame = QFrame()
        actions_frame.setStyleSheet("""
            QFrame {
                background-color: white;
                border: 1px solid #e0e0e0;
                border-radius: 8px;
                padding: 10px;
            }
        """)
        
        actions_layout = QVBoxLayout(actions_frame)
        
        # Başlık
        title = QLabel("⚡ Hızlı İşlemler")
        title.setStyleSheet("""
            font-size: 14px; 
            font-weight: bold; 
            color: #333333; 
            margin-bottom: 8px;
        """)
        actions_layout.addWidget(title)
        
        # Butonlar
        buttons_layout = QHBoxLayout()
        buttons_layout.setSpacing(8)
        
        quick_buttons = [
            ("🚗 Araç Ekle", self.quick_add_vehicle, "#3498db"),
            ("👨‍✈️ Sürücü Ekle", self.quick_add_driver, "#e74c3c"),
            ("🔧 Bakım Ekle", self.quick_add_maintenance, "#f39c12"),
            ("⛽ Yakıt Ekle", self.quick_add_fuel, "#27ae60"),
            ("📊 Rapor Oluştur", self.quick_create_report, "#9b59b6")
        ]
        
        for text, callback, color in quick_buttons:
            btn = QPushButton(text)
            btn.setStyleSheet(f"""
                QPushButton {{
                    background-color: {color};
                    color: white;
                    border: none;
                    border-radius: 6px;
                    padding: 8px 12px;
                    font-size: 11px;
                    font-weight: bold;
                }}
                QPushButton:hover {{
                    background-color: {color}dd;
                }}
                QPushButton:pressed {{
                    background-color: {color}aa;
                }}
            """)
            btn.clicked.connect(callback)
            buttons_layout.addWidget(btn)
        
        actions_layout.addLayout(buttons_layout)
        parent_layout.addWidget(actions_frame)
    
    def quick_add_vehicle(self):
        """Hızlı araç ekleme"""
        # Ana pencereye sinyal gönder
        if hasattr(self.parent(), 'show_panel'):
            self.parent().show_panel('araclar')
    
    def quick_add_driver(self):
        """Hızlı sürücü ekleme"""
        if hasattr(self.parent(), 'show_panel'):
            self.parent().show_panel('suruculer')
    
    def quick_add_maintenance(self):
        """Hızlı bakım ekleme"""
        if hasattr(self.parent(), 'show_panel'):
            self.parent().show_panel('bakimlar')
    
    def quick_add_fuel(self):
        """Hızlı yakıt ekleme"""
        if hasattr(self.parent(), 'show_panel'):
            self.parent().show_panel('yakitlar')
    
    def quick_create_report(self):
        """Hızlı rapor oluşturma"""
        if hasattr(self.parent(), 'show_panel'):
            self.parent().show_panel('raporlar')
    
    def quick_backup_data(self):
        """Hızlı veri yedekleme"""
        if hasattr(self.parent(), 'backup_data'):
            self.parent().backup_data()
    
    def create_summary_cards(self, parent_layout):
        """Yan yana özet kartları - Her biri ayrı başlıklı"""
        # Ana container
        summary_container = QFrame()
        summary_container.setStyleSheet("""
            QFrame {
                background-color: #f8f9fa;
                border: 1px solid #dee2e6;
                border-radius: 8px;
                padding: 15px;
            }
        """)
        
        container_layout = QVBoxLayout(summary_container)
        container_layout.setSpacing(15)
        container_layout.setContentsMargins(15, 15, 15, 15)
        
        # Ana başlık
        title_label = QLabel("📊 ÖZET BİLGİLER")
        title_label.setStyleSheet("""
            font-size: 16px;
                font-weight: bold;
            color: #2c3e50;
            margin-bottom: 10px;
        """)
        title_label.setAlignment(Qt.AlignmentFlag.AlignCenter)
        container_layout.addWidget(title_label)
        
        # Kartlar için yatay layout
        cards_layout = QHBoxLayout()
        cards_layout.setSpacing(0)
        cards_layout.setContentsMargins(0, 0, 0, 0)
        
        # Kart konfigürasyonları
        card_configs = [
            ("🚗", "Araçlar", "#3498db", "araclar"),
            ("💰", "Bakım & Onarım", "#e74c3c", "bakimlar"),
            ("⛽", "Yakıt Harcamaları", "#f39c12", "yakitlar"),
            ("👨‍✈️", "Sürücüler", "#27ae60", "suruculer")
        ]
        
        self.summary_cards = {}
        
        for i, (icon, title, color, data_type) in enumerate(card_configs):
            card = self.create_summary_card(icon, title, color, data_type)
            self.summary_cards[title] = card
            cards_layout.addWidget(card)
            
            # Son kart değilse dikey çizgi ekle
            if i < len(card_configs) - 1:
                separator = QFrame()
                separator.setFrameShape(QFrame.Shape.VLine)
                separator.setStyleSheet("""
                    QFrame {
                        background-color: #dee2e6;
                        max-width: 1px;
                        min-width: 1px;
                    }
                """)
                cards_layout.addWidget(separator)
        
        container_layout.addLayout(cards_layout)
        parent_layout.addWidget(summary_container)
    
    def create_summary_card(self, icon, title, color, data_type):
        """Tek bir özet kartı oluştur"""
        card = QFrame()
        card.setStyleSheet(f"""
            QFrame {{
                background-color: white;
                border: 1px solid #dee2e6;
                border-radius: 6px;
                padding: 12px;
                min-width: 200px;
                max-width: 250px;
            }}
        """)
        
        layout = QVBoxLayout(card)
        layout.setSpacing(8)
        layout.setContentsMargins(12, 12, 12, 12)
        
        # Başlık ve Yıl seçimi yan yana - basit düzen
        header_layout = QHBoxLayout()
        header_layout.setSpacing(8)
        header_layout.setContentsMargins(0, 0, 0, 5)
        
        # Başlık
        title_label = QLabel(f"{icon} {title}")
        title_label.setStyleSheet("""
            font-size: 13px;
            font-weight: bold;
            color: #2c3e50;
        """)
        header_layout.addWidget(title_label)
        header_layout.addStretch()
        
        # Yıl seçimi - sadece dropdown
        year_combo = QComboBox()
        year_combo.setStyleSheet("""
            QComboBox {
                background-color: white;
                border: 1px solid #ccc;
                border-radius: 2px;
                padding: 1px 2px;
                font-size: 8px;
                min-width: 35px;
                max-width: 45px;
                height: 16px;
            }
        """)
        current_year = datetime.now().year
        years = [str(y) for y in range(current_year, current_year-5, -1)]
        year_combo.addItems(years)
        year_combo.setCurrentText(str(current_year))
        
        header_layout.addWidget(year_combo)
        layout.addLayout(header_layout)
        
        # Toplam satırı (sadece bakım kartı için)
        total_label = QLabel("")
        total_label.setStyleSheet("""
            font-size: 11px;
            font-weight: bold;
            color: #2c3e50;
            padding: 3px 0px;
            border-top: 1px solid #dee2e6;
            margin-top: 2px;
        """)
        total_label.setAlignment(Qt.AlignmentFlag.AlignCenter)
        total_label.setVisible(False)  # Başlangıçta gizli
        layout.addWidget(total_label)
        
        # Veri listesi
        data_list = QListWidget()
        data_list.setStyleSheet("""
            QListWidget {
                background-color: transparent;
                border: none;
                font-size: 11px;
                padding: 0px;
            }
            QListWidget::item {
                padding: 4px 0px;
                border: none;
                color: #495057;
            }
        """)
        data_list.setMaximumHeight(120)
        layout.addWidget(data_list)
        
        # Kart referanslarını sakla
        card.data_list = data_list
        card.year_combo = year_combo
        card.data_type = data_type
        card.total_label = total_label
        
        # Yıl değişikliğini dinle
        year_combo.currentTextChanged.connect(lambda: self.update_card_data(card))
        
        # İlk veriyi yükle
        QTimer.singleShot(100, lambda: self.update_card_data(card))
        
        return card
    
    def update_card_data(self, card):
        """Kart verilerini güncelle"""
        try:
            year = int(card.year_combo.currentText())
            data_type = card.data_type
            card.data_list.clear()
            
            def tr_money(val):
                return f"{val:,.2f}".replace(",", "X").replace(".", ",").replace("X", ".") + " TL"
            
            if data_type == "araclar":
                # Toplam satırını gizle
                card.total_label.setVisible(False)
                
                # Araçlar verisi
                vehicles_data = self.data_manager.load_data('araclar')
                if not vehicles_data.empty:
                    active_vehicles = vehicles_data[vehicles_data['Durum'] == 'Aktif']
                    total_active = len(active_vehicles)
                    
                    if 'Eklenme Tarihi' in vehicles_data.columns:
                        vehicles_data['Eklenme Tarihi'] = pd.to_datetime(vehicles_data['Eklenme Tarihi'], errors='coerce')
                        year_vehicles = vehicles_data[vehicles_data['Eklenme Tarihi'].dt.year == year]
                        year_count = len(year_vehicles)
                    else:
                        year_count = 0
                    
                    card.data_list.addItem(f"Aktif: {total_active}")
                    card.data_list.addItem(f"{year}: {year_count} eklendi")
                else:
                    card.data_list.addItem("Veri bulunamadı")
            
            elif data_type == "bakimlar":
                # Bakım & Onarım verisi
                maintenance_data = self.data_manager.load_data('bakimlar')
                giderler_data = self.data_manager.load_data('giderler')
                
                # Debug bilgileri
                print(f"Bakım verisi yüklendi: {len(maintenance_data)} satır")
                if not maintenance_data.empty:
                    print(f"Bakım sütunları: {list(maintenance_data.columns)}")
                    print(f"Bakım verisi örnek: {maintenance_data.head()}")
                
                # Bakım & Onarım giderleri (bakimlar.xlsx dosyasından)
                total_maintenance = 0
                if not maintenance_data.empty and 'Bakım Tarihi' in maintenance_data.columns:
                    maintenance_data['Bakım Tarihi'] = pd.to_datetime(maintenance_data['Bakım Tarihi'], errors='coerce')
                    year_maintenance = maintenance_data[maintenance_data['Bakım Tarihi'].dt.year == year]
                    print(f"Seçili yıl ({year}) için bakım verisi: {len(year_maintenance)} satır")
                    total_maintenance = year_maintenance['Tutar'].sum() if 'Tutar' in year_maintenance.columns else 0
                    print(f"Bakım toplam tutar: {total_maintenance}")
                else:
                    print("Bakım verisi boş veya 'Bakım Tarihi' sütunu yok")
                
                # Kasko/Muayene/Sigorta giderleri (giderler.xlsx dosyasından)
                kasko_total = 0
                muayene_total = 0
                sigorta_total = 0
                
                if not giderler_data.empty and 'Tarih' in giderler_data.columns:
                    giderler_data['Tarih'] = pd.to_datetime(giderler_data['Tarih'], errors='coerce')
                    year_giderler = giderler_data[giderler_data['Tarih'].dt.year == year]
                    
                    # Gider türlerine göre ayrıştır
                    kasko_data = year_giderler[year_giderler['Gider Türü'] == 'Kasko']
                    muayene_data = year_giderler[year_giderler['Gider Türü'] == 'Muayene']
                    sigorta_data = year_giderler[year_giderler['Gider Türü'] == 'Sigorta']
                    
                    kasko_total = kasko_data['Tutar'].sum() if 'Tutar' in kasko_data.columns else 0
                    muayene_total = muayene_data['Tutar'].sum() if 'Tutar' in muayene_data.columns else 0
                    sigorta_total = sigorta_data['Tutar'].sum() if 'Tutar' in sigorta_data.columns else 0
                
                # Toplam hesapla ve göster
                total_all = total_maintenance + kasko_total + muayene_total + sigorta_total
                card.total_label.setText(f"Toplam: {tr_money(total_all)}")
                card.total_label.setVisible(True)
                
                # Bakım & Onarım bilgileri
                card.data_list.addItem("Bakım & Onarım:")
                card.data_list.addItem(f"  {tr_money(total_maintenance)}")
                
                # Kasko/Muayene/Sigorta bilgileri
                card.data_list.addItem("")
                card.data_list.addItem("Sigorta Giderleri:")
                if kasko_total > 0:
                    card.data_list.addItem(f"  Kasko: {tr_money(kasko_total)}")
                if muayene_total > 0:
                    card.data_list.addItem(f"  Muayene: {tr_money(muayene_total)}")
                if sigorta_total > 0:
                    card.data_list.addItem(f"  Sigorta: {tr_money(sigorta_total)}")
            
            elif data_type == "yakitlar":
                # Yakıt verisi
                fuel_data = self.data_manager.load_data('yakitlar')
                if not fuel_data.empty and 'Tarih' in fuel_data.columns:
                    fuel_data['Tarih'] = pd.to_datetime(fuel_data['Tarih'], errors='coerce')
                    year_fuel = fuel_data[fuel_data['Tarih'].dt.year == year]
                    
                    # Yakıt türü sütununu bul
                    fuel_type_column = None
                    for col in year_fuel.columns:
                        if 'yakıt' in col.lower() or 'tür' in col.lower() or 'tip' in col.lower() or 'ürün' in col.lower():
                            fuel_type_column = col
                            break
                    
                    total_fuel_cost = 0
                    total_fuel_liter = 0
                    
                    if fuel_type_column and not year_fuel.empty:
                        # Yakıt türlerine göre grupla
                        fuel_types = year_fuel[fuel_type_column].value_counts()
                        
                        for fuel_type, count in fuel_types.items():
                            type_data = year_fuel[year_fuel[fuel_type_column] == fuel_type]
                            type_liter = type_data['Litre'].sum() if 'Litre' in type_data.columns else 0
                            type_cost = type_data['Tutar'].sum() if 'Tutar' in type_data.columns else 0
                            
                            total_fuel_cost += type_cost
                            total_fuel_liter += type_liter
                            
                            card.data_list.addItem(f"{fuel_type}:")
                            card.data_list.addItem(f"  {type_liter:.1f}L - {tr_money(type_cost)}")
                    else:
                        # Yakıt türü sütunu yoksa toplam bilgileri göster
                        total_fuel_cost = year_fuel['Tutar'].sum() if 'Tutar' in year_fuel.columns else 0
                        total_fuel_liter = year_fuel['Litre'].sum() if 'Litre' in year_fuel.columns else 0
                        
                        card.data_list.addItem(f"Toplam Tutar: {tr_money(total_fuel_cost)}")
                        card.data_list.addItem(f"Toplam Litre: {total_fuel_liter:.1f} L")
                    
                    # Toplam satırını göster
                    card.total_label.setText(f"Toplam: {tr_money(total_fuel_cost)} ({total_fuel_liter:.1f}L)")
                    card.total_label.setVisible(True)
                else:
                    card.data_list.addItem("Veri bulunamadı")
                    card.total_label.setVisible(False)
            
            elif data_type == "suruculer":
                # Toplam satırını gizle
                card.total_label.setVisible(False)
                
                # Sürücüler verisi
                drivers_data = self.data_manager.load_data('suruculer')
                if not drivers_data.empty:
                    total_drivers = len(drivers_data)
                    active_drivers = drivers_data[drivers_data['Durum'] == 'Aktif']
                    total_active = len(active_drivers)
                    
                    card.data_list.addItem(f"Toplam: {total_drivers}")
                    card.data_list.addItem(f"Aktif: {total_active}")
                else:
                    card.data_list.addItem("Veri bulunamadı")
            
        except Exception as e:
            print(f"Kart güncelleme hatası: {e}")
            card.data_list.addItem("❌ Hata oluştu")
    

    

    
    def create_info_card(self, title, value, color, with_fuel_table=False):
        """Küçük ve kompakt özet kartı oluşturur, yıl seçici sağ üstte."""
        card = QFrame()
        card.setFrameStyle(QFrame.Box)
        card.setStyleSheet(f"""
            QFrame {{
                background-color: {color};
                border-radius: 5px;
                padding: 7px;
                color: white;
            }}
        """)
        layout = QVBoxLayout(card)
        layout.setAlignment(Qt.AlignmentFlag.AlignCenter)
        layout.setSpacing(6)

        # Başlık ve yıl seçici aynı satırda
        header_layout = QHBoxLayout()
        header_layout.setContentsMargins(0, 0, 0, 0)
        header_layout.setSpacing(3)
        title_label = QLabel(title)
        title_label.setAlignment(Qt.AlignmentFlag.AlignCenter)
        title_label.setStyleSheet("font-size: 12px; font-weight: bold;")
        header_layout.addWidget(title_label, alignment=Qt.AlignmentFlag.AlignLeft | Qt.AlignmentFlag.AlignVCenter)
        year_combo = None
        if title in ["💰 Araçların Yıl İçindeki Masrafları", "⛽ Yıllık Yakıt"]:
            year_combo = QComboBox()
            year_combo.setStyleSheet("background-color: white; color: #2c3e50; font-weight: bold; border-radius: 4px; padding: 1px 6px; font-size: 10px;")
            current_year = datetime.now().year
            years = [str(y) for y in range(current_year, current_year-10, -1)]
            year_combo.addItems(years)
            year_combo.setCurrentText(str(current_year))
            if title == "💰 Araçların Yıl İçindeki Masrafları":
                self.maintenance_year_combo = year_combo
                year_combo.currentTextChanged.connect(self.update_maintenance_card)
                # İlk değeri yükle
                QTimer.singleShot(100, self.update_maintenance_card)
            else:
                self.fuel_year_combo = year_combo
                year_combo.currentTextChanged.connect(self.update_fuel_details_label)
                # İlk değeri yükle
                QTimer.singleShot(100, self.update_fuel_details_label)
            header_layout.addWidget(year_combo, alignment=Qt.AlignmentFlag.AlignRight | Qt.AlignmentFlag.AlignVCenter)
        else:
            header_layout.addStretch()
        layout.addLayout(header_layout)

        # Ana değer
        value_label = QLabel(value)
        value_label.setAlignment(Qt.AlignmentFlag.AlignCenter)
        value_label.setStyleSheet("font-size: 15px; font-weight: bold; margin-bottom: 2px;")
        value_label.setObjectName("value_label")
        layout.addWidget(value_label)

        # Alt bilgi veya detay
        if with_fuel_table:
            self.fuel_details_label = QLabel()
            self.fuel_details_label.setAlignment(Qt.AlignmentFlag.AlignCenter)
            self.fuel_details_label.setStyleSheet("font-size: 10px; color: #eaeaea; margin-top: 2px;")
            layout.addWidget(self.fuel_details_label)
        elif title == "🚗 Tüm Araçlar":
            self.vehicle_status_sub = QLabel()
            self.vehicle_status_sub.setAlignment(Qt.AlignmentFlag.AlignCenter)
            self.vehicle_status_sub.setStyleSheet("font-size: 10px; color: #eaeaea; margin-top: 2px;")
            layout.addWidget(self.vehicle_status_sub)
        elif title == "💰 Araçların Yıl İçindeki Masrafları":
            self.maintenance_cost_sub = QLabel()
            self.maintenance_cost_sub.setAlignment(Qt.AlignmentFlag.AlignCenter)
            self.maintenance_cost_sub.setStyleSheet("font-size: 10px; color: #eaeaea; margin-top: 2px;")
            layout.addWidget(self.maintenance_cost_sub)
        return card
    
    def create_upcoming_events(self, parent_layout):
        """Yaklaşan işlemler"""
        group = QGroupBox("📅 Yaklaşan İşlemler")
        group.setStyleSheet("""
            QGroupBox {
                font-weight: bold;
                border: 2px solid #27ae60;
                border-radius: 8px;
                margin-top: 10px;
                padding-top: 10px;
                background-color: #f0f8f0;
            }
            QGroupBox::title {
                subcontrol-origin: margin;
                left: 10px;
                padding: 0 8px 0 8px;
                color: #27ae60;
            }
        """)
        
        layout = QVBoxLayout(group)
        
        self.upcoming_table = QTableWidget()
        self.upcoming_table.setColumnCount(4)
        self.upcoming_table.setHorizontalHeaderLabels(["Tarih", "Tür", "Araç", "Açıklama"])
        self.upcoming_table.horizontalHeader().setStretchLastSection(True)
        self.upcoming_table.setAlternatingRowColors(True)
        
        layout.addWidget(self.upcoming_table)
        parent_layout.addWidget(group)
    
    def create_recent_activities(self, parent_layout):
        """Son aktiviteler - sağ taraf"""
        activities_frame = QFrame()
        activities_frame.setStyleSheet("""
            QFrame {
                background-color: white;
                border: 1px solid #e0e0e0;
                border-radius: 8px;
                padding: 10px;
            }
        """)
        activities_frame.setMinimumWidth(250)
        activities_frame.setMaximumWidth(280)
        
        activities_layout = QVBoxLayout(activities_frame)
        
        # Başlık
        title = QLabel("📋 Son Aktiviteler")
        title.setStyleSheet("""
            font-size: 12px; 
                font-weight: bold;
            color: #333333; 
            margin-bottom: 6px;
        """)
        activities_layout.addWidget(title)
        
        # Aktivite listesi
        self.activities_list = QListWidget()
        self.activities_list.setStyleSheet("""
            QListWidget {
                border: none;
                background-color: transparent;
                font-size: 9px;
            }
            QListWidget::item {
                padding: 3px;
                border-bottom: 1px solid #f0f0f0;
            }
        """)
        
        # Örnek aktiviteler
        sample_activities = [
            "🚗 34ABC123 eklendi (2 saat önce)",
            "🔧 34DEF456 bakım (1 gün önce)",
            "⛽ 34GHI789 yakıt (2 gün önce)",
            "👨‍✈️ Ahmet Yılmaz güncellendi (3 gün önce)",
            "📊 Aylık rapor (1 hafta önce)"
        ]
        
        for activity in sample_activities:
            item = QListWidgetItem(activity)
            self.activities_list.addItem(item)
        
        activities_layout.addWidget(self.activities_list)
        
        parent_layout.addWidget(activities_frame)
    
    def create_expiring_documents(self, parent_layout):
        """Yaklaşan işlemler - sol taraf"""
        expiring_frame = QFrame()
        expiring_frame.setStyleSheet("""
            QFrame {
                background-color: white;
                border: 1px solid #e0e0e0;
                border-radius: 8px;
                padding: 12px;
            }
        """)
        expiring_frame.setMinimumWidth(400)
        
        expiring_layout = QVBoxLayout(expiring_frame)
        
        # Başlık
        title = QLabel("⏰ Yaklaşan İşlemler")
        title.setStyleSheet("""
            font-size: 14px; 
                font-weight: bold;
            color: #333333; 
            margin-bottom: 8px;
        """)
        expiring_layout.addWidget(title)
        
        # Yan yana 3 sütun layout
        cards_layout = QHBoxLayout()
        cards_layout.setSpacing(8)
        
        # Muayene kartı
        muayene_frame = QFrame()
        muayene_frame.setStyleSheet("""
            QFrame {
                border: 2px solid #f39c12;
                border-radius: 6px;
                background-color: #fef9e7;
                padding: 8px;
            }
        """)
        muayene_layout = QVBoxLayout(muayene_frame)
        
        muayene_title = QLabel("🔧 Muayenesi Yaklaşan Araçlar")
        muayene_title.setStyleSheet("""
            font-size: 8px; 
            font-weight: bold; 
            color: #f39c12;
            text-align: center;
            padding: 2px;
            max-height: 20px;
        """)
        
        self.muayene_list = QListWidget()
        self.muayene_list.setMaximumHeight(140)
        self.muayene_list.setStyleSheet("""
            QListWidget {
                border: none;
                background-color: transparent;
                font-size: 8px;
            }
        """)
        
        muayene_layout.addWidget(muayene_title)
        muayene_layout.addWidget(self.muayene_list)
        
        # Sigorta kartı
        sigorta_frame = QFrame()
        sigorta_frame.setStyleSheet("""
            QFrame {
                border: 2px solid #e74c3c;
                border-radius: 6px;
                background-color: #fdf2f2;
                padding: 8px;
            }
        """)
        sigorta_layout = QVBoxLayout(sigorta_frame)
        
        sigorta_title = QLabel("🛡️ Sigortası Yaklaşan Araçlar")
        sigorta_title.setStyleSheet("""
            font-size: 8px; 
            font-weight: bold; 
            color: #e74c3c;
            text-align: center;
            padding: 2px;
            max-height: 20px;
        """)
        
        self.sigorta_list = QListWidget()
        self.sigorta_list.setMaximumHeight(140)
        self.sigorta_list.setStyleSheet("""
            QListWidget {
                border: none;
                background-color: transparent;
                font-size: 8px;
            }
        """)
        
        sigorta_layout.addWidget(sigorta_title)
        sigorta_layout.addWidget(self.sigorta_list)
        
        # Kasko kartı
        kasko_frame = QFrame()
        kasko_frame.setStyleSheet("""
            QFrame {
                border: 2px solid #9b59b6;
                border-radius: 6px;
                background-color: #f8f4fd;
                padding: 8px;
            }
        """)
        kasko_layout = QVBoxLayout(kasko_frame)
        
        kasko_title = QLabel("🛡️ Kaskosu Yaklaşan Araçlar")
        kasko_title.setStyleSheet("""
            font-size: 8px; 
            font-weight: bold; 
            color: #9b59b6;
            text-align: center;
            padding: 2px;
            max-height: 20px;
        """)
        
        self.kasko_list = QListWidget()
        self.kasko_list.setMaximumHeight(140)
        self.kasko_list.setStyleSheet("""
            QListWidget {
                border: none;
                background-color: transparent;
                font-size: 8px;
            }
        """)
        
        kasko_layout.addWidget(kasko_title)
        kasko_layout.addWidget(self.kasko_list)
        
        # Kartları yan yana ekle
        cards_layout.addWidget(muayene_frame)
        cards_layout.addWidget(sigorta_frame)
        cards_layout.addWidget(kasko_frame)
        
        expiring_layout.addLayout(cards_layout)
        
        parent_layout.addWidget(expiring_frame)
    
    def load_data(self):
        """Veri yükle"""
        try:
            # Araç verileri
            vehicles = self.data_manager.load_data('araclar')
            vehicle_count = len(vehicles) if not vehicles.empty else 0
            
            # Araç durumlarına göre dağılım
            vehicle_statuses = {}
            if not vehicles.empty:
                vehicle_statuses = vehicles['Durum'].value_counts().to_dict()
            
            # Durum dağılımını metin olarak oluştur
            status_text = f"Toplam: {vehicle_count}"
            for status, count in vehicle_statuses.items():
                status_text += f"\n{status}: {count}"
            
            # Masraf detaylarını hesapla
            maintenance = self.data_manager.load_data('bakimlar')
            expenses = self.data_manager.load_data('giderler')
            current_year = datetime.now().year
            
            # Bakım onarım maliyetleri
            maintenance_cost = 0
            if not maintenance.empty:
                for _, record in maintenance.iterrows():
                    try:
                        bakim_tarihi = str(record.get('Bakım Tarihi', ''))
                        if bakim_tarihi and len(bakim_tarihi) >= 4:
                            bakim_yili = int(bakim_tarihi[:4])
                            if bakim_yili == current_year:
                                tutar_str = str(record.get('Tutar', '0'))
                                tutar_str = tutar_str.replace('TL', '').replace('₺', '').replace(',', '').strip()
                                try:
                                    tutar = float(tutar_str)
                                    maintenance_cost += tutar
                                except:
                                    pass
                    except:
                        pass
            
            # Kasko/Muayene/Sigorta giderleri
            insurance_cost = 0
            if not expenses.empty:
                for _, record in expenses.iterrows():
                    try:
                        gider_turu = str(record.get('Gider Türü', '')).strip()
                        # Tam eşleşme kontrolü
                        if gider_turu in ['Kasko', 'Muayene', 'Sigorta']:
                            tarih = str(record.get('Tarih', ''))
                            if tarih and len(tarih) >= 4:
                                gider_yili = int(tarih[:4])
                                if gider_yili == current_year:
                                    tutar_str = str(record.get('Tutar', '0'))
                                    tutar_str = tutar_str.replace('TL', '').replace('₺', '').replace(',', '').strip()
                                    try:
                                        tutar = float(tutar_str)
                                        insurance_cost += tutar
                                    except:
                                        pass
                    except:
                        pass
            
            # NaN kontrolü
            if pd.isna(maintenance_cost) or maintenance_cost is None:
                maintenance_cost = 0
            if pd.isna(insurance_cost) or insurance_cost is None:
                insurance_cost = 0
                
            # Toplam masraf
            total_cost = maintenance_cost + insurance_cost
            
            # Yakıt giderlerini hesapla
            expenses = self.data_manager.load_data('giderler')
            fuel_cost = 0
            
            if not expenses.empty:
                for _, record in expenses.iterrows():
                    try:
                        # Gider türünü kontrol et
                        gider_turu = str(record.get('Gider Türü', '')).lower()
                        if 'yakıt' in gider_turu or 'benzin' in gider_turu or 'dizel' in gider_turu:
                            # Tarihi kontrol et
                            tarih = str(record.get('Tarih', ''))
                            if tarih and len(tarih) >= 4:
                                gider_yili = int(tarih[:4])
                                if gider_yili == current_year:
                                    # Tutarı hesapla
                                    tutar_str = str(record.get('Tutar', '0'))
                                    tutar_str = tutar_str.replace('TL', '').replace('₺', '').replace(',', '').strip()
                                    try:
                                        tutar = float(tutar_str)
                                        fuel_cost += tutar
                                    except:
                                        pass
                    except:
                        pass
            
            # Yakıt maliyeti NaN kontrolü
            if pd.isna(fuel_cost) or fuel_cost is None:
                fuel_cost = 0
            
            # Süresi Yaklaşan İşlemleri yükle
            self.load_expiring_documents()
            
            # Kartları güncelle (varsa)
            if hasattr(self, 'maintenance_cost_card'):
                maintenance_cost_label = self.maintenance_cost_card.findChild(QLabel, "value_label")
            else:
                maintenance_cost_label = None
                
            if hasattr(self, 'fuel_cost_card'):
                fuel_cost_label = self.fuel_cost_card.findChild(QLabel, "value_label")
            else:
                fuel_cost_label = None
            
            if maintenance_cost_label:
                # Masraf detaylarını oluştur
                def tr_money(val):
                    if pd.isna(val) or val is None:
                        return "0,00"
                    try:
                        return f"{float(val):,.2f}".replace(',', 'X').replace('.', ',').replace('X', '.')
                    except:
                        return "0,00"
                cost_text = f"Toplam: {tr_money(total_cost)} TL\n"
                cost_text += f"Bakım/Onarım: {tr_money(maintenance_cost)} TL\n"
                cost_text += f"Kasko/Muayene/Sigorta: {tr_money(insurance_cost)} TL"
                maintenance_cost_label.setText(cost_text)
            if fuel_cost_label:
                def tr_money(val):
                    if pd.isna(val) or val is None:
                        return "0,00"
                    try:
                        return f"{float(val):,.2f}".replace(',', 'X').replace('.', ',').replace('X', '.')
                    except:
                        return "0,00"
                fuel_cost_label.setText(f"{tr_money(fuel_cost)} TL")
            
            # Yıllık yakıt detayını güncelle (sadece combo box'lar oluşturulduktan sonra)
            if hasattr(self, 'fuel_year_combo') and hasattr(self, 'maintenance_year_combo'):
                self.update_fuel_details_label()
                self.update_maintenance_card()
            
        except Exception as e:
            print(f"Dashboard veri yükleme hatası: {e}")
    
    def update_maintenance_card(self):
        try:
            year = int(self.maintenance_year_combo.currentText()) if hasattr(self, 'maintenance_year_combo') else datetime.now().year
            maintenance = self.data_manager.load_data('bakimlar')
            expenses = self.data_manager.load_data('giderler')
            maintenance_cost = 0
            insurance_cost = 0
            # Bakım/onarım
            if not maintenance.empty:
                for _, record in maintenance.iterrows():
                    try:
                        bakim_tarihi = str(record.get('Bakım Tarihi', ''))
                        if bakim_tarihi and len(bakim_tarihi) >= 4:
                            bakim_yili = int(bakim_tarihi[-4:]) if bakim_tarihi[-4:].isdigit() else int(bakim_tarihi[:4])
                            if bakim_yili == year:
                                tutar_str = str(record.get('Tutar', '0')).replace('TL', '').replace('₺', '').replace(',', '').strip()
                                try:
                                    tutar = float(tutar_str)
                                    maintenance_cost += tutar
                                except:
                                    pass
                    except:
                        pass
            # Kasko/Muayene/Sigorta
            if not expenses.empty:
                for _, record in expenses.iterrows():
                    try:
                        gider_turu = str(record.get('Gider Türü', '')).strip()
                        # Tam eşleşme kontrolü
                        if gider_turu in ['Kasko', 'Muayene', 'Sigorta']:
                            tarih = str(record.get('Tarih', '')).strip()
                            if not tarih or len(tarih) < 4:
                                tarih = str(record.get('Bitiş', '')).strip()
                            gider_yili = None
                            if tarih and len(tarih) >= 4:
                                # Yıl başta mı sonda mı kontrol et
                                if tarih[:4].isdigit():
                                    gider_yili = int(tarih[:4])
                                elif tarih[-4:].isdigit():
                                    gider_yili = int(tarih[-4:])
                            if gider_yili == year:
                                tutar_str = str(record.get('Tutar', '0')).replace('TL', '').replace('₺', '').replace(',', '').strip()
                                try:
                                    tutar = float(tutar_str)
                                except:
                                    tutar = 0
                                insurance_cost += tutar
                    except:
                        pass
            # NaN ve None değerlerini kontrol et
            if pd.isna(maintenance_cost) or maintenance_cost is None:
                maintenance_cost = 0
            if pd.isna(insurance_cost) or insurance_cost is None:
                insurance_cost = 0
                
            total_cost = maintenance_cost + insurance_cost
            
            def tr_money(val):
                if pd.isna(val) or val is None:
                    return "0,00"
                try:
                    return f"{float(val):,.2f}".replace(',', 'X').replace('.', ',').replace('X', '.')
                except:
                    return "0,00"
                    
            cost_text = f"Toplam: {tr_money(total_cost)} TL\n"
            cost_text += f"Bakım/Onarım: {tr_money(maintenance_cost)} TL\n"
            cost_text += f"Kasko/Muayene/Sigorta: {tr_money(insurance_cost)} TL"
            
            # Ana değer label'ını güncelle
            maintenance_cost_label = self.maintenance_cost_card.findChild(QLabel, "value_label")
            if maintenance_cost_label:
                maintenance_cost_label.setText(cost_text)
            
            # Alt bilgi label'ını da güncelle
            if hasattr(self, 'maintenance_cost_sub'):
                self.maintenance_cost_sub.setText(f"{year} yılı toplam masrafı")
        except Exception as e:
            maintenance_cost_label = self.maintenance_cost_card.findChild(QLabel, "value_label")
            if maintenance_cost_label:
                maintenance_cost_label.setText("")
            if hasattr(self, 'maintenance_cost_sub'):
                self.maintenance_cost_sub.setText("")
    
    def update_fuel_details_label(self):
        try:
            year = int(self.fuel_year_combo.currentText()) if hasattr(self, 'fuel_year_combo') else datetime.now().year
            yakitlar = self.data_manager.load_data('yakitlar')
            summary = {"Benzin": {"litre": 0, "tutar": 0}, "Motorin": {"litre": 0, "tutar": 0}}
            if not yakitlar.empty:
                for _, record in yakitlar.iterrows():
                    try:
                        yakit_tipi = str(record.get('Yakıt Tipi', ''))
                        tarih = str(record.get('Tarih', ''))
                        litre = float(str(record.get('Litre', '0')).replace(',', '.'))
                        tutar = float(str(record.get('Tutar', '0')).replace(',', '.'))
                        if tarih and len(tarih) >= 4:
                            yili = int(tarih[:4])
                            if yili == year and yakit_tipi in summary:
                                summary[yakit_tipi]["litre"] += litre
                                summary[yakit_tipi]["tutar"] += tutar
                    except:
                        pass
            def tr_money(val):
                if pd.isna(val) or val is None:
                    return "0,00"
                try:
                    return f"{float(val):,.2f}".replace(',', 'X').replace('.', ',').replace('X', '.')
                except:
                    return "0,00"
                    
            lines = []
            for yakit_tipi in ["Benzin", "Motorin"]:
                litre = summary[yakit_tipi]["litre"]
                tutar = summary[yakit_tipi]["tutar"]
                # NaN kontrolü
                if pd.isna(litre) or litre is None:
                    litre = 0
                if pd.isna(tutar) or tutar is None:
                    tutar = 0
                lines.append(f"{yakit_tipi}: {litre:,.2f} L / {tr_money(tutar)} TL")
            
            # Detay label'ını güncelle
            if hasattr(self, 'fuel_details_label'):
                self.fuel_details_label.setText("\n".join(lines))
            
            # Ana değer label'ını güncelle
            fuel_cost_label = self.fuel_cost_card.findChild(QLabel, "value_label")
            if fuel_cost_label:
                toplam = summary["Benzin"]["tutar"] + summary["Motorin"]["tutar"]
                fuel_cost_label.setText(f"{tr_money(toplam)} TL")
        except Exception as e:
            if hasattr(self, 'fuel_details_label'):
                self.fuel_details_label.setText("")
            fuel_cost_label = self.fuel_cost_card.findChild(QLabel, "value_label")
            if fuel_cost_label:
                fuel_cost_label.setText("")
    
    def load_expiring_documents(self):
        """Süresi Yaklaşan İşlemleri yükle"""
        try:
            vehicles = self.data_manager.load_data('araclar')
            current_date = datetime.now().date()
            
            # Listeleri temizle
            self.muayene_list.clear()
            self.sigorta_list.clear()
            self.kasko_list.clear()
            
            if vehicles.empty:
                self.muayene_list.addItem("Muayenesi yaklaşan araç yok")
                self.sigorta_list.addItem("Sigortası yaklaşan araç yok")
                self.kasko_list.addItem("Kaskosu yaklaşan araç yok")
                return
            
            muayene_count = 0
            sigorta_count = 0
            kasko_count = 0
            
            # Geçmiş belgeler için sayaç
            expired_muayene = 0
            expired_sigorta = 0
            expired_kasko = 0
            
            for _, vehicle in vehicles.iterrows():
                plaka = vehicle.get('Plaka', '')
                if not plaka:
                    continue
                
                # Muayene kontrolü
                muayene_date = self._parse_date(vehicle.get('Muayene Tarihi', ''))
                if muayene_date:
                    days_left = (muayene_date - current_date).days
                    
                    # Geçmiş tarih kontrolü - sürekli uyarı
                    if days_left < 0:
                        expired_muayene += 1
                        days_text = f"{abs(days_left)} gün geçmiş"
                        item_text = f"🚨 {plaka} - {days_text} - YENİ TARİH GİRİN!"
                        
                        list_item = QListWidgetItem(item_text)
                        list_item.setBackground(QColor("#ffcdd2"))  # Kırmızı arka plan
                        list_item.setForeground(QColor("#d32f2f"))  # Kırmızı yazı
                        
                        self.muayene_list.addItem(list_item)
                    # Yaklaşan tarih kontrolü
                    elif 0 <= days_left <= 30:
                        muayene_count += 1
                        days_text = f"{days_left} gün" if days_left > 0 else "BUGÜN"
                        item_text = f"{plaka} - {days_text}"
                        
                        list_item = QListWidgetItem(item_text)
                        if days_left <= 7:
                            list_item.setBackground(QColor("#ffebee"))
                        elif days_left <= 15:
                            list_item.setBackground(QColor("#fff3e0"))
                        else:
                            list_item.setBackground(QColor("#f1f8e9"))
                        
                        self.muayene_list.addItem(list_item)
                
                # Sigorta kontrolü
                sigorta_date = self._parse_date(vehicle.get('Sigorta Bitiş', ''))
                if sigorta_date:
                    days_left = (sigorta_date - current_date).days
                    
                    # Geçmiş tarih kontrolü - sürekli uyarı
                    if days_left < 0:
                        expired_sigorta += 1
                        days_text = f"{abs(days_left)} gün geçmiş"
                        item_text = f"🚨 {plaka} - {days_text} - YENİ TARİH GİRİN!"
                        
                        list_item = QListWidgetItem(item_text)
                        list_item.setBackground(QColor("#ffcdd2"))  # Kırmızı arka plan
                        list_item.setForeground(QColor("#d32f2f"))  # Kırmızı yazı
                        
                        self.sigorta_list.addItem(list_item)
                    # Yaklaşan tarih kontrolü
                    elif 0 <= days_left <= 30:
                        sigorta_count += 1
                        days_text = f"{days_left} gün" if days_left > 0 else "BUGÜN"
                        item_text = f"{plaka} - {days_text}"
                        
                        list_item = QListWidgetItem(item_text)
                        if days_left <= 7:
                            list_item.setBackground(QColor("#ffebee"))
                        elif days_left <= 15:
                            list_item.setBackground(QColor("#fff3e0"))
                        else:
                            list_item.setBackground(QColor("#f1f8e9"))
                        
                        self.sigorta_list.addItem(list_item)
                
                # Kasko kontrolü
                kasko_date = self._parse_date(vehicle.get('Kasko Bitiş', ''))
                if kasko_date:
                    days_left = (kasko_date - current_date).days
                    
                    # Geçmiş tarih kontrolü - sürekli uyarı
                    if days_left < 0:
                        expired_kasko += 1
                        days_text = f"{abs(days_left)} gün geçmiş"
                        item_text = f"🚨 {plaka} - {days_text} - YENİ TARİH GİRİN!"
                        
                        list_item = QListWidgetItem(item_text)
                        list_item.setBackground(QColor("#ffcdd2"))  # Kırmızı arka plan
                        list_item.setForeground(QColor("#d32f2f"))  # Kırmızı yazı
                        
                        self.kasko_list.addItem(list_item)
                    # Yaklaşan tarih kontrolü
                    elif 0 <= days_left <= 30:
                        kasko_count += 1
                        days_text = f"{days_left} gün" if days_left > 0 else "BUGÜN"
                        item_text = f"{plaka} - {days_text}"
                        
                        list_item = QListWidgetItem(item_text)
                        if days_left <= 7:
                            list_item.setBackground(QColor("#ffebee"))
                        elif days_left <= 15:
                            list_item.setBackground(QColor("#fff3e0"))
                        else:
                            list_item.setBackground(QColor("#f1f8e9"))
                        
                        self.kasko_list.addItem(list_item)
            
            # Eğer hiç belge yoksa mesaj göster
            if muayene_count == 0 and expired_muayene == 0:
                self.muayene_list.addItem("Muayenesi yaklaşan araç yok")
            if sigorta_count == 0 and expired_sigorta == 0:
                self.sigorta_list.addItem("Sigortası yaklaşan araç yok")
            if kasko_count == 0 and expired_kasko == 0:
                self.kasko_list.addItem("Kaskosu yaklaşan araç yok")
                
        except Exception as e:
            print(f"Belge yükleme hatası: {e}")
            self.muayene_list.addItem("Yükleme hatası")
            self.sigorta_list.addItem("Yükleme hatası")
            self.kasko_list.addItem("Yükleme hatası")
    
    def _parse_date(self, date_str):
        """Tarih string'ini parse et - Türkçe format (dd.mm.yyyy) için dayfirst=True"""
        try:
            # NaN, None, boş string, float NaN kontrolü
            if pd.isna(date_str) or date_str is None or str(date_str).strip() == '' or str(date_str).lower() == 'nan':
                return None
            
            # String'e çevir ve boşlukları temizle
            date_str = str(date_str).strip()
            if not date_str or date_str.lower() == 'nan':
                return None
                
            return pd.to_datetime(date_str, dayfirst=True).date()
        except Exception as e:
            print(f"Dashboard tarih parse hatası: {date_str} - {str(e)}")
            return None
    
    def create_fuel_summary(self, parent_layout):
        group = QGroupBox("Yakıt Tipine Göre Yıllık Tüketim ve Tutar")
        group.setStyleSheet("""
            QGroupBox {
                font-weight: bold;
                border: 2px solid #f39c12;
                border-radius: 5px;
                margin-top: 10px;
                padding-top: 10px;
            }
        """)
        layout = QVBoxLayout(group)
        self.fuel_table = QTableWidget()
        self.fuel_table.setColumnCount(3)
        self.fuel_table.setHorizontalHeaderLabels(["Yakıt Tipi", "Toplam Litre", "Toplam Tutar (TL)"])
        self.fuel_table.horizontalHeader().setStretchLastSection(True)
        self.fuel_table.setAlternatingRowColors(True)
        layout.addWidget(self.fuel_table)
        parent_layout.addWidget(group)
    
    def load_fuel_summary(self):
        try:
            yakitlar = self.data_manager.load_data('yakitlar')
            current_year = datetime.now().year
            summary = {}
            if not yakitlar.empty:
                for _, record in yakitlar.iterrows():
                    try:
                        yakit_tipi = str(record.get('Yakıt Tipi', 'Diğer'))
                        tarih = str(record.get('Tarih', ''))
                        litre = float(str(record.get('Litre', '0')).replace(',', '.'))
                        tutar = float(str(record.get('Tutar', '0')).replace(',', '.'))
                        if tarih and len(tarih) >= 4:
                            year = int(tarih[:4])
                            if year == current_year:
                                if yakit_tipi not in summary:
                                    summary[yakit_tipi] = {'litre': 0, 'tutar': 0}
                                summary[yakit_tipi]['litre'] += litre
                                summary[yakit_tipi]['tutar'] += tutar
                    except:
                        pass
            self.fuel_table.setRowCount(len(summary))
            for row, (yakit_tipi, vals) in enumerate(summary.items()):
                self.fuel_table.setItem(row, 0, QTableWidgetItem(yakit_tipi))
                self.fuel_table.setItem(row, 1, QTableWidgetItem(f"{vals['litre']:.2f}"))
                self.fuel_table.setItem(row, 2, QTableWidgetItem(f"{vals['tutar']:.2f}"))
        except Exception as e:
            print(f"Yakıt özet yükleme hatası: {e}")

# Diğer panel sınıfları burada tanımlanacak...
class VehiclesPanel(QWidget):
    def __init__(self, data_manager):
        super().__init__()
        self.data_manager = data_manager
        self.init_ui()
        self.load_vehicles()
        self.load_deleted_vehicles()  # Silinen araçları da yükle
    
    def init_ui(self):
        layout = QVBoxLayout(self)
        layout.setContentsMargins(20, 20, 20, 20)
        layout.setSpacing(20)
        
        # Başlık
        header_layout = QHBoxLayout()
        title = QLabel("Araç Yönetimi")
        title.setStyleSheet("font-size: 24px; font-weight: bold; color: #2c3e50;")
        header_layout.addWidget(title)
        header_layout.addStretch()
        layout.addLayout(header_layout)
        
        # Tab Widget oluştur
        self.tab_widget = QTabWidget()
        self.tab_widget.setStyleSheet("""
            QTabWidget::pane {
                border: 1px solid #bdc3c7;
                border-radius: 5px;
                background-color: white;
            }
            QTabBar::tab {
                background-color: #ecf0f1;
                color: #2c3e50;
                padding: 10px 20px;
                margin-right: 2px;
                border-top-left-radius: 5px;
                border-top-right-radius: 5px;
            font-weight: bold; 
            }
            QTabBar::tab:selected {
                background-color: #3498db;
                color: white;
            }
            QTabBar::tab:hover {
                background-color: #2980b9;
                color: white;
            }
        """)
        
        # Aktif Araçlar sekmesi
        self.create_active_vehicles_tab()
        
        # Silinen Araçlar sekmesi
        self.create_deleted_vehicles_tab()
        
        layout.addWidget(self.tab_widget)
        
        # İlk yükleme için timer kullan
        QTimer.singleShot(100, self.load_deleted_vehicles)
    
    def create_active_vehicles_tab(self):
        """Aktif araçlar sekmesi"""
        active_tab = QWidget()
        layout = QVBoxLayout(active_tab)
        layout.setContentsMargins(10, 10, 10, 10)
        layout.setSpacing(15)
        
        # Butonlar
        btn_layout = QHBoxLayout()
        btn_add = QPushButton("➕ Araç Ekle")
        btn_add.setStyleSheet("""
            QPushButton {
                background-color: #27ae60;
                color: white;
                border: none;
                padding: 10px 20px;
                border-radius: 5px;
                font-weight: bold;
            }
            QPushButton:hover {
                background-color: #229954;
            }
        """)
        btn_add.clicked.connect(self.add_vehicle)
        
        btn_edit = QPushButton("✏️ Düzenle")
        btn_edit.setStyleSheet("""
            QPushButton {
                background-color: #3498db;
                color: white;
                border: none;
                padding: 10px 20px;
                border-radius: 5px;
                font-weight: bold;
            }
            QPushButton:hover {
                background-color: #2980b9;
            }
        """)
        btn_edit.clicked.connect(self.edit_vehicle)
        
        btn_delete = QPushButton("🗑️ Sil")
        btn_delete.setStyleSheet("""
            QPushButton {
                background-color: #e74c3c;
                color: white;
                border: none;
                padding: 10px 20px;
                border-radius: 5px;
                font-weight: bold;
            }
            QPushButton:hover {
                background-color: #c0392b;
            }
        """)
        btn_delete.clicked.connect(self.delete_vehicle)
        
        btn_export = QPushButton("📤 Excel'e Aktar")
        btn_export.setStyleSheet("""
            QPushButton {
                background-color: #f1c40f;
                color: #2c3e50;
                border: none;
                padding: 10px 20px;
                border-radius: 5px;
                font-weight: bold;
            }
            QPushButton:hover {
                background-color: #f39c12;
            }
        """)
        btn_export.clicked.connect(lambda: export_table_to_excel(self.model, self.proxy_model, self))
        
        btn_layout.addWidget(btn_add)
        btn_layout.addWidget(btn_edit)
        btn_layout.addWidget(btn_delete)
        btn_layout.addWidget(btn_export)
        btn_layout.addStretch()
        layout.addLayout(btn_layout)
        
        # Arama kutusu
        self.search_box = QLineEdit()
        self.search_box.setPlaceholderText("Aktif araçlarda ara...")
        self.search_box.textChanged.connect(self.filter_vehicles)
        layout.addWidget(self.search_box)
        
        # QTableView ve model
        self.table_view = QTableView()
        self.model = QStandardItemModel()
        self.proxy_model = QSortFilterProxyModel()
        self.proxy_model.setSourceModel(self.model)
        self.proxy_model.setFilterCaseSensitivity(Qt.CaseInsensitive)
        self.table_view.setModel(self.proxy_model)
        self.table_view.setSortingEnabled(True)
        self.table_view.setSelectionBehavior(QTableView.SelectRows)
        self.table_view.setSelectionMode(QTableView.SingleSelection)
        self.table_view.horizontalHeader().setStretchLastSection(True)
        
        # Satır numaralarını gizle
        self.table_view.verticalHeader().setVisible(False)
        
        layout.addWidget(self.table_view)
        
        self.tab_widget.addTab(active_tab, "🚗 Aktif Araçlar")
    
    def create_deleted_vehicles_tab(self):
        """Silinen araçlar sekmesi"""
        deleted_tab = QWidget()
        layout = QVBoxLayout(deleted_tab)
        layout.setContentsMargins(10, 10, 10, 10)
        layout.setSpacing(15)
        
        # Filtreler
        filter_layout = QHBoxLayout()
        
        # Tarih filtresi
        date_label = QLabel("Tarih Aralığı:")
        self.deleted_start_date = QDateEdit()
        self.deleted_start_date.setDate(QDate.currentDate().addMonths(-6))
        self.deleted_end_date = QDateEdit()
        self.deleted_end_date.setDate(QDate.currentDate())
        
        # Sebep filtresi
        reason_label = QLabel("Silme Sebebi:")
        self.deleted_reason_filter = QComboBox()
        self.deleted_reason_filter.addItems(["Tümü", "Hurda", "Satış", "Kaza", "Değişim", "Devir", "Diğer"])
        
        # Arama kutusu
        self.deleted_search_box = QLineEdit()
        self.deleted_search_box.setPlaceholderText("Silinen araçlarda ara...")
        
        filter_layout.addWidget(date_label)
        filter_layout.addWidget(self.deleted_start_date)
        filter_layout.addWidget(QLabel("-"))
        filter_layout.addWidget(self.deleted_end_date)
        filter_layout.addWidget(reason_label)
        filter_layout.addWidget(self.deleted_reason_filter)
        filter_layout.addWidget(self.deleted_search_box)
        filter_layout.addStretch()
        layout.addLayout(filter_layout)
        
        # Butonlar
        btn_layout = QHBoxLayout()
        
        btn_restore = QPushButton("🔄 Geri Yükle")
        btn_restore.setStyleSheet("""
            QPushButton {
                background-color: #27ae60;
                color: white;
                border: none;
                padding: 10px 20px;
                border-radius: 5px;
                font-weight: bold;
            }
            QPushButton:hover {
                background-color: #229954;
            }
        """)
        btn_restore.clicked.connect(self.restore_vehicle)
        
        btn_export_deleted = QPushButton("📤 Excel'e Aktar")
        btn_export_deleted.setStyleSheet("""
            QPushButton {
                background-color: #f1c40f;
                color: #2c3e50;
                border: none;
                padding: 10px 20px;
                border-radius: 5px;
            font-weight: bold; 
            }
            QPushButton:hover {
                background-color: #f39c12;
            }
        """)
        btn_export_deleted.clicked.connect(self.export_deleted_vehicles)
        
        btn_report = QPushButton("📊 Silme Raporu")
        btn_report.setStyleSheet("""
            QPushButton {
                background-color: #9b59b6;
                color: white;
                border: none;
                padding: 10px 20px;
                border-radius: 5px;
                font-weight: bold;
            }
            QPushButton:hover {
                background-color: #8e44ad;
            }
        """)
        btn_report.clicked.connect(self.generate_deletion_report)
        
        btn_layout.addWidget(btn_restore)
        btn_layout.addWidget(btn_export_deleted)
        btn_layout.addWidget(btn_report)
        btn_layout.addStretch()
        layout.addLayout(btn_layout)
        
        # Silinen araçlar tablosu
        self.deleted_table_view = QTableView()
        self.deleted_model = QStandardItemModel()
        self.deleted_proxy_model = QSortFilterProxyModel()
        self.deleted_proxy_model.setSourceModel(self.deleted_model)
        self.deleted_proxy_model.setFilterCaseSensitivity(Qt.CaseInsensitive)
        self.deleted_table_view.setModel(self.deleted_proxy_model)
        self.deleted_table_view.setSortingEnabled(True)
        self.deleted_table_view.setSelectionBehavior(QTableView.SelectRows)
        self.deleted_table_view.setSelectionMode(QTableView.SingleSelection)
        self.deleted_table_view.horizontalHeader().setStretchLastSection(True)
        
        # Satır numaralarını gizle
        self.deleted_table_view.verticalHeader().setVisible(False)
        
        layout.addWidget(self.deleted_table_view)
        
        # Filtre bağlantıları
        self.deleted_start_date.dateChanged.connect(self.load_deleted_vehicles)
        self.deleted_end_date.dateChanged.connect(self.load_deleted_vehicles)
        self.deleted_reason_filter.currentTextChanged.connect(self.load_deleted_vehicles)
        self.deleted_search_box.textChanged.connect(self.filter_deleted_vehicles)
        
        # Test butonu ekle
        btn_test = QPushButton("🔄 Test Yükle")
        btn_test.setStyleSheet("""
            QPushButton {
                background-color: #9b59b6;
                color: white;
                border: none;
                padding: 5px 10px;
                border-radius: 3px;
                font-size: 10px;
            }
            QPushButton:hover {
                background-color: #8e44ad;
            }
        """)
        btn_test.clicked.connect(self.load_deleted_vehicles)
        btn_layout.addWidget(btn_test)
        
        self.tab_widget.addTab(deleted_tab, "🗑️ Silinen Araçlar")
    def load_vehicles(self):
        try:
            vehicles = self.data_manager.load_data('araclar')
            columns = [
                "Plaka", "Marka", "Model", "Yıl", "Şasi No", "Ruhsat Seri No", "Yakıt Tipi", "Araç Tipi", "Son KM", "Durum", "Sigorta Bitiş", "Kasko Var mı?", "Kasko Bitiş", "Muayene Tarihi", "Birim", "Hizmet Kişisi"
            ]
            self.model.clear()
            self.model.setHorizontalHeaderLabels(columns)
            for _, vehicle in vehicles.iterrows():
                row = []
                for idx, col in enumerate(columns):
                    val = str(vehicle.get(col, ''))
                    item = QStandardItem(val)
                    if col in ["Sigorta Bitiş", "Kasko Bitiş", "Muayene Tarihi"] and val:
                        try:
                            date_val = pd.to_datetime(val, dayfirst=False)
                            days_left = (date_val - datetime.now()).days
                            if days_left > 60:
                                item.setBackground(QColor("#b6fcb6"))  # Yeşil
                            elif 30 < days_left <= 60:
                                item.setBackground(QColor("#fff9b1"))  # Sarı
                            elif 0 <= days_left <= 30:
                                item.setBackground(QColor("#ffb3b3"))  # Kırmızı
                            else:
                                item.setBackground(QColor("#e0e0e0"))  # Gri
                        except:
                            pass
                    row.append(item)
                self.model.appendRow(row)
        except Exception as e:
            print(f"Araç yükleme hatası: {e}")
    def filter_vehicles(self, text):
        self.proxy_model.setFilterWildcard(text)
    def add_vehicle(self):
        dialog = VehicleDialog(self)
        if dialog.exec_() == QDialog.Accepted:
            self.load_vehicles()
    def edit_vehicle(self):
        index = self.table_view.currentIndex()
        if index.isValid():
            row = self.proxy_model.mapToSource(index).row()
            plaka = self.model.item(row, 0).text()
            vehicles = self.data_manager.load_data('araclar')
            vehicle = vehicles[vehicles['Plaka'] == plaka]
            if not vehicle.empty:
                dialog = VehicleDialog(self, vehicle.iloc[0].to_dict())
                if dialog.exec_() == QDialog.Accepted:
                    self.load_vehicles()
        else:
            QMessageBox.warning(self, "Uyarı", "Lütfen düzenlenecek aracı seçin.")
    def delete_vehicle(self):
        index = self.table_view.currentIndex()
        if index.isValid():
            row = self.proxy_model.mapToSource(index).row()
            plaka = self.model.item(row, 0).text()
            
            # Gelişmiş silme dialog'u
            dialog = VehicleDeletionDialog(self, plaka)
            if dialog.exec_() == QDialog.Accepted:
                try:
                    # Araç verilerini al
                    vehicles = self.data_manager.load_data('araclar')
                    vehicle_data = vehicles[vehicles['Plaka'] == plaka]
                    
                    if not vehicle_data.empty:
                        vehicle_info = vehicle_data.iloc[0].to_dict()
                        
                        # Silinen araçlar tablosuna ekle
                        self.add_to_deleted_vehicles(vehicle_info, dialog.get_deletion_reason(), dialog.get_deletion_notes())
                        
                        # Ana tablodan çıkar
                        vehicles = vehicles[vehicles['Plaka'] != plaka]
                        self.data_manager.save_data('araclar', vehicles)
                        
                        # İlgili kayıtları güncelle (bakım, yakıt vs.)
                        self.update_related_records(plaka)
                        
                        self.load_vehicles()
                        self.load_deleted_vehicles()  # Silinen araçlar sekmesini güncelle
                        
                        QMessageBox.information(self, "Başarılı", 
                            f"'{plaka}' plakalı araç başarıyla silindi.\n"
                            f"Sebep: {dialog.get_deletion_reason()}")
                    else:
                        QMessageBox.warning(self, "Uyarı", "Araç bulunamadı.")
                        
                except Exception as e:
                    QMessageBox.critical(self, "Hata", f"Araç silme hatası: {str(e)}")
        else:
            QMessageBox.warning(self, "Uyarı", "Lütfen silinecek aracı seçin.")
    
    def add_to_deleted_vehicles(self, vehicle_info, reason, notes):
        """Silinen araçları ayrı tabloya ekle"""
        try:
            deleted_vehicles = self.data_manager.load_data('silinen_araclar')
            
            # Yeni silinen araç kaydı
            deleted_record = {
                'ID': len(deleted_vehicles) + 1 if not deleted_vehicles.empty else 1,
                'Plaka': vehicle_info.get('Plaka', ''),
                'Marka': vehicle_info.get('Marka', ''),
                'Model': vehicle_info.get('Model', ''),
                'Yıl': vehicle_info.get('Yıl', ''),
                'Şasi No': vehicle_info.get('Şasi No', ''),
                'Motor No': vehicle_info.get('Motor No', ''),
                'Araç Tipi': vehicle_info.get('Araç Tipi', ''),
                'Son KM': vehicle_info.get('Son KM', ''),
                'Durum': 'Silindi',
                'Sigorta Bitiş': vehicle_info.get('Sigorta Bitiş', ''),
                'Muayene Tarihi': vehicle_info.get('Muayene Tarihi', ''),
                'Birim': vehicle_info.get('Birim', ''),
                'Hizmet Kişisi': vehicle_info.get('Hizmet Kişisi', ''),
                'Fotoğraf Yolu': vehicle_info.get('Fotoğraf Yolu', ''),
                'Evrak Yolu': vehicle_info.get('Evrak Yolu', ''),
                'KM Geçmişi': vehicle_info.get('KM Geçmişi', ''),
                'Oluşturma Tarihi': vehicle_info.get('Oluşturma Tarihi', ''),
                'Güncelleme Tarihi': vehicle_info.get('Güncelleme Tarihi', ''),
                'Silme Tarihi': datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
                'Silme Sebebi': reason,
                'Silme Notları': notes,
                'Silen Kullanıcı': 'Admin'  # Gelecekte kullanıcı sistemi eklenebilir
            }
            
            new_df = pd.DataFrame([deleted_record])
            deleted_vehicles = pd.concat([deleted_vehicles, new_df], ignore_index=True)
            
            self.data_manager.save_data('silinen_araclar', deleted_vehicles)
            
        except Exception as e:
            logger.error(f"Silinen araç ekleme hatası: {str(e)}", e)
    
    def update_related_records(self, plaka):
        """İlgili kayıtları güncelle (bakım, yakıt, ceza vs.)"""
        try:
            # Bakım kayıtlarını güncelle
            maintenance = self.data_manager.load_data('bakimlar')
            if not maintenance.empty:
                maintenance.loc[maintenance['Araç Plakası'] == plaka, 'Durum'] = 'Araç Silindi'
                self.data_manager.save_data('bakimlar', maintenance)
            
            # Yakıt kayıtlarını güncelle
            fuel = self.data_manager.load_data('yakitlar')
            if not fuel.empty:
                fuel.loc[fuel['Araç Plakası'] == plaka, 'Durum'] = 'Araç Silindi'
                self.data_manager.save_data('yakitlar', fuel)
            
            # Sürücü atamalarını güncelle
            drivers = self.data_manager.load_data('suruculer')
            if not drivers.empty:
                drivers.loc[drivers['Atanan Araç'] == plaka, 'Atanan Araç'] = ''
                drivers.loc[drivers['Atanan Araç'] == plaka, 'Atama Tarihi'] = ''
                self.data_manager.save_data('suruculer', drivers)
                
        except Exception as e:
            logger.error(f"İlgili kayıt güncelleme hatası: {str(e)}", e)
    
    def load_deleted_vehicles(self):
        """Silinen araçları yükle"""
        try:
            print("🔍 Silinen araçlar yükleniyor...")
            deleted_vehicles = self.data_manager.load_data('silinen_araclar')
            print(f"📊 Yüklenen kayıt sayısı: {len(deleted_vehicles)}")
            
            # Filtreleri uygula
            start_date = self.deleted_start_date.date().toPyDate()
            end_date = self.deleted_end_date.date().toPyDate()
            selected_reason = self.deleted_reason_filter.currentText()
            
            print(f"📅 Tarih aralığı: {start_date} - {end_date}")
            print(f"🏷️ Seçilen sebep: {selected_reason}")
            
            if not deleted_vehicles.empty:
                # Tarih filtresi - Güvenli tarih dönüşümü
                try:
                    # Silme Tarihi sütununu datetime'a çevir
                    deleted_vehicles['Silme Tarihi'] = pd.to_datetime(deleted_vehicles['Silme Tarihi'], errors='coerce')
                    
                    # Geçerli tarihleri filtrele
                    valid_dates = deleted_vehicles['Silme Tarihi'].notna()
                    deleted_vehicles = deleted_vehicles[valid_dates]
                    
                    # Tarih aralığı filtresi
                    if not deleted_vehicles.empty:
                        deleted_vehicles = deleted_vehicles[
                            (deleted_vehicles['Silme Tarihi'].dt.date >= start_date) &
                            (deleted_vehicles['Silme Tarihi'].dt.date <= end_date)
                        ]
                except Exception as date_error:
                    logger.error(f"Tarih filtreleme hatası: {str(date_error)}")
                    # Tarih filtresi olmadan devam et
                
                # Sebep filtresi
                if selected_reason != "Tümü":
                    deleted_vehicles = deleted_vehicles[
                        deleted_vehicles['Silme Sebebi'] == selected_reason
                    ]
            
            # Tablo sütunları
            columns = [
                "Plaka", "Marka", "Model", "Yıl", "Son KM", 
                "Silme Tarihi", "Silme Sebebi", "Silen Kullanıcı"
            ]
            
            print(f"📋 Tablo sütunları: {columns}")
            print(f"📊 Filtrelenmiş kayıt sayısı: {len(deleted_vehicles)}")
            
            self.deleted_model.clear()
            self.deleted_model.setHorizontalHeaderLabels(columns)
            
            for _, vehicle in deleted_vehicles.iterrows():
                row = []
                for col in columns:
                    value = str(vehicle.get(col, ''))
                    # Tarih formatını düzenle
                    if col == "Silme Tarihi" and pd.notna(vehicle.get(col, '')):
                        try:
                            date_value = pd.to_datetime(vehicle.get(col, ''))
                            value = date_value.strftime("%d.%m.%Y %H:%M")
                        except:
                            value = str(vehicle.get(col, ''))
                    item = QStandardItem(value)
                    row.append(item)
                self.deleted_model.appendRow(row)
            
        except Exception as e:
            logger.error(f"Silinen araçlar yükleme hatası: {str(e)}", e)
            # Hata durumunda boş tablo göster
            self.deleted_model.clear()
            columns = [
                "Plaka", "Marka", "Model", "Yıl", "Son KM", 
                "Silme Tarihi", "Silme Sebebi", "Silen Kullanıcı"
            ]
            self.deleted_model.setHorizontalHeaderLabels(columns)
    
    def filter_deleted_vehicles(self, text):
        """Silinen araçlarda arama"""
        self.deleted_proxy_model.setFilterWildcard(text)
    
    def restore_vehicle(self):
        """Silinen aracı geri yükle"""
        index = self.deleted_table_view.currentIndex()
        if index.isValid():
            row = self.deleted_proxy_model.mapToSource(index).row()
            plaka = self.deleted_model.item(row, 0).text()
            
            reply = QMessageBox.question(self, "Geri Yükleme Onayı", 
                f"'{plaka}' plakalı aracı geri yüklemek istediğinizden emin misiniz?",
                QMessageBox.Yes | QMessageBox.No)
            
            if reply == QMessageBox.Yes:
                try:
                    # Silinen araçlar tablosundan al
                    deleted_vehicles = self.data_manager.load_data('silinen_araclar')
                    vehicle_data = deleted_vehicles[deleted_vehicles['Plaka'] == plaka]
                    
                    if not vehicle_data.empty:
                        vehicle_info = vehicle_data.iloc[0].to_dict()
                        
                        # Ana tabloya geri ekle
                        vehicles = self.data_manager.load_data('araclar')
                        
                        restored_vehicle = {
                            'ID': len(vehicles) + 1 if not vehicles.empty else 1,
                            'Plaka': vehicle_info.get('Plaka', ''),
                            'Marka': vehicle_info.get('Marka', ''),
                            'Model': vehicle_info.get('Model', ''),
                            'Yıl': vehicle_info.get('Yıl', ''),
                            'Şasi No': vehicle_info.get('Şasi No', ''),
                            'Motor No': vehicle_info.get('Motor No', ''),
                            'Araç Tipi': vehicle_info.get('Araç Tipi', ''),
                            'Son KM': vehicle_info.get('Son KM', ''),
                            'Durum': 'Aktif',
                            'Sigorta Bitiş': vehicle_info.get('Sigorta Bitiş', ''),
                            'Muayene Tarihi': vehicle_info.get('Muayene Tarihi', ''),
                            'Birim': vehicle_info.get('Birim', ''),
                            'Hizmet Kişisi': vehicle_info.get('Hizmet Kişisi', ''),
                            'Fotoğraf Yolu': vehicle_info.get('Fotoğraf Yolu', ''),
                            'Evrak Yolu': vehicle_info.get('Evrak Yolu', ''),
                            'KM Geçmişi': vehicle_info.get('KM Geçmişi', ''),
                            'Oluşturma Tarihi': vehicle_info.get('Oluşturma Tarihi', ''),
                            'Güncelleme Tarihi': datetime.now().strftime("%Y-%m-%d %H:%M:%S")
                        }
                        
                        new_df = pd.DataFrame([restored_vehicle])
                        vehicles = pd.concat([vehicles, new_df], ignore_index=True)
                        self.data_manager.save_data('araclar', vehicles)
                        
                        # Silinen araçlar tablosundan çıkar
                        deleted_vehicles = deleted_vehicles[deleted_vehicles['Plaka'] != plaka]
                        self.data_manager.save_data('silinen_araclar', deleted_vehicles)
                        
                        self.load_vehicles()
                        self.load_deleted_vehicles()
                        
                        QMessageBox.information(self, "Başarılı", 
                            f"'{plaka}' plakalı araç başarıyla geri yüklendi.")
                    else:
                        QMessageBox.warning(self, "Uyarı", "Araç bulunamadı.")
                        
                except Exception as e:
                    QMessageBox.critical(self, "Hata", f"Araç geri yükleme hatası: {str(e)}")
        else:
            QMessageBox.warning(self, "Uyarı", "Lütfen geri yüklenecek aracı seçin.")
    
    def export_deleted_vehicles(self):
        """Silinen araçları Excel'e aktar"""
        try:
            export_table_to_excel(self.deleted_model, self.deleted_proxy_model, self)
        except Exception as e:
            QMessageBox.critical(self, "Hata", f"Export hatası: {str(e)}")
    
    def generate_deletion_report(self):
        """Silme raporu oluştur"""
        try:
            deleted_vehicles = self.data_manager.load_data('silinen_araclar')
            
            if deleted_vehicles.empty:
                QMessageBox.information(self, "Bilgi", "Silinen araç bulunmuyor.")
                return
            
            # İstatistikler
            total_deleted = len(deleted_vehicles)
            reasons = deleted_vehicles['Silme Sebebi'].value_counts()
            
            # Rapor içeriği
            report_content = f"""
SİLİNEN ARAÇLAR RAPORU
========================

Toplam Silinen Araç: {total_deleted}

SİLME SEBEPLERİ:
"""
            
            for reason, count in reasons.items():
                percentage = (count / total_deleted) * 100
                report_content += f"• {reason}: {count} araç (%{percentage:.1f})\n"
            
            report_content += f"""

DETAYLI LİSTE:
"""
            
            for _, vehicle in deleted_vehicles.iterrows():
                report_content += f"""
Plaka: {vehicle.get('Plaka', '')}
Marka/Model: {vehicle.get('Marka', '')} {vehicle.get('Model', '')}
Silme Tarihi: {vehicle.get('Silme Tarihi', '')}
Silme Sebebi: {vehicle.get('Silme Sebebi', '')}
Son KM: {vehicle.get('Son KM', '')}
---
"""
            
            # Rapor dialog'u
            self.show_deletion_report_dialog(report_content)
            
        except Exception as e:
            QMessageBox.critical(self, "Hata", f"Rapor oluşturma hatası: {str(e)}")
    
    def show_deletion_report_dialog(self, content):
        """Silme raporu dialog'u"""
        dialog = QDialog(self)
        dialog.setWindowTitle("Silinen Araçlar Raporu")
        dialog.setMinimumSize(600, 400)
        dialog.setModal(True)
        
        layout = QVBoxLayout(dialog)
        
        # Rapor içeriği
        text_edit = QTextEdit()
        text_edit.setPlainText(content)
        text_edit.setReadOnly(True)
        layout.addWidget(text_edit)
        
        # Butonlar
        btn_layout = QHBoxLayout()
        
        btn_save = QPushButton("💾 Kaydet")
        btn_save.clicked.connect(lambda: self.save_deletion_report(content))
        
        btn_close = QPushButton("❌ Kapat")
        btn_close.clicked.connect(dialog.accept)
        
        btn_layout.addWidget(btn_save)
        btn_layout.addWidget(btn_close)
        layout.addLayout(btn_layout)
        
        dialog.exec_()
    
    def save_deletion_report(self, content):
        """Silme raporunu kaydet"""
        try:
            file_path, _ = QFileDialog.getSaveFileName(
                self, "Raporu Kaydet", "", "Metin Dosyası (*.txt)")
            
            if file_path:
                with open(file_path, 'w', encoding='utf-8') as f:
                    f.write(content)
                
                QMessageBox.information(self, "Başarılı", "Rapor başarıyla kaydedildi.")
        except Exception as e:
            QMessageBox.critical(self, "Hata", f"Rapor kaydetme hatası: {str(e)}")

class VehicleDeletionDialog(QDialog):
    """Araç silme dialog'u"""
    
    def __init__(self, parent=None, plaka=""):
        super().__init__(parent)
        self.plaka = plaka
        self.deletion_reason = ""
        self.deletion_notes = ""
        self.init_ui()
    
    def init_ui(self):
        self.setWindowTitle("Araç Silme Onayı")
        self.setFixedSize(500, 400)
        self.setModal(True)
        
        layout = QVBoxLayout(self)
        layout.setSpacing(20)
        
        # Uyarı mesajı
        warning_label = QLabel("⚠️ DİKKAT: Bu işlem geri alınamaz!")
        warning_label.setStyleSheet("""
            QLabel {
            color: #e74c3c;
                font-size: 16px;
                font-weight: bold;
                padding: 10px;
                background-color: #fdf2f2;
                border: 2px solid #e74c3c;
                border-radius: 5px;
            }
        """)
        warning_label.setAlignment(Qt.AlignmentFlag.AlignCenter)
        layout.addWidget(warning_label)
        
        # Araç bilgileri
        vehicle_info = QLabel(f"Silinecek Araç: {self.plaka}")
        vehicle_info.setStyleSheet("font-size: 14px; font-weight: bold; color: #2c3e50;")
        vehicle_info.setAlignment(Qt.AlignmentFlag.AlignCenter)
        layout.addWidget(vehicle_info)
        
        # Silme sebebi
        reason_group = QGroupBox("Silme Sebebi")
        reason_layout = QVBoxLayout(reason_group)
        
        self.reason_combo = QComboBox()
        self.reason_combo.addItems([
            "Hurda",
            "Satış", 
            "Kaza",
            "Değişim",
            "Devir",
            "Diğer"
        ])
        self.reason_combo.currentTextChanged.connect(self.on_reason_changed)
        reason_layout.addWidget(self.reason_combo)
        
        layout.addWidget(reason_group)
        
        # Açıklama
        notes_group = QGroupBox("Açıklama (Opsiyonel)")
        notes_layout = QVBoxLayout(notes_group)
        
        self.notes_text = QTextEdit()
        self.notes_text.setPlaceholderText("Silme sebebi hakkında detaylı açıklama yazabilirsiniz...")
        self.notes_text.setMaximumHeight(100)
        notes_layout.addWidget(self.notes_text)
        
        layout.addWidget(notes_group)
        
        # Onay checkbox'ı
        self.confirm_checkbox = QCheckBox("Bu işlemi gerçekleştirmek istediğimi onaylıyorum")
        self.confirm_checkbox.setStyleSheet("font-weight: bold; color: #e74c3c;")
        layout.addWidget(self.confirm_checkbox)
        
        # Butonlar
        btn_layout = QHBoxLayout()
        
        btn_delete = QPushButton("🗑️ Sil")
        btn_delete.setStyleSheet("""
            QPushButton {
                background-color: #e74c3c;
                color: white;
                border: none;
                padding: 10px 20px;
                border-radius: 5px;
                font-weight: bold;
            }
            QPushButton:hover {
                background-color: #c0392b;
            }
            QPushButton:disabled {
                background-color: #bdc3c7;
                color: #7f8c8d;
            }
        """)
        btn_delete.clicked.connect(self.accept)
        btn_delete.setEnabled(False)
        self.confirm_checkbox.toggled.connect(btn_delete.setEnabled)
        
        btn_cancel = QPushButton("❌ İptal")
        btn_cancel.setStyleSheet("""
            QPushButton {
                background-color: #95a5a6;
                color: white;
                border: none;
                padding: 10px 20px;
                border-radius: 5px;
                font-weight: bold;
            }
            QPushButton:hover {
                background-color: #7f8c8d;
            }
        """)
        btn_cancel.clicked.connect(self.reject)
        
        btn_layout.addWidget(btn_delete)
        btn_layout.addWidget(btn_cancel)
        layout.addLayout(btn_layout)
    
    def on_reason_changed(self, reason):
        """Silme sebebi değiştiğinde"""
        self.deletion_reason = reason
    
    def get_deletion_reason(self):
        """Silme sebebini döndür"""
        return self.deletion_reason
    
    def get_deletion_notes(self):
        """Silme notlarını döndür"""
        return self.notes_text.toPlainText()

class VehicleDialog(QDialog):
    """Araç ekleme/düzenleme dialog"""
    
    def __init__(self, parent=None, vehicle_data=None):
        super().__init__(parent)
        self.vehicle_data = vehicle_data
        self.data_manager = parent.data_manager if parent else DataManager()
        self.selected_photo_path = None
        self.init_ui()
        if vehicle_data:
            self.load_vehicle_data(vehicle_data)
    
    def init_ui(self):
        self.setWindowTitle("Araç Ekle" if not self.vehicle_data else "Araç Düzenle")
        self.setMinimumWidth(500)
        self.setModal(True)
        
        layout = QFormLayout(self)
        layout.setSpacing(15)
        
        # Form alanları
        self.plaka = QLineEdit()
        self.marka = QLineEdit()
        self.model = QLineEdit()
        self.yil = QLineEdit()
        self.sasi = QLineEdit()
        self.ruhsat_seri_no = QLineEdit()
        self.ruhsat_seri_no.setPlaceholderText("Ruhsat Seri Numarası")
        
        self.yakit_tipi = QComboBox()
        self.yakit_tipi.addItems(["Benzin", "Dizel", "LPG", "Elektrik", "Hibrit", "Diğer"])
        
        self.arac_tipi = QComboBox()
        self.arac_tipi.addItems(["Binek", "SUV", "Elektrikli", "Minibüs", "Pick-up", "Ticari", "Diğer"])
        
        self.son_km = QLineEdit()
        
        self.durum = QComboBox()
        self.durum.addItems(["Aktif", "Pasif", "Hurda", "Satıldı"])
        
        self.sigorta = QDateEdit()
        self.sigorta.setCalendarPopup(True)
        self.sigorta.setDate(QDate.currentDate())
        
        self.kasko_var = QComboBox()
        self.kasko_var.addItems(["Yok", "Var"])
        
        self.kasko = QDateEdit()
        self.kasko.setCalendarPopup(True)
        self.kasko.setDate(QDate.currentDate())
        self.kasko.setEnabled(False)
        self.kasko_var.currentTextChanged.connect(self.kasko_durum_kontrol)
        
        # Kasko seçimi için açıklama
        kasko_info = QLabel("Kasko sigortası var mı? Varsa tarih girin.")
        kasko_info.setStyleSheet("color: #7f8c8d; font-size: 11px; font-style: italic;")
        
        self.muayene = QDateEdit()
        self.muayene.setCalendarPopup(True)
        self.muayene.setDate(QDate.currentDate())
        
        # Birim ve hizmet bilgileri
        self.birim = QLineEdit()
        self.birim.setPlaceholderText("Örn: İdari İşler, Teknik Servis, Güvenlik")
        
        self.hizmet_kisi = QLineEdit()
        self.hizmet_kisi.setPlaceholderText("Örn: Ahmet Yılmaz, Mehmet Demir")
        
        # Fotoğraf bölümü
        photo_group = QGroupBox("📸 Araç Fotoğrafı")
        photo_group.setStyleSheet("""
            QGroupBox {
                font-weight: bold;
                border: 2px solid #3498db;
                border-radius: 5px;
                margin-top: 10px;
                padding-top: 10px;
            }
        """)
        
        photo_layout = QVBoxLayout(photo_group)
        
        # Fotoğraf görüntüleme alanı
        self.photo_label = QLabel()
        self.photo_label.setMinimumSize(200, 150)
        self.photo_label.setMaximumSize(300, 200)
        self.photo_label.setStyleSheet("""
            QLabel {
                border: 2px dashed #bdc3c7;
                border-radius: 5px;
                background-color: #f8f9fa;
            }
        """)
        self.photo_label.setAlignment(Qt.AlignmentFlag.AlignCenter)
        self.photo_label.setText("Fotoğraf seçilmedi\n\n📷 Fotoğraf eklemek için tıklayın")
        
        # Fotoğraf seçme butonu
        btn_photo_layout = QHBoxLayout()
        btn_select_photo = QPushButton("📷 Fotoğraf Seç")
        btn_select_photo.setStyleSheet("""
            QPushButton {
                background-color: #3498db;
                color: white;
                border: none;
                padding: 8px 15px;
                border-radius: 5px;
                font-weight: bold;
            }
            QPushButton:hover {
                background-color: #2980b9;
            }
        """)
        btn_select_photo.clicked.connect(self.select_photo)
        
        btn_clear_photo = QPushButton("🗑️ Temizle")
        btn_clear_photo.setStyleSheet("""
            QPushButton {
                background-color: #e74c3c;
                color: white;
                border: none;
                padding: 8px 15px;
                border-radius: 5px;
                font-weight: bold;
            }
            QPushButton:hover {
                background-color: #c0392b;
            }
        """)
        btn_clear_photo.clicked.connect(self.clear_photo)
        
        btn_photo_layout.addWidget(btn_select_photo)
        btn_photo_layout.addWidget(btn_clear_photo)
        
        photo_layout.addWidget(self.photo_label)
        photo_layout.addLayout(btn_photo_layout)
        
        # Form'a ekle
        layout.addRow("Plaka:", self.plaka)
        layout.addRow("Marka:", self.marka)
        layout.addRow("Model:", self.model)
        layout.addRow("Yıl:", self.yil)
        layout.addRow("Şasi No:", self.sasi)
        layout.addRow("Ruhsat Seri No:", self.ruhsat_seri_no)
        layout.addRow("Yakıt Tipi:", self.yakit_tipi)
        layout.addRow("Araç Tipi:", self.arac_tipi)
        layout.addRow("Son KM:", self.son_km)
        layout.addRow("Durum:", self.durum)
        layout.addRow("Sigorta Bitiş:", self.sigorta)
        layout.addRow("Kasko Var mı?:", self.kasko_var)
        layout.addRow("", kasko_info)  # Açıklama satırı
        layout.addRow("Kasko Bitiş:", self.kasko)
        layout.addRow("Muayene Tarihi:", self.muayene)
        layout.addRow("Birim:", self.birim)
        layout.addRow("Hizmet Kişisi:", self.hizmet_kisi)
        layout.addRow(photo_group)
        
        # Butonlar
        btn_layout = QHBoxLayout()
        btn_save = QPushButton("Kaydet")
        btn_save.clicked.connect(self.save_vehicle)
        btn_cancel = QPushButton("İptal")
        btn_cancel.clicked.connect(self.reject)
        
        btn_layout.addWidget(btn_save)
        btn_layout.addWidget(btn_cancel)
        layout.addRow(btn_layout)
    
    def kasko_durum_kontrol(self, text):
        """Kasko durum kontrolü"""
        self.kasko.setEnabled(text == "Var")
    
    def select_photo(self):
        """Fotoğraf seç"""
        file_path, _ = QFileDialog.getOpenFileName(
            self, "Araç Fotoğrafı Seç", "", 
            "Resim Dosyaları (*.jpg *.jpeg *.png *.bmp *.gif);;Tüm Dosyalar (*)"
        )
        
        if file_path:
            try:
                # Fotoğrafı yükle ve boyutlandır
                pixmap = QPixmap(file_path)
                if not pixmap.isNull():
                    # Fotoğrafı label boyutuna sığdır
                    scaled_pixmap = pixmap.scaled(
                        self.photo_label.size(), 
                        Qt.AspectRatioMode.KeepAspectRatio, 
                        Qt.TransformationMode.SmoothTransformation
                    )
                    self.photo_label.setPixmap(scaled_pixmap)
                    self.photo_label.setStyleSheet("""
                        QLabel {
                            border: 2px solid #27ae60;
                            border-radius: 5px;
                            background-color: white;
                        }
                    """)
                    # Fotoğraf yolunu sakla
                    self.selected_photo_path = file_path
                else:
                    QMessageBox.warning(self, "Hata", "Fotoğraf yüklenemedi.")
            except Exception as e:
                QMessageBox.critical(self, "Hata", f"Fotoğraf yükleme hatası: {str(e)}")
    
    def clear_photo(self):
        """Fotoğrafı temizle"""
        self.photo_label.clear()
        self.photo_label.setText("Fotoğraf seçilmedi\n\n📷 Fotoğraf eklemek için tıklayın")
        self.photo_label.setStyleSheet("""
            QLabel {
                border: 2px dashed #bdc3c7;
                border-radius: 5px;
                background-color: #f8f9fa;
            }
        """)
        self.selected_photo_path = None
    
    def load_vehicle_data(self, vehicle_data):
        """Araç verilerini yükle"""
        self.plaka.setText(str(vehicle_data.get('Plaka', '')))
        self.marka.setText(str(vehicle_data.get('Marka', '')))
        self.model.setText(str(vehicle_data.get('Model', '')))
        self.yil.setText(str(vehicle_data.get('Yıl', '')))
        self.sasi.setText(str(vehicle_data.get('Şasi No', '')))
        self.ruhsat_seri_no.setText(str(vehicle_data.get('Ruhsat Seri No', '')))
        self.yakit_tipi.setCurrentText(str(vehicle_data.get('Yakıt Tipi', 'Benzin')))
        self.arac_tipi.setCurrentText(str(vehicle_data.get('Araç Tipi', 'Binek')))
        self.son_km.setText(str(vehicle_data.get('Son KM', '')))
        self.durum.setCurrentText(str(vehicle_data.get('Durum', 'Aktif')))
        
        # Tarihleri yükle
        try:
            sigorta_date = QDate.fromString(str(vehicle_data.get('Sigorta Bitiş', '')), "dd.MM.yyyy")
            if sigorta_date.isValid():
                self.sigorta.setDate(sigorta_date)
        except:
            pass
        
        self.kasko_var.setCurrentText(str(vehicle_data.get('Kasko Var mı?', 'Yok')))
        
        try:
            kasko_date = QDate.fromString(str(vehicle_data.get('Kasko Bitiş', '')), "dd.MM.yyyy")
            if kasko_date.isValid():
                self.kasko.setDate(kasko_date)
        except:
            pass
        
        try:
            muayene_date = QDate.fromString(str(vehicle_data.get('Muayene Tarihi', '')), "dd.MM.yyyy")
            if muayene_date.isValid():
                self.muayene.setDate(muayene_date)
        except:
            pass
        
        # Birim ve hizmet bilgilerini yükle
        self.birim.setText(str(vehicle_data.get('Birim', '')))
        self.hizmet_kisi.setText(str(vehicle_data.get('Hizmet Kişisi', '')))
        
        # Fotoğrafı yükle
        photo_path = vehicle_data.get('Fotoğraf Yolu', '')
        plaka = vehicle_data.get('Plaka', '').strip()
        
        # Önce veri tabanındaki fotoğraf yolunu kontrol et
        if isinstance(photo_path, str) and photo_path and os.path.exists(photo_path):
            try:
                pixmap = QPixmap(photo_path)
                if not pixmap.isNull():
                    scaled_pixmap = pixmap.scaled(
                        self.photo_label.size(), 
                        Qt.AspectRatioMode.KeepAspectRatio, 
                        Qt.TransformationMode.SmoothTransformation
                    )
                    self.photo_label.setPixmap(scaled_pixmap)
                    self.photo_label.setStyleSheet("""
                        QLabel {
                            border: 2px solid #27ae60;
                            border-radius: 5px;
                            background-color: white;
                        }
                    """)
                    self.selected_photo_path = photo_path
                    return
            except Exception as e:
                print(f"Fotoğraf yükleme hatası: {e}")
        
        # Veri tabanında fotoğraf yolu yoksa, plaka ile eşleşen dosyayı ara
        if plaka:
            photo_dir = "veri/arac_fotograflari"
            possible_extensions = ['.png', '.jpg', '.jpeg']
            
            for ext in possible_extensions:
                possible_path = os.path.join(photo_dir, f"{plaka}{ext}")
                if os.path.exists(possible_path):
                    try:
                        pixmap = QPixmap(possible_path)
                        if not pixmap.isNull():
                            scaled_pixmap = pixmap.scaled(
                                self.photo_label.size(), 
                                Qt.AspectRatioMode.KeepAspectRatio, 
                                Qt.TransformationMode.SmoothTransformation
                            )
                            self.photo_label.setPixmap(scaled_pixmap)
                            self.photo_label.setStyleSheet("""
                                QLabel {
                                    border: 2px solid #27ae60;
                                    border-radius: 5px;
                                    background-color: white;
                                }
                            """)
                            self.selected_photo_path = possible_path
                            return
                    except Exception as e:
                        print(f"Otomatik fotoğraf yükleme hatası: {e}")
        
        # Hiçbir fotoğraf bulunamazsa varsayılan görünümü göster
        self.photo_label.clear()
        self.photo_label.setText("Fotoğraf seçilmedi\n\n📷 Fotoğraf eklemek için tıklayın")
        self.photo_label.setStyleSheet("""
            QLabel {
                border: 2px dashed #bdc3c7;
                border-radius: 5px;
                background-color: #f8f9fa;
            }
        """)
        self.selected_photo_path = None
    
    def save_vehicle(self):
        """Araç kaydet"""
        if not self.plaka.text().strip():
            QMessageBox.warning(self, "Uyarı", "Plaka alanı zorunludur.")
            return

        # Fotoğraf yolu belirleme
        foto_path = ""
        if self.selected_photo_path:  # Kullanıcı yeni fotoğraf seçtiyse
            foto_path = save_vehicle_photo(self.selected_photo_path, self.plaka.text().strip())
        elif self.vehicle_data and self.vehicle_data.get('Fotoğraf Yolu', ''):
            foto_path = self.vehicle_data.get('Fotoğraf Yolu', '')
        else:
            foto_path = ""

        try:
            vehicles = self.data_manager.load_data('araclar')
            
            # Yeni araç verisi
            new_vehicle = {
                'ID': len(vehicles) + 1 if not vehicles.empty else 1,
                'Plaka': self.plaka.text().strip(),
                'Marka': self.marka.text().strip(),
                'Model': self.model.text().strip(),
                'Yıl': self.yil.text().strip(),
                'Şasi No': self.sasi.text().strip(),
                'Ruhsat Seri No': self.ruhsat_seri_no.text().strip(),
                'Yakıt Tipi': self.yakit_tipi.currentText(),
                'Araç Tipi': self.arac_tipi.currentText(),
                'Son KM': self.son_km.text().strip(),
                'Durum': self.durum.currentText(),
                'Sigorta Bitiş': self.sigorta.date().toString("dd.MM.yyyy"),
                'Kasko Var mı?': self.kasko_var.currentText(),
                'Kasko Bitiş': self.kasko.date().toString("dd.MM.yyyy") if self.kasko_var.currentText() == "Var" else "",
                'Muayene Tarihi': self.muayene.date().toString("dd.MM.yyyy"),
                'Birim': self.birim.text().strip(),
                'Hizmet Kişisi': self.hizmet_kisi.text().strip(),
                'Fotoğraf Yolu': foto_path,
                'Evrak Yolu': "",
                'KM Geçmişi': "",
                'Oluşturma Tarihi': datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
                'Güncelleme Tarihi': datetime.now().strftime("%Y-%m-%d %H:%M:%S")
            }
            
            if self.vehicle_data:
                # Düzenleme modu
                vehicles = vehicles[vehicles['Plaka'] != self.vehicle_data['Plaka']]
            
            # Yeni veriyi ekle
            new_df = pd.DataFrame([new_vehicle])
            vehicles = pd.concat([vehicles, new_df], ignore_index=True)
            
            # Kaydet
            if self.data_manager.save_data('araclar', vehicles):
                # --- Kasko/Muayene/Sigorta giderlerini otomatik ekle ---
                giderler = self.data_manager.load_data('giderler')
                plaka = self.plaka.text().strip()
                # Sigorta
                sigorta_bitis = self.sigorta.date().toString("dd.MM.yyyy")
                if sigorta_bitis:
                    mevcut = giderler[(giderler['Araç Plakası'] == plaka) & (giderler['Gider Türü'] == 'Sigorta') & (giderler['Bitiş'] == sigorta_bitis)]
                    if mevcut.empty:
                        yeni = {'Araç Plakası': plaka, 'Gider Türü': 'Sigorta', 'Başlangıç': '', 'Bitiş': sigorta_bitis, 'Şirket': '', 'Tutar': ''}
                        giderler = pd.concat([giderler, pd.DataFrame([yeni])], ignore_index=True)
                # Kasko
                if self.kasko_var.currentText() == 'Var':
                    kasko_bitis = self.kasko.date().toString("dd.MM.yyyy")
                    if kasko_bitis:
                        mevcut = giderler[(giderler['Araç Plakası'] == plaka) & (giderler['Gider Türü'] == 'Kasko') & (giderler['Bitiş'] == kasko_bitis)]
                        if mevcut.empty:
                            yeni = {'Araç Plakası': plaka, 'Gider Türü': 'Kasko', 'Başlangıç': '', 'Bitiş': kasko_bitis, 'Şirket': '', 'Tutar': ''}
                            giderler = pd.concat([giderler, pd.DataFrame([yeni])], ignore_index=True)
                # Muayene
                muayene_bitis = self.muayene.date().toString("dd.MM.yyyy")
                if muayene_bitis:
                    mevcut = giderler[(giderler['Araç Plakası'] == plaka) & (giderler['Gider Türü'] == 'Muayene') & (giderler['Bitiş'] == muayene_bitis)]
                    if mevcut.empty:
                        yeni = {'Araç Plakası': plaka, 'Gider Türü': 'Muayene', 'Başlangıç': '', 'Bitiş': muayene_bitis, 'Şirket': '', 'Tutar': ''}
                        giderler = pd.concat([giderler, pd.DataFrame([yeni])], ignore_index=True)
                self.data_manager.save_data('giderler', giderler)
                # --- Son ---
                QMessageBox.information(self, "Başarılı", 
                    "Araç güncellendi." if self.vehicle_data else "Araç eklendi.")
                self.accept()
            else:
                QMessageBox.critical(self, "Hata", "Kaydetme hatası.")
                
        except Exception as e:
            QMessageBox.critical(self, "Hata", f"Kaydetme hatası: {str(e)}")

class DriversPanel(QWidget):
    def __init__(self, data_manager):
        super().__init__()
        self.data_manager = data_manager
        self.init_ui()
        self.load_drivers()
    
    def init_ui(self):
        layout = QVBoxLayout(self)
        layout.setContentsMargins(20, 20, 20, 20)
        layout.setSpacing(20)
        
        # Başlık ve butonlar
        header_layout = QHBoxLayout()
        
        title = QLabel("Sürücü Yönetimi")
        title.setStyleSheet("font-size: 24px; font-weight: bold; color: #2c3e50;")
        header_layout.addWidget(title)
        
        header_layout.addStretch()
        
        # Butonlar
        btn_add = QPushButton("➕ Sürücü Ekle")
        btn_add.setStyleSheet("""
            QPushButton {
                background-color: #27ae60;
                color: white;
                border: none;
                padding: 10px 20px;
                border-radius: 5px;
                font-weight: bold;
            }
            QPushButton:hover {
                background-color: #229954;
            }
        """)
        btn_add.clicked.connect(self.add_driver)
        
        btn_edit = QPushButton("✏️ Düzenle")
        btn_edit.setStyleSheet("""
            QPushButton {
                background-color: #3498db;
                color: white;
                border: none;
                padding: 10px 20px;
                border-radius: 5px;
                font-weight: bold;
            }
            QPushButton:hover {
                background-color: #2980b9;
            }
        """)
        btn_edit.clicked.connect(self.edit_driver)
        
        btn_delete = QPushButton("🗑️ Sil")
        btn_delete.setStyleSheet("""
            QPushButton {
                background-color: #e74c3c;
                color: white;
                border: none;
                padding: 10px 20px;
                border-radius: 5px;
                font-weight: bold;
            }
            QPushButton:hover {
                background-color: #c0392b;
            }
        """)
        btn_delete.clicked.connect(self.delete_driver)
        
        header_layout.addWidget(btn_add)
        header_layout.addWidget(btn_edit)
        header_layout.addWidget(btn_delete)
        
        layout.addLayout(header_layout)
        
        # Tablo
        self.drivers_table = QTableWidget()
        self.drivers_table.setColumnCount(7)
        self.drivers_table.setHorizontalHeaderLabels([
            "Ad Soyad", "TC Kimlik", "Telefon", "Ehliyet Sınıfı", 
            "Ehliyet Tarihi", "Atanan Araç", "Durum"
        ])
        self.drivers_table.horizontalHeader().setStretchLastSection(True)
        self.drivers_table.setAlternatingRowColors(True)
        self.drivers_table.setSelectionBehavior(QAbstractItemView.SelectRows)
        self.drivers_table.setSelectionMode(QAbstractItemView.SingleSelection)
        
        # Satır numaralarını gizle
        self.drivers_table.verticalHeader().setVisible(False)
        
        layout.addWidget(self.drivers_table)
    
    def load_drivers(self):
        """Sürücüleri yükle"""
        try:
            drivers = self.data_manager.load_data('suruculer')
            self.drivers_table.setRowCount(0)  # Önce tabloyu temizle
            self.drivers_table.setRowCount(len(drivers))
            
            for row, (_, driver) in enumerate(drivers.iterrows()):
                self.drivers_table.setItem(row, 0, QTableWidgetItem(str(driver.get('Ad Soyad', ''))))
                self.drivers_table.setItem(row, 1, QTableWidgetItem(str(driver.get('TC Kimlik', ''))))
                self.drivers_table.setItem(row, 2, QTableWidgetItem(str(driver.get('Telefon', ''))))
                self.drivers_table.setItem(row, 3, QTableWidgetItem(str(driver.get('Ehliyet Sınıfı', ''))))
                self.drivers_table.setItem(row, 4, QTableWidgetItem(str(driver.get('Ehliyet Tarihi', ''))))
                self.drivers_table.setItem(row, 5, QTableWidgetItem(str(driver.get('Atanan Araç', ''))))
                self.drivers_table.setItem(row, 6, QTableWidgetItem(str(driver.get('Durum', ''))))
        except Exception as e:
            print(f"Sürücü yükleme hatası: {e}")
    
    def add_driver(self):
        """Sürücü ekle"""
        dialog = DriverDialog(self)
        if dialog.exec_() == QDialog.Accepted:
            self.load_drivers()
    
    def edit_driver(self):
        """Sürücü düzenle"""
        current_row = self.drivers_table.currentRow()
        if current_row >= 0:
            tc_kimlik = self.drivers_table.item(current_row, 1).text()
            drivers = self.data_manager.load_data('suruculer')
            driver = drivers[drivers['TC Kimlik'] == tc_kimlik]
            if not driver.empty:
                dialog = DriverDialog(self, driver.iloc[0].to_dict())
                if dialog.exec_() == QDialog.Accepted:
                    self.load_drivers()
        else:
            QMessageBox.warning(self, "Uyarı", "Lütfen düzenlenecek sürücüyü seçin.")
    
    def delete_driver(self):
        """Sürücü sil"""
        current_row = self.drivers_table.currentRow()
        if current_row >= 0:
            tc_kimlik = self.drivers_table.item(current_row, 1).text()
            ad_soyad = self.drivers_table.item(current_row, 0).text()
            reply = QMessageBox.question(self, "Onay", 
                f"'{ad_soyad}' isimli sürücüyü silmek istediğinizden emin misiniz?",
                QMessageBox.Yes | QMessageBox.No)
            
            if reply == QMessageBox.Yes:
                try:
                    drivers = self.data_manager.load_data('suruculer')
                    drivers = drivers[drivers['TC Kimlik'].astype(str) != str(tc_kimlik)]
                    success = self.data_manager.save_data('suruculer', drivers)
                    if success:
                        self.load_drivers()
                        QMessageBox.information(self, "Başarılı", "Sürücü silindi.")
                    else:
                        QMessageBox.critical(self, "Hata", "Dosya kaydedilemedi! (Excel dosyası açık olabilir veya başka bir hata oluştu.)")
                except Exception as e:
                    QMessageBox.critical(self, "Hata", f"Sürücü silme hatası: {str(e)}")
        else:
            QMessageBox.warning(self, "Uyarı", "Lütfen silinecek sürücüyü seçin.")

class DriverDialog(QDialog):
    """Sürücü ekleme/düzenleme dialog"""
    
    def __init__(self, parent=None, driver_data=None):
        super().__init__(parent)
        self.driver_data = driver_data
        self.data_manager = parent.data_manager if parent else DataManager()
        self.init_ui()
        if driver_data:
            self.load_driver_data(driver_data)
    
    def init_ui(self):
        self.setWindowTitle("Sürücü Ekle" if not self.driver_data else "Sürücü Düzenle")
        self.setMinimumWidth(500)
        self.setModal(True)
        
        layout = QFormLayout(self)
        layout.setSpacing(15)
        
        # Form alanları
        self.ad_soyad = QLineEdit()
        self.tc_kimlik = QLineEdit()
        self.telefon = QLineEdit()
        
        self.ehliyet_sinifi = QComboBox()
        self.ehliyet_sinifi.addItems(["B", "C", "D", "E", "F", "G", "H"])
        
        self.ehliyet_tarihi = QDateEdit()
        self.ehliyet_tarihi.setCalendarPopup(True)
        self.ehliyet_tarihi.setDate(QDate.currentDate())
        
        self.atanan_arac = QComboBox()
        self.load_available_vehicles()
        
        self.durum = QComboBox()
        self.durum.addItems(["Aktif", "Pasif", "İzinli", "İşten Ayrıldı"])
        
        # Form'a ekle
        layout.addRow("Ad Soyad:", self.ad_soyad)
        layout.addRow("TC Kimlik:", self.tc_kimlik)
        layout.addRow("Telefon:", self.telefon)
        layout.addRow("Ehliyet Sınıfı:", self.ehliyet_sinifi)
        layout.addRow("Ehliyet Tarihi:", self.ehliyet_tarihi)
        layout.addRow("Atanan Araç:", self.atanan_arac)
        layout.addRow("Durum:", self.durum)
        
        # Butonlar
        btn_layout = QHBoxLayout()
        btn_save = QPushButton("Kaydet")
        btn_save.clicked.connect(self.save_driver)
        btn_cancel = QPushButton("İptal")
        btn_cancel.clicked.connect(self.reject)
        
        btn_layout.addWidget(btn_save)
        btn_layout.addWidget(btn_cancel)
        layout.addRow(btn_layout)
    
    def load_available_vehicles(self):
        """Mevcut araçları yükle"""
        try:
            vehicles = self.data_manager.load_data('araclar')
            self.atanan_arac.clear()
            self.atanan_arac.addItem("Atanmamış")
            
            if not vehicles.empty:
                for _, vehicle in vehicles.iterrows():
                    if vehicle.get('Durum') == 'Aktif':
                        plaka = vehicle.get('Plaka', '')
                        marka = vehicle.get('Marka', '')
                        model = vehicle.get('Model', '')
                        self.atanan_arac.addItem(f"{plaka} - {marka} {model}")
        except Exception as e:
            print(f"Araç yükleme hatası: {e}")
    
    def load_driver_data(self, driver_data):
        """Sürücü verilerini yükle"""
        self.ad_soyad.setText(str(driver_data.get('Ad Soyad', '')))
        self.tc_kimlik.setText(str(driver_data.get('TC Kimlik', '')))
        self.telefon.setText(str(driver_data.get('Telefon', '')))
        self.ehliyet_sinifi.setCurrentText(str(driver_data.get('Ehliyet Sınıfı', 'B')))
        
        # Tarihleri yükle
        try:
            ehliyet_date = QDate.fromString(str(driver_data.get('Ehliyet Tarihi', '')), "dd.MM.yyyy")
            if ehliyet_date.isValid():
                self.ehliyet_tarihi.setDate(ehliyet_date)
        except:
            pass
        
        self.atanan_arac.setCurrentText(str(driver_data.get('Atanan Araç', 'Atanmamış')))
        self.durum.setCurrentText(str(driver_data.get('Durum', 'Aktif')))
    
    def save_driver(self):
        """Sürücü kaydet"""
        if not self.ad_soyad.text().strip():
            QMessageBox.warning(self, "Uyarı", "Ad Soyad alanı zorunludur.")
            return
        
        if not self.tc_kimlik.text().strip():
            QMessageBox.warning(self, "Uyarı", "TC Kimlik alanı zorunludur.")
            return
        
        try:
            drivers = self.data_manager.load_data('suruculer')
            
            # Yeni sürücü verisi
            new_driver = {
                'ID': len(drivers) + 1 if not drivers.empty else 1,
                'Ad Soyad': self.ad_soyad.text().strip(),
                'TC Kimlik': self.tc_kimlik.text().strip(),
                'Telefon': self.telefon.text().strip(),
                'Ehliyet Sınıfı': self.ehliyet_sinifi.currentText(),
                'Ehliyet Tarihi': self.ehliyet_tarihi.date().toString("dd.MM.yyyy"),
                'Atanan Araç': self.atanan_arac.currentText(),
                'Atama Tarihi': datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
                'Durum': self.durum.currentText(),
                'Oluşturma Tarihi': datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
                'Güncelleme Tarihi': datetime.now().strftime("%Y-%m-%d %H:%M:%S")
            }
            
            if self.driver_data:
                # Düzenleme modu
                drivers = drivers[drivers['TC Kimlik'] != self.driver_data['TC Kimlik']]
            
            # Yeni veriyi ekle
            new_df = pd.DataFrame([new_driver])
            drivers = pd.concat([drivers, new_df], ignore_index=True)
            
            # Kaydet
            if self.data_manager.save_data('suruculer', drivers):
                QMessageBox.information(self, "Başarılı", 
                    "Sürücü güncellendi." if self.driver_data else "Sürücü eklendi.")
                self.accept()
            else:
                QMessageBox.critical(self, "Hata", "Kaydetme hatası.")
                
        except Exception as e:
            QMessageBox.critical(self, "Hata", f"Kaydetme hatası: {str(e)}")

class MaintenancePanel(QWidget):
    def __init__(self, data_manager):
        super().__init__()
        self.data_manager = data_manager
        self.init_ui()
        self.load_maintenance()
    
    def init_ui(self):
        layout = QVBoxLayout(self)
        layout.setContentsMargins(20, 20, 20, 20)
        layout.setSpacing(20)
        header_layout = QHBoxLayout()
        title = QLabel("Bakım & Onarımlar")
        title.setStyleSheet("font-size: 24px; font-weight: bold; color: #2c3e50;")
        header_layout.addWidget(title)
        header_layout.addStretch()
        btn_add = QPushButton("➕ Bakım Ekle")
        btn_add.setStyleSheet("""
            QPushButton {
                background-color: #27ae60;
                color: white;
                border: none;
                padding: 10px 20px;
                border-radius: 5px;
                font-weight: bold;
            }
            QPushButton:hover {
                background-color: #229954;
            }
        """)
        btn_add.clicked.connect(self.add_maintenance)
        btn_edit = QPushButton("✏️ Düzenle")
        btn_edit.setStyleSheet("""
            QPushButton {
                background-color: #3498db;
                color: white;
                border: none;
                padding: 10px 20px;
                border-radius: 5px;
                font-weight: bold;
            }
            QPushButton:hover {
                background-color: #2980b9;
            }
        """)
        btn_edit.clicked.connect(self.edit_maintenance)
        btn_delete = QPushButton("🗑️ Sil")
        btn_delete.setStyleSheet("""
            QPushButton {
                background-color: #e74c3c;
                color: white;
                border: none;
                padding: 10px 20px;
                border-radius: 5px;
                font-weight: bold;
            }
            QPushButton:hover {
                background-color: #c0392b;
            }
        """)
        btn_delete.clicked.connect(self.delete_maintenance)
        btn_export = QPushButton("📤 Excel'e Aktar")
        btn_export.setStyleSheet("""
            QPushButton {
                background-color: #f1c40f;
                color: #2c3e50;
                border: none;
                padding: 10px 20px;
                border-radius: 5px;
                font-weight: bold;
            }
            QPushButton:hover {
                background-color: #f39c12;
            }
        """)
        btn_export.clicked.connect(lambda: export_table_to_excel(self.model, self.proxy_model, self))
        header_layout.addWidget(btn_add)
        header_layout.addWidget(btn_edit)
        header_layout.addWidget(btn_delete)
        header_layout.addWidget(btn_export)
        layout.addLayout(header_layout)
        self.search_box = QLineEdit()
        self.search_box.setPlaceholderText("Bakımlarda ara...")
        self.search_box.textChanged.connect(self.filter_maintenance)
        layout.addWidget(self.search_box)
        self.table_view = QTableView()
        self.model = QStandardItemModel()
        self.proxy_model = QSortFilterProxyModel()
        self.proxy_model.setSourceModel(self.model)
        self.proxy_model.setFilterCaseSensitivity(Qt.CaseInsensitive)
        self.table_view.setModel(self.proxy_model)
        self.table_view.setSortingEnabled(True)
        self.table_view.setSelectionBehavior(QTableView.SelectRows)
        self.table_view.setSelectionMode(QTableView.SingleSelection)
        self.table_view.horizontalHeader().setStretchLastSection(True)
        
        # Satır numaralarını gizle
        self.table_view.verticalHeader().setVisible(False)
        
        layout.addWidget(self.table_view)
    def load_maintenance(self):
        try:
            maintenance = self.data_manager.load_data('bakimlar')
            columns = [
                "Araç Plakası", "Bakım Tarihi", "Bakıma Girdiği KM", "İşlem Türü", "Servis Adı", "Tutar", "Açıklama"
            ]
            self.model.clear()
            self.model.setHorizontalHeaderLabels(columns)
            for _, record in maintenance.iterrows():
                row = [
                    QStandardItem(str(record.get("Araç Plakası", ''))),
                    QStandardItem(str(record.get("Bakım Tarihi", ''))),
                    QStandardItem(str(record.get("KM", ''))),
                    QStandardItem(str(record.get("İşlem Türü", ''))),
                    QStandardItem(str(record.get("Servis Adı", ''))),
                    QStandardItem(str(record.get("Tutar", ''))),
                    QStandardItem(str(record.get("Açıklama", '')))
                ]
                self.model.appendRow(row)
        except Exception as e:
            print(f"Bakım yükleme hatası: {e}")
    def filter_maintenance(self, text):
        self.proxy_model.setFilterWildcard(text)
    def add_maintenance(self):
        dialog = MaintenanceDialog(self)
        if dialog.exec_() == QDialog.Accepted:
            self.load_maintenance()
    def edit_maintenance(self):
        index = self.table_view.currentIndex()
        if index.isValid():
            row = self.proxy_model.mapToSource(index).row()
            plaka = self.model.item(row, 0).text()
            tarih = self.model.item(row, 1).text()
            maintenance = self.data_manager.load_data('bakimlar')
            maintenance_record = maintenance[(maintenance['Araç Plakası'] == plaka) & (maintenance['Bakım Tarihi'] == tarih)]
            if not maintenance_record.empty:
                dialog = MaintenanceDialog(self, maintenance_record.iloc[0].to_dict())
                if dialog.exec_() == QDialog.Accepted:
                    self.load_maintenance()
        else:
            QMessageBox.warning(self, "Uyarı", "Lütfen düzenlenecek bakım kaydını seçin.")
    def delete_maintenance(self):
        index = self.table_view.currentIndex()
        if index.isValid():
            row = self.proxy_model.mapToSource(index).row()
            plaka = self.model.item(row, 0).text()
            tarih = self.model.item(row, 1).text()
            reply = QMessageBox.question(self, "Onay", f"'{plaka}' plakalı aracın '{tarih}' tarihli bakım kaydını silmek istediğinizden emin misiniz?", QMessageBox.Yes | QMessageBox.No)
            if reply == QMessageBox.Yes:
                try:
                    maintenance = self.data_manager.load_data('bakimlar')
                    maintenance = maintenance[~((maintenance['Araç Plakası'] == plaka) & (maintenance['Bakım Tarihi'] == tarih))]
                    self.data_manager.save_data('bakimlar', maintenance)
                    self.load_maintenance()
                    QMessageBox.information(self, "Başarılı", "Bakım kaydı silindi.")
                except Exception as e:
                    QMessageBox.critical(self, "Hata", f"Bakım silme hatası: {str(e)}")
        else:
            QMessageBox.warning(self, "Uyarı", "Lütfen silinecek bakım kaydını seçin.")

class MaintenanceDialog(QDialog):
    """Bakım ekleme/düzenleme dialog"""
    
    def __init__(self, parent=None, maintenance_data=None):
        super().__init__(parent)
        self.maintenance_data = maintenance_data
        self.data_manager = parent.data_manager if parent else DataManager()
        self.init_ui()
        if maintenance_data:
            self.load_maintenance_data(maintenance_data)
    
    def init_ui(self):
        self.setWindowTitle("Bakım Ekle" if not self.maintenance_data else "Bakım Düzenle")
        self.setMinimumWidth(500)
        self.setModal(True)
        
        layout = QFormLayout(self)
        layout.setSpacing(15)
        
        # Form alanları
        self.arac_plakasi = QComboBox()
        self.load_vehicles()
        
        self.bakim_tarihi = QDateEdit()
        self.bakim_tarihi.setCalendarPopup(True)
        self.bakim_tarihi.setDate(QDate.currentDate())
        
        self.km = QLineEdit()
        
        self.islem_turu = QComboBox()
        self.islem_turu.addItems([
            "Yağ Değişimi", "Fren Bakımı", "Lastik Değişimi", "Motor Bakımı",
            "Elektrik Sistemi", "Klima Bakımı", "Kaporta Boya", "Diğer"
        ])
        
        self.servis_adi = QLineEdit()
        
        self.tutar = QLineEdit()
        self.tutar.setPlaceholderText("0.00")
        
        self.aciklama = QTextEdit()
        self.aciklama.setMaximumHeight(100)
        
        # Form'a ekle
        layout.addRow("Araç Plakası:", self.arac_plakasi)
        layout.addRow("Bakım Tarihi:", self.bakim_tarihi)
        layout.addRow("KM:", self.km)
        layout.addRow("İşlem Türü:", self.islem_turu)
        layout.addRow("Servis Adı:", self.servis_adi)
        layout.addRow("Tutar:", self.tutar)
        layout.addRow("Açıklama:", self.aciklama)
        
        # Butonlar
        btn_layout = QHBoxLayout()
        btn_save = QPushButton("Kaydet")
        btn_save.clicked.connect(self.save_maintenance)
        btn_cancel = QPushButton("İptal")
        btn_cancel.clicked.connect(self.reject)
        
        btn_layout.addWidget(btn_save)
        btn_layout.addWidget(btn_cancel)
        layout.addRow(btn_layout)
    
    def load_vehicles(self):
        """Mevcut araçları yükle"""
        try:
            vehicles = self.data_manager.load_data('araclar')
            self.arac_plakasi.clear()
            
            if not vehicles.empty:
                for _, vehicle in vehicles.iterrows():
                    plaka = vehicle.get('Plaka', '')
                    marka = vehicle.get('Marka', '')
                    model = vehicle.get('Model', '')
                    self.arac_plakasi.addItem(f"{plaka} - {marka} {model}")
        except Exception as e:
            print(f"Araç yükleme hatası: {e}")
    
    def load_maintenance_data(self, maintenance_data):
        """Bakım verilerini yükle"""
        # Araç plakasını ayarla
        arac_text = f"{maintenance_data.get('Araç Plakası', '')} - "
        vehicles = self.data_manager.load_data('araclar')
        if not vehicles.empty:
            vehicle = vehicles[vehicles['Plaka'] == maintenance_data.get('Araç Plakası', '')]
            if not vehicle.empty:
                marka = vehicle.iloc[0].get('Marka', '')
                model = vehicle.iloc[0].get('Model', '')
                arac_text += f"{marka} {model}"
        
        # ComboBox'ta araç seç
        index = self.arac_plakasi.findText(arac_text)
        if index >= 0:
            self.arac_plakasi.setCurrentIndex(index)
        
        # Tarihi yükle
        try:
            bakim_date = QDate.fromString(str(maintenance_data.get('Bakım Tarihi', '')), "dd.MM.yyyy")
            if bakim_date.isValid():
                self.bakim_tarihi.setDate(bakim_date)
        except:
            pass
        
        # Diğer alanları yükle
        self.km.setText(str(maintenance_data.get('KM', '')))
        
        # İşlem türünü ayarla
        islem_index = self.islem_turu.findText(str(maintenance_data.get('İşlem Türü', '')))
        if islem_index >= 0:
            self.islem_turu.setCurrentIndex(islem_index)
        
        self.servis_adi.setText(str(maintenance_data.get('Servis Adı', '')))
        self.tutar.setText(str(maintenance_data.get('Tutar', '')))
        self.aciklama.setPlainText(str(maintenance_data.get('Açıklama', '')))
    
    def save_maintenance(self):
        """Bakım kaydet"""
        if not self.arac_plakasi.currentText():
            QMessageBox.warning(self, "Uyarı", "Lütfen bir araç seçin.")
            return
        
        try:
            maintenance = self.data_manager.load_data('bakimlar')
            
            # Araç plakasını ayır
            arac_text = self.arac_plakasi.currentText()
            plaka = arac_text.split(' - ')[0] if ' - ' in arac_text else arac_text
            
            # Yeni bakım verisi
            new_maintenance = {
                'ID': len(maintenance) + 1 if not maintenance.empty else 1,
                'Araç Plakası': plaka,
                'Bakım Tarihi': self.bakim_tarihi.date().toString("dd.MM.yyyy"),
                'KM': self.km.text().strip(),
                'İşlem Türü': self.islem_turu.currentText(),
                'Servis Adı': self.servis_adi.text().strip(),
                'Tutar': self.tutar.text().strip(),
                'Açıklama': self.aciklama.toPlainText().strip(),
                'Fatura Dosyası': "",
                'Oluşturma Tarihi': datetime.now().strftime("%Y-%m-%d %H:%M:%S")
            }
            
            if self.maintenance_data:
                # Düzenleme modu - eski kaydı sil
                maintenance = maintenance[
                    ~((maintenance['Araç Plakası'] == self.maintenance_data['Araç Plakası']) & 
                      (maintenance['Bakım Tarihi'] == self.maintenance_data['Bakım Tarihi']) &
                      (maintenance['İşlem Türü'] == self.maintenance_data['İşlem Türü']))
                ]
            
            # Yeni veriyi ekle
            new_df = pd.DataFrame([new_maintenance])
            maintenance = pd.concat([maintenance, new_df], ignore_index=True)
            
            # Kaydet
            if self.data_manager.save_data('bakimlar', maintenance):
                QMessageBox.information(self, "Başarılı", 
                    "Bakım kaydı güncellendi." if self.maintenance_data else "Bakım kaydı eklendi.")
                self.accept()
            else:
                QMessageBox.critical(self, "Hata", "Kaydetme hatası.")
                
        except Exception as e:
            QMessageBox.critical(self, "Hata", f"Kaydetme hatası: {str(e)}")

class YakıtlarPanel(QWidget):
    def __init__(self, data_manager):
        super().__init__()
        self.data_manager = data_manager
        self.init_ui()
        self.load_yakitlar()
    
    def init_ui(self):
        layout = QVBoxLayout(self)
        layout.setContentsMargins(20, 20, 20, 20)
        layout.setSpacing(20)
        header_layout = QHBoxLayout()
        title = QLabel("Yakıtlar")
        title.setStyleSheet("font-size: 24px; font-weight: bold; color: #2c3e50;")
        header_layout.addWidget(title)
        header_layout.addStretch()
        btn_add = QPushButton("➕ Yakıt Ekle")
        btn_add.setStyleSheet("""
            QPushButton {
                background-color: #27ae60;
                color: white;
                border: none;
                padding: 10px 20px;
                border-radius: 5px;
                font-weight: bold;
            }
            QPushButton:hover {
                background-color: #229954;
            }
        """)
        btn_add.clicked.connect(self.add_yakit)
        btn_edit = QPushButton("✏️ Düzenle")
        btn_edit.setStyleSheet("""
            QPushButton {
                background-color: #3498db;
                color: white;
                border: none;
                padding: 10px 20px;
                border-radius: 5px;
                font-weight: bold;
            }
            QPushButton:hover {
                background-color: #2980b9;
            }
        """)
        btn_edit.clicked.connect(self.edit_yakit)
        btn_delete = QPushButton("🗑️ Sil")
        btn_delete.setStyleSheet("""
            QPushButton {
                background-color: #e74c3c;
                color: white;
                border: none;
                padding: 10px 20px;
                border-radius: 5px;
                font-weight: bold;
            }
            QPushButton:hover {
                background-color: #c0392b;
            }
        """)
        btn_delete.clicked.connect(self.delete_yakit)
        btn_export = QPushButton("📤 Excel'e Aktar")
        btn_export.setStyleSheet("""
            QPushButton {
                background-color: #f1c40f;
                color: #2c3e50;
                border: none;
                padding: 10px 20px;
                border-radius: 5px;
                font-weight: bold;
            }
            QPushButton:hover {
                background-color: #f39c12;
            }
        """)
        btn_export.clicked.connect(lambda: export_table_to_excel(self.model, self.proxy_model, self))
        
        # Excel import butonu
        btn_import = QPushButton("📥 Excel'den Yükle")
        btn_import.setStyleSheet("""
            QPushButton {
                background-color: #9b59b6;
                color: white;
                border: none;
                padding: 10px 20px;
                border-radius: 5px;
                font-weight: bold;
            }
            QPushButton:hover {
                background-color: #8e44ad;
            }
        """)
        btn_import.clicked.connect(self.import_fuel_excel)
        
        header_layout.addWidget(btn_add)
        header_layout.addWidget(btn_edit)
        header_layout.addWidget(btn_delete)
        header_layout.addWidget(btn_export)
        header_layout.addWidget(btn_import)
        layout.addLayout(header_layout)
        self.search_box = QLineEdit()
        self.search_box.setPlaceholderText("Yakıtlarda ara...")
        self.search_box.textChanged.connect(self.filter_yakitlar)
        layout.addWidget(self.search_box)
        self.table_view = QTableView()
        self.model = QStandardItemModel()
        self.proxy_model = QSortFilterProxyModel()
        self.proxy_model.setSourceModel(self.model)
        self.proxy_model.setFilterCaseSensitivity(Qt.CaseInsensitive)
        self.table_view.setModel(self.proxy_model)
        self.table_view.setSortingEnabled(True)
        self.table_view.setSelectionBehavior(QTableView.SelectRows)
        self.table_view.setSelectionMode(QTableView.SingleSelection)
        self.table_view.horizontalHeader().setStretchLastSection(True)
        
        # Satır numaralarını gizle
        self.table_view.verticalHeader().setVisible(False)
        
        layout.addWidget(self.table_view)
    def load_yakitlar(self):
        try:
            yakitlar = self.data_manager.load_data('yakitlar')
            columns = [
                "Araç Plakası", "Yakıt Tipi", "Tarih", "Litre", "Tutar", "Bayi", "Açıklama"
            ]
            self.model.clear()
            self.model.setHorizontalHeaderLabels(columns)
            
            def tr_money(val):
                try:
                    val = float(str(val).replace(",", ".").replace("₺", "").replace("TL", ""))
                    return f"{val:,.2f} ₺".replace(",", "X").replace(".", ",").replace("X", ".")
                except:
                    return str(val)
            
            for _, yakit in yakitlar.iterrows():
                row = []
                for col in columns:
                    if col == "Tutar":
                        row.append(QStandardItem(tr_money(yakit.get("Tutar", 0))))
                    else:
                        row.append(QStandardItem(str(yakit.get(col, ''))))
                self.model.appendRow(row)
        except Exception as e:
            print(f"Yakıt yükleme hatası: {e}")
    def filter_yakitlar(self, text):
        self.proxy_model.setFilterWildcard(text)
    def add_yakit(self):
        dialog = YakitDialog(self)
        if dialog.exec_() == QDialog.Accepted:
            self.load_yakitlar()
    def edit_yakit(self):
        index = self.table_view.currentIndex()
        if index.isValid():
            row = self.proxy_model.mapToSource(index).row()
            plaka = self.model.item(row, 0).text()
            tarih = self.model.item(row, 2).text()
            yakitlar = self.data_manager.load_data('yakitlar')
            yakit = yakitlar[(yakitlar['Araç Plakası'] == plaka) & (yakitlar['Tarih'] == tarih)]
            if not yakit.empty:
                dialog = YakitDialog(self, yakit.iloc[0].to_dict())
                if dialog.exec_() == QDialog.Accepted:
                    self.load_yakitlar()
        else:
            QMessageBox.warning(self, "Uyarı", "Lütfen düzenlenecek yakıt kaydını seçin.")
    def delete_yakit(self):
        index = self.table_view.currentIndex()
        if index.isValid():
            row = self.proxy_model.mapToSource(index).row()
            plaka = self.model.item(row, 0).text()
            tarih = self.model.item(row, 2).text()
            reply = QMessageBox.question(self, "Onay", f"'{plaka}' plakalı aracın '{tarih}' tarihli yakıt kaydını silmek istediğinizden emin misiniz?", QMessageBox.Yes | QMessageBox.No)
            if reply == QMessageBox.Yes:
                try:
                    yakitlar = self.data_manager.load_data('yakitlar')
                    yakitlar = yakitlar[~((yakitlar['Araç Plakası'] == plaka) & (yakitlar['Tarih'] == tarih))]
                    self.data_manager.save_data('yakitlar', yakitlar)
                    self.load_yakitlar()
                    QMessageBox.information(self, "Başarılı", "Yakıt kaydı silindi.")
                except Exception as e:
                    QMessageBox.critical(self, "Hata", f"Yakıt silme hatası: {str(e)}")
        else:
            QMessageBox.warning(self, "Uyarı", "Lütfen silinecek yakıt kaydını seçin.")
    
    def import_fuel_excel(self):
        """Excel'den yakıt verilerini yükle"""
        try:
            # Önce şablon indirme seçeneği sun
            reply = QMessageBox.question(self, "Yakıt Veri Yükleme", 
                "Yakıt verilerini yüklemek için:\n\n"
                "1️⃣ **Şablon İndir**: Boş Excel şablonu indirip doldurun\n"
                "2️⃣ **Dosya Seç**: Dolu Excel dosyasını seçin\n\n"
                "Şablon indirmek ister misiniz?",
                QMessageBox.Yes | QMessageBox.No | QMessageBox.Cancel)
            
            if reply == QMessageBox.Yes:
                self.download_fuel_template()
                return
            elif reply == QMessageBox.Cancel:
                return
            
            # Dosya seç
            file_path, _ = QFileDialog.getOpenFileName(
                self, "Yakıt Excel Dosyası Seç", "", 
                "Excel Dosyaları (*.xlsx *.xls)"
            )
            
            if not file_path:
                return
            
            # Excel dosyasını oku
            df = pd.read_excel(file_path)
            
            # Sütun isimlerini kontrol et ve normalize et
            column_mapping = {
                'Plaka': 'Araç Plakası',
                'Araç Plakası': 'Araç Plakası',
                'Yakıt Tipi': 'Yakıt Tipi',
                'Tarih': 'Tarih',
                'Litre': 'Litre',
                'Tutar': 'Tutar',
                'Bayi': 'Bayi',
                'İstasyon': 'Bayi',
                'Açıklama': 'Açıklama',
                'Not': 'Açıklama'
            }
            
            # Sütun isimlerini normalize et
            df.columns = [column_mapping.get(col, col) for col in df.columns]
            
            # Gerekli sütunları kontrol et
            required_columns = ['Araç Plakası', 'Tarih', 'Litre', 'Tutar']
            missing_columns = [col for col in required_columns if col not in df.columns]
            
            if missing_columns:
                QMessageBox.critical(self, "Hata", 
                    f"Excel dosyasında gerekli sütunlar eksik:\n{', '.join(missing_columns)}\n\n"
                    f"Beklenen sütunlar: {', '.join(required_columns)}")
                return
            
            # Veri temizleme ve formatlama
            processed_data = []
            
            for _, row in df.iterrows():
                try:
                    # Plaka kontrolü
                    plaka = str(row.get('Araç Plakası', '')).strip()
                    if not plaka or plaka.lower() in ['nan', 'none', '']:
                        continue
                    
                    # Tarih formatını kontrol et ve normalize et
                    tarih = row.get('Tarih', '')
                    if pd.isna(tarih):
                        continue
                    
                    # Tarih formatını kontrol et
                    if isinstance(tarih, str):
                        # String tarih formatlarını kontrol et
                        tarih_formats = ['%d.%m.%Y', '%d/%m/%Y', '%Y-%m-%d', '%d-%m-%Y']
                        parsed_date = None
                        for fmt in tarih_formats:
                            try:
                                parsed_date = datetime.strptime(tarih, fmt)
                                break
                            except:
                                continue
                        
                        if parsed_date:
                            tarih = parsed_date.strftime('%d.%m.%Y')
                        else:
                            continue  # Tarih parse edilemezse satırı atla
                    elif isinstance(tarih, datetime):
                        tarih = tarih.strftime('%d.%m.%Y')
                    else:
                        continue
                    
                    # Litre kontrolü
                    litre = row.get('Litre', 0)
                    if pd.isna(litre):
                        litre = 0
                    try:
                        litre = float(str(litre).replace(',', '.'))
                    except:
                        litre = 0
                    
                    # Tutar kontrolü ve formatlama
                    tutar = row.get('Tutar', 0)
                    if pd.isna(tutar):
                        tutar = 0
                    
                    # Tutar formatını temizle (₺, TL, virgül, nokta)
                    if isinstance(tutar, str):
                        tutar = str(tutar).replace('₺', '').replace('TL', '').replace(' ', '')
                        tutar = tutar.replace(',', '.')
                    
                    try:
                        tutar = float(tutar)
                    except:
                        tutar = 0
                    
                    # Diğer alanlar
                    yakit_tipi = str(row.get('Yakıt Tipi', 'Benzin')).strip()
                    if not yakit_tipi or yakit_tipi.lower() in ['nan', 'none']:
                        yakit_tipi = 'Benzin'
                    
                    bayi = str(row.get('Bayi', '')).strip()
                    if not bayi or bayi.lower() in ['nan', 'none']:
                        bayi = ''
                    
                    aciklama = str(row.get('Açıklama', '')).strip()
                    if not aciklama or aciklama.lower() in ['nan', 'none']:
                        aciklama = ''
                    
                    # Yeni kayıt oluştur
                    new_record = {
                        'ID': len(processed_data) + 1,
                        'Araç Plakası': plaka,
                        'Yakıt Tipi': yakit_tipi,
                        'Tarih': tarih,
                        'Litre': litre,
                        'Tutar': tutar,
                        'Bayi': bayi,
                        'Açıklama': aciklama,
                        'Oluşturma Tarihi': datetime.now().strftime('%Y-%m-%d %H:%M:%S')
                    }
                    
                    processed_data.append(new_record)
                    
                except Exception as e:
                    print(f"Satır işleme hatası: {e}")
                    continue
            
            if not processed_data:
                QMessageBox.warning(self, "Uyarı", "İşlenebilir veri bulunamadı.")
                return
            
            # Mevcut yakıt verilerini yükle
            existing_fuel = self.data_manager.load_data('yakitlar')
            
            # Yeni verileri ekle
            new_df = pd.DataFrame(processed_data)
            combined_fuel = pd.concat([existing_fuel, new_df], ignore_index=True)
            
            # Verileri kaydet
            self.data_manager.save_data('yakitlar', combined_fuel)
            
            # Tabloyu yenile
            self.load_yakitlar()
            
            QMessageBox.information(self, "Başarılı", 
                f"{len(processed_data)} adet yakıt kaydı başarıyla yüklendi.\n\n"
                f"Toplam kayıt sayısı: {len(combined_fuel)}")
            
        except Exception as e:
            QMessageBox.critical(self, "Hata", f"Excel yükleme hatası: {str(e)}")
    
    def download_fuel_template(self):
        """Yakıt veri şablonu indir"""
        try:
            # Kayıt klasörü seç
            folder = QFileDialog.getExistingDirectory(self, "Şablon Kayıt Klasörü Seç")
            if not folder:
                return
            
            # Şablon dosya adı
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"Yakit_Veri_Sablonu_{timestamp}.xlsx"
            filepath = os.path.join(folder, filename)
            
            # Şablon verisi oluştur
            template_data = {
                'Araç Plakası': ['55ABC123', '55DEF456', '55GHI789', '', ''],
                'Yakıt Tipi': ['Benzin', 'Dizel', 'LPG', '', ''],
                'Tarih': ['23.08.2025', '24.08.2025', '25.08.2025', '', ''],
                'Litre': [45.5, 52.0, 38.5, '', ''],
                'Tutar': [1250.50, 1450.00, 1050.75, '', ''],
                'Bayi': ['Shell', 'BP', 'Petrol Ofisi', '', ''],
                'Açıklama': ['Tam doldurma', 'Yarı doldurma', 'Acil yakıt', '', '']
            }
            
            # DataFrame oluştur
            df = pd.DataFrame(template_data)
            
            # Excel dosyası oluştur
            with pd.ExcelWriter(filepath, engine='openpyxl') as writer:
                # Ana veri sayfası
                df.to_excel(writer, sheet_name='Yakıt Verileri', index=False)
                
                # Talimat sayfası
                instructions = {
                    'Sütun': [
                        'Araç Plakası',
                        'Yakıt Tipi', 
                        'Tarih',
                        'Litre',
                        'Tutar',
                        'Bayi',
                        'Açıklama'
                    ],
                    'Açıklama': [
                        'Araç plakası (zorunlu)',
                        'Benzin, Dizel, LPG, Elektrik, Hibrit, Diğer',
                        'Tarih formatı: dd.mm.yyyy (zorunlu)',
                        'Yakıt miktarı litre (zorunlu)',
                        'Tutar TL (zorunlu)',
                        'İstasyon/bayi adı (opsiyonel)',
                        'Ek açıklama (opsiyonel)'
                    ],
                    'Örnek': [
                        '55ABC123',
                        'Benzin',
                        '23.08.2025',
                        '45.5',
                        '1250.50',
                        'Shell',
                        'Tam doldurma'
                    ]
                }
                
                instruction_df = pd.DataFrame(instructions)
                instruction_df.to_excel(writer, sheet_name='Kullanım Talimatları', index=False)
                
                # Excel dosyasını formatla
                workbook = writer.book
                
                # Ana sayfa formatlaması
                worksheet = writer.sheets['Yakıt Verileri']
                
                # Başlık satırını kalın yap
                for col in range(1, len(df.columns) + 1):
                    cell = worksheet.cell(row=1, column=col)
                    cell.font = openpyxl.styles.Font(bold=True)
                    cell.fill = openpyxl.styles.PatternFill(start_color="CCCCCC", end_color="CCCCCC", fill_type="solid")
                
                # Sütun genişliklerini ayarla
                column_widths = [15, 12, 12, 10, 12, 15, 20]
                for i, width in enumerate(column_widths):
                    worksheet.column_dimensions[openpyxl.utils.get_column_letter(i + 1)].width = width
                
                # Talimat sayfası formatlaması
                instruction_worksheet = writer.sheets['Kullanım Talimatları']
                
                # Başlık ekle
                instruction_worksheet.insert_rows(1)
                instruction_worksheet['A1'] = "📋 YAKIT VERİ ŞABLONU KULLANIM TALİMATLARI"
                instruction_worksheet['A1'].font = openpyxl.styles.Font(bold=True, size=14)
                instruction_worksheet.merge_cells('A1:C1')
                
                # Alt başlık
                instruction_worksheet.insert_rows(2)
                instruction_worksheet['A2'] = "Bu şablonu doldurduktan sonra 'Yakıtlar' menüsünden 'Excel'den Yükle' butonunu kullanarak verileri sisteme aktarabilirsiniz."
                instruction_worksheet['A2'].font = openpyxl.styles.Font(size=10, color="666666")
                instruction_worksheet.merge_cells('A2:C2')
                
                # Sütun genişliklerini ayarla
                instruction_worksheet.column_dimensions['A'].width = 20
                instruction_worksheet.column_dimensions['B'].width = 40
                instruction_worksheet.column_dimensions['C'].width = 25
            
            QMessageBox.information(self, "Başarılı", 
                f"Yakıt veri şablonu başarıyla oluşturuldu!\n\n"
                f"Dosya: {filename}\n"
                f"Konum: {folder}\n\n"
                "Şablonu doldurduktan sonra tekrar 'Excel'den Yükle' butonunu kullanın.")
            
        except Exception as e:
            QMessageBox.critical(self, "Hata", f"Şablon oluşturma hatası: {str(e)}")

# Yakıt ekleme/düzenleme dialog
class YakitDialog(QDialog):
    def __init__(self, parent=None, yakit_data=None):
        super().__init__(parent)
        self.yakit_data = yakit_data
        self.data_manager = parent.data_manager if parent else DataManager()
        self.init_ui()
        if yakit_data:
            self.load_yakit_data(yakit_data)
    
    def init_ui(self):
        self.setWindowTitle("Yakıt Ekle" if not self.yakit_data else "Yakıt Düzenle")
        self.setMinimumWidth(500)
        self.setModal(True)
        layout = QFormLayout(self)
        layout.setSpacing(15)
        self.arac_plakasi = QComboBox()
        self.load_vehicles()
        self.yakit_tipi = QComboBox()
        self.yakit_tipi.addItems(["Benzin", "Dizel", "LPG", "Elektrik", "Hibrit", "Diğer"])
        self.tarih = QDateEdit()
        self.tarih.setCalendarPopup(True)
        self.tarih.setDate(QDate.currentDate())
        self.litre = QLineEdit()
        self.litre.setPlaceholderText("0.00")
        self.tutar = QLineEdit()
        self.tutar.setPlaceholderText("0.00")
        self.bayi = QLineEdit()
        self.bayi.setPlaceholderText("Bayi/İstasyon Adı")
        self.aciklama = QTextEdit()
        self.aciklama.setMaximumHeight(100)
        layout.addRow("Araç Plakası:", self.arac_plakasi)
        layout.addRow("Yakıt Tipi:", self.yakit_tipi)
        layout.addRow("Tarih:", self.tarih)
        layout.addRow("Litre:", self.litre)
        layout.addRow("Tutar:", self.tutar)
        layout.addRow("Bayi:", self.bayi)
        layout.addRow("Açıklama:", self.aciklama)
        btn_layout = QHBoxLayout()
        btn_save = QPushButton("Kaydet")
        btn_save.clicked.connect(self.save_yakit)
        btn_cancel = QPushButton("İptal")
        btn_cancel.clicked.connect(self.reject)
        btn_layout.addWidget(btn_save)
        btn_layout.addWidget(btn_cancel)
        layout.addRow(btn_layout)
    def load_vehicles(self):
        try:
            vehicles = self.data_manager.load_data('araclar')
            self.arac_plakasi.clear()
            if not vehicles.empty:
                for _, vehicle in vehicles.iterrows():
                    plaka = vehicle.get('Plaka', '')
                    marka = vehicle.get('Marka', '')
                    model = vehicle.get('Model', '')
                    self.arac_plakasi.addItem(f"{plaka} - {marka} {model}")
        except Exception as e:
            print(f"Araç yükleme hatası: {e}")
    def load_yakit_data(self, yakit_data):
        self.arac_plakasi.setCurrentText(str(yakit_data.get('Araç Plakası', '')))
        self.yakit_tipi.setCurrentText(str(yakit_data.get('Yakıt Tipi', 'Benzin')))
        try:
            tarih = QDate.fromString(str(yakit_data.get('Tarih', '')), "yyyy-MM-dd")
            if tarih.isValid():
                self.tarih.setDate(tarih)
        except:
            pass
        self.litre.setText(str(yakit_data.get('Litre', '')))
        self.tutar.setText(str(yakit_data.get('Tutar', '')))
        self.bayi.setText(str(yakit_data.get('Bayi', '')))
        self.aciklama.setPlainText(str(yakit_data.get('Açıklama', '')))
    def save_yakit(self):
        if not self.arac_plakasi.currentText():
            QMessageBox.warning(self, "Uyarı", "Lütfen bir araç seçin.")
            return
        if not self.litre.text().strip():
            QMessageBox.warning(self, "Uyarı", "Litre alanı zorunludur.")
            return
        if not self.tutar.text().strip():
            QMessageBox.warning(self, "Uyarı", "Tutar alanı zorunludur.")
            return
        try:
            yakitlar = self.data_manager.load_data('yakitlar')
            arac_text = self.arac_plakasi.currentText()
            plaka = arac_text.split(' - ')[0] if ' - ' in arac_text else arac_text
            new_yakit = {
                'ID': len(yakitlar) + 1 if not yakitlar.empty else 1,
                'Araç Plakası': plaka,
                'Yakıt Tipi': self.yakit_tipi.currentText(),
                'Tarih': self.tarih.date().toString("yyyy-MM-dd"),
                'Litre': self.litre.text().strip(),
                'Tutar': self.tutar.text().strip(),
                'Bayi': self.bayi.text().strip(),
                'Açıklama': self.aciklama.toPlainText().strip(),
                'Oluşturma Tarihi': datetime.now().strftime("%Y-%m-%d %H:%M:%S")
            }
            if self.yakit_data:
                yakitlar = yakitlar[~((yakitlar['Araç Plakası'] == self.yakit_data['Araç Plakası']) & (yakitlar['Tarih'] == self.yakit_data['Tarih']))]
            new_df = pd.DataFrame([new_yakit])
            yakitlar = pd.concat([yakitlar, new_df], ignore_index=True)
            if self.data_manager.save_data('yakitlar', yakitlar):
                QMessageBox.information(self, "Başarılı", "Yakıt kaydı eklendi.")
                self.accept()
            else:
                QMessageBox.critical(self, "Hata", "Kaydetme hatası.")
        except Exception as e:
            QMessageBox.critical(self, "Hata", f"Kaydetme hatası: {str(e)}")

class ExpensesPanel(QWidget):
    def __init__(self, data_manager):
        super().__init__()
        self.data_manager = data_manager
        self.init_ui()
        self.load_expenses()
    
    def init_ui(self):
        layout = QVBoxLayout(self)
        layout.setContentsMargins(20, 20, 20, 20)
        layout.setSpacing(20)
        
        # Başlık ve butonlar
        header_layout = QHBoxLayout()
        
        title = QLabel("Yakıt & Giderler")
        title.setStyleSheet("font-size: 24px; font-weight: bold; color: #2c3e50;")
        header_layout.addWidget(title)
        
        header_layout.addStretch()
        
        # Butonlar
        btn_add = QPushButton("➕ Gider Ekle")
        btn_add.setStyleSheet("""
            QPushButton {
                background-color: #27ae60;
                color: white;
                border: none;
                padding: 10px 20px;
                border-radius: 5px;
                font-weight: bold;
            }
            QPushButton:hover {
                background-color: #229954;
            }
        """)
        btn_add.clicked.connect(self.add_expense)
        
        btn_edit = QPushButton("✏️ Düzenle")
        btn_edit.setStyleSheet("""
            QPushButton {
                background-color: #3498db;
                color: white;
                border: none;
                padding: 10px 20px;
                border-radius: 5px;
                font-weight: bold;
            }
            QPushButton:hover {
                background-color: #2980b9;
            }
        """)
        btn_edit.clicked.connect(self.edit_expense)
        
        btn_delete = QPushButton("🗑️ Sil")
        btn_delete.setStyleSheet("""
            QPushButton {
                background-color: #e74c3c;
                color: white;
                border: none;
                padding: 10px 20px;
                border-radius: 5px;
                font-weight: bold;
            }
            QPushButton:hover {
                background-color: #c0392b;
            }
        """)
        btn_delete.clicked.connect(self.delete_expense)
        
        header_layout.addWidget(btn_add)
        header_layout.addWidget(btn_edit)
        header_layout.addWidget(btn_delete)
        
        layout.addLayout(header_layout)
        
        # Tablo
        self.expenses_table = QTableWidget()
        self.expenses_table.setColumnCount(6)
        self.expenses_table.setHorizontalHeaderLabels([
            "Araç Plakası", "Gider Türü", "Tarih", "Tutar", "Açıklama", "Fiş"
        ])
        self.expenses_table.horizontalHeader().setStretchLastSection(True)
        self.expenses_table.setAlternatingRowColors(True)
        self.expenses_table.setSelectionBehavior(QAbstractItemView.SelectRows)
        self.expenses_table.setSelectionMode(QAbstractItemView.SingleSelection)
        
        # Satır numaralarını gizle
        self.expenses_table.verticalHeader().setVisible(False)
        
        layout.addWidget(self.expenses_table)
    
    def load_expenses(self):
        """Giderleri yükle"""
        try:
            expenses = self.data_manager.load_data('giderler')
            self.expenses_table.setRowCount(len(expenses))
            
            for row, (_, expense) in enumerate(expenses.iterrows()):
                self.expenses_table.setItem(row, 0, QTableWidgetItem(str(expense.get('Araç Plakası', ''))))
                self.expenses_table.setItem(row, 1, QTableWidgetItem(str(expense.get('Gider Türü', ''))))
                self.expenses_table.setItem(row, 2, QTableWidgetItem(str(expense.get('Tarih', ''))))
                self.expenses_table.setItem(row, 3, QTableWidgetItem(str(expense.get('Tutar', ''))))
                self.expenses_table.setItem(row, 4, QTableWidgetItem(str(expense.get('Açıklama', ''))))
                self.expenses_table.setItem(row, 5, QTableWidgetItem(str(expense.get('Fiş Dosyası', ''))))
        except Exception as e:
            print(f"Gider yükleme hatası: {e}")
    
    def add_expense(self):
        """Gider ekle"""
        dialog = ExpenseDialog(self)
        if dialog.exec_() == QDialog.Accepted:
            self.load_expenses()
    
    def edit_expense(self):
        """Gider düzenle"""
        current_row = self.expenses_table.currentRow()
        if current_row >= 0:
            QMessageBox.information(self, "Bilgi", "Gider düzenleme özelliği geliştiriliyor...")
        else:
            QMessageBox.warning(self, "Uyarı", "Lütfen düzenlenecek gideri seçin.")
    
    def delete_expense(self):
        """Gider sil"""
        current_row = self.expenses_table.currentRow()
        if current_row >= 0:
            plaka = self.expenses_table.item(current_row, 0).text()
            gider_turu = self.expenses_table.item(current_row, 1).text()
            tarih = self.expenses_table.item(current_row, 2).text()
            reply = QMessageBox.question(self, "Onay", 
                f"'{plaka}' plakalı aracın '{gider_turu}' giderini silmek istediğinizden emin misiniz?",
                QMessageBox.Yes | QMessageBox.No)
            
            if reply == QMessageBox.Yes:
                try:
                    expenses = self.data_manager.load_data('giderler')
                    # Gider kaydını sil
                    expenses = expenses[
                        ~((expenses['Araç Plakası'] == plaka) & 
                          (expenses['Gider Türü'] == gider_turu) &
                          (expenses['Tarih'] == tarih))
                    ]
                    self.data_manager.save_data('giderler', expenses)
                    self.load_expenses()
                    QMessageBox.information(self, "Başarılı", "Gider silindi.")
                except Exception as e:
                    QMessageBox.critical(self, "Hata", f"Gider silme hatası: {str(e)}")
        else:
            QMessageBox.warning(self, "Uyarı", "Lütfen silinecek gideri seçin.")

class TrafficPanel(QWidget):
    def __init__(self, data_manager):
        super().__init__()
        self.data_manager = data_manager
        self.init_ui()
        self.load_traffic_data()
    
    def init_ui(self):
        layout = QVBoxLayout(self)
        layout.setContentsMargins(20, 20, 20, 20)
        layout.setSpacing(20)
        
        # Başlık
        title = QLabel("🚨 Trafik Cezaları & Kazalar")
        title.setStyleSheet("font-size: 24px; font-weight: bold; color: #2c3e50;")
        layout.addWidget(title)
        
        # Tab widget oluştur
        self.tab_widget = QTabWidget()
        
        # Cezalar tab'ı
        self.create_penalties_tab()
        
        # Kazalar tab'ı
        self.create_accidents_tab()
        
        layout.addWidget(self.tab_widget)
    
    def create_penalties_tab(self):
        """Cezalar tab'ı oluştur"""
        penalties_widget = QWidget()
        layout = QVBoxLayout(penalties_widget)
        
        # Başlık ve butonlar
        header_layout = QHBoxLayout()
        
        title = QLabel("📋 Trafik Cezaları")
        title.setStyleSheet("font-size: 18px; font-weight: bold; color: #e74c3c;")
        header_layout.addWidget(title)
        
        header_layout.addStretch()
        
        # Butonlar
        btn_add = QPushButton("➕ Ceza Ekle")
        btn_add.setStyleSheet("""
            QPushButton {
                background-color: #e74c3c;
                color: white;
                border: none;
                padding: 10px 20px;
                border-radius: 5px;
                font-weight: bold;
            }
            QPushButton:hover {
                background-color: #c0392b;
            }
        """)
        btn_add.clicked.connect(self.add_penalty)
        
        btn_edit = QPushButton("✏️ Düzenle")
        btn_edit.setStyleSheet("""
            QPushButton {
                background-color: #3498db;
                color: white;
                border: none;
                padding: 10px 20px;
                border-radius: 5px;
                font-weight: bold;
            }
            QPushButton:hover {
                background-color: #2980b9;
            }
        """)
        btn_edit.clicked.connect(self.edit_penalty)
        
        btn_delete = QPushButton("🗑️ Sil")
        btn_delete.setStyleSheet("""
            QPushButton {
                background-color: #95a5a6;
                color: white;
                border: none;
                padding: 10px 20px;
                border-radius: 5px;
                font-weight: bold;
            }
            QPushButton:hover {
                background-color: #7f8c8d;
            }
        """)
        btn_delete.clicked.connect(self.delete_penalty)
        
        header_layout.addWidget(btn_add)
        header_layout.addWidget(btn_edit)
        header_layout.addWidget(btn_delete)
        
        layout.addLayout(header_layout)
        
        # Tablo
        self.penalties_table = QTableWidget()
        self.penalties_table.setColumnCount(8)
        self.penalties_table.setHorizontalHeaderLabels([
            "Araç Plakası", "Sürücü", "Ceza Tarihi", "Ceza Türü", 
            "Ceza Tutarı", "Ceza Yeri", "Ödeme Durumu", "Ceza Nedeni"
        ])
        self.penalties_table.horizontalHeader().setStretchLastSection(True)
        self.penalties_table.setAlternatingRowColors(True)
        self.penalties_table.setSelectionBehavior(QAbstractItemView.SelectRows)
        self.penalties_table.setSelectionMode(QAbstractItemView.SingleSelection)
        
        # Satır numaralarını gizle
        self.penalties_table.verticalHeader().setVisible(False)
        
        layout.addWidget(self.penalties_table)
        
        self.tab_widget.addTab(penalties_widget, "🚨 Cezalar")
    
    def create_accidents_tab(self):
        """Kazalar tab'ı oluştur"""
        accidents_widget = QWidget()
        layout = QVBoxLayout(accidents_widget)
        
        # Başlık ve butonlar
        header_layout = QHBoxLayout()
        
        title = QLabel("🚗 Kazalar")
        title.setStyleSheet("font-size: 18px; font-weight: bold; color: #e67e22;")
        header_layout.addWidget(title)
        
        header_layout.addStretch()
        
        # Butonlar
        btn_add = QPushButton("➕ Kaza Ekle")
        btn_add.setStyleSheet("""
            QPushButton {
                background-color: #e67e22;
                color: white;
                border: none;
                padding: 10px 20px;
                border-radius: 5px;
                font-weight: bold;
            }
            QPushButton:hover {
                background-color: #d35400;
            }
        """)
        btn_add.clicked.connect(self.add_accident)
        
        btn_edit = QPushButton("✏️ Düzenle")
        btn_edit.setStyleSheet("""
            QPushButton {
                background-color: #3498db;
                color: white;
                border: none;
                padding: 10px 20px;
                border-radius: 5px;
                font-weight: bold;
            }
            QPushButton:hover {
                background-color: #2980b9;
            }
        """)
        btn_edit.clicked.connect(self.edit_accident)
        
        btn_delete = QPushButton("🗑️ Sil")
        btn_delete.setStyleSheet("""
            QPushButton {
                background-color: #95a5a6;
                color: white;
                border: none;
                padding: 10px 20px;
                border-radius: 5px;
                font-weight: bold;
            }
            QPushButton:hover {
                background-color: #7f8c8d;
            }
        """)
        btn_delete.clicked.connect(self.delete_accident)
        
        header_layout.addWidget(btn_add)
        header_layout.addWidget(btn_edit)
        header_layout.addWidget(btn_delete)
        
        layout.addLayout(header_layout)
        
        # Tablo
        self.accidents_table = QTableWidget()
        self.accidents_table.setColumnCount(8)
        self.accidents_table.setHorizontalHeaderLabels([
            "Araç Plakası", "Sürücü", "Kaza Tarihi", "Kaza Yeri", 
            "Kaza Türü", "Hasar Durumu", "Hasar Tutarı", "Sigorta Şirketi"
        ])
        self.accidents_table.horizontalHeader().setStretchLastSection(True)
        self.accidents_table.setAlternatingRowColors(True)
        
        # Satır numaralarını gizle
        self.accidents_table.verticalHeader().setVisible(False)
        self.accidents_table.setSelectionBehavior(QAbstractItemView.SelectRows)
        self.accidents_table.setSelectionMode(QAbstractItemView.SingleSelection)
        
        layout.addWidget(self.accidents_table)
        
        self.tab_widget.addTab(accidents_widget, "🚗 Kazalar")
    
    def load_traffic_data(self):
        """Trafik verilerini yükle"""
        self.load_penalties()
        self.load_accidents()
    
    def load_penalties(self):
        """Cezaları yükle"""
        try:
            penalties = self.data_manager.load_data('cezalar')
            self.penalties_table.setRowCount(len(penalties))
            
            for row, (_, penalty) in enumerate(penalties.iterrows()):
                self.penalties_table.setItem(row, 0, QTableWidgetItem(str(penalty.get('Araç Plakası', ''))))
                self.penalties_table.setItem(row, 1, QTableWidgetItem(str(penalty.get('Sürücü', ''))))
                self.penalties_table.setItem(row, 2, QTableWidgetItem(str(penalty.get('Ceza Tarihi', ''))))
                self.penalties_table.setItem(row, 3, QTableWidgetItem(str(penalty.get('Ceza Türü', ''))))
                self.penalties_table.setItem(row, 4, QTableWidgetItem(str(penalty.get('Ceza Tutarı', ''))))
                self.penalties_table.setItem(row, 5, QTableWidgetItem(str(penalty.get('Ceza Yeri', ''))))
                self.penalties_table.setItem(row, 6, QTableWidgetItem(str(penalty.get('Ödeme Durumu', ''))))
                self.penalties_table.setItem(row, 7, QTableWidgetItem(str(penalty.get('Ceza Nedeni', ''))))
        except Exception as e:
            print(f"Ceza yükleme hatası: {e}")
    
    def load_accidents(self):
        """Kazaları yükle"""
        try:
            accidents = self.data_manager.load_data('kazalar')
            self.accidents_table.setRowCount(len(accidents))
            
            for row, (_, accident) in enumerate(accidents.iterrows()):
                self.accidents_table.setItem(row, 0, QTableWidgetItem(str(accident.get('Araç Plakası', ''))))
                self.accidents_table.setItem(row, 1, QTableWidgetItem(str(accident.get('Sürücü', ''))))
                self.accidents_table.setItem(row, 2, QTableWidgetItem(str(accident.get('Kaza Tarihi', ''))))
                self.accidents_table.setItem(row, 3, QTableWidgetItem(str(accident.get('Kaza Yeri', ''))))
                self.accidents_table.setItem(row, 4, QTableWidgetItem(str(accident.get('Kaza Türü', ''))))
                self.accidents_table.setItem(row, 5, QTableWidgetItem(str(accident.get('Hasar Durumu', ''))))
                self.accidents_table.setItem(row, 6, QTableWidgetItem(str(accident.get('Hasar Tutarı', ''))))
                self.accidents_table.setItem(row, 7, QTableWidgetItem(str(accident.get('Sigorta Şirketi', ''))))
        except Exception as e:
            print(f"Kaza yükleme hatası: {e}")
    
    def add_penalty(self):
        """Ceza ekle"""
        dialog = PenaltyDialog(self)
        if dialog.exec_() == QDialog.Accepted:
            self.load_penalties()
    
    def edit_penalty(self):
        """Ceza düzenle"""
        current_row = self.penalties_table.currentRow()
        if current_row >= 0:
            QMessageBox.information(self, "Bilgi", "Ceza düzenleme özelliği geliştiriliyor...")
        else:
            QMessageBox.warning(self, "Uyarı", "Lütfen düzenlenecek cezayı seçin.")
    
    def delete_penalty(self):
        """Ceza sil"""
        current_row = self.penalties_table.currentRow()
        if current_row >= 0:
            plaka = self.penalties_table.item(current_row, 0).text()
            ceza_tarihi = self.penalties_table.item(current_row, 2).text()
            reply = QMessageBox.question(self, "Onay", 
                f"'{plaka}' plakalı aracın '{ceza_tarihi}' tarihli cezasını silmek istediğinizden emin misiniz?",
                QMessageBox.Yes | QMessageBox.No)
            
            if reply == QMessageBox.Yes:
                try:
                    penalties = self.data_manager.load_data('cezalar')
                    penalties = penalties[
                        ~((penalties['Araç Plakası'] == plaka) & 
                          (penalties['Ceza Tarihi'] == ceza_tarihi))
                    ]
                    self.data_manager.save_data('cezalar', penalties)
                    self.load_penalties()
                    QMessageBox.information(self, "Başarılı", "Ceza silindi.")
                except Exception as e:
                    QMessageBox.critical(self, "Hata", f"Ceza silme hatası: {str(e)}")
        else:
            QMessageBox.warning(self, "Uyarı", "Lütfen silinecek cezayı seçin.")
    
    def add_accident(self):
        """Kaza ekle"""
        dialog = AccidentDialog(self)
        if dialog.exec_() == QDialog.Accepted:
            self.load_accidents()
    
    def edit_accident(self):
        """Kaza düzenle"""
        current_row = self.accidents_table.currentRow()
        if current_row >= 0:
            QMessageBox.information(self, "Bilgi", "Kaza düzenleme özelliği geliştiriliyor...")
        else:
            QMessageBox.warning(self, "Uyarı", "Lütfen düzenlenecek kazayı seçin.")
    
    def delete_accident(self):
        """Kaza sil"""
        current_row = self.accidents_table.currentRow()
        if current_row >= 0:
            plaka = self.accidents_table.item(current_row, 0).text()
            kaza_tarihi = self.accidents_table.item(current_row, 2).text()
            reply = QMessageBox.question(self, "Onay", 
                f"'{plaka}' plakalı aracın '{kaza_tarihi}' tarihli kaza kaydını silmek istediğinizden emin misiniz?",
                QMessageBox.Yes | QMessageBox.No)
            
            if reply == QMessageBox.Yes:
                try:
                    accidents = self.data_manager.load_data('kazalar')
                    accidents = accidents[
                        ~((accidents['Araç Plakası'] == plaka) & 
                          (accidents['Kaza Tarihi'] == kaza_tarihi))
                    ]
                    self.data_manager.save_data('kazalar', accidents)
                    self.load_accidents()
                    QMessageBox.information(self, "Başarılı", "Kaza kaydı silindi.")
                except Exception as e:
                    QMessageBox.critical(self, "Hata", f"Kaza silme hatası: {str(e)}")
        else:
            QMessageBox.warning(self, "Uyarı", "Lütfen silinecek kaza kaydını seçin.")

class PenaltyDialog(QDialog):
    """Ceza ekleme dialog"""
    
    def __init__(self, parent=None):
        super().__init__(parent)
        self.data_manager = parent.data_manager if parent else DataManager()
        self.init_ui()
    
    def init_ui(self):
        self.setWindowTitle("Ceza Ekle")
        self.setMinimumWidth(500)
        self.setModal(True)
        
        layout = QFormLayout(self)
        layout.setSpacing(15)
        
        # Form alanları
        self.arac_plakasi = QComboBox()
        self.load_vehicles()
        
        self.surucu = QComboBox()
        self.load_drivers()
        
        self.ceza_tarihi = QDateEdit()
        self.ceza_tarihi.setCalendarPopup(True)
        self.ceza_tarihi.setDate(QDate.currentDate())
        
        self.ceza_turu = QComboBox()
        self.ceza_turu.addItems([
            "Hız İhlali", "Kırmızı Işık İhlali", "Park İhlali", "Emniyet Kemeri",
            "Cep Telefonu", "Alkollü Araç Kullanma", "Trafik İşareti İhlali", "Diğer"
        ])
        
        self.ceza_tutari = QLineEdit()
        self.ceza_tutari.setPlaceholderText("0.00")
        
        self.ceza_yeri = QLineEdit()
        self.ceza_yeri.setPlaceholderText("Örn: İstanbul, Kadıköy")
        
        self.ceza_nedeni = QTextEdit()
        self.ceza_nedeni.setMaximumHeight(100)
        self.ceza_nedeni.setPlaceholderText("Ceza nedenini detaylı açıklayın...")
        
        self.odeme_durumu = QComboBox()
        self.odeme_durumu.addItems(["Ödenmedi", "Ödendi", "İndirimli Ödendi"])
        
        self.odeme_tarihi = QDateEdit()
        self.odeme_tarihi.setCalendarPopup(True)
        self.odeme_tarihi.setDate(QDate.currentDate())
        self.odeme_tarihi.setEnabled(False)
        self.odeme_durumu.currentTextChanged.connect(self.odeme_durum_kontrol)
        
        # Form'a ekle
        layout.addRow("Araç Plakası:", self.arac_plakasi)
        layout.addRow("Sürücü:", self.surucu)
        layout.addRow("Ceza Tarihi:", self.ceza_tarihi)
        layout.addRow("Ceza Türü:", self.ceza_turu)
        layout.addRow("Ceza Tutarı:", self.ceza_tutari)
        layout.addRow("Ceza Yeri:", self.ceza_yeri)
        layout.addRow("Ceza Nedeni:", self.ceza_nedeni)
        layout.addRow("Ödeme Durumu:", self.odeme_durumu)
        layout.addRow("Ödeme Tarihi:", self.odeme_tarihi)
        
        # Butonlar
        btn_layout = QHBoxLayout()
        btn_save = QPushButton("Kaydet")
        btn_save.clicked.connect(self.save_penalty)
        btn_cancel = QPushButton("İptal")
        btn_cancel.clicked.connect(self.reject)
        
        btn_layout.addWidget(btn_save)
        btn_layout.addWidget(btn_cancel)
        layout.addRow(btn_layout)
    
    def load_vehicles(self):
        """Mevcut araçları yükle"""
        try:
            vehicles = self.data_manager.load_data('araclar')
            self.arac_plakasi.clear()
            
            if not vehicles.empty:
                for _, vehicle in vehicles.iterrows():
                    plaka = vehicle.get('Plaka', '')
                    marka = vehicle.get('Marka', '')
                    model = vehicle.get('Model', '')
                    self.arac_plakasi.addItem(f"{plaka} - {marka} {model}")
        except Exception as e:
            print(f"Araç yükleme hatası: {e}")
    
    def load_drivers(self):
        """Mevcut sürücüleri yükle"""
        try:
            drivers = self.data_manager.load_data('suruculer')
            self.surucu.clear()
            
            if not drivers.empty:
                for _, driver in drivers.iterrows():
                    ad_soyad = driver.get('Ad Soyad', '')
                    tc_kimlik = driver.get('TC Kimlik', '')
                    self.surucu.addItem(f"{ad_soyad} - {tc_kimlik}")
        except Exception as e:
            print(f"Sürücü yükleme hatası: {e}")
    
    def odeme_durum_kontrol(self, text):
        """Ödeme durum kontrolü"""
        self.odeme_tarihi.setEnabled(text == "Ödendi" or text == "İndirimli Ödendi")
    
    def save_penalty(self):
        """Ceza kaydet"""
        if not self.arac_plakasi.currentText():
            QMessageBox.warning(self, "Uyarı", "Lütfen bir araç seçin.")
            return
        
        try:
            penalties = self.data_manager.load_data('cezalar')
            
            # Araç plakasını ayır
            arac_text = self.arac_plakasi.currentText()
            plaka = arac_text.split(' - ')[0] if ' - ' in arac_text else arac_text
            
            # Sürücü adını ayır
            surucu_text = self.surucu.currentText()
            surucu = surucu_text.split(' - ')[0] if ' - ' in surucu_text else surucu_text
            
            # Yeni ceza verisi
            new_penalty = {
                'ID': len(penalties) + 1 if not penalties.empty else 1,
                'Araç Plakası': plaka,
                'Sürücü': surucu,
                'Ceza Tarihi': self.ceza_tarihi.date().toString("dd.MM.yyyy"),
                'Ceza Türü': self.ceza_turu.currentText(),
                'Ceza Tutarı': self.ceza_tutari.text().strip(),
                'Ceza Yeri': self.ceza_yeri.text().strip(),
                'Ceza Nedeni': self.ceza_nedeni.toPlainText().strip(),
                'Ödeme Durumu': self.odeme_durumu.currentText(),
                'Ödeme Tarihi': self.odeme_tarihi.date().toString("yyyy-MM-dd") if self.odeme_durumu.currentText() in ["Ödendi", "İndirimli Ödendi"] else "",
                'Ceza Dosyası': "",
                'Oluşturma Tarihi': datetime.now().strftime("%Y-%m-%d %H:%M:%S")
            }
            
            # Yeni veriyi ekle
            new_df = pd.DataFrame([new_penalty])
            penalties = pd.concat([penalties, new_df], ignore_index=True)
            
            # Kaydet
            if self.data_manager.save_data('cezalar', penalties):
                QMessageBox.information(self, "Başarılı", "Ceza kaydı eklendi.")
                self.accept()
            else:
                QMessageBox.critical(self, "Hata", "Kaydetme hatası.")
                
        except Exception as e:
            QMessageBox.critical(self, "Hata", f"Kaydetme hatası: {str(e)}")

class AccidentDialog(QDialog):
    """Kaza ekleme dialog"""
    
    def __init__(self, parent=None):
        super().__init__(parent)
        self.data_manager = parent.data_manager if parent else DataManager()
        self.init_ui()
    
    def init_ui(self):
        self.setWindowTitle("Kaza Ekle")
        self.setMinimumWidth(500)
        self.setModal(True)
        
        layout = QFormLayout(self)
        layout.setSpacing(15)
        
        # Form alanları
        self.arac_plakasi = QComboBox()
        self.load_vehicles()
        
        self.surucu = QComboBox()
        self.load_drivers()
        
        self.kaza_tarihi = QDateEdit()
        self.kaza_tarihi.setCalendarPopup(True)
        self.kaza_tarihi.setDate(QDate.currentDate())
        
        self.kaza_yeri = QLineEdit()
        self.kaza_yeri.setPlaceholderText("Örn: İstanbul, Kadıköy")
        
        self.kaza_turu = QComboBox()
        self.kaza_turu.addItems([
            "Çarpışma", "Devrilme", "Yaya Kazası", "Hayvan Kazası", 
            "Çukur Kazası", "Hava Koşulları", "Fren Arızası", "Diğer"
        ])
        
        self.hasar_durumu = QComboBox()
        self.hasar_durumu.addItems([
            "Hafif Hasar", "Orta Hasar", "Ağır Hasar", "Total Kayıp"
        ])
        
        self.hasar_tutari = QLineEdit()
        self.hasar_tutari.setPlaceholderText("0.00")
        
        self.sigorta_sirketi = QLineEdit()
        self.sigorta_sirketi.setPlaceholderText("Örn: Anadolu Sigorta")
        
        self.sigorta_dosya_no = QLineEdit()
        self.sigorta_dosya_no.setPlaceholderText("Sigorta dosya numarası")
        
        self.kaza_aciklamasi = QTextEdit()
        self.kaza_aciklamasi.setMaximumHeight(100)
        self.kaza_aciklamasi.setPlaceholderText("Kaza detaylarını açıklayın...")
        
        # Fotoğraf bölümü
        photo_group = QGroupBox("📷 Kaza Fotoğrafı")
        photo_layout = QVBoxLayout(photo_group)
        
        self.photo_label = QLabel("Fotoğraf seçilmedi")
        self.photo_label.setMinimumSize(200, 150)
        self.photo_label.setMaximumSize(300, 200)
        self.photo_label.setStyleSheet("""
            QLabel {
                border: 2px dashed #bdc3c7;
                border-radius: 5px;
                background-color: #f8f9fa;
                color: #6c757d;
                font-size: 12px;
            }
        """)
        self.photo_label.setAlignment(Qt.AlignmentFlag.AlignCenter)
        
        photo_btn_layout = QHBoxLayout()
        btn_select_photo = QPushButton("📁 Fotoğraf Seç")
        btn_select_photo.clicked.connect(self.select_photo)
        btn_clear_photo = QPushButton("🗑️ Temizle")
        btn_clear_photo.clicked.connect(self.clear_photo)
        
        photo_btn_layout.addWidget(btn_select_photo)
        photo_btn_layout.addWidget(btn_clear_photo)
        
        photo_layout.addWidget(self.photo_label)
        photo_layout.addLayout(photo_btn_layout)
        
        self.selected_photo_path = ""
        
        # Form'a ekle
        layout.addRow("Araç Plakası:", self.arac_plakasi)
        layout.addRow("Sürücü:", self.surucu)
        layout.addRow("Kaza Tarihi:", self.kaza_tarihi)
        layout.addRow("Kaza Yeri:", self.kaza_yeri)
        layout.addRow("Kaza Türü:", self.kaza_turu)
        layout.addRow("Hasar Durumu:", self.hasar_durumu)
        layout.addRow("Hasar Tutarı:", self.hasar_tutari)
        layout.addRow("Sigorta Şirketi:", self.sigorta_sirketi)
        layout.addRow("Sigorta Dosya No:", self.sigorta_dosya_no)
        layout.addRow("Kaza Açıklaması:", self.kaza_aciklamasi)
        layout.addRow(photo_group)
        
        # Butonlar
        btn_layout = QHBoxLayout()
        btn_save = QPushButton("Kaydet")
        btn_save.clicked.connect(self.save_accident)
        btn_cancel = QPushButton("İptal")
        btn_cancel.clicked.connect(self.reject)
        
        btn_layout.addWidget(btn_save)
        btn_layout.addWidget(btn_cancel)
        layout.addRow(btn_layout)
    
    def load_vehicles(self):
        """Mevcut araçları yükle"""
        try:
            vehicles = self.data_manager.load_data('araclar')
            self.arac_plakasi.clear()
            
            if not vehicles.empty:
                for _, vehicle in vehicles.iterrows():
                    plaka = vehicle.get('Plaka', '')
                    marka = vehicle.get('Marka', '')
                    model = vehicle.get('Model', '')
                    self.arac_plakasi.addItem(f"{plaka} - {marka} {model}")
        except Exception as e:
            print(f"Araç yükleme hatası: {e}")
    
    def load_drivers(self):
        """Mevcut sürücüleri yükle"""
        try:
            drivers = self.data_manager.load_data('suruculer')
            self.surucu.clear()
            
            if not drivers.empty:
                for _, driver in drivers.iterrows():
                    ad_soyad = driver.get('Ad Soyad', '')
                    tc_kimlik = driver.get('TC Kimlik', '')
                    self.surucu.addItem(f"{ad_soyad} - {tc_kimlik}")
        except Exception as e:
            print(f"Sürücü yükleme hatası: {e}")
    
    def select_photo(self):
        """Fotoğraf seç"""
        file_path, _ = QFileDialog.getOpenFileName(
            self, "Kaza Fotoğrafı Seç", "", 
            "Resim Dosyaları (*.jpg *.jpeg *.png *.bmp *.gif);;Tüm Dosyalar (*)"
        )
        
        if file_path:
            try:
                pixmap = QPixmap(file_path)
                if not pixmap.isNull():
                    scaled_pixmap = pixmap.scaled(
                        self.photo_label.size(), 
                        Qt.AspectRatioMode.KeepAspectRatio, 
                        Qt.TransformationMode.SmoothTransformation
                    )
                    self.photo_label.setPixmap(scaled_pixmap)
                    self.photo_label.setStyleSheet("""
                        QLabel {
                            border: 2px solid #27ae60;
                            border-radius: 5px;
                            background-color: white;
                        }
                    """)
                    self.selected_photo_path = file_path
                else:
                    QMessageBox.warning(self, "Hata", "Seçilen dosya geçerli bir resim değil.")
            except Exception as e:
                QMessageBox.critical(self, "Hata", f"Fotoğraf yükleme hatası: {str(e)}")
    
    def clear_photo(self):
        """Fotoğrafı temizle"""
        self.photo_label.clear()
        self.photo_label.setText("Fotoğraf seçilmedi")
        self.photo_label.setStyleSheet("""
            QLabel {
                border: 2px dashed #bdc3c7;
                border-radius: 5px;
                background-color: #f8f9fa;
                color: #6c757d;
                font-size: 12px;
            }
        """)
        self.selected_photo_path = ""
    
    def save_accident(self):
        """Kaza kaydet"""
        if not self.arac_plakasi.currentText():
            QMessageBox.warning(self, "Uyarı", "Lütfen bir araç seçin.")
            return
        
        try:
            accidents = self.data_manager.load_data('kazalar')
            
            # Araç plakasını ayır
            arac_text = self.arac_plakasi.currentText()
            plaka = arac_text.split(' - ')[0] if ' - ' in arac_text else arac_text
            
            # Sürücü adını ayır
            surucu_text = self.surucu.currentText()
            surucu = surucu_text.split(' - ')[0] if ' - ' in surucu_text else surucu_text
            
            # Yeni kaza verisi
            new_accident = {
                'ID': len(accidents) + 1 if not accidents.empty else 1,
                'Araç Plakası': plaka,
                'Sürücü': surucu,
                'Kaza Tarihi': self.kaza_tarihi.date().toString("yyyy-MM-dd"),
                'Kaza Yeri': self.kaza_yeri.text().strip(),
                'Kaza Türü': self.kaza_turu.currentText(),
                'Hasar Durumu': self.hasar_durumu.currentText(),
                'Hasar Tutarı': self.hasar_tutari.text().strip(),
                'Sigorta Şirketi': self.sigorta_sirketi.text().strip(),
                'Sigorta Dosya No': self.sigorta_dosya_no.text().strip(),
                'Kaza Açıklaması': self.kaza_aciklamasi.toPlainText().strip(),
                'Kaza Dosyası': self.selected_photo_path if self.selected_photo_path else "",
                'Oluşturma Tarihi': datetime.now().strftime("%Y-%m-%d %H:%M:%S")
            }
            
            # Yeni veriyi ekle
            new_df = pd.DataFrame([new_accident])
            accidents = pd.concat([accidents, new_df], ignore_index=True)
            
            # Kaydet
            if self.data_manager.save_data('kazalar', accidents):
                QMessageBox.information(self, "Başarılı", "Kaza kaydı eklendi.")
                self.accept()
            else:
                QMessageBox.critical(self, "Hata", "Kaydetme hatası.")
                
        except Exception as e:
            QMessageBox.critical(self, "Hata", f"Kaydetme hatası: {str(e)}")

class ExpenseDialog(QDialog):
    """Gider ekleme dialog"""
    
    def __init__(self, parent=None):
        super().__init__(parent)
        self.data_manager = parent.data_manager if parent else DataManager()
        self.init_ui()
    
    def init_ui(self):
        self.setWindowTitle("Gider Ekle")
        self.setMinimumWidth(500)
        self.setModal(True)
        
        layout = QFormLayout(self)
        layout.setSpacing(15)
        
        # Form alanları
        self.arac_plakasi = QComboBox()
        self.load_vehicles()
        
        self.gider_turu = QComboBox()
        self.gider_turu.addItems([
            "Yakıt", "Otoyol", "Otopark", "Temizlik", "Yıkama", "Diğer"
        ])
        
        self.tarih = QDateEdit()
        self.tarih.setCalendarPopup(True)
        self.tarih.setDate(QDate.currentDate())
        
        self.tutar = QLineEdit()
        self.tutar.setPlaceholderText("0.00")
        
        self.aciklama = QTextEdit()
        self.aciklama.setMaximumHeight(100)
        
        # Form'a ekle
        layout.addRow("Araç Plakası:", self.arac_plakasi)
        layout.addRow("Gider Türü:", self.gider_turu)
        layout.addRow("Tarih:", self.tarih)
        layout.addRow("Tutar:", self.tutar)
        layout.addRow("Açıklama:", self.aciklama)
        
        # Butonlar
        btn_layout = QHBoxLayout()
        btn_save = QPushButton("Kaydet")
        btn_save.clicked.connect(self.save_expense)
        btn_cancel = QPushButton("İptal")
        btn_cancel.clicked.connect(self.reject)
        
        btn_layout.addWidget(btn_save)
        btn_layout.addWidget(btn_cancel)
        layout.addRow(btn_layout)
    
    def load_vehicles(self):
        """Mevcut araçları yükle"""
        try:
            vehicles = self.data_manager.load_data('araclar')
            self.arac_plakasi.clear()
            
            if not vehicles.empty:
                for _, vehicle in vehicles.iterrows():
                    plaka = vehicle.get('Plaka', '')
                    marka = vehicle.get('Marka', '')
                    model = vehicle.get('Model', '')
                    self.arac_plakasi.addItem(f"{plaka} - {marka} {model}")
        except Exception as e:
            print(f"Araç yükleme hatası: {e}")
    
    def save_expense(self):
        """Gider kaydet"""
        if not self.arac_plakasi.currentText():
            QMessageBox.warning(self, "Uyarı", "Lütfen bir araç seçin.")
            return
        
        try:
            expenses = self.data_manager.load_data('giderler')
            
            # Araç plakasını ayır
            arac_text = self.arac_plakasi.currentText()
            plaka = arac_text.split(' - ')[0] if ' - ' in arac_text else arac_text
            
            # Yeni gider verisi
            new_expense = {
                'ID': len(expenses) + 1 if not expenses.empty else 1,
                'Araç Plakası': plaka,
                'Gider Türü': self.gider_turu.currentText(),
                'Tarih': self.tarih.date().toString("yyyy-MM-dd"),
                'Tutar': self.tutar.text().strip(),
                'Açıklama': self.aciklama.toPlainText().strip(),
                'Fiş Dosyası': "",
                'Oluşturma Tarihi': datetime.now().strftime("%Y-%m-%d %H:%M:%S")
            }
            
            # Yeni veriyi ekle
            new_df = pd.DataFrame([new_expense])
            expenses = pd.concat([expenses, new_df], ignore_index=True)
            
            # Kaydet
            if self.data_manager.save_data('giderler', expenses):
                QMessageBox.information(self, "Başarılı", "Gider eklendi.")
                self.accept()
            else:
                QMessageBox.critical(self, "Hata", "Kaydetme hatası.")
                
        except Exception as e:
            QMessageBox.critical(self, "Hata", f"Kaydetme hatası: {str(e)}")
            

class ReportsPanel(QWidget):
    def __init__(self, data_manager):
        super().__init__()
        self.data_manager = data_manager
        self.init_ui()
        self.load_reports()

    @staticmethod
    def tr_money(val):
        try:
            return f"{float(val):,.2f} TL".replace(",", "_").replace(".", ",").replace("_", ".")
        except:
            return str(val)
    
    def init_ui(self):
        layout = QVBoxLayout(self)
        # Başlık
        title = QLabel("Raporlar & İstatistikler")
        title.setStyleSheet("font-size: 24px; font-weight: bold; color: #2c3e50;")
        layout.addWidget(title)
        # Rapor seçenekleri
        report_group = QGroupBox("Rapor Türleri")
        report_group.setStyleSheet("""
            QGroupBox {
                font-weight: bold;
                border: 2px solid #bdc3c7;
                border-radius: 5px;
                margin-top: 10px;
                padding-top: 10px;
            }
        """)
        
        report_layout = QGridLayout(report_group)
        
        # Rapor butonları
        btn_vehicle_info_report = QPushButton("📄 Araç Bilgi Formu")
        btn_vehicle_info_report.setStyleSheet("""
            QPushButton {
                background-color: #8e44ad;
                color: white;
                border: none;
                padding: 15px;
                border-radius: 5px;
                font-weight: bold;
                font-size: 14px;
            }
            QPushButton:hover {
                background-color: #6c3483;
            }
        """)
        btn_vehicle_info_report.clicked.connect(self.create_vehicle_info_report)
        report_layout.addWidget(btn_vehicle_info_report, 1, 1)  # Uygun bir yere ekleyin
        
        btn_vehicle_report = QPushButton("🚗 Araç Raporu")
        btn_vehicle_report.setStyleSheet("""
            QPushButton {
                background-color: #3498db;
                color: white;
                border: none;
                padding: 15px;
                border-radius: 5px;
                font-weight: bold;
                font-size: 14px;
            }
            QPushButton:hover {
                background-color: #2980b9;
            }
        """)
        btn_vehicle_report.clicked.connect(self.generate_vehicle_report)
        
        btn_maintenance_report = QPushButton("🔧 Bakım Raporu")
        btn_maintenance_report.setStyleSheet("""
            QPushButton {
                background-color: #e67e22;
                color: white;
                border: none;
                padding: 15px;
                border-radius: 5px;
                font-weight: bold;
                font-size: 14px;
            }
            QPushButton:hover {
                background-color: #d35400;
            }
        """)
        btn_maintenance_report.clicked.connect(self.generate_maintenance_report)
        
        btn_expense_report = QPushButton("💰 Gider Raporu")
        btn_expense_report.setStyleSheet("""
            QPushButton {
                background-color: #27ae60;
                color: white;
                border: none;
                padding: 15px;
                border-radius: 5px;
                font-weight: bold;
                font-size: 14px;
            }
            QPushButton:hover {
                background-color: #229954;
            }
        """)
        btn_expense_report.clicked.connect(self.generate_expense_report)
        
        # Butonları layout'a ekle
        report_layout.addWidget(btn_vehicle_report, 0, 0)
        report_layout.addWidget(btn_maintenance_report, 0, 1)
        report_layout.addWidget(btn_expense_report, 1, 0)
        report_group.setLayout(report_layout)
        layout.addWidget(report_group)
        
        # İstatistikler
        stats_group = QGroupBox("Özet İstatistikler")
        stats_group.setStyleSheet("""
            QGroupBox {
                font-weight: bold;
                border: 2px solid #bdc3c7;
                border-radius: 5px;
                margin-top: 10px;
                padding-top: 10px;
            }
        """)
        
        stats_layout = QVBoxLayout(stats_group)
        
        self.stats_text = QTextEdit()
        self.stats_text.setMaximumHeight(200)
        self.stats_text.setReadOnly(True)
        stats_layout.addWidget(self.stats_text)
        
        layout.addWidget(stats_group)
    
    def load_reports(self):
        """Rapor istatistiklerini yükle"""
        try:
            # Araç sayısı
            vehicles = self.data_manager.load_data('araclar')
            vehicle_count = len(vehicles) if not vehicles.empty else 0
            
            # Bakım sayısı
            maintenance = self.data_manager.load_data('bakimlar')
            maintenance_count = len(maintenance) if not maintenance.empty else 0
            
            # Gider sayısı
            expenses = self.data_manager.load_data('giderler')
            expense_count = len(expenses) if not expenses.empty else 0
            
            # Sürücü sayısı
            drivers = self.data_manager.load_data('suruculer')
            driver_count = len(drivers) if not drivers.empty else 0
            
            # Toplam gider
            total_expense = 0
            if not expenses.empty:
                for _, expense in expenses.iterrows():
                    amount = expense.get('Tutar', '0')
                    try:
                        total_expense += float(amount)
                    except:
                        pass
            
            # İstatistik metni oluştur
            stats_text = f"""
📊 ÖZET İSTATİSTİKLER

🚗 Toplam Araç: {vehicle_count}
👨‍💼 Toplam Sürücü: {driver_count}
🔧 Toplam Bakım: {maintenance_count}
💰 Toplam Gider: {total_expense:.2f} ₺
📋 Toplam Gider Kaydı: {expense_count}

💡 Raporlar panelinden detaylı raporlar oluşturabilirsiniz.
            """
            
            self.stats_text.setText(stats_text.strip())
            
        except Exception as e:
            self.stats_text.setText(f"İstatistik yükleme hatası: {str(e)}")
    
    def create_vehicle_info_report(self):
        from PyQt5.QtWidgets import QInputDialog, QMessageBox, QFileDialog
        import os
        import pandas as pd
        import math
        from datetime import datetime

        plaka, ok = QInputDialog.getText(self, "Plaka Girin", "Raporunu almak istediğiniz plakayı girin:")
        if not ok or not plaka:
            return

        # Araç verisini bul
        vehicles = self.data_manager.load_data('araclar')
        arac = vehicles[vehicles['Plaka'].astype(str).str.strip().str.upper() == plaka.strip().upper()]
        if arac.empty:
            QMessageBox.warning(self, "Uyarı", f"{plaka} plakalı araç bulunamadı.")
            return
        arac = arac.iloc[0]

        # Fotoğraf yolu
        foto_path = arac.get('Evrak Yolu', '') or arac.get('Fotoğraf Yolu', '')
        if not isinstance(foto_path, str) or not foto_path or (isinstance(foto_path, float) and math.isnan(foto_path)) or not os.path.exists(str(foto_path)):
            foto_path = "araba_icon.png"

        # Bakım verilerini al
        bakimlar = self.data_manager.load_data('bakimlar')
        giderler = self.data_manager.load_data('giderler')
        
        # Bu araca ait bakım ve giderleri filtrele
        arac_bakimlar = bakimlar[bakimlar['Araç Plakası'].astype(str).str.strip().str.upper() == plaka.strip().upper()]
        arac_giderler = giderler[giderler['Araç Plakası'].astype(str).str.strip().str.upper() == plaka.strip().upper()]

        # Dinamik yıl hesaplama - Araçın filoda olduğu yılları bul
        current_year = datetime.now().year
        all_years = set()
        
        # Bakım tarihlerinden yılları çıkar
        for tarih in arac_bakimlar['Bakım Tarihi'].astype(str):
            try:
                if tarih and tarih != 'nan':
                    year = int(tarih[:4])  # İlk 4 karakter yıl
                    all_years.add(year)
            except:
                pass
        
        # Gider tarihlerinden yılları çıkar
        for tarih in arac_giderler['Tarih'].astype(str):
            try:
                if tarih and tarih != 'nan':
                    year = int(tarih[:4])  # İlk 4 karakter yıl
                    all_years.add(year)
            except:
                pass
        
        # Araç yılından günümüze kadar olan yılları ekle
        try:
            arac_yili = int(str(arac.get('Yıl', current_year))[:4])
            for year in range(arac_yili, current_year + 1):
                all_years.add(year)
        except:
            pass
        
        # En az son 3 yıl olsun
        if len(all_years) < 3:
            for year in range(current_year-2, current_year+1):
                all_years.add(year)
        
        # Yılları sırala
        years_list = sorted(list(all_years))
        yearly_summary = {}
        
        for year in years_list:
            yearly_summary[year] = {
                'bakim_sayisi': 0,
                'bakim_tutari': 0,
                'gider_sayisi': 0,
                'gider_tutari': 0
            }
            
            # O yılın bakımları
            year_bakimlar = arac_bakimlar[arac_bakimlar['Bakım Tarihi'].astype(str).str.contains(str(year))]
            yearly_summary[year]['bakim_sayisi'] = len(year_bakimlar)
            yearly_summary[year]['bakim_tutari'] = year_bakimlar['Tutar'].astype(float).sum() if not year_bakimlar.empty else 0
            
            # O yılın giderleri
            year_giderler = arac_giderler[arac_giderler['Tarih'].astype(str).str.contains(str(year))]
            yearly_summary[year]['gider_sayisi'] = len(year_giderler)
            yearly_summary[year]['gider_tutari'] = year_giderler['Tutar'].astype(float).sum() if not year_giderler.empty else 0
        
        # Sadece veri içeren yılları filtrele
        years_with_data = []
        for year in years_list:
            summary = yearly_summary[year]
            if summary['bakim_sayisi'] > 0 or summary['gider_sayisi'] > 0:
                years_with_data.append(year)
        
        # Eğer hiç veri yoksa, en az son 3 yılı göster
        if not years_with_data:
            current_year = datetime.now().year
            for year in range(current_year-2, current_year+1):
                years_with_data.append(year)
                if year not in yearly_summary:
                    yearly_summary[year] = {
                        'bakim_sayisi': 0,
                        'bakim_tutari': 0,
                        'gider_sayisi': 0,
                        'gider_tutari': 0
                    }

        # Format seçimi - Butonlu dialog
        format_choice = self.show_format_selection_dialog()
        if not format_choice:
            return

        # Rapor verilerini hazırla
        rapor_data = []
        
        # 1. Araç Bilgileri - Başlık satırı kaldırıldı
        # Fotoğraf varsa icon, yoksa boş
        foto_icon = "🚗" if foto_path and foto_path != "araba_icon.png" and os.path.exists(foto_path) else "🚗"
        
        # Son güncel KM bilgilerini al ve karşılaştır
        son_km_araclar = arac.get('Son KM', '')
        son_km_araclar = "" if str(son_km_araclar).lower() == "nan" else son_km_araclar
        
        # Bakımlar tablosundan son güncel KM bilgisini al
        arac_bakimlar = bakimlar[bakimlar['Araç Plakası'].astype(str).str.strip().str.upper() == plaka.strip().upper()]
        son_km_bakimlar = 0
        if not arac_bakimlar.empty:
            # KM alanındaki değerleri sayısal değerlere çevir ve en yüksek olanı al
            km_values = arac_bakimlar['KM'].dropna()
            if not km_values.empty:
                try:
                    # Sayısal olmayan değerleri filtrele
                    numeric_km_values = []
                    for km_val in km_values:
                        try:
                            numeric_km = float(str(km_val).replace(',', '').replace(' ', ''))
                            numeric_km_values.append(numeric_km)
                        except:
                            continue
                    
                    if numeric_km_values:
                        son_km_bakimlar = max(numeric_km_values)
                except:
                    son_km_bakimlar = 0
        
        # İki KM değerini karşılaştır ve en büyük olanı al
        guncel_km = 0
        try:
            if son_km_araclar and str(son_km_araclar).replace(',', '').replace(' ', '').replace('.', '').isdigit():
                km_araclar = float(str(son_km_araclar).replace(',', '').replace(' ', ''))
                guncel_km = max(guncel_km, km_araclar)
        except:
            pass
        
        guncel_km = max(guncel_km, son_km_bakimlar)
        
        # Güncel KM değerini formatla
        if guncel_km > 0:
            guncel_km_str = f"{guncel_km:,.0f}".replace(",", ".")
        else:
            guncel_km_str = "KM bilgisi bulunamadı"
        
        for key, label in [
            ('Plaka', 'PLAKA'),
            ('Marka', 'MARKA'),
            ('Model', 'MODEL'),
            ('Yıl', 'YIL'),
            ('Şasi No', 'ŞASİ NO'),
            ('Sigorta Bitiş', 'SİGORTA BİTİŞ'),
            ('Kasko Bitiş', 'KASKO BİTİŞ'),
            ('Muayene Bitiş', 'MUAYENE BİTİŞ'),
            ('Birim', 'BİRİM'),
            ('Şoför', 'ŞOFÖR'),
            ('Hizmet', 'HİZMET')
        ]:
            value = arac.get(key, '')
            value = "" if str(value).lower() == "nan" else value
            rapor_data.append(["", label, value])
        
        # Güncel KM bilgisini ekle
        rapor_data.append(["", "GÜNCEL KM", guncel_km_str])
        
        # 2. Boş satır
        rapor_data.append(["", "", ""])
        
        # 3. Yıllık Bakım ve Gider Özeti - Kompakt Tablo
        rapor_data.append(["", "YILLIK BAKIM VE GİDER ÖZETİ", ""])
        
        # Tablo başlığı - Boş sütun kaldırıldı, gider sayısı ve gider tutarı sütunları kaldırıldı
        rapor_data.append(["Yıl", "Bakım Sayısı", "Bakım Tutarı", "Toplam Tutar"])
        
        # Türkçe para birimi formatı için yardımcı fonksiyon
        def tr_money(val):
            try:
                return f"{val:,.2f}".replace(",", "X").replace(".", ",").replace("X", ".") + " TL"
            except:
                return f"{val} TL"
        
        # Yıllık veriler - Boş sütun kaldırıldı, gider sayısı ve gider tutarı sütunları kaldırıldı
        # 0 değerli yıllar gizlendi
        for year in years_with_data:
            summary = yearly_summary[year]
            # Sadece bakım sayısı 0'dan büyük olan yılları göster
            if summary['bakim_sayisi'] > 0:
                rapor_data.append([
                    f"{year}",
                    f"{summary['bakim_sayisi']}",
                    tr_money(summary['bakim_tutari']),
                    tr_money(summary['bakim_tutari'])  # Toplam tutar artık sadece bakım tutarı
                ])
        
        # Toplam satırı - Sadece gösterilen yıllar için hesaplama (0 değerli yıllar hariç)
        total_bakim_sayisi = sum(yearly_summary[year]['bakim_sayisi'] for year in years_with_data if yearly_summary[year]['bakim_sayisi'] > 0)
        total_bakim_tutari = sum(yearly_summary[year]['bakim_tutari'] for year in years_with_data if yearly_summary[year]['bakim_sayisi'] > 0)
        
        # Sadece veri varsa toplam satırını ekle - Boş satır kaldırıldı
        if total_bakim_sayisi > 0:
            rapor_data.append([
                "TOPLAM",
                f"{total_bakim_sayisi}",
                tr_money(total_bakim_tutari),
                tr_money(total_bakim_tutari)  # Toplam tutar artık sadece bakım tutarı
            ])

        # DataFrame oluştur - Dinamik sütun sayısı
        max_cols = max(len(row) for row in rapor_data)
        columns = [f"Sütun_{i+1}" for i in range(max_cols)]
        df = pd.DataFrame(rapor_data, columns=columns)  # type: ignore

        success_count = 0
        total_formats = 1 if format_choice != "Tümü" else 2

        # Word oluştur
        if format_choice in ["Word (.docx)", "Tümü"]:
            try:
                from docx import Document
                from docx.shared import Inches, Pt, Cm, RGBColor
                from docx.enum.text import WD_ALIGN_PARAGRAPH
                from docx.oxml import parse_xml
                import os
                import pandas as pd
                file_path, _ = QFileDialog.getSaveFileName(self, "Word Olarak Kaydet", f"{plaka}_bilgi_formu.docx", "Word Dosyası (*.docx)")
                if file_path:
                    doc = Document()
                    section = doc.sections[0]
                    section.left_margin = Cm(1)
                    section.right_margin = Cm(1)
                    section.top_margin = Cm(1)
                    section.bottom_margin = Cm(1)
                    plaka = plaka.upper()
                    # Fotoğraf
                    foto_path = os.path.join("veri", "arac_fotograflari", f"{plaka}.jpg")
                    if not os.path.exists(foto_path):
                        foto_path = os.path.join("veri", "arac_fotograflari", f"{plaka}.png")
                    if not os.path.exists(foto_path):
                        foto_path = "araba_icon.png"
                    if os.path.exists(foto_path):
                        try:
                            doc.add_picture(foto_path, width=Inches(2.0))
                            last_paragraph = doc.paragraphs[-1]
                            last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        except Exception as e:
                            print("Word fotoğraf ekleme hatası:", repr(e))
                    # Başlık
                    title = doc.add_heading(f'{plaka} PLAKALI ARAÇ BİLGİ FORMU', 0)
                    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    doc.add_paragraph("")
                    # Araç Bilgileri Tablosu (4 sütun, PDF ile aynı genişlik ve font)
                    arac_bilgi_pairs = []
                    for row in rapor_data:
                        if len(row) > 1 and row[1] == "YILLIK BAKIM VE GİDER ÖZETİ":
                            break
                        if len(row) >= 3:
                            baslik = row[1]
                            deger = row[2]
                            if baslik in ["KASKO BİTİŞ", "SİGORTA BİTİŞ", "MUAYENE BİTİŞ"] and deger:
                                try:
                                    tarih = pd.to_datetime(deger, dayfirst=True)
                                    deger = tarih.strftime("%d.%m.%Y")
                                except:
                                    pass
                            arac_bilgi_pairs.append((baslik, deger))
                    rows_4sutun = []
                    for i in range(0, len(arac_bilgi_pairs), 2):
                        left = arac_bilgi_pairs[i]
                        right = arac_bilgi_pairs[i+1] if i+1 < len(arac_bilgi_pairs) else ("", "")
                        rows_4sutun.append([left[0], left[1], right[0], right[1]])
                    table1 = doc.add_table(rows=len(rows_4sutun), cols=4)
                    table1.style = 'Table Grid'
                    table1.allow_autofit = False
                    for i, row_data in enumerate(rows_4sutun):
                        for j in range(4):
                            val = row_data[j] if j < len(row_data) else ""
                            cell = table1.cell(i, j)
                            cell.text = str(val)
                            for paragraph in cell.paragraphs:
                                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                for run in paragraph.runs:
                                    run.font.name = 'Calibri'
                                    run.font.size = Pt(10)
                            cell._tc.get_or_add_tcPr().append(
                                parse_xml(r'<w:tcW xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:w="2692" w:type="dxa"/>')
                            )
                    doc.add_paragraph("")
                    # Yıllık Özet Tablosu (4 sütun, başlık renkli, PDF ile aynı)
                    yillik_ozet_rows = []
                    yillik_ozet_basladi = False
                    for row in rapor_data:
                        if len(row) > 1 and row[1] == "YILLIK BAKIM VE GİDER ÖZETİ":
                            yillik_ozet_basladi = True
                            continue
                        if yillik_ozet_basladi:
                            yillik_ozet_rows.append(row)
                    table2 = doc.add_table(rows=len(yillik_ozet_rows), cols=4)
                    table2.style = 'Table Grid'
                    table2.allow_autofit = False
                    for i, row_data in enumerate(yillik_ozet_rows):
                        for j in range(4):
                            val = row_data[j] if j < len(row_data) else ""
                            cell = table2.cell(i, j)
                            cell.text = str(val)
                            for paragraph in cell.paragraphs:
                                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                for run in paragraph.runs:
                                    run.font.name = 'Calibri'
                                    run.font.size = Pt(10)
                            cell._tc.get_or_add_tcPr().append(
                                parse_xml(r'<w:tcW xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:w="2692" w:type="dxa"/>')
                            )
                        # Başlık satırını renkli ve kalın yap
                        if i == 0:
                            for cell in table2.rows[i].cells:
                                cell_par = cell.paragraphs[0]
                                for run in cell_par.runs:
                                    run.bold = True
                                    run.font.size = Pt(11)
                                    run.font.color.rgb = RGBColor(255,255,255)  # Beyaz
                                cell._tc.get_or_add_tcPr().append(
                                    parse_xml(r'<w:shd xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:fill="A23B72"/>')
                                )
                        # Toplam satırını renkli ve kalın yap
                        if row_data[0] == "TOPLAM":
                            for cell in table2.rows[i].cells:
                                cell_par = cell.paragraphs[0]
                                for run in cell_par.runs:
                                    run.bold = True
                                    run.font.size = Pt(11)
                                    run.font.color.rgb = RGBColor(255,255,255)  # Beyaz
                                cell._tc.get_or_add_tcPr().append(
                                    parse_xml(r'<w:shd xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:fill="F18F01"/>')
                                )
                    doc.save(file_path)
                    success_count += 1
            except ImportError:
                QMessageBox.warning(self, "Kütüphane Hatası", "Word oluşturmak için python-docx kütüphanesi gerekli. Lütfen 'pip install python-docx' komutunu çalıştırın.")
            except Exception as e:
                QMessageBox.warning(self, "Word Hatası", f"Word dosyası oluşturulamadı: {str(e)}")


        # PDF oluştur
        if format_choice in ["PDF (.pdf)", "Tümü"]:
            try:
                from reportlab.lib.pagesizes import A4
                from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer, Image
                from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
                from reportlab.lib.units import inch, cm
                from reportlab.lib import colors
                from reportlab.pdfbase import pdfmetrics
                from reportlab.pdfbase.ttfonts import TTFont
                from reportlab.pdfbase.pdfmetrics import stringWidth
                
                
                # Türkçe karakter desteği için font yükleme
                try:
                    # Windows'ta varsayılan fontları dene
                    font_paths = [
                        "C:/Windows/Fonts/arial.ttf",
                        "C:/Windows/Fonts/calibri.ttf", 
                        "C:/Windows/Fonts/tahoma.ttf",
                        "C:/Windows/Fonts/verdana.ttf"
                    ]
                    
                    turkish_font = None
                    for font_path in font_paths:
                        if os.path.exists(font_path):
                            try:
                                pdfmetrics.registerFont(TTFont('TurkishFont', font_path))
                                turkish_font = 'TurkishFont'
                                break
                            except:
                                continue
                    
                    # Font bulunamazsa varsayılan font kullan
                    if not turkish_font:
                        turkish_font = 'Helvetica'
                except:
                    turkish_font = 'Helvetica'
                
                file_path, _ = QFileDialog.getSaveFileName(self, "PDF Olarak Kaydet", f"{plaka}_bilgi_formu.pdf", "PDF Dosyası (*.pdf)")
                if file_path:
                    # A4 boyutu ayarları
                    doc = SimpleDocTemplate(file_path, pagesize=(A4),
                                            leftMargin=1*cm, rightMargin=1*cm,
                                            topMargin=1*cm, bottomMargin=1*cm)
                    story = []
                    
                    # Fotoğraf yolu belirle
                    foto_path = os.path.join("veri", "arac_fotograflari", f"{plaka}.jpg")
                    if not os.path.exists(foto_path):
                        foto_path = os.path.join("veri", "arac_fotograflari", f"{plaka}.png")
                    if not os.path.exists(foto_path):
                        foto_path = "araba_icon.png"

                    # Fotoğrafı ekle
                    if os.path.exists(foto_path):
                        try:
                            img = Image(foto_path, width=150, height=150)
                            img.hAlign = 'CENTER'
                            story.append(img)
                            story.append(Spacer(1, 10))
                        except Exception as e:
                            print("PDF fotoğraf ekleme hatası:", e)
                    # Başlık
                    styles = getSampleStyleSheet()
                    title_style = ParagraphStyle(
                        'CustomTitle',
                        parent=styles['Heading1'],
                        fontSize=18,
                        spaceAfter=20,
                        alignment=1,  # Center
                        fontName=turkish_font
                    )
                    title = Paragraph(f'{plaka} PLAKALI ARAÇ BİLGİ FORMU', title_style)
                    story.append(title)

                    
                    # Araç fotoğrafını başlığın altına ekle
                    foto_path = os.path.join("veri", "arac_fotograflari", f"{plaka}.jpg")
                    if not os.path.exists(foto_path):
                        foto_path = os.path.join("veri", "arac_fotograflari", f"{plaka}.png")
                    if not os.path.exists(foto_path):
                        foto_path = "araba_icon.png"
                    
                    
                    # İki ayrı tablo oluştur - Araç bilgileri ve yıllık özet için
                    story.append(Spacer(1, 10))
                    
                    # 1. Araç Bilgileri Tablosu (3 sütunlu)
                    arac_bilgileri_data = []
                    yillik_ozet_data = []
                    current_section = "arac_bilgileri"
                    
                    for row in rapor_data:
                        if len(row) >= 2 and row[1] == "YILLIK BAKIM VE GİDER ÖZETİ":
                            current_section = "yillik_ozet"
                            continue
                        
                        if current_section == "arac_bilgileri":
                            if len(row) >= 3:
                                arac_bilgileri_data.append([row[0], row[1], row[2]])
                            elif len(row) == 2:
                                arac_bilgileri_data.append([row[0], row[1], ""])
                            elif len(row) == 1:
                                arac_bilgileri_data.append([row[0], "", ""])
                        elif current_section == "yillik_ozet":
                            if len(row) >= 4:
                                # 4 sütunlu veriyi al (boş sütun kaldırıldı, gider sayısı ve gider tutarı sütunları kaldırıldı)
                                yillik_ozet_data.append(row[:4])
                            elif len(row) >= 2:
                                # 2 sütunlu satırları 4 sütuna genişlet
                                expanded_row = [""] * 4
                                for i, val in enumerate(row):
                                    if i < 4:
                                        expanded_row[i] = val
                                yillik_ozet_data.append(expanded_row)
                    
                     # --- 5 Sütunlu Araç Bilgileri Tablosu ---
                    # Bilgi başlıkları ve değerleri hazırlanıyor
                    from reportlab.platypus import Paragraph
                    from reportlab.lib.styles import ParagraphStyle
                    cell_style = ParagraphStyle('cell', fontName=turkish_font, fontSize=9)
                    bilgi_listesi = []
                    for row in arac_bilgileri_data:
                        if len(row) >= 3:
                            baslik = row[1]
                            deger = row[2]
                            if baslik in ["KASKO BİTİŞ", "SİGORTA BİTİŞ", "MUAYENE BİTİŞ"] and deger:
                                try:
                                    tarih = str(deger).split()[0]
                                    if '-' in tarih:
                                        yil, ay, gun = tarih.split('-')
                                        deger = f"{gun}.{ay}.{yil}"
                                    elif '/' in tarih:
                                        gun, ay, yil = tarih.split('/')
                                        deger = f"{gun}.{ay}.{yil}"
                                    else:
                                        deger = tarih
                                except:
                                    pass
                            bilgi_listesi.append((baslik, deger))
                    rows = []
                    for i in range(0, len(bilgi_listesi), 2):
                        b1, v1 = bilgi_listesi[i]
                        if i+1 < len(bilgi_listesi):
                            b2, v2 = bilgi_listesi[i+1]
                        else:
                            b2, v2 = "", ""
                        # Metinleri Paragraph ile sığdır - 4 sütunlu yapı
                        row = [Paragraph(str(b1), cell_style), Paragraph(str(v1), cell_style), Paragraph(str(b2), cell_style), Paragraph(str(v2), cell_style)]
                        rows.append(row)
                    # A4 sayfası kullanılabilir genişlik: 21cm - 2cm kenar boşlukları = 19cm
                    col_widths = [4.75*cm, 4.75*cm, 4.75*cm, 4.75*cm]  # Toplam 19 cm (sayfa genişliğine göre)
                    table = Table(rows, colWidths=col_widths, hAlign='CENTER')
                    table_style = TableStyle([
                        ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                        ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
                        ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
                        ('FONTNAME', (0, 0), (-1, -1), turkish_font),
                        ('FONTSIZE', (0, 0), (-1, -1), 9),
                        ('BACKGROUND', (0, 0), (0, len(rows)-1), colors.HexColor('#F7F7F7')),
                        ('BACKGROUND', (2, 0), (2, len(rows)-1), colors.HexColor('#F7F7F7')),
                        ('BACKGROUND', (1, 0), (1, len(rows)-1), colors.white),
                        ('BACKGROUND', (3, 0), (3, len(rows)-1), colors.white),
                    ])
                    table.setStyle(table_style)
                    story.append(table)
                    story.append(Spacer(1, 20))

                    # --- Eski tabloyu kaldırdık, yeni tablo yukarıda eklendi ---
                    
                    # 2. Yıllık Özet Tablosu (4 sütunlu - Boş sütun kaldırıldı, gider sayısı ve gider tutarı sütunları kaldırıldı)
                    if yillik_ozet_data:
                        yillik_col_widths = [4.75*cm, 4.75*cm, 4.75*cm, 4.75*cm]  # 4 sütun - Toplam 19 cm (sayfa genişliğine göre)
                        yillik_table = Table(yillik_ozet_data, colWidths=yillik_col_widths)
                        
                        # Yıllık özet tablo stilleri - Düzgün hizalanmış kenarlıklar
                        yillik_table_style = TableStyle([
                            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                            ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
                            ('GRID', (0, 0), (-1, -1), 1, colors.black),
                            ('BOTTOMPADDING', (0, 0), (-1, -1), 8),
                            ('TOPPADDING', (0, 0), (-1, -1), 8),
                            ('LEFTPADDING', (0, 0), (-1, -1), 6),
                            ('RIGHTPADDING', (0, 0), (-1, -1), 6),
                            ('FONTNAME', (0, 0), (-1, -1), turkish_font),
                            ('FONTSIZE', (0, 0), (-1, -1), 9),
                            ('FONTNAME', (0, 0), (-1, 0), turkish_font),  # Başlık satırı
                            ('FONTSIZE', (0, 0), (-1, 0), 10),
                            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#A23B72')),  # Başlık arka planı
                            ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),  # Başlık yazı rengi
                            ('FONTNAME', (0, -1), (-1, -1), turkish_font),  # Toplam satırı
                            ('FONTSIZE', (0, -1), (-1, -1), 10),
                            ('BACKGROUND', (0, -1), (-1, -1), colors.HexColor('#F18F01')),  # Toplam arka planı
                            ('TEXTCOLOR', (0, -1), (-1, -1), colors.white),  # Toplam yazı rengi
                        ])
                        
                        # Özel satır stilleri uygula
                        for i, row_data in enumerate(yillik_ozet_data):
                            if len(row_data) >= 2:
                                bilgi = row_data[1]
                                
                                # Ana başlık satırı
                                if bilgi == "YILLIK BAKIM VE GİDER ÖZETİ":
                                    yillik_table_style.add('BACKGROUND', (0, i), (-1, i), colors.HexColor('#2E86AB'))
                                    yillik_table_style.add('TEXTCOLOR', (0, i), (-1, i), colors.white)
                                    yillik_table_style.add('FONTSIZE', (0, i), (-1, i), 11)
                                    yillik_table_style.add('FONTNAME', (0, i), (-1, i), turkish_font)
                                
                                # Tablo başlık satırı
                                elif bilgi == "Yıl":
                                    yillik_table_style.add('BACKGROUND', (0, i), (-1, i), colors.HexColor('#A23B72'))
                                    yillik_table_style.add('TEXTCOLOR', (0, i), (-1, i), colors.white)
                                    yillik_table_style.add('FONTSIZE', (0, i), (-1, i), 10)
                                    yillik_table_style.add('FONTNAME', (0, i), (-1, i), turkish_font)
                                
                                # Toplam satırı
                                elif bilgi == "TOPLAM":
                                    yillik_table_style.add('BACKGROUND', (0, i), (-1, i), colors.HexColor('#F18F01'))
                                    yillik_table_style.add('TEXTCOLOR', (0, i), (-1, i), colors.white)
                                    yillik_table_style.add('FONTSIZE', (0, i), (-1, i), 10)
                                    yillik_table_style.add('FONTNAME', (0, i), (-1, i), turkish_font)
                                
                                # Yıl veri satırları
                                elif bilgi.isdigit():
                                    yillik_table_style.add('BACKGROUND', (0, i), (-1, i), colors.HexColor('#F7F7F7'))
                                    yillik_table_style.add('TEXTCOLOR', (0, i), (-1, i), colors.black)
                                    yillik_table_style.add('FONTSIZE', (0, i), (-1, i), 9)
                                    yillik_table_style.add('FONTNAME', (0, i), (-1, i), turkish_font)
                        
                        yillik_table.setStyle(yillik_table_style)
                        story.append(yillik_table)
                    
                    # Tablolar zaten yukarıda oluşturuldu ve story'ye eklendi
                    doc.build(story)
                    success_count += 1
            except ImportError:
                QMessageBox.warning(self, "Kütüphane Hatası", "PDF oluşturmak için reportlab kütüphanesi gerekli. Lütfen 'pip install reportlab' komutunu çalıştırın.")
            except Exception as e:
                QMessageBox.warning(self, "PDF Hatası", f"PDF dosyası oluşturulamadı: {str(e)}")

        # Başarı mesajı
        if success_count == total_formats:
            QMessageBox.information(self, "Başarılı", f"Araç bilgi formu {success_count} formatında başarıyla kaydedildi.")
        elif success_count > 0:
            QMessageBox.information(self, "Kısmi Başarı", f"{success_count}/{total_formats} formatında kaydedildi. Bazı formatlar oluşturulamadı.")
        else:
            QMessageBox.warning(self, "Hata", "Hiçbir format oluşturulamadı.")
    
    def generate_vehicle_report(self):
        """Araç raporu oluştur"""
        try:
            vehicles = self.data_manager.load_data('araclar')
            if vehicles.empty:
                QMessageBox.information(self, "Bilgi", "Araç verisi bulunamadı.")
                return
            
            report_text = "🚗 ARAÇ RAPORU\n"
            report_text += "=" * 50 + "\n\n"
            
            for _, vehicle in vehicles.iterrows():
                report_text += f"Plaka: {vehicle.get('Plaka', 'N/A')}\n"
                report_text += f"Marka/Model: {vehicle.get('Marka', '')} {vehicle.get('Model', '')}\n"
                report_text += f"Yıl: {vehicle.get('Yıl', 'N/A')}\n"
                report_text += f"Durum: {vehicle.get('Durum', 'N/A')}\n"
                report_text += f"Son KM: {vehicle.get('Son KM', 'N/A')}\n"
                report_text += f"Muayene: {vehicle.get('Muayene Tarihi', 'N/A')}\n"
                report_text += "-" * 30 + "\n"
            
            self.show_report_dialog("Araç Raporu", report_text)
            
        except Exception as e:
            QMessageBox.critical(self, "Hata", f"Rapor oluşturma hatası: {str(e)}")
    
    
    def tr_money(self, val):
        try:
            return f"{float(val):,.2f} TL".replace(",", "_").replace(".", ",").replace("_", ".")
        except:
            return str(val)

    def generate_maintenance_report(self):
        """Bakım raporunu araçlara ve tarih aralığına göre tablo olarak oluşturur ve Word/PDF olarak kaydeder."""
        try:
            maintenance = self.data_manager.load_data('bakimlar')
            vehicles = self.data_manager.load_data('araclar')
            vehicles = vehicles[vehicles['Durum'].str.lower() == 'aktif']
            aktif_plakalar = set(vehicles['Plaka'])
            maintenance = maintenance[maintenance['Araç Plakası'].isin(aktif_plakalar)]
            if maintenance.empty:
                QMessageBox.information(self, "Bilgi", "Bakım verisi bulunamadı.")
                return
            if vehicles.empty:
                QMessageBox.information(self, "Bilgi", "Araç verisi bulunamadı.")
                return

            # Tarih sütunlarını bul
            maintenance['Bakım Tarihi'] = pd.to_datetime(maintenance['Bakım Tarihi'], errors='coerce', dayfirst=True)
            maintenance = maintenance.dropna(subset=['Bakım Tarihi'])
            
            # Tarih aralığı seçim dialog'u
            date_range = self.show_date_range_selection_dialog("Bakım Raporu için Tarih Aralığı Seçimi")
            if not date_range:
                return
            
            start_date, end_date = date_range
            
            # Seçilen tarih aralığına göre filtrele
            maintenance = maintenance[
                (maintenance['Bakım Tarihi'] >= pd.Timestamp(start_date)) & 
                (maintenance['Bakım Tarihi'] <= pd.Timestamp(end_date))
            ]
            
            if maintenance.empty:
                QMessageBox.information(self, "Bilgi", "Seçilen tarih aralığında bakım verisi bulunamadı.")
                return
                
            # Yıl bilgilerini al
            maintenance['Yıl'] = maintenance['Bakım Tarihi'].dt.year
            years = sorted(maintenance['Yıl'].unique())

            # Araç bilgilerini al
            vehicle_info = vehicles.set_index('Plaka')[['Marka', 'Model', 'Yıl']]

            # Pivot tablo: satır=plaka, sütun=yıl, değer=toplam bakım tutarı
            pivot = maintenance.pivot_table(index='Araç Plakası', columns='Yıl', values='Tutar', aggfunc=lambda x: pd.to_numeric(x, errors='coerce').sum(), fill_value=0)
            # Araç bilgilerini ekle
            pivot = pivot.merge(vehicle_info, left_index=True, right_index=True, how='left')
            # Sütun sırası: Marka, Model, Yıl, [yıllar...], Toplam
            cols = ['Marka', 'Model', 'Yıl'] + list(years)
            pivot = pivot[cols]
            pivot['Toplam'] = pivot[years].sum(axis=1)
            # Sıralama
            pivot = pivot.sort_values(by='Toplam', ascending=False)
            # Sıfır olmayanlar
            pivot = pivot[pivot['Toplam'] > 0]
            if pivot.empty:
                QMessageBox.information(self, "Bilgi", "Bakım gideri olan araç bulunamadı.")
                return

            # Tabloyu rapor formatına çevir
            rapor_data = []
            header = ['No', 'Plaka', 'Marka', 'Model', 'Yıl'] + [str(y) for y in years] + ['Toplam']
            rapor_data.append(header)
            for idx, (plaka, row) in enumerate(pivot.iterrows(), 1):
                rapor_data.append([
                    idx,
                    plaka,
                    row['Marka'],
                    row['Model'],
                    str(row['Yıl'])[:4] if pd.notna(row['Yıl']) else '',
                    *[self.tr_money(row[y]) for y in years],
                    self.tr_money(row['Toplam'])
                ])
            total_row = ['GENEL TOPLAM', '', '', '', '']
            for y in years:
                total_row.append(self.tr_money(pivot[y].sum()))
            total_row.append(self.tr_money(pivot['Toplam'].sum()))
            rapor_data.append(total_row)

            # Kullanıcıdan format seçimi al
            format_choice = self.show_format_selection_dialog()
            if not format_choice:
                return
            self.create_maintenance_report_files(rapor_data, format_choice, years)
        except Exception as e:
            QMessageBox.critical(self, "Hata", f"Bakım raporu oluşturulamadı: {str(e)}")

    def create_maintenance_report_files(self, rapor_data, format_choice, selected_years=None):
        """Bakım raporunu Word/PDF olarak kaydeder."""
        import os
        import pandas as pd
        from PyQt5.QtWidgets import QFileDialog, QMessageBox
        # Word
        if format_choice in ["Word (.docx)", "Tümü"]:
            try:
                from docx import Document
                from docx.shared import Pt, Cm, RGBColor
                from docx.enum.text import WD_ALIGN_PARAGRAPH
                from docx.oxml import parse_xml
                from docx.enum.section import WD_ORIENT
                file_path, _ = QFileDialog.getSaveFileName(self, "Word Olarak Kaydet", f"bakim_raporu_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx", "Word Dosyası (*.docx)")
                if file_path:
                    doc = Document()
                    section = doc.sections[0]
                    section.orientation = WD_ORIENT.LANDSCAPE
                    new_width, new_height = section.page_height, section.page_width
                    section.page_width = new_width
                    section.page_height = new_height
                    section.left_margin = Cm(1)
                    section.right_margin = Cm(1)
                    section.top_margin = Cm(1)
                    section.bottom_margin = Cm(1)
                    # Başlık oluştur
                    if selected_years:
                        years_text = ", ".join(map(str, selected_years))
                        title_text = f'Yıllara Göre Araç Bakım/Onarım Raporu ({years_text})'
                    else:
                        title_text = 'Yıllara Göre Araç Bakım/Onarım Raporu'
                    title = doc.add_heading(title_text, 0)
                    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    doc.add_paragraph("")
                    table = doc.add_table(rows=len(rapor_data), cols=len(rapor_data[0]))
                    table.style = 'Table Grid'
                    table.allow_autofit = False
                    for i, row_data in enumerate(rapor_data):
                        for j, val in enumerate(row_data):
                            cell = table.cell(i, j)
                            cell.text = str(val)
                            for paragraph in cell.paragraphs:
                                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                for run in cell.paragraphs[0].runs:
                                    run.font.name = 'Calibri'
                                    run.font.size = Pt(10)
                            cell._tc.get_or_add_tcPr().append(
                                parse_xml(r'<w:tcW xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:w="2000" w:type="dxa"/>')
                            )
                        # Başlık satırı
                        if i == 0:
                            for cell in table.rows[i].cells:
                                cell_par = cell.paragraphs[0]
                                for run in cell_par.runs:
                                    run.bold = True
                                    run.font.size = Pt(11)
                                    run.font.color.rgb = RGBColor(255,255,255)
                                cell._tc.get_or_add_tcPr().append(
                                    parse_xml(r'<w:shd xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:fill="A23B72"/>')
                                )
                        # Genel toplam satırı
                        if row_data[0] == "GENEL TOPLAM":
                            for cell in table.rows[i].cells:
                                cell_par = cell.paragraphs[0]
                                for run in cell_par.runs:
                                    run.bold = True
                                    run.font.size = Pt(11)
                                    run.font.color.rgb = RGBColor(255,255,255)
                                cell._tc.get_or_add_tcPr().append(
                                    parse_xml(r'<w:shd xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:fill="F18F01"/>')
                                )
                    doc.save(file_path)
            except ImportError:
                QMessageBox.warning(self, "Kütüphane Hatası", "Word oluşturmak için python-docx kütüphanesi gerekli. Lütfen 'pip install python-docx' komutunu çalıştırın.")
            except Exception as e:
                QMessageBox.warning(self, "Word Hatası", f"Word dosyası oluşturulamadı: {str(e)}")
        # PDF
        if format_choice in ["PDF (.pdf)", "Tümü"]:
            try:
                from reportlab.lib.pagesizes import A4, landscape
                from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
                from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
                from reportlab.lib.units import cm
                from reportlab.lib import colors
                from reportlab.pdfbase import pdfmetrics
                from reportlab.pdfbase.ttfonts import TTFont
                # Türkçe font desteği
                try:
                    font_paths = [
                        "C:/Windows/Fonts/arial.ttf",
                        "C:/Windows/Fonts/calibri.ttf",
                        "C:/Windows/Fonts/tahoma.ttf",
                        "C:/Windows/Fonts/verdana.ttf"
                    ]
                    turkish_font = None
                    for font_path in font_paths:
                        if os.path.exists(font_path):
                            try:
                                pdfmetrics.registerFont(TTFont('TurkishFont', font_path))
                                turkish_font = 'TurkishFont'
                                break
                            except:
                                continue
                    if not turkish_font:
                        turkish_font = 'Helvetica'
                except:
                    turkish_font = 'Helvetica'
                file_path, _ = QFileDialog.getSaveFileName(self, "PDF Olarak Kaydet", f"bakim_raporu_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf", "PDF Dosyası (*.pdf)")
                if file_path:
                    doc = SimpleDocTemplate(file_path, pagesize=landscape(A4),
                                          leftMargin=1*cm, rightMargin=1*cm,
                                          topMargin=1*cm, bottomMargin=1*cm)
                    story = []
                    styles = getSampleStyleSheet()
                    title_style = ParagraphStyle('title', parent=styles['Heading1'], alignment=1, fontName=turkish_font, fontSize=16, textColor=colors.HexColor('#A23B72'))
                    # Başlık oluştur
                    if selected_years:
                        years_text = ", ".join(map(str, selected_years))
                        title_text = f'Yıllara Göre Araç Bakım/Onarım Raporu ({years_text})'
                    else:
                        title_text = 'Yıllara Göre Araç Bakım/Onarım Raporu'
                    story.append(Paragraph(title_text, title_style))
                    story.append(Spacer(1, 12))
                    # Tablo
                    table_data = rapor_data
                    table = Table(table_data, repeatRows=1)
                    table_style = TableStyle([
                        ('BACKGROUND', (0,0), (-1,0), colors.HexColor('#A23B72')),
                        ('TEXTCOLOR', (0,0), (-1,0), colors.white),
                        ('ALIGN', (0,0), (-1,-1), 'CENTER'),
                        ('FONTNAME', (0,0), (-1,-1), turkish_font),
                        ('FONTSIZE', (0,0), (-1,-1), 9),
                        ('GRID', (0,0), (-1,-1), 0.5, colors.grey),
                        ('BACKGROUND', (0,-1), (-1,-1), colors.HexColor('#F18F01')),
                        ('TEXTCOLOR', (0,-1), (-1,-1), colors.white),
                        ('FONTNAME', (0,-1), (-1,-1), turkish_font),
                        ('FONTSIZE', (0,-1), (-1,-1), 10),
                        ('ALIGN', (0,-1), (-1,-1), 'CENTER'),
                        ('VALIGN', (0,0), (-1,-1), 'MIDDLE'),
                    ])
                    table.setStyle(table_style)
                    story.append(table)
                    doc.build(story)
            except ImportError:
                QMessageBox.warning(self, "Kütüphane Hatası", "PDF oluşturmak için reportlab kütüphanesi gerekli. Lütfen 'pip install reportlab' komutunu çalıştırın.")
            except Exception as e:
                QMessageBox.warning(self, "PDF Hatası", f"PDF dosyası oluşturulamadı: {str(e)}")



    
    def generate_expense_report(self):
        """Gider raporu oluştur"""
        try:
            # Tüm gider verilerini yükle
            expenses = self.data_manager.load_data('giderler')
            maintenance = self.data_manager.load_data('bakimlar')
            fuel = self.data_manager.load_data('yakitlar')
            
            if expenses.empty and maintenance.empty and fuel.empty:
                QMessageBox.information(self, "Bilgi", "Gider verisi bulunamadı.")
                return
            
            # Tarih sütunlarını datetime'a çevir
            if not expenses.empty and 'Tarih' in expenses.columns:
                expenses['Tarih'] = pd.to_datetime(expenses['Tarih'], errors='coerce')
            
            if not maintenance.empty and 'Bakım Tarihi' in maintenance.columns:
                maintenance['Bakım Tarihi'] = pd.to_datetime(maintenance['Bakım Tarihi'], errors='coerce')
            
            if not fuel.empty and 'Tarih' in fuel.columns:
                fuel['Tarih'] = pd.to_datetime(fuel['Tarih'], errors='coerce')
            
            # Tarih aralığı seçim dialog'u
            date_range = self.show_date_range_selection_dialog("Gider Raporu için Tarih Aralığı Seçimi")
            if not date_range:
                return
            
            start_date, end_date = date_range
            
            # Seçilen tarih aralığına göre filtrele
            if not expenses.empty:
                expenses = expenses[
                    (expenses['Tarih'] >= pd.Timestamp(start_date)) & 
                    (expenses['Tarih'] <= pd.Timestamp(end_date))
                ]
            
            if not maintenance.empty:
                maintenance = maintenance[
                    (maintenance['Bakım Tarihi'] >= pd.Timestamp(start_date)) & 
                    (maintenance['Bakım Tarihi'] <= pd.Timestamp(end_date))
                ]
            
            if not fuel.empty:
                fuel = fuel[
                    (fuel['Tarih'] >= pd.Timestamp(start_date)) & 
                    (fuel['Tarih'] <= pd.Timestamp(end_date))
                ]
            
            # Gider türü seçim dialog'u
            expense_types = self.show_expense_type_selection_dialog()
            if not expense_types:
                return
            
            # Yıl bilgilerini al
            if not expenses.empty:
                expenses['Yıl'] = expenses['Tarih'].dt.year
            if not maintenance.empty:
                maintenance['Yıl'] = maintenance['Bakım Tarihi'].dt.year
            if not fuel.empty:
                fuel['Yıl'] = fuel['Tarih'].dt.year
            
            # Mevcut yılları topla
            available_years = set()
            if not expenses.empty:
                available_years.update(expenses['Yıl'].dropna().unique())
            if not maintenance.empty:
                available_years.update(maintenance['Yıl'].dropna().unique())
            if not fuel.empty:
                available_years.update(fuel['Yıl'].dropna().unique())
            
            selected_years = sorted(list(available_years))
            
            # Format seçimi
            format_choice = self.show_format_selection_dialog()
            if not format_choice:
                return
            
            # Rapor verilerini hazırla
            rapor_data = {
                'expenses': expenses,
                'maintenance': maintenance,
                'fuel': fuel,
                'selected_years': selected_years,
                'expense_types': expense_types
            }
            
            # Rapor dosyalarını oluştur
            self.create_expense_report_files(rapor_data, format_choice, selected_years)
            
        except Exception as e:
            QMessageBox.critical(self, "Hata", f"Rapor oluşturma hatası: {str(e)}")
    

    
    def show_report_dialog(self, title, content):
        """Rapor dialog göster"""
        dialog = QDialog(self)
        dialog.setWindowTitle(title)
        dialog.setMinimumSize(600, 400)
        dialog.setModal(True)
        
        layout = QVBoxLayout(dialog)
        
        # Rapor içeriği
        text_edit = QTextEdit()
        text_edit.setPlainText(content)
        text_edit.setReadOnly(True)
        layout.addWidget(text_edit)
        
        # Butonlar
        btn_layout = QHBoxLayout()
        btn_save = QPushButton("Kaydet")
        btn_save.clicked.connect(lambda: self.save_report(title, content))
        btn_close = QPushButton("Kapat")
        btn_close.clicked.connect(dialog.accept)
        
        btn_layout.addWidget(btn_save)
        btn_layout.addWidget(btn_close)
        layout.addLayout(btn_layout)
        
        dialog.exec_()
    
    def save_report(self, title, content):
        """Raporu dosyaya kaydet"""
        try:
            filename, _ = QFileDialog.getSaveFileName(
                self, "Raporu Kaydet", 
                f"{title}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.txt",
                "Metin Dosyası (*.txt)"
            )
            
            if filename:
                with open(filename, 'w', encoding='utf-8') as f:
                    f.write(content)
                QMessageBox.information(self, "Başarılı", f"Rapor kaydedildi:\n{filename}")
        except Exception as e:
            QMessageBox.critical(self, "Hata", f"Kaydetme hatası: {str(e)}")
    
    def show_format_selection_dialog(self):
        """Format seçimi için butonlu dialog göster"""
        from PyQt5.QtWidgets import QDialog, QVBoxLayout, QHBoxLayout, QPushButton, QLabel
        from PyQt5.QtCore import Qt
        
        dialog = QDialog(self)
        dialog.setWindowTitle("Format Seçin")
        dialog.setFixedSize(500, 250)
        dialog.setModal(True)
        
        layout = QVBoxLayout(dialog)
        layout.setSpacing(20)
        
        # Başlık
        title = QLabel("Hangi formatta kaydetmek istiyorsunuz?")
        title.setStyleSheet("font-size: 16px; font-weight: bold; color: #2c3e50;")
        title.setAlignment(Qt.AlignCenter)
        layout.addWidget(title)
        
        # Butonlar için layout
        button_layout = QHBoxLayout()
        button_layout.setSpacing(20)
        
        # Word butonu
        btn_word = QPushButton("📄 Word (.docx)")
        btn_word.setStyleSheet("""
            QPushButton {
                background-color: #3498db;
                color: white;
                border: none;
                padding: 12px 20px;
                border-radius: 8px;
                font-size: 12px;
                font-weight: bold;
                min-width: 100px;
                max-width: 100px;
            }
            QPushButton:hover {
                background-color: #2980b9;
            }
        """)
        btn_word.clicked.connect(lambda: self.select_format(dialog, "Word (.docx)"))
        
        # PDF butonu
        btn_pdf = QPushButton("📋 PDF (.pdf)")
        btn_pdf.setStyleSheet("""
            QPushButton {
                background-color: #e74c3c;
                color: white;
                border: none;
                padding: 12px 20px;
                border-radius: 8px;
                font-size: 12px;
                font-weight: bold;
                min-width: 100px;
                max-width: 100px;
            }
            QPushButton:hover {
                background-color: #c0392b;
            }
        """)
        btn_pdf.clicked.connect(lambda: self.select_format(dialog, "PDF (.pdf)"))
        
        # Tümü butonu
        btn_all = QPushButton("📚 Tümü")
        btn_all.setStyleSheet("""
            QPushButton {
                background-color: #27ae60;
                color: white;
                border: none;
                padding: 12px 20px;
                border-radius: 8px;
                font-size: 12px;
                font-weight: bold;
                min-width: 100px;
                max-width: 100px;
            }
            QPushButton:hover {
                background-color: #229954;
            }
        """)
        btn_all.clicked.connect(lambda: self.select_format(dialog, "Tümü"))
        
        # İptal butonu
        btn_cancel = QPushButton("❌ İptal")
        btn_cancel.setStyleSheet("""
            QPushButton {
                background-color: #95a5a6;
                color: white;
                border: none;
                padding: 12px 20px;
                border-radius: 8px;
                font-size: 12px;
                font-weight: bold;
                min-width: 100px;
                max-width: 100px;
            }
            QPushButton:hover {
                background-color: #7f8c8d;
            }
        """)
        btn_cancel.clicked.connect(dialog.reject)
        
        button_layout.addWidget(btn_word)
        button_layout.addWidget(btn_pdf)
        button_layout.addWidget(btn_all)
        button_layout.addWidget(btn_cancel)
        
        layout.addLayout(button_layout)
        
        # Dialog'u göster ve sonucu döndür
        result = dialog.exec_()
        if result == QDialog.Accepted:
            return getattr(dialog, 'selected_format', None)
        return None
    
    def select_format(self, dialog, format_choice):
        """Format seçimi yapıldığında çağrılır"""
        dialog.selected_format = format_choice
        dialog.accept()
    
    def show_date_range_selection_dialog(self, title="Tarih Aralığı Seçimi"):
        """Özel tarih aralığı seçimi dialog'u"""
        dialog = QDialog(self)
        dialog.setWindowTitle(title)
        dialog.setModal(True)
        dialog.setMinimumWidth(400)
        
        layout = QVBoxLayout(dialog)
        
        # Başlık
        title_label = QLabel("Hangi tarihler arası rapor almak istiyorsunuz?")
        title_label.setStyleSheet("""
            font-size: 14px;
            font-weight: bold;
            color: #2c3e50;
            margin-bottom: 10px;
        """)
        layout.addWidget(title_label)
        
        # Tarih seçimi
        date_layout = QHBoxLayout()
        
        # Başlangıç tarihi
        start_layout = QVBoxLayout()
        start_label = QLabel("Başlangıç Tarihi:")
        start_label.setStyleSheet("font-weight: bold; color: #34495e;")
        self.start_date_edit = QDateEdit()
        self.start_date_edit.setCalendarPopup(True)
        self.start_date_edit.setDate(QDate.currentDate().addYears(-1))  # Varsayılan 1 yıl önce
        self.start_date_edit.setStyleSheet("""
            QDateEdit {
                padding: 8px;
                border: 2px solid #bdc3c7;
                border-radius: 5px;
                font-size: 12px;
            }
            QDateEdit:focus {
                border-color: #3498db;
            }
        """)
        start_layout.addWidget(start_label)
        start_layout.addWidget(self.start_date_edit)
        
        # Bitiş tarihi
        end_layout = QVBoxLayout()
        end_label = QLabel("Bitiş Tarihi:")
        end_label.setStyleSheet("font-weight: bold; color: #34495e;")
        self.end_date_edit = QDateEdit()
        self.end_date_edit.setCalendarPopup(True)
        self.end_date_edit.setDate(QDate.currentDate())  # Varsayılan bugün
        self.end_date_edit.setStyleSheet("""
            QDateEdit {
                padding: 8px;
                border: 2px solid #bdc3c7;
                border-radius: 5px;
                font-size: 12px;
            }
            QDateEdit:focus {
                border-color: #3498db;
            }
        """)
        end_layout.addWidget(end_label)
        end_layout.addWidget(self.end_date_edit)
        
        date_layout.addLayout(start_layout)
        date_layout.addLayout(end_layout)
        layout.addLayout(date_layout)
        
        # Hızlı seçim butonları
        quick_layout = QHBoxLayout()
        
        btn_last_month = QPushButton("Son 1 Ay")
        btn_last_month.setStyleSheet("""
            QPushButton {
                background-color: #3498db;
                color: white;
                border: none;
                padding: 8px 12px;
                border-radius: 4px;
                font-size: 11px;
            }
            QPushButton:hover {
                background-color: #2980b9;
            }
        """)
        btn_last_month.clicked.connect(lambda: self.set_date_range(-1, 0))
        
        btn_last_3_months = QPushButton("Son 3 Ay")
        btn_last_3_months.setStyleSheet("""
            QPushButton {
                background-color: #e67e22;
                color: white;
                border: none;
                padding: 8px 12px;
                border-radius: 4px;
                font-size: 11px;
            }
            QPushButton:hover {
                background-color: #d35400;
            }
        """)
        btn_last_3_months.clicked.connect(lambda: self.set_date_range(-3, 0))
        
        btn_last_6_months = QPushButton("Son 6 Ay")
        btn_last_6_months.setStyleSheet("""
            QPushButton {
                background-color: #27ae60;
                color: white;
                border: none;
                padding: 8px 12px;
                border-radius: 4px;
                font-size: 11px;
            }
            QPushButton:hover {
                background-color: #229954;
            }
        """)
        btn_last_6_months.clicked.connect(lambda: self.set_date_range(-6, 0))
        
        btn_last_year = QPushButton("Son 1 Yıl")
        btn_last_year.setStyleSheet("""
            QPushButton {
                background-color: #9b59b6;
                color: white;
                border: none;
                padding: 8px 12px;
                border-radius: 4px;
                font-size: 11px;
            }
            QPushButton:hover {
                background-color: #8e44ad;
            }
        """)
        btn_last_year.clicked.connect(lambda: self.set_date_range(-12, 0))
        
        quick_layout.addWidget(btn_last_month)
        quick_layout.addWidget(btn_last_3_months)
        quick_layout.addWidget(btn_last_6_months)
        quick_layout.addWidget(btn_last_year)
        layout.addLayout(quick_layout)
        
        # Butonlar
        button_layout = QHBoxLayout()
        
        btn_ok = QPushButton("Tamam")
        btn_ok.setStyleSheet("""
            QPushButton {
                background-color: #27ae60;
                color: white;
                border: none;
                padding: 10px 20px;
                border-radius: 5px;
                font-weight: bold;
                font-size: 12px;
            }
            QPushButton:hover {
                background-color: #229954;
            }
        """)
        btn_ok.clicked.connect(dialog.accept)
        
        btn_cancel = QPushButton("İptal")
        btn_cancel.setStyleSheet("""
            QPushButton {
                background-color: #e74c3c;
                color: white;
                border: none;
                padding: 10px 20px;
                border-radius: 5px;
                font-weight: bold;
                font-size: 12px;
            }
            QPushButton:hover {
                background-color: #c0392b;
            }
        """)
        btn_cancel.clicked.connect(dialog.reject)
        
        button_layout.addWidget(btn_ok)
        button_layout.addWidget(btn_cancel)
        layout.addLayout(button_layout)
        
        # Dialog'u çalıştır
        if dialog.exec_() == QDialog.Accepted:
            start_date = self.start_date_edit.date().toPyDate()
            end_date = self.end_date_edit.date().toPyDate()
            return start_date, end_date
        else:
            return None
    
    def set_date_range(self, months_start, months_end):
        """Hızlı tarih aralığı ayarla"""
        current_date = QDate.currentDate()
        start_date = current_date.addMonths(months_start)
        end_date = current_date.addMonths(months_end)
        
        self.start_date_edit.setDate(start_date)
        self.end_date_edit.setDate(end_date)
    
    def show_year_selection_dialog(self, available_years):
        """Yıl seçim dialog'u gösterir"""
        dialog = QDialog(self)
        dialog.setWindowTitle("Yıl Seçimi")
        dialog.setModal(True)
        dialog.setFixedSize(400, 350)
        
        layout = QVBoxLayout(dialog)
        
        # Başlık
        title_label = QLabel("Bakım Raporu için Yıl Seçimi")
        title_label.setStyleSheet("font-size: 14px; font-weight: bold; margin-bottom: 10px;")
        title_label.setAlignment(Qt.AlignmentFlag.AlignCenter)
        layout.addWidget(title_label)
        
        # Açıklama
        desc_label = QLabel("Rapora dahil edilecek yılları seçin:")
        desc_label.setStyleSheet("font-size: 12px; margin-bottom: 10px;")
        layout.addWidget(desc_label)
        
        # Scroll area içinde checkbox'lar
        scroll_area = QScrollArea()
        scroll_area.setWidgetResizable(True)
        scroll_area.setMaximumHeight(200)
        scroll_area.setStyleSheet("""
            QScrollArea {
                border: 2px solid #bdc3c7;
                border-radius: 5px;
                background-color: white;
            }
            QScrollBar:vertical {
                background-color: #ecf0f1;
                width: 12px;
                border-radius: 6px;
            }
            QScrollBar::handle:vertical {
                background-color: #bdc3c7;
                border-radius: 6px;
                min-height: 20px;
            }
            QScrollBar::handle:vertical:hover {
                background-color: #95a5a6;
            }
        """)
        
        # Checkbox'lar için widget
        checkbox_widget = QWidget()
        checkbox_layout = QVBoxLayout(checkbox_widget)
        checkbox_layout.setSpacing(8)
        checkbox_layout.setContentsMargins(10, 10, 10, 10)
        
        # Onay kutucuklarını oluştur
        self.year_checkboxes = {}
        for year in available_years:
            checkbox = QCheckBox(str(year))
            checkbox.setChecked(True)  # Varsayılan olarak seçili
            checkbox.setStyleSheet("""
                QCheckBox {
                    font-size: 13px;
                    spacing: 8px;
                }
                QCheckBox::indicator {
                    width: 16px;
                    height: 16px;
                }
                QCheckBox::indicator:unchecked {
                    border: 1px solid #95a5a6;
                    background-color: white;
                    border-radius: 2px;
                }
                QCheckBox::indicator:checked {
                    border: 1px solid #2c3e50;
                    background-color: #2c3e50;
                    border-radius: 2px;
                }
                QCheckBox::indicator:hover {
                    border: 1px solid #3498db;
                }
            """)
            self.year_checkboxes[year] = checkbox
            checkbox_layout.addWidget(checkbox)
        
        scroll_area.setWidget(checkbox_widget)
        layout.addWidget(scroll_area)
        
        # Butonlar
        button_layout = QHBoxLayout()
        
        select_all_btn = QPushButton("Tümünü Seç")
        select_all_btn.setStyleSheet("""
            QPushButton {
                background-color: #27ae60;
                color: white;
                border: none;
                padding: 8px 16px;
                border-radius: 4px;
                font-weight: bold;
            }
            QPushButton:hover {
                background-color: #229954;
            }
        """)
        select_all_btn.clicked.connect(lambda: self.select_all_years_checkboxes())
        
        clear_all_btn = QPushButton("Seçimi Temizle")
        clear_all_btn.setStyleSheet("""
            QPushButton {
                background-color: #e74c3c;
                color: white;
                border: none;
                padding: 8px 16px;
                border-radius: 4px;
                font-weight: bold;
            }
            QPushButton:hover {
                background-color: #c0392b;
            }
        """)
        clear_all_btn.clicked.connect(lambda: self.clear_all_years_checkboxes())
        
        button_layout.addWidget(select_all_btn)
        button_layout.addWidget(clear_all_btn)
        button_layout.addStretch()
        
        layout.addLayout(button_layout)
        
        # Tamam/İptal butonları
        button_layout2 = QHBoxLayout()
        
        ok_btn = QPushButton("Tamam")
        ok_btn.setStyleSheet("""
            QPushButton {
                background-color: #3498db;
                color: white;
                border: none;
                padding: 8px 20px;
                border-radius: 4px;
                font-weight: bold;
            }
            QPushButton:hover {
                background-color: #2980b9;
            }
        """)
        ok_btn.clicked.connect(dialog.accept)
        
        cancel_btn = QPushButton("İptal")
        cancel_btn.setStyleSheet("""
            QPushButton {
                background-color: #95a5a6;
                color: white;
                border: none;
                padding: 8px 20px;
                border-radius: 4px;
                font-weight: bold;
            }
            QPushButton:hover {
                background-color: #7f8c8d;
            }
        """)
        cancel_btn.clicked.connect(dialog.reject)
        
        button_layout2.addStretch()
        button_layout2.addWidget(ok_btn)
        button_layout2.addWidget(cancel_btn)
        
        layout.addLayout(button_layout2)
        
        # Dialog'u göster
        if dialog.exec_() == QDialog.Accepted:
            selected_years = []
            for year, checkbox in self.year_checkboxes.items():
                if checkbox.isChecked():
                    selected_years.append(year)
            
            if not selected_years:
                QMessageBox.warning(self, "Uyarı", "En az bir yıl seçmelisiniz!")
                return None
            return selected_years
        return None
    
    def select_all_years_checkboxes(self):
        """Tüm yıl checkbox'larını seçer"""
        for checkbox in self.year_checkboxes.values():
            checkbox.setChecked(True)
    
    def clear_all_years_checkboxes(self):
        """Tüm yıl checkbox'larının seçimini temizler"""
        for checkbox in self.year_checkboxes.values():
            checkbox.setChecked(False)
    
    def show_expense_type_selection_dialog(self):
        """Gider türü seçim dialog'u göster"""
        dialog = QDialog(self)
        dialog.setWindowTitle("Gider Türü Seçimi")
        dialog.setFixedSize(500, 400)
        dialog.setModal(True)
        
        layout = QVBoxLayout(dialog)
        layout.setSpacing(20)
        
        # Başlık
        title = QLabel("Hangi gider türlerini dahil etmek istiyorsunuz?")
        title.setStyleSheet("font-size: 16px; font-weight: bold; color: #2c3e50;")
        title.setAlignment(Qt.AlignmentFlag.AlignCenter)
        layout.addWidget(title)
        
        # Scroll area
        scroll = QScrollArea()
        scroll.setWidgetResizable(True)
        scroll.setStyleSheet("""
            QScrollArea {
                border: 2px solid #bdc3c7;
                border-radius: 8px;
                background-color: white;
            }
            QScrollBar:vertical {
                background-color: #ecf0f1;
                width: 12px;
                border-radius: 6px;
            }
            QScrollBar::handle:vertical {
                background-color: #bdc3c7;
                border-radius: 6px;
                min-height: 20px;
            }
            QScrollBar::handle:vertical:hover {
                background-color: #95a5a6;
            }
        """)
        
        # Checkbox'lar için widget
        checkbox_widget = QWidget()
        checkbox_layout = QVBoxLayout(checkbox_widget)
        checkbox_layout.setSpacing(10)
        
        # Gider türleri
        expense_types = {
            'genel_giderler': '💰 Genel Giderler (Kasko, Muayene, Sigorta vb.)',
            'bakim_onarim': '🔧 Bakım/Onarım Giderleri',
            'yakit_giderleri': '⛽ Yakıt Giderleri'
        }
        
        self.expense_type_checkboxes = {}
        for key, label in expense_types.items():
            checkbox = QCheckBox(label)
            checkbox.setChecked(True)  # Varsayılan olarak hepsi seçili
            checkbox.setStyleSheet("""
                QCheckBox {
                    font-size: 13px;
                    spacing: 8px;
                    padding: 8px 0px;
                }
                QCheckBox::indicator {
                    width: 16px;
                    height: 16px;
                }
                QCheckBox::indicator:unchecked {
                    border: 1px solid #95a5a6;
                    background-color: white;
                    border-radius: 2px;
                }
                QCheckBox::indicator:checked {
                    border: 1px solid #2c3e50;
                    background-color: #2c3e50;
                    border-radius: 2px;
                }
                QCheckBox::indicator:hover {
                    border: 1px solid #3498db;
                }
            """)
            self.expense_type_checkboxes[key] = checkbox
            checkbox_layout.addWidget(checkbox)
        
        scroll.setWidget(checkbox_widget)
        layout.addWidget(scroll)
        
        # Butonlar
        btn_layout = QHBoxLayout()
        
        btn_select_all = QPushButton("Tümünü Seç")
        btn_select_all.clicked.connect(self.select_all_expense_types)
        btn_select_all.setStyleSheet("""
            QPushButton {
                background-color: #27ae60;
                color: white;
                border: none;
                padding: 10px 20px;
                border-radius: 6px;
                font-weight: bold;
            }
            QPushButton:hover {
                background-color: #229954;
            }
        """)
        
        btn_clear_all = QPushButton("Seçimi Temizle")
        btn_clear_all.clicked.connect(self.clear_all_expense_types)
        btn_clear_all.setStyleSheet("""
            QPushButton {
                background-color: #e74c3c;
                color: white;
                border: none;
                padding: 10px 20px;
                border-radius: 6px;
                font-weight: bold;
            }
            QPushButton:hover {
                background-color: #c0392b;
            }
        """)
        
        btn_ok = QPushButton("Tamam")
        btn_ok.clicked.connect(dialog.accept)
        btn_ok.setStyleSheet("""
            QPushButton {
                background-color: #3498db;
                color: white;
                border: none;
                padding: 10px 20px;
                border-radius: 6px;
                font-weight: bold;
            }
            QPushButton:hover {
                background-color: #2980b9;
            }
        """)
        
        btn_cancel = QPushButton("İptal")
        btn_cancel.clicked.connect(dialog.reject)
        btn_cancel.setStyleSheet("""
            QPushButton {
                background-color: #95a5a6;
                color: white;
                border: none;
                padding: 10px 20px;
                border-radius: 6px;
                font-weight: bold;
            }
            QPushButton:hover {
                background-color: #7f8c8d;
            }
        """)
        
        btn_layout.addWidget(btn_select_all)
        btn_layout.addWidget(btn_clear_all)
        btn_layout.addStretch()
        btn_layout.addWidget(btn_ok)
        btn_layout.addWidget(btn_cancel)
        
        layout.addLayout(btn_layout)
        
        if dialog.exec_() == QDialog.Accepted:
            selected_types = []
            for key, checkbox in self.expense_type_checkboxes.items():
                if checkbox.isChecked():
                    selected_types.append(key)
            return selected_types
        else:
            return None
    
    def select_all_expense_types(self):
        """Tüm gider türü checkbox'larını seçer"""
        for checkbox in self.expense_type_checkboxes.values():
            checkbox.setChecked(True)
    
    def clear_all_expense_types(self):
        """Tüm gider türü checkbox'larının seçimini temizler"""
        for checkbox in self.expense_type_checkboxes.values():
            checkbox.setChecked(False)
    
    def create_expense_report_files(self, rapor_data, format_choice, selected_years=None):
        """Gider raporu dosyalarını oluştur"""
        try:
            expenses = rapor_data['expenses']
            maintenance = rapor_data['maintenance']
            fuel = rapor_data['fuel']
            expense_types = rapor_data['expense_types']
            
            # Araç bilgilerini yükle
            vehicles = self.data_manager.load_data('araclar')
            
            # Dosya kaydetme yeri seç
            if format_choice == 'word':
                default_filename = f"Gider_Raporu_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
                filename, _ = QFileDialog.getSaveFileName(
                    self, "Gider Raporunu Kaydet", 
                    default_filename,
                    "Word Dosyası (*.docx)"
                )
            else:  # pdf
                default_filename = f"Gider_Raporu_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf"
                filename, _ = QFileDialog.getSaveFileName(
                    self, "Gider Raporunu Kaydet", 
                    default_filename,
                    "PDF Dosyası (*.pdf)"
                )
            
            if not filename:
                return  # Kullanıcı iptal etti
            
            # Dosya yolunu al
            report_dir = os.path.dirname(filename)
            os.makedirs(report_dir, exist_ok=True)
            
            if format_choice == 'word':
                # Word raporu oluştur
                from docx import Document
                from docx.shared import Inches, Pt
                from docx.enum.text import WD_ALIGN_PARAGRAPH
                
                doc = Document()
                
                # Başlık
                if selected_years:
                    years_text = ", ".join(map(str, selected_years))
                    title_text = f'Yıllara Göre Araç Gider Raporu ({years_text})'
                else:
                    title_text = 'Yıllara Göre Araç Gider Raporu'
                title = doc.add_heading(title_text, 0)
                title.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                # Rapor bilgileri
                doc.add_paragraph(f"Rapor Tarihi: {datetime.now().strftime('%d.%m.%Y %H:%M')}")
                doc.add_paragraph(f"Dahil Edilen Gider Türleri: {', '.join(expense_types)}")
                doc.add_paragraph("")
                
                # Genel giderler
                if 'genel_giderler' in expense_types and not expenses.empty:
                    doc.add_heading('💰 Genel Giderler', level=1)
                    total_general = 0
                    
                    for _, expense in expenses.iterrows():
                        amount = expense.get('Tutar', 0)
                        try:
                            amount = float(amount)
                            total_general += amount
                        except:
                            amount = 0
                        
                        # Araç bilgilerini bul
                        vehicle_info = vehicles[vehicles['Plaka'] == expense.get('Araç Plakası', '')]
                        if not vehicle_info.empty:
                            vehicle_info = vehicle_info.iloc[0]
                            vehicle_text = f"{expense.get('Araç Plakası', 'N/A')} - {vehicle_info.get('Marka', '')} {vehicle_info.get('Model', '')}"
                        else:
                            vehicle_text = expense.get('Araç Plakası', 'N/A')
                        
                        p = doc.add_paragraph()
                        p.add_run(f"Araç: ").bold = True
                        p.add_run(f"{vehicle_text}\n")
                        p.add_run(f"Gider Türü: ").bold = True
                        p.add_run(f"{expense.get('Gider Türü', 'N/A')}\n")
                        p.add_run(f"Tarih: ").bold = True
                        p.add_run(f"{expense.get('Tarih', 'N/A')}\n")
                        p.add_run(f"Tutar: ").bold = True
                        p.add_run(f"{self.tr_money(amount)} TL\n")
                        p.add_run(f"Açıklama: ").bold = True
                        p.add_run(f"{expense.get('Açıklama', 'N/A')}\n")
                        doc.add_paragraph("")
                    
                    doc.add_paragraph(f"Genel Giderler Toplamı: {self.tr_money(total_general)} TL")
                    doc.add_paragraph("")
                
                # Bakım/Onarım giderleri
                if 'bakim_onarim' in expense_types and not maintenance.empty:
                    doc.add_heading('🔧 Bakım/Onarım Giderleri', level=1)
                    total_maintenance = 0
                    
                    for _, maint in maintenance.iterrows():
                        amount = maint.get('Maliyet', 0)
                        try:
                            amount = float(amount)
                            total_maintenance += amount
                        except:
                            amount = 0
                        
                        # Araç bilgilerini bul
                        vehicle_info = vehicles[vehicles['Plaka'] == maint.get('Araç Plakası', '')]
                        if not vehicle_info.empty:
                            vehicle_info = vehicle_info.iloc[0]
                            vehicle_text = f"{maint.get('Araç Plakası', 'N/A')} - {vehicle_info.get('Marka', '')} {vehicle_info.get('Model', '')}"
                        else:
                            vehicle_text = maint.get('Araç Plakası', 'N/A')
                        
                        p = doc.add_paragraph()
                        p.add_run(f"Araç: ").bold = True
                        p.add_run(f"{vehicle_text}\n")
                        p.add_run(f"Bakım Türü: ").bold = True
                        p.add_run(f"{maint.get('Bakım Türü', 'N/A')}\n")
                        p.add_run(f"Bakım Tarihi: ").bold = True
                        p.add_run(f"{maint.get('Bakım Tarihi', 'N/A')}\n")
                        p.add_run(f"Maliyet: ").bold = True
                        p.add_run(f"{self.tr_money(amount)} TL\n")
                        p.add_run(f"Açıklama: ").bold = True
                        p.add_run(f"{maint.get('Açıklama', 'N/A')}\n")
                        doc.add_paragraph("")
                    
                    doc.add_paragraph(f"Bakım/Onarım Toplamı: {self.tr_money(total_maintenance)} TL")
                    doc.add_paragraph("")
                
                # Yakıt giderleri
                if 'yakit_giderleri' in expense_types and not fuel.empty:
                    doc.add_heading('⛽ Yakıt Giderleri', level=1)
                    total_fuel = 0
                    
                    for _, fuel_record in fuel.iterrows():
                        amount = fuel_record.get('Tutar', 0)
                        try:
                            amount = float(amount)
                            total_fuel += amount
                        except:
                            amount = 0
                        
                        # Araç bilgilerini bul
                        vehicle_info = vehicles[vehicles['Plaka'] == fuel_record.get('Araç Plakası', '')]
                        if not vehicle_info.empty:
                            vehicle_info = vehicle_info.iloc[0]
                            vehicle_text = f"{fuel_record.get('Araç Plakası', 'N/A')} - {vehicle_info.get('Marka', '')} {vehicle_info.get('Model', '')}"
                        else:
                            vehicle_text = fuel_record.get('Araç Plakası', 'N/A')
                        
                        p = doc.add_paragraph()
                        p.add_run(f"Araç: ").bold = True
                        p.add_run(f"{vehicle_text}\n")
                        p.add_run(f"Yakıt Türü: ").bold = True
                        p.add_run(f"{fuel_record.get('Yakıt Türü', 'N/A')}\n")
                        p.add_run(f"Tarih: ").bold = True
                        p.add_run(f"{fuel_record.get('Tarih', 'N/A')}\n")
                        p.add_run(f"Litre: ").bold = True
                        p.add_run(f"{fuel_record.get('Litre', 'N/A')} L\n")
                        p.add_run(f"Tutar: ").bold = True
                        p.add_run(f"{self.tr_money(amount)} TL\n")
                        doc.add_paragraph("")
                    
                    doc.add_paragraph(f"Yakıt Giderleri Toplamı: {self.tr_money(total_fuel)} TL")
                    doc.add_paragraph("")
                
                # Genel toplam
                total_all = 0
                if 'genel_giderler' in expense_types and not expenses.empty:
                    total_all += expenses['Tutar'].astype(float).sum()
                if 'bakim_onarim' in expense_types and not maintenance.empty:
                    total_all += maintenance['Maliyet'].astype(float).sum()
                if 'yakit_giderleri' in expense_types and not fuel.empty:
                    total_all += fuel['Tutar'].astype(float).sum()
                
                doc.add_heading('📊 GENEL TOPLAM', level=1)
                doc.add_paragraph(f"Seçilen yıllar için toplam gider: {self.tr_money(total_all)} TL")
                
                # Dosyayı kaydet
                doc.save(filename)
                QMessageBox.information(self, "Başarılı", f"Word raporu oluşturuldu:\n{filename}")
                
            elif format_choice == 'pdf':
                # PDF raporu oluştur
                from reportlab.lib.pagesizes import A4
                from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
                from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
                from reportlab.lib.units import inch
                from reportlab.lib import colors
                from reportlab.pdfbase import pdfmetrics
                from reportlab.pdfbase.ttfonts import TTFont
                
                # Türkçe font desteği
                try:
                    pdfmetrics.registerFont(TTFont('DejaVuSans', 'DejaVuSans.ttf'))
                    font_name = 'DejaVuSans'
                except:
                    font_name = 'Helvetica'
                
                doc = SimpleDocTemplate(filename, pagesize=A4)
                story = []
                
                # Stil tanımları
                styles = getSampleStyleSheet()
                title_style = ParagraphStyle(
                    'CustomTitle',
                    parent=styles['Heading1'],
                    fontSize=18,
                    spaceAfter=30,
                    alignment=1,  # Center
                    fontName=font_name
                )
                
                heading_style = ParagraphStyle(
                    'CustomHeading',
                    parent=styles['Heading2'],
                    fontSize=14,
                    spaceAfter=12,
                    fontName=font_name
                )
                
                normal_style = ParagraphStyle(
                    'CustomNormal',
                    parent=styles['Normal'],
                    fontSize=10,
                    spaceAfter=6,
                    fontName=font_name
                )
                
                # Başlık
                if selected_years:
                    years_text = ", ".join(map(str, selected_years))
                    title_text = f'Yıllara Göre Araç Gider Raporu ({years_text})'
                else:
                    title_text = 'Yıllara Göre Araç Gider Raporu'
                story.append(Paragraph(title_text, title_style))
                
                # Rapor bilgileri
                story.append(Paragraph(f"Rapor Tarihi: {datetime.now().strftime('%d.%m.%Y %H:%M')}", normal_style))
                story.append(Paragraph(f"Dahil Edilen Gider Türleri: {', '.join(expense_types)}", normal_style))
                story.append(Spacer(1, 20))
                
                # Genel giderler
                if 'genel_giderler' in expense_types and not expenses.empty:
                    story.append(Paragraph('💰 Genel Giderler', heading_style))
                    total_general = 0
                    
                    for _, expense in expenses.iterrows():
                        amount = expense.get('Tutar', 0)
                        try:
                            amount = float(amount)
                            total_general += amount
                        except:
                            amount = 0
                        
                        # Araç bilgilerini bul
                        vehicle_info = vehicles[vehicles['Plaka'] == expense.get('Araç Plakası', '')]
                        if not vehicle_info.empty:
                            vehicle_info = vehicle_info.iloc[0]
                            vehicle_text = f"{expense.get('Araç Plakası', 'N/A')} - {vehicle_info.get('Marka', '')} {vehicle_info.get('Model', '')}"
                        else:
                            vehicle_text = expense.get('Araç Plakası', 'N/A')
                        
                        story.append(Paragraph(f"<b>Araç:</b> {vehicle_text}", normal_style))
                        story.append(Paragraph(f"<b>Gider Türü:</b> {expense.get('Gider Türü', 'N/A')}", normal_style))
                        story.append(Paragraph(f"<b>Tarih:</b> {expense.get('Tarih', 'N/A')}", normal_style))
                        story.append(Paragraph(f"<b>Tutar:</b> {self.tr_money(amount)} TL", normal_style))
                        story.append(Paragraph(f"<b>Açıklama:</b> {expense.get('Açıklama', 'N/A')}", normal_style))
                        story.append(Spacer(1, 10))
                    
                    story.append(Paragraph(f"<b>Genel Giderler Toplamı: {self.tr_money(total_general)} TL</b>", normal_style))
                    story.append(Spacer(1, 20))
                
                # Bakım/Onarım giderleri
                if 'bakim_onarim' in expense_types and not maintenance.empty:
                    story.append(Paragraph('🔧 Bakım/Onarım Giderleri', heading_style))
                    total_maintenance = 0
                    
                    for _, maint in maintenance.iterrows():
                        amount = maint.get('Maliyet', 0)
                        try:
                            amount = float(amount)
                            total_maintenance += amount
                        except:
                            amount = 0
                        
                        # Araç bilgilerini bul
                        vehicle_info = vehicles[vehicles['Plaka'] == maint.get('Araç Plakası', '')]
                        if not vehicle_info.empty:
                            vehicle_info = vehicle_info.iloc[0]
                            vehicle_text = f"{maint.get('Araç Plakası', 'N/A')} - {vehicle_info.get('Marka', '')} {vehicle_info.get('Model', '')}"
                        else:
                            vehicle_text = maint.get('Araç Plakası', 'N/A')
                        
                        story.append(Paragraph(f"<b>Araç:</b> {vehicle_text}", normal_style))
                        story.append(Paragraph(f"<b>Bakım Türü:</b> {maint.get('Bakım Türü', 'N/A')}", normal_style))
                        story.append(Paragraph(f"<b>Bakım Tarihi:</b> {maint.get('Bakım Tarihi', 'N/A')}", normal_style))
                        story.append(Paragraph(f"<b>Maliyet:</b> {self.tr_money(amount)} TL", normal_style))
                        story.append(Paragraph(f"<b>Açıklama:</b> {maint.get('Açıklama', 'N/A')}", normal_style))
                        story.append(Spacer(1, 10))
                    
                    story.append(Paragraph(f"<b>Bakım/Onarım Toplamı: {self.tr_money(total_maintenance)} TL</b>", normal_style))
                    story.append(Spacer(1, 20))
                
                # Yakıt giderleri
                if 'yakit_giderleri' in expense_types and not fuel.empty:
                    story.append(Paragraph('⛽ Yakıt Giderleri', heading_style))
                    total_fuel = 0
                    
                    for _, fuel_record in fuel.iterrows():
                        amount = fuel_record.get('Tutar', 0)
                        try:
                            amount = float(amount)
                            total_fuel += amount
                        except:
                            amount = 0
                        
                        # Araç bilgilerini bul
                        vehicle_info = vehicles[vehicles['Plaka'] == fuel_record.get('Araç Plakası', '')]
                        if not vehicle_info.empty:
                            vehicle_info = vehicle_info.iloc[0]
                            vehicle_text = f"{fuel_record.get('Araç Plakası', 'N/A')} - {vehicle_info.get('Marka', '')} {vehicle_info.get('Model', '')}"
                        else:
                            vehicle_text = fuel_record.get('Araç Plakası', 'N/A')
                        
                        story.append(Paragraph(f"<b>Araç:</b> {vehicle_text}", normal_style))
                        story.append(Paragraph(f"<b>Yakıt Türü:</b> {fuel_record.get('Yakıt Türü', 'N/A')}", normal_style))
                        story.append(Paragraph(f"<b>Tarih:</b> {fuel_record.get('Tarih', 'N/A')}", normal_style))
                        story.append(Paragraph(f"<b>Litre:</b> {fuel_record.get('Litre', 'N/A')} L", normal_style))
                        story.append(Paragraph(f"<b>Tutar:</b> {self.tr_money(amount)} TL", normal_style))
                        story.append(Spacer(1, 10))
                    
                    story.append(Paragraph(f"<b>Yakıt Giderleri Toplamı: {self.tr_money(total_fuel)} TL</b>", normal_style))
                    story.append(Spacer(1, 20))
                
                # Genel toplam
                total_all = 0
                if 'genel_giderler' in expense_types and not expenses.empty:
                    total_all += expenses['Tutar'].astype(float).sum()
                if 'bakim_onarim' in expense_types and not maintenance.empty:
                    total_all += maintenance['Maliyet'].astype(float).sum()
                if 'yakit_giderleri' in expense_types and not fuel.empty:
                    total_all += fuel['Tutar'].astype(float).sum()
                
                story.append(Paragraph('📊 GENEL TOPLAM', heading_style))
                story.append(Paragraph(f"<b>Seçilen yıllar için toplam gider: {self.tr_money(total_all)} TL</b>", normal_style))
                
                # PDF oluştur
                doc.build(story)
                QMessageBox.information(self, "Başarılı", f"PDF raporu oluşturuldu:\n{filename}")
                
        except Exception as e:
            QMessageBox.critical(self, "Hata", f"Rapor oluşturma hatası: {str(e)}")

class CalendarPanel(QWidget):
    def __init__(self, data_manager):
        super().__init__()
        self.data_manager = data_manager
        self.init_ui()
        self.load_reminders()
    
    def init_ui(self):
        layout = QVBoxLayout(self)
        layout.setContentsMargins(20, 20, 20, 20)
        layout.setSpacing(20)
        
        # Başlık ve butonlar
        header_layout = QHBoxLayout()
        
        title = QLabel("Takvim & Hatırlatmalar")
        title.setStyleSheet("font-size: 24px; font-weight: bold; color: #2c3e50;")
        header_layout.addWidget(title)
        
        header_layout.addStretch()
        
        # Butonlar
        btn_add = QPushButton("➕ Hatırlatma Ekle")
        btn_add.setStyleSheet("""
            QPushButton {
                background-color: #27ae60;
                color: white;
                border: none;
                padding: 10px 20px;
                border-radius: 5px;
                font-weight: bold;
            }
            QPushButton:hover {
                background-color: #229954;
            }
        """)
        btn_add.clicked.connect(self.add_reminder)
        
        btn_edit = QPushButton("✏️ Düzenle")
        btn_edit.setStyleSheet("""
            QPushButton {
                background-color: #3498db;
                color: white;
                border: none;
                padding: 10px 20px;
                border-radius: 5px;
                font-weight: bold;
            }
            QPushButton:hover {
                background-color: #2980b9;
            }
        """)
        btn_edit.clicked.connect(self.edit_reminder)
        
        btn_delete = QPushButton("🗑️ Sil")
        btn_delete.setStyleSheet("""
            QPushButton {
                background-color: #e74c3c;
                color: white;
                border: none;
                padding: 10px 20px;
                border-radius: 5px;
                font-weight: bold;
            }
            QPushButton:hover {
                background-color: #c0392b;
            }
        """)
        btn_delete.clicked.connect(self.delete_reminder)
        
        header_layout.addWidget(btn_add)
        header_layout.addWidget(btn_edit)
        header_layout.addWidget(btn_delete)
        
        layout.addLayout(header_layout)
        
        # Yaklaşan hatırlatmalar
        upcoming_group = QGroupBox("Yaklaşan Hatırlatmalar")
        upcoming_group.setStyleSheet("""
            QGroupBox {
                font-weight: bold;
                border: 2px solid #bdc3c7;
                border-radius: 5px;
                margin-top: 10px;
                padding-top: 10px;
            }
        """)
        
        upcoming_layout = QVBoxLayout(upcoming_group)
        
        self.upcoming_list = QListWidget()
        self.upcoming_list.setMaximumHeight(150)
        upcoming_layout.addWidget(self.upcoming_list)
        
        layout.addWidget(upcoming_group)
        
        # Tüm hatırlatmalar
        all_group = QGroupBox("Tüm Hatırlatmalar")
        all_group.setStyleSheet("""
            QGroupBox {
                font-weight: bold;
                border: 2px solid #bdc3c7;
                border-radius: 5px;
                margin-top: 10px;
                padding-top: 10px;
            }
        """)
        
        all_layout = QVBoxLayout(all_group)
        
        self.reminders_table = QTableWidget()
        self.reminders_table.setColumnCount(5)
        self.reminders_table.setHorizontalHeaderLabels([
            "Başlık", "Açıklama", "Tarih", "Tür", "Durum"
        ])
        self.reminders_table.horizontalHeader().setStretchLastSection(True)
        self.reminders_table.setAlternatingRowColors(True)
        self.reminders_table.setSelectionBehavior(QAbstractItemView.SelectRows)
        self.reminders_table.setSelectionMode(QAbstractItemView.SingleSelection)
        
        # Satır numaralarını gizle
        self.reminders_table.verticalHeader().setVisible(False)
        
        all_layout.addWidget(self.reminders_table)
        layout.addWidget(all_group)
    
    def load_reminders(self):
        """Hatırlatmaları yükle"""
        try:
            reminders = self.data_manager.load_data('hatirlatmalar')
            self.reminders_table.setRowCount(len(reminders))
            
            # Yaklaşan hatırlatmaları temizle
            self.upcoming_list.clear()
            
            current_date = datetime.now().date()
            
            for row, (_, reminder) in enumerate(reminders.iterrows()):
                # Tablo için
                self.reminders_table.setItem(row, 0, QTableWidgetItem(str(reminder.get('Başlık', ''))))
                self.reminders_table.setItem(row, 1, QTableWidgetItem(str(reminder.get('Açıklama', ''))))
                self.reminders_table.setItem(row, 2, QTableWidgetItem(str(reminder.get('Tarih', ''))))
                self.reminders_table.setItem(row, 3, QTableWidgetItem(str(reminder.get('Tür', ''))))
                self.reminders_table.setItem(row, 4, QTableWidgetItem(str(reminder.get('Durum', ''))))
                
                # Yaklaşan hatırlatmalar için
                try:
                    reminder_date = datetime.strptime(str(reminder.get('Tarih', '')), "%Y-%m-%d").date()
                    if reminder_date >= current_date and reminder.get('Durum') != 'Tamamlandı':
                        days_left = (reminder_date - current_date).days
                        item_text = f"{reminder.get('Başlık', '')} - {reminder_date.strftime('%d.%m.%Y')} ({days_left} gün kaldı)"
                        self.upcoming_list.addItem(item_text)
                except:
                    pass
                    
        except Exception as e:
            print(f"Hatırlatma yükleme hatası: {e}")
    
    def add_reminder(self):
        """Hatırlatma ekle"""
        dialog = ReminderDialog(self)
        if dialog.exec_() == QDialog.Accepted:
            self.load_reminders()
    
    def edit_reminder(self):
        """Hatırlatma düzenle"""
        current_row = self.reminders_table.currentRow()
        if current_row >= 0:
            baslik = self.reminders_table.item(current_row, 0).text()
            reminders = self.data_manager.load_data('hatirlatmalar')
            reminder = reminders[reminders['Başlık'] == baslik]
            if not reminder.empty:
                dialog = ReminderDialog(self, reminder.iloc[0].to_dict())
                if dialog.exec_() == QDialog.Accepted:
                    self.load_reminders()
        else:
            QMessageBox.warning(self, "Uyarı", "Lütfen düzenlenecek hatırlatmayı seçin.")
    
    def delete_reminder(self):
        """Hatırlatma sil"""
        current_row = self.reminders_table.currentRow()
        if current_row >= 0:
            baslik = self.reminders_table.item(current_row, 0).text()
            reply = QMessageBox.question(self, "Onay", 
                f"'{baslik}' hatırlatmasını silmek istediğinizden emin misiniz?",
                QMessageBox.Yes | QMessageBox.No)
            
            if reply == QMessageBox.Yes:
                try:
                    reminders = self.data_manager.load_data('hatirlatmalar')
                    reminders = reminders[reminders['Başlık'] != baslik]
                    self.data_manager.save_data('hatirlatmalar', reminders)
                    self.load_reminders()
                    QMessageBox.information(self, "Başarılı", "Hatırlatma silindi.")
                except Exception as e:
                    QMessageBox.critical(self, "Hata", f"Hatırlatma silme hatası: {str(e)}")
        else:
            QMessageBox.warning(self, "Uyarı", "Lütfen silinecek hatırlatmayı seçin.")

class ReminderDialog(QDialog):
    """Hatırlatma ekleme/düzenleme dialog"""
    
    def __init__(self, parent=None, reminder_data=None):
        super().__init__(parent)
        self.reminder_data = reminder_data
        self.data_manager = parent.data_manager if parent else DataManager()
        self.init_ui()
        if reminder_data:
            self.load_reminder_data(reminder_data)
    
    def init_ui(self):
        self.setWindowTitle("Hatırlatma Ekle" if not self.reminder_data else "Hatırlatma Düzenle")
        self.setMinimumWidth(500)
        self.setModal(True)
        
        layout = QFormLayout(self)
        layout.setSpacing(15)
        
        # Form alanları
        self.baslik = QLineEdit()
        
        self.aciklama = QTextEdit()
        self.aciklama.setMaximumHeight(100)
        
        self.tarih = QDateEdit()
        self.tarih.setCalendarPopup(True)
        self.tarih.setDate(QDate.currentDate())
        
        self.tur = QComboBox()
        self.tur.addItems([
            "Muayene", "Sigorta", "Bakım", "Kasko", "Vergi", "Diğer"
        ])
        
        self.durum = QComboBox()
        self.durum.addItems(["Bekliyor", "Tamamlandı", "İptal"])
        
        # Form'a ekle
        layout.addRow("Başlık:", self.baslik)
        layout.addRow("Açıklama:", self.aciklama)
        layout.addRow("Tarih:", self.tarih)
        layout.addRow("Tür:", self.tur)
        layout.addRow("Durum:", self.durum)
        
        # Butonlar
        btn_layout = QHBoxLayout()
        btn_save = QPushButton("Kaydet")
        btn_save.clicked.connect(self.save_reminder)
        btn_cancel = QPushButton("İptal")
        btn_cancel.clicked.connect(self.reject)
        
        btn_layout.addWidget(btn_save)
        btn_layout.addWidget(btn_cancel)
        layout.addRow(btn_layout)
    
    def load_reminder_data(self, reminder_data):
        """Hatırlatma verilerini yükle"""
        self.baslik.setText(str(reminder_data.get('Başlık', '')))
        self.aciklama.setPlainText(str(reminder_data.get('Açıklama', '')))
        
        # Tarihi yükle
        try:
            reminder_date = QDate.fromString(str(reminder_data.get('Tarih', '')), "yyyy-MM-dd")
            if reminder_date.isValid():
                self.tarih.setDate(reminder_date)
        except:
            pass
        
        self.tur.setCurrentText(str(reminder_data.get('Tür', 'Diğer')))
        self.durum.setCurrentText(str(reminder_data.get('Durum', 'Bekliyor')))
    
    def save_reminder(self):
        """Hatırlatma kaydet"""
        if not self.baslik.text().strip():
            QMessageBox.warning(self, "Uyarı", "Başlık alanı zorunludur.")
            return
        
        try:
            reminders = self.data_manager.load_data('hatirlatmalar')
            
            # Yeni hatırlatma verisi
            new_reminder = {
                'ID': len(reminders) + 1 if not reminders.empty else 1,
                'Başlık': self.baslik.text().strip(),
                'Açıklama': self.aciklama.toPlainText().strip(),
                'Tarih': self.tarih.date().toString("yyyy-MM-dd"),
                'Tür': self.tur.currentText(),
                'Durum': self.durum.currentText(),
                'Oluşturma Tarihi': datetime.now().strftime("%Y-%m-%d %H:%M:%S")
            }
            
            if self.reminder_data:
                # Düzenleme modu
                reminders = reminders[reminders['Başlık'] != self.reminder_data['Başlık']]
            
            # Yeni veriyi ekle
            new_df = pd.DataFrame([new_reminder])
            reminders = pd.concat([reminders, new_df], ignore_index=True)
            
            # Kaydet
            if self.data_manager.save_data('hatirlatmalar', reminders):
                QMessageBox.information(self, "Başarılı", 
                    "Hatırlatma güncellendi." if self.reminder_data else "Hatırlatma eklendi.")
                self.accept()
            else:
                QMessageBox.critical(self, "Hata", "Kaydetme hatası.")
                
        except Exception as e:
            QMessageBox.critical(self, "Hata", f"Kaydetme hatası: {str(e)}")

class DocumentsPanel(QWidget):
    def __init__(self, data_manager):
        super().__init__()
        self.data_manager = data_manager
        self.init_ui()
        self.load_documents()
    
    def init_ui(self):
        layout = QVBoxLayout(self)
        layout.setContentsMargins(20, 20, 20, 20)
        layout.setSpacing(20)
        
        # Başlık ve butonlar
        header_layout = QHBoxLayout()
        
        title = QLabel("Belge Yönetimi")
        title.setStyleSheet("font-size: 24px; font-weight: bold; color: #2c3e50;")
        header_layout.addWidget(title)
        
        header_layout.addStretch()
        
        # Butonlar
        btn_add = QPushButton("📁 Belge Ekle")
        btn_add.setStyleSheet("""
            QPushButton {
                background-color: #27ae60;
                color: white;
                border: none;
                padding: 10px 20px;
                border-radius: 5px;
                font-weight: bold;
            }
            QPushButton:hover {
                background-color: #229954;
            }
        """)
        btn_add.clicked.connect(self.add_document)
        
        btn_view = QPushButton("👁️ Görüntüle")
        btn_view.setStyleSheet("""
            QPushButton {
                background-color: #3498db;
                color: white;
                border: none;
                padding: 10px 20px;
                border-radius: 5px;
                font-weight: bold;
            }
            QPushButton:hover {
                background-color: #2980b9;
            }
        """)
        btn_view.clicked.connect(self.view_document)
        
        btn_delete = QPushButton("🗑️ Sil")
        btn_delete.setStyleSheet("""
            QPushButton {
                background-color: #e74c3c;
                color: white;
                border: none;
                padding: 10px 20px;
                border-radius: 5px;
                font-weight: bold;
            }
            QPushButton:hover {
                background-color: #c0392b;
            }
        """)
        btn_delete.clicked.connect(self.delete_document)
        
        header_layout.addWidget(btn_add)
        header_layout.addWidget(btn_view)
        header_layout.addWidget(btn_delete)
        
        layout.addLayout(header_layout)
        
        # Belge listesi
        self.documents_table = QTableWidget()
        self.documents_table.setColumnCount(5)
        self.documents_table.setHorizontalHeaderLabels([
            "Araç Plakası", "Belge Türü", "Dosya Adı", "Yükleme Tarihi", "Açıklama"
        ])
        self.documents_table.horizontalHeader().setStretchLastSection(True)
        self.documents_table.setAlternatingRowColors(True)
        self.documents_table.setSelectionBehavior(QAbstractItemView.SelectRows)
        self.documents_table.setSelectionMode(QAbstractItemView.SingleSelection)
        
        # Satır numaralarını gizle
        self.documents_table.verticalHeader().setVisible(False)
        
        layout.addWidget(self.documents_table)
    
    def load_documents(self):
        """Belgeleri yükle"""
        try:
            documents = self.data_manager.load_data('belgeler')
            self.documents_table.setRowCount(len(documents))
            
            for row, (_, document) in enumerate(documents.iterrows()):
                self.documents_table.setItem(row, 0, QTableWidgetItem(str(document.get('Araç Plakası', ''))))
                self.documents_table.setItem(row, 1, QTableWidgetItem(str(document.get('Belge Türü', ''))))
                self.documents_table.setItem(row, 2, QTableWidgetItem(str(document.get('Dosya Adı', ''))))
                self.documents_table.setItem(row, 3, QTableWidgetItem(str(document.get('Yükleme Tarihi', ''))))
                self.documents_table.setItem(row, 4, QTableWidgetItem(str(document.get('Açıklama', ''))))
        except Exception as e:
            print(f"Belge yükleme hatası: {e}")
    
    def add_document(self):
        """Belge ekle"""
        dialog = DocumentDialog(self)
        if dialog.exec_() == QDialog.Accepted:
            self.load_documents()
    
    def view_document(self):
        """Belge görüntüle"""
        current_row = self.documents_table.currentRow()
        if current_row >= 0:
            dosya_adi = self.documents_table.item(current_row, 2).text()
            dosya_yolu = self.documents_table.item(current_row, 2).text()  # Basit implementasyon
            
            if dosya_yolu and os.path.exists(dosya_yolu):
                try:
                    os.startfile(dosya_yolu)  # Windows için
                except:
                    QMessageBox.information(self, "Bilgi", f"Dosya açılamadı: {dosya_yolu}")
            else:
                QMessageBox.warning(self, "Uyarı", "Dosya bulunamadı.")
        else:
            QMessageBox.warning(self, "Uyarı", "Lütfen görüntülenecek belgeyi seçin.")
    
    def delete_document(self):
        """Belge sil"""
        current_row = self.documents_table.currentRow()
        if current_row >= 0:
            dosya_adi = self.documents_table.item(current_row, 2).text()
            reply = QMessageBox.question(self, "Onay", 
                f"'{dosya_adi}' belgesini silmek istediğinizden emin misiniz?",
                QMessageBox.Yes | QMessageBox.No)
            
            if reply == QMessageBox.Yes:
                try:
                    documents = self.data_manager.load_data('belgeler')
                    documents = documents[documents['Dosya Adı'] != dosya_adi]
                    self.data_manager.save_data('belgeler', documents)
                    self.load_documents()
                    QMessageBox.information(self, "Başarılı", "Belge silindi.")
                except Exception as e:
                    QMessageBox.critical(self, "Hata", f"Belge silme hatası: {str(e)}")
        else:
            QMessageBox.warning(self, "Uyarı", "Lütfen silinecek belgeyi seçin.")

class DocumentDialog(QDialog):
    """Belge ekleme dialog"""
    
    def __init__(self, parent=None):
        super().__init__(parent)
        self.data_manager = parent.data_manager if parent else DataManager()
        self.selected_file_path = ""
        self.init_ui()
    
    def init_ui(self):
        self.setWindowTitle("Belge Ekle")
        self.setMinimumWidth(500)
        self.setModal(True)
        
        layout = QFormLayout(self)
        layout.setSpacing(15)
        
        # Form alanları
        self.arac_plakasi = QComboBox()
        self.load_vehicles()
        
        self.belge_turu = QComboBox()
        self.belge_turu.addItems([
            "Ruhsat", "Sigorta", "Muayene", "Kasko", "Fatura", "Diğer"
        ])
        
        # Dosya seçimi
        file_layout = QHBoxLayout()
        self.file_path_label = QLabel("Dosya seçilmedi")
        self.file_path_label.setStyleSheet("color: #7f8c8d; font-style: italic;")
        
        btn_browse = QPushButton("Dosya Seç")
        btn_browse.clicked.connect(self.browse_file)
        
        file_layout.addWidget(self.file_path_label)
        file_layout.addWidget(btn_browse)
        
        self.aciklama = QTextEdit()
        self.aciklama.setMaximumHeight(100)
        
        # Form'a ekle
        layout.addRow("Araç Plakası:", self.arac_plakasi)
        layout.addRow("Belge Türü:", self.belge_turu)
        layout.addRow("Dosya:", file_layout)
        layout.addRow("Açıklama:", self.aciklama)
        
        # Butonlar
        btn_layout = QHBoxLayout()
        btn_save = QPushButton("Kaydet")
        btn_save.clicked.connect(self.save_document)
        btn_cancel = QPushButton("İptal")
        btn_cancel.clicked.connect(self.reject)
        
        btn_layout.addWidget(btn_save)
        btn_layout.addWidget(btn_cancel)
        layout.addRow(btn_layout)
    
    def load_vehicles(self):
        """Mevcut araçları yükle"""
        try:
            vehicles = self.data_manager.load_data('araclar')
            self.arac_plakasi.clear()
            
            if not vehicles.empty:
                for _, vehicle in vehicles.iterrows():
                    plaka = vehicle.get('Plaka', '')
                    marka = vehicle.get('Marka', '')
                    model = vehicle.get('Model', '')
                    self.arac_plakasi.addItem(f"{plaka} - {marka} {model}")
        except Exception as e:
            print(f"Araç yükleme hatası: {e}")
    
    def browse_file(self):
        """Dosya seç"""
        file_path, _ = QFileDialog.getOpenFileName(
            self, "Belge Seç", "", 
            "Tüm Dosyalar (*);;PDF Dosyaları (*.pdf);;Resim Dosyaları (*.jpg *.jpeg *.png);;Word Dosyaları (*.doc *.docx)"
        )
        
        if file_path:
            self.selected_file_path = file_path
            self.file_path_label.setText(os.path.basename(file_path))
            self.file_path_label.setStyleSheet("color: #2c3e50; font-weight: bold;")
    
    def save_document(self):
        """Belge kaydet"""
        if not self.arac_plakasi.currentText():
            QMessageBox.warning(self, "Uyarı", "Lütfen bir araç seçin.")
            return
        
        if not self.selected_file_path:
            QMessageBox.warning(self, "Uyarı", "Lütfen bir dosya seçin.")
            return
        
        try:
            documents = self.data_manager.load_data('belgeler')
            
            # Araç plakasını ayır
            arac_text = self.arac_plakasi.currentText()
            plaka = arac_text.split(' - ')[0] if ' - ' in arac_text else arac_text
            
            # Belgeler klasörünü oluştur
            belgeler_dir = os.path.join(Config.DATA_DIR, "belgeler")
            if not os.path.exists(belgeler_dir):
                os.makedirs(belgeler_dir)
            
            # Orijinal dosya adını al
            orijinal_dosya_adi = os.path.basename(self.selected_file_path)
            dosya_uzantisi = os.path.splitext(orijinal_dosya_adi)[1]
            
            # Yeni dosya adını oluştur: plaka_tarih_uzanti
            tarih_str = datetime.now().strftime("%Y%m%d_%H%M%S")
            yeni_dosya_adi = f"{plaka}_{tarih_str}{dosya_uzantisi}"
            
            # Yeni dosya yolu
            yeni_dosya_yolu = os.path.join(belgeler_dir, yeni_dosya_adi)
            
            # Dosyayı belgeler klasörüne kopyala
            import shutil
            shutil.copy2(self.selected_file_path, yeni_dosya_yolu)
            
            # Yeni belge verisi
            new_document = {
                'ID': len(documents) + 1 if not documents.empty else 1,
                'Araç Plakası': plaka,
                'Belge Türü': self.belge_turu.currentText(),
                'Dosya Adı': orijinal_dosya_adi,
                'Dosya Yolu': yeni_dosya_yolu,
                'Yükleme Tarihi': datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
                'Açıklama': self.aciklama.toPlainText().strip()
            }
            
            # Yeni veriyi ekle
            new_df = pd.DataFrame([new_document])
            documents = pd.concat([documents, new_df], ignore_index=True)
            
            # Kaydet
            if self.data_manager.save_data('belgeler', documents):
                QMessageBox.information(self, "Başarılı", f"Belge başarıyla eklendi.\nDosya: {yeni_dosya_yolu}")
                self.accept()
            else:
                QMessageBox.critical(self, "Hata", "Kaydetme hatası.")
                
        except Exception as e:
            QMessageBox.critical(self, "Hata", f"Kaydetme hatası: {str(e)}")

class SettingsPanel(QWidget):
    def __init__(self, data_manager):
        super().__init__()
        self.data_manager = data_manager
        self.init_ui()
    
    def init_ui(self):
        layout = QVBoxLayout(self)
        layout.setContentsMargins(20, 20, 20, 20)
        layout.setSpacing(20)
        
        # Başlık
        title = QLabel("Ayarlar")
        title.setStyleSheet("font-size: 24px; font-weight: bold; color: #2c3e50;")
        layout.addWidget(title)
        
        # Uygulama bilgileri
        app_group = QGroupBox("Uygulama Bilgileri")
        app_group.setStyleSheet("""
            QGroupBox {
                font-weight: bold;
                border: 2px solid #bdc3c7;
                border-radius: 5px;
                margin-top: 10px;
                padding-top: 10px;
            }
        """)
        
        app_layout = QFormLayout(app_group)
        
        app_name_label = QLabel(Config.APP_NAME)
        app_version_label = QLabel(Config.VERSION)
        app_developer_label = QLabel(Config.DEVELOPER)
        app_email_label = QLabel(Config.EMAIL)
        
        app_layout.addRow("Uygulama Adı:", app_name_label)
        app_layout.addRow("Versiyon:", app_version_label)
        app_layout.addRow("Geliştirici:", app_developer_label)
        app_layout.addRow("E-posta:", app_email_label)
        
        layout.addWidget(app_group)
        
        # Veri yönetimi
        data_group = QGroupBox("Veri Yönetimi")
        data_group.setStyleSheet("""
            QGroupBox {
                font-weight: bold;
                border: 2px solid #bdc3c7;
                border-radius: 5px;
                margin-top: 10px;
                padding-top: 10px;
            }
        """)
        
        data_layout = QVBoxLayout(data_group)
        
        # Veri yedekleme
        btn_backup = QPushButton("💾 Veri Yedekle")
        btn_backup.setStyleSheet("""
            QPushButton {
                background-color: #3498db;
                color: white;
                border: none;
                padding: 15px;
                border-radius: 5px;
                font-weight: bold;
                font-size: 14px;
            }
            QPushButton:hover {
                background-color: #2980b9;
            }
        """)
        btn_backup.clicked.connect(self.backup_data)
        
        # Excel export
        btn_export = QPushButton("📊 Excel'e Aktar")
        btn_export.setStyleSheet("""
            QPushButton {
                background-color: #27ae60;
                color: white;
                border: none;
                padding: 15px;
                border-radius: 5px;
                font-weight: bold;
                font-size: 14px;
            }
            QPushButton:hover {
                background-color: #229954;
            }
        """)
        btn_export.clicked.connect(self.export_to_excel)
        
        # Excel şablon indirme
        btn_templates = QPushButton("📋 Excel Şablonları İndir")
        btn_templates.setStyleSheet("""
            QPushButton {
                background-color: #f39c12;
                color: white;
                border: none;
                padding: 15px;
                border-radius: 5px;
                font-weight: bold;
                font-size: 14px;
            }
            QPushButton:hover {
                background-color: #e67e22;
            }
        """)
        btn_templates.clicked.connect(self.download_excel_templates)
        
        data_layout.addWidget(btn_backup)
        data_layout.addWidget(btn_export)
        data_layout.addWidget(btn_templates)
        
        layout.addWidget(data_group)
        
        # Sistem bilgileri
        system_group = QGroupBox("Sistem Bilgileri")
        system_group.setStyleSheet("""
            QGroupBox {
                font-weight: bold;
                border: 2px solid #bdc3c7;
                border-radius: 5px;
                margin-top: 10px;
                padding-top: 10px;
            }
        """)
        
        system_layout = QFormLayout(system_group)
        
        # Veri klasörü
        data_dir_label = QLabel(Config.DATA_DIR)
        system_layout.addRow("Veri Klasörü:", data_dir_label)
        
        # Yedek klasörü
        backup_dir_label = QLabel(Config.BACKUP_DIR)
        system_layout.addRow("Yedek Klasörü:", backup_dir_label)
        
        # Log klasörü
        log_dir_label = QLabel(Config.LOG_DIR)
        system_layout.addRow("Log Klasörü:", log_dir_label)
        
        layout.addWidget(system_group)
        
        # Güncelleme ayarları
        update_group = QGroupBox("🔄 Güncelleme Ayarları")
        update_group.setStyleSheet("""
            QGroupBox {
                font-weight: bold;
                border: 2px solid #bdc3c7;
                border-radius: 5px;
                margin-top: 10px;
                padding-top: 10px;
            }
        """)
        
        update_layout = QVBoxLayout(update_group)
        
        # Güncelleme kontrolü
        btn_check_update = QPushButton("🔍 Güncelleme Kontrol Et")
        btn_check_update.setStyleSheet("""
            QPushButton {
                background-color: #3498db;
                color: white;
                border: none;
                padding: 15px;
                border-radius: 5px;
                font-weight: bold;
                font-size: 14px;
            }
            QPushButton:hover {
                background-color: #2980b9;
            }
        """)
        btn_check_update.clicked.connect(self.check_for_updates)
        
        # GitHub repository bilgisi
        repo_info = QLabel(f"Repository: {Config.GITHUB_REPO_URL}")
        repo_info.setStyleSheet("font-size: 12px; color: #7f8c8d; padding: 5px;")
        repo_info.setWordWrap(True)
        
        # Güncelleme durumu
        self.update_status_label = QLabel("Son kontrol: Henüz kontrol edilmedi")
        self.update_status_label.setStyleSheet("font-size: 12px; color: #7f8c8d; padding: 5px;")
        
        update_layout.addWidget(btn_check_update)
        update_layout.addWidget(repo_info)
        update_layout.addWidget(self.update_status_label)
        
        layout.addWidget(update_group)
        

        
        # Alt boşluk
        layout.addStretch()
    
    def backup_data(self):
        """Veri yedekle"""
        try:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            backup_dir = os.path.join(Config.BACKUP_DIR, f"yedek_{timestamp}")
            os.makedirs(backup_dir, exist_ok=True)
            
            for data_type, filename in Config.EXCEL_FILES.items():
                source = os.path.join(Config.DATA_DIR, filename)
                if os.path.exists(source):
                    shutil.copy2(source, backup_dir)
            
            QMessageBox.information(self, "Başarılı", f"Veriler yedeklendi:\n{backup_dir}")
        except Exception as e:
            QMessageBox.critical(self, "Hata", f"Yedekleme hatası: {str(e)}")
    
    def export_to_excel(self):
        """Excel'e aktar"""
        try:
            folder = QFileDialog.getExistingDirectory(self, "Excel dosyalarını kaydet")
            if folder:
                for data_type, filename in Config.EXCEL_FILES.items():
                    source = os.path.join(Config.DATA_DIR, filename)
                    if os.path.exists(source):
                        dest = os.path.join(folder, filename)
                        shutil.copy2(source, dest)
                
                QMessageBox.information(self, "Başarılı", f"Excel dosyaları kaydedildi:\n{folder}")
        except Exception as e:
            QMessageBox.critical(self, "Hata", f"Export hatası: {str(e)}")
    

    
    def download_excel_templates(self):
        """Excel şablonlarını indir"""
        try:
            # Template manager oluştur
            template_manager = ExcelTemplateManager(self.data_manager)
            
            # Dialog'u göster
            dialog = ExcelTemplateDownloadDialog(self, template_manager)
            dialog.exec_()
            
        except Exception as e:
            QMessageBox.critical(self, "Hata", f"Şablon indirme hatası: {str(e)}")
    
    def check_for_updates(self):
        """Güncelleme kontrolü"""
        try:
            # Ana pencereye erişim
            main_window = self.window()
            if hasattr(main_window, 'check_for_updates'):
                # Güncelleme kontrolünü başlat
                main_window.check_for_updates()
                
                # Durum etiketini güncelle
                self.update_status_label.setText(f"Son kontrol: {datetime.now().strftime('%d.%m.%Y %H:%M')}")
                self.update_status_label.setStyleSheet("font-size: 12px; color: #27ae60; padding: 5px;")
                
        except Exception as e:
            QMessageBox.critical(self, "Hata", f"Güncelleme kontrolü hatası: {str(e)}")

class ExcelTemplateDownloadDialog(QDialog):
    """Excel şablon indirme dialog'u"""
    
    def __init__(self, parent=None, template_manager=None):
        super().__init__(parent)
        self.template_manager = template_manager
        self.selected_templates = []
        self.init_ui()
    
    def init_ui(self):
        self.setWindowTitle("Excel Şablon İndirme")
        self.setFixedSize(600, 500)
        self.setWindowFlags(self.windowFlags() & ~Qt.WindowContextHelpButtonHint)
        
        layout = QVBoxLayout()
        
        # Başlık
        title_label = QLabel("📥 Excel Şablonları İndir")
        title_label.setStyleSheet("font-size: 16px; font-weight: bold; margin-bottom: 15px; color: #2196f3;")
        layout.addWidget(title_label)
        
        # Açıklama
        desc_label = QLabel("Veri yüklemesi için kullanabileceğiniz Excel şablonlarını indirin. Bu şablonları doldurup uygulamaya yükleyebilirsiniz.")
        desc_label.setStyleSheet("font-size: 12px; color: #666; margin-bottom: 15px;")
        desc_label.setWordWrap(True)
        layout.addWidget(desc_label)
        
        # Şablon listesi
        self.template_list = QListWidget()
        self.template_list.setStyleSheet("""
            QListWidget {
                border: 2px solid #ddd;
                border-radius: 8px;
                padding: 10px;
                background-color: white;
                font-size: 12px;
            }
            QListWidget::item {
                padding: 15px;
                margin: 3px;
                border-radius: 6px;
                background-color: #f8f9fa;
                border: 1px solid #e9ecef;
            }
            QListWidget::item:hover {
                background-color: #e3f2fd;
                border-color: #2196f3;
            }
            QListWidget::item:selected {
                background-color: #2196f3;
                color: white;
                border-color: #1976d2;
            }
        """)
        
        if self.template_manager:
            for data_type, info in self.template_manager.template_descriptions.items():
                item = QListWidgetItem()
                item.setText(f"📋 {info['title']}\n{info['description']}")
                item.setData(Qt.UserRole, data_type)
                item.setFlags(item.flags() | Qt.ItemIsUserCheckable)
                item.setCheckState(Qt.Unchecked)
                self.template_list.addItem(item)
        
        layout.addWidget(self.template_list)
        
        # Seçim butonları
        selection_layout = QHBoxLayout()
        
        select_all_btn = QPushButton("Tümünü Seç")
        select_all_btn.setStyleSheet("""
            QPushButton {
                background-color: #ff9800;
                color: white;
                border: none;
                padding: 8px 15px;
                border-radius: 5px;
                font-size: 12px;
                font-weight: bold;
            }
            QPushButton:hover {
                background-color: #f57c00;
            }
        """)
        select_all_btn.clicked.connect(self.select_all_templates)
        
        clear_all_btn = QPushButton("Seçimi Temizle")
        clear_all_btn.setStyleSheet("""
            QPushButton {
                background-color: #9e9e9e;
                color: white;
                border: none;
                padding: 8px 15px;
                border-radius: 5px;
                font-size: 12px;
                font-weight: bold;
            }
            QPushButton:hover {
                background-color: #757575;
            }
        """)
        clear_all_btn.clicked.connect(self.clear_all_templates)
        
        selection_layout.addWidget(select_all_btn)
        selection_layout.addWidget(clear_all_btn)
        selection_layout.addStretch()
        
        layout.addLayout(selection_layout)
        
        # Butonlar
        button_layout = QHBoxLayout()
        
        cancel_btn = QPushButton("İptal")
        cancel_btn.setStyleSheet("""
            QPushButton {
                background-color: #f44336;
                color: white;
                border: none;
                padding: 10px 20px;
                border-radius: 6px;
                font-size: 13px;
                font-weight: bold;
            }
            QPushButton:hover {
                background-color: #d32f2f;
            }
        """)
        cancel_btn.clicked.connect(self.reject)
        
        download_btn = QPushButton("Şablonları İndir")
        download_btn.setStyleSheet("""
            QPushButton {
                background-color: #4caf50;
                color: white;
                border: none;
                padding: 10px 20px;
                border-radius: 6px;
                font-size: 13px;
                font-weight: bold;
            }
            QPushButton:hover {
                background-color: #388e3c;
            }
        """)
        download_btn.clicked.connect(self.download_templates)
        
        button_layout.addWidget(cancel_btn)
        button_layout.addWidget(download_btn)
        
        layout.addLayout(button_layout)
        
        self.setLayout(layout)
    
    def select_all_templates(self):
        """Tüm şablonları seç"""
        for i in range(self.template_list.count()):
            item = self.template_list.item(i)
            item.setCheckState(Qt.Checked)
    
    def clear_all_templates(self):
        """Tüm seçimleri temizle"""
        for i in range(self.template_list.count()):
            item = self.template_list.item(i)
            item.setCheckState(Qt.Unchecked)
    
    def download_templates(self):
        """Seçili şablonları indir"""
        selected_templates = []
        for i in range(self.template_list.count()):
            item = self.template_list.item(i)
            if item.checkState() == Qt.Checked:
                selected_templates.append(item.data(Qt.UserRole))
        
        if not selected_templates:
            QMessageBox.warning(self, "Uyarı", "Lütfen en az bir şablon seçin.")
            return
        
        try:
            # Klasör seç
            output_dir = QFileDialog.getExistingDirectory(self, "Şablonları Kaydetmek İçin Klasör Seçin")
            if not output_dir:
                return
            
            # Şablonları oluştur
            created_files = []
            for template_type in selected_templates:
                filename = f"{template_type}_sablonu.xlsx"
                filepath = os.path.join(output_dir, filename)
                self.template_manager.create_template(template_type, filepath)
                created_files.append(filepath)
            
            # Başarı mesajı
            QMessageBox.information(
                self, 
                "Başarılı", 
                f"{len(created_files)} şablon başarıyla indirildi!\n\n"
                f"Konum: {output_dir}\n\n"
                "Bu şablonları doldurup uygulamaya yükleyebilirsiniz."
            )
            
            self.accept()
            
        except Exception as e:
            QMessageBox.critical(self, "Hata", f"Şablon indirme hatası: {str(e)}")






# =============================================================================
# SPLASH SCREEN SİSTEMİ
# =============================================================================

class SplashScreen(QSplashScreen):
    """Başlangıç ekranı"""
    
    def __init__(self, pixmap=None):
        if pixmap is None:
            # Varsayılan splash screen oluştur
            pixmap = self.create_default_splash()
        
        super().__init__(pixmap)
        self.setWindowFlags(Qt.WindowStaysOnTopHint | Qt.FramelessWindowHint)
        
        # Başlangıç mesajı
        self.showMessage("Araç Filo Yönetim Sistemi Başlatılıyor...", 
                        Qt.AlignmentFlag.AlignBottom | Qt.AlignmentFlag.AlignCenter, 
                        QColor(255, 255, 255))
    
    def create_default_splash(self):
        """Varsayılan splash screen oluştur"""
        # 800x600 boyutunda splash screen
        pixmap = QPixmap(800, 600)
        pixmap.fill(QColor(37, 99, 235))  # Mavi arka plan
        
        painter = QPainter(pixmap)
        
        # Gradient arka plan
        gradient = QLinearGradient(0, 0, 0, 600)
        gradient.setColorAt(0, QColor(37, 99, 235))  # Mavi
        gradient.setColorAt(1, QColor(59, 130, 246))  # Açık mavi
        painter.fillRect(0, 0, 800, 600, gradient)
        
        # Başlık
        title_font = QFont("Arial", 32, QFont.Weight.Bold)
        painter.setFont(title_font)
        painter.setPen(QColor(255, 255, 255))
        painter.drawText(400, 200, "🚗 Araç Filo Yönetim Sistemi")
        
        # Alt başlık
        subtitle_font = QFont("Arial", 16)
        painter.setFont(subtitle_font)
        painter.setPen(QColor(255, 255, 255))
        painter.drawText(400, 250, "v22.07.24.01")
        
        # Araç ikonları (basit çizimler)
        self.draw_car_icons(painter)
        
        # Alt bilgi
        info_font = QFont("Arial", 12)
        painter.setFont(info_font)
        painter.setPen(QColor(255, 255, 255))
        painter.drawText(400, 550, "Ertuğrul Yazılım © 2024")
        
        painter.end()
        return pixmap
    
    def draw_car_icons(self, painter):
        """Basit araç ikonları çiz"""
        # Araç gövdesi (basit dikdörtgenler)
        cars = [
            (150, 350, 80, 40),  # Sol araç
            (350, 350, 80, 40),  # Orta araç
            (550, 350, 80, 40),  # Sağ araç
        ]
        
        for x, y, w, h in cars:
            # Araç gövdesi
            painter.fillRect(x, y, w, h, QColor(255, 255, 255))
            # Tekerlekler
            painter.fillRect(x + 10, y + h, 15, 8, QColor(50, 50, 50))
            painter.fillRect(x + w - 25, y + h, 15, 8, QColor(50, 50, 50))
            # Ön cam
            painter.fillRect(x + 5, y + 5, 20, 15, QColor(200, 220, 255))

class SplashThread(QThread):
    """Splash screen için ayrı thread"""
    finished = pyqtSignal()
    
    def __init__(self, duration=5):
        super().__init__()
        self.duration = duration
    
    def run(self):
        time.sleep(self.duration)
        self.finished.emit()
    
    def stop(self):
        """Thread'i güvenli şekilde durdur"""
        self.quit()
        self.wait()

def show_splash_screen(app, main_window_class, splash_image_path=None):
    """Başlangıç ekranını göster"""
    
    # PyInstaller ile paketlenmiş dosya yolu
    if getattr(sys, 'frozen', False):
        # EXE dosyası olarak çalışıyor
        base_path = sys._MEIPASS
    else:
        # Python script olarak çalışıyor
        base_path = os.path.dirname(os.path.abspath(__file__))
    
    # Varsayılan splash screen dosyası
    if splash_image_path is None:
        splash_image_path = os.path.join(base_path, "baslangic.jpg")
    
    print(f"🔍 Splash screen dosyası aranıyor: {splash_image_path}")
    print(f"📁 Dosya mevcut mu: {os.path.exists(splash_image_path)}")
    
    # Splash screen oluştur
    if splash_image_path and os.path.exists(splash_image_path):
        print("✅ Splash screen dosyası bulundu, yükleniyor...")
        pixmap = QPixmap(splash_image_path)
        if not pixmap.isNull():
            print(f"✅ Görsel yüklendi: {pixmap.width()}x{pixmap.height()}")
            
            # Görseli %50 küçült
            original_width = pixmap.width()
            original_height = pixmap.height()
            new_width = int(original_width * 0.5)
            new_height = int(original_height * 0.5)
            
            scaled_pixmap = pixmap.scaled(new_width, new_height, Qt.KeepAspectRatio, Qt.SmoothTransformation)
            print(f"📏 Görsel küçültüldü: {new_width}x{new_height}")
            
            splash = SplashScreen(scaled_pixmap)
        else:
            print("❌ Görsel yüklenemedi, varsayılan kullanılıyor")
            splash = SplashScreen()
    else:
        print("❌ Splash screen dosyası bulunamadı, varsayılan kullanılıyor")
        splash = SplashScreen()
    
    # Splash screen'i göster
    print("🚀 Splash screen gösteriliyor...")
    splash.show()
    app.processEvents()
    
    # Ana pencereyi oluştur (arka planda)
    print("🏗️ Ana pencere hazırlanıyor...")
    main_window = main_window_class()
    
    # 1 saniye bekle
    print("⏱️ 1 saniye bekleniyor...")
    splash_thread = SplashThread(1)
    splash_thread.finished.connect(lambda: finish_splash(splash, main_window, splash_thread))
    splash_thread.start()
    
    return main_window

def finish_splash(splash, main_window, splash_thread):
    """Splash screen'i kapat ve ana pencereyi göster"""
    # Thread'i güvenli şekilde durdur
    splash_thread.stop()
    
    # Splash screen'i kapat
    splash.finish(main_window)
    main_window.show()
    main_window.raise_()
    main_window.activateWindow()

# =============================================================================
# GÜNCELLEME SİSTEMİ TEST FONKSİYONU
# =============================================================================

def test_update_system():
    """Güncelleme sistemini test et"""
    print("🔄 Güncelleme sistemi test ediliyor...")
    
    # Test versiyonları
    test_versions = [
        ("22.07.24.01", "22.07.24.02", True),   # Yeni sürüm mevcut
        ("22.07.24.02", "22.07.24.01", False),  # Eski sürüm
        ("22.07.24.01", "22.07.24.01", False),  # Aynı sürüm
        ("22.07.24.01", "22.08.01.01", True),   # Büyük güncelleme
        ("22.08.01.01", "22.07.24.01", False),  # Geri dönüş
        ("22.07.24.01", "22.07.24.01-beta", True),  # Beta sürüm
        ("22.07.24.01-beta", "22.07.24.01", True),  # Stable sürüm
    ]
    
    checker = UpdateChecker()
    
    for current, latest, expected in test_versions:
        result = checker.compare_versions(latest, current)
        status = "✅" if result == expected else "❌"
        print(f"{status} {current} -> {latest}: {result} (beklenen: {expected})")
    
    print("✅ Güncelleme sistemi test tamamlandı!")

# =============================================================================
# ANA UYGULAMA
# =============================================================================

if __name__ == "__main__":
    # Test modu kontrolü
    if len(sys.argv) > 1 and sys.argv[1] == "--test-update":
        test_update_system()
        sys.exit(0)
    
    app = QApplication(sys.argv)
    app.setApplicationName(Config.APP_NAME)
    app.setApplicationVersion(Config.VERSION)

    def exception_hook(exctype, value, tb):
        error_msg = f"Hata: {exctype.__name__}: {value}"
        logger.critical(f"Beklenmeyen hata: {error_msg}", exc=value)
        
        # Hata detaylarını log dosyasına yaz
        try:
            with open(os.path.join(Config.LOG_DIR, "error_traceback.log"), 'a', encoding='utf-8') as f:
                f.write(f"\n{'='*50}\n")
                f.write(f"Tarih: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n")
                f.write(f"Hata: {exctype.__name__}: {value}\n")
                f.write("Traceback:\n")
                traceback.print_tb(tb, file=f)
                f.write(f"{'='*50}\n")
        except Exception as e:
            print(f"Traceback log hatası: {e}")
        
        # Kullanıcıya hata mesajı göster
        if QApplication.activeWindow():
            QMessageBox.critical(QApplication.activeWindow(), "Kritik Hata", 
                               f"Beklenmeyen bir hata oluştu:\n{error_msg}\n\n"
                               f"Detaylar log dosyasına kaydedildi.")
    sys.excepthook = exception_hook

    # Splash screen ile başlat
    window = show_splash_screen(app, MainWindow)
    
    sys.exit(app.exec_())